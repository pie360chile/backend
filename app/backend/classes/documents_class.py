import uuid
import re
import os
import platform
import subprocess
import tempfile
from pathlib import Path
import fitz  # PyMuPDF
from docx import Document
from typing import Optional, List, Dict, Any
from sqlalchemy.orm import Session
from datetime import datetime, date
from app.backend.db.models import DocumentModel, BirthCertificateDocumentModel, HealthEvaluationModel, FolderModel

# pypdf imports (opcional, solo si está instalado)
try:
    from pypdf import PdfReader, PdfWriter
    from pypdf.generic import NameObject, NumberObject, BooleanObject, TextStringObject
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False
    NameObject = None
    NumberObject = None
    BooleanObject = None
    TextStringObject = None

# ReportLab imports (opcional, solo si está instalado)
try:
    from reportlab.lib.pagesizes import A4, letter
    from reportlab.lib.units import inch, cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image, KeepTogether
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY, TA_RIGHT
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# Win32 imports (opcional, solo Windows)
try:
    import win32com.client
    WIN32_AVAILABLE = True
except ImportError:
    WIN32_AVAILABLE = False

# docx imports adicionales
try:
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.ns import qn
    DOCX_EXTRA_AVAILABLE = True
except ImportError:
    DOCX_EXTRA_AVAILABLE = False


class DocumentsClass:
    def __init__(self, db: Session):
        self.db = db

    @staticmethod
    def generate_document_pdf(
        document_id: int,
        document_data: Dict[str, Any],
        db: Optional[Session] = None,
        template_path: Optional[str] = None,
        tag_replacements: Optional[Dict[str, str]] = None,
        output_directory: str = "files/system/students"
    ) -> Dict[str, Any]:
        """
        Función general para generar PDFs de cualquier documento.
        Dependiendo de si hay template_path:
        - Si hay template_path: Sustituye datos en el template PDF usando tag_replacements
        - Si no hay template_path: Crea el PDF desde cero usando ReportLab
        
        Args:
            document_id: ID del documento
            document_data: Diccionario con los datos del documento
            db: Sesión de base de datos (opcional)
            template_path: Ruta del archivo PDF template (opcional, si es None se crea desde cero)
            tag_replacements: Diccionario con mapeo de etiquetas a valores (opcional, se genera automáticamente si es None)
            output_directory: Directorio donde se guardará el archivo generado
        
        Returns:
            dict: Diccionario con status, message, filename y file_path
        """
        try:
            # Si es documento 19 (Estado de Avance PAI - progress_status_individual_support), usar función específica
            if document_id == 19:
                return DocumentsClass._generate_progress_status_individual_support_from_scratch(
                    document_id=document_id,
                    ps_data=document_data,
                    db=db,
                    output_directory=output_directory
                )
            
            # Si es documento 18 (Estado de Avance - progress_status_students), usar función específica
            if document_id == 18:
                return DocumentsClass._generate_progress_status_from_scratch(
                    document_id=document_id,
                    progress_status_data=document_data,
                    db=db,
                    output_directory=output_directory
                )
            
            # Si es documento 20 (CESP - Plan de Acompañamiento Emocional y Conductual), usar función específica
            if document_id == 20:
                return DocumentsClass._generate_cesp_from_scratch(
                    document_id=document_id,
                    cesp_data=document_data,
                    db=db,
                    output_directory=output_directory
                )
            
            # Si es documento 22 (Plan de Apoyo Individual), usar función específica
            if document_id == 22:
                return DocumentsClass._generate_individual_support_plan_from_scratch(
                    document_id=document_id,
                    isp_data=document_data,
                    db=db,
                    output_directory=output_directory
                )
            
            # Si es documento 8 (Informe Fonoaudiológico), usar función específica
            if document_id == 8:
                return DocumentsClass._generate_fonoaudiological_report_from_scratch(
                    document_id=document_id,
                    report_data=document_data,
                    db=db,
                    output_directory=output_directory
                )
            
            # Si es documento 9 (Informe fonoaudiológico IDTEL), usar función específica
            if document_id == 9:
                return DocumentsClass._generate_idtel_report_from_scratch(
                    document_id=document_id,
                    report_data=document_data,
                    db=db,
                    output_directory=output_directory
                )

            # Informe de evaluación psicomotriz (detección por nombre en catálogo `documents`)
            if db is not None:
                _inf = (
                    db.query(DocumentModel)
                    .filter(
                        DocumentModel.id == document_id,
                        DocumentModel.deleted_date.is_(None),
                    )
                    .first()
                )
                if _inf and (_inf.document or "").strip() == "Informe de evaluación psicomotriz":
                    return DocumentsClass._generate_psychomotor_evaluation_report_from_scratch(
                        document_id=document_id,
                        report_data=document_data,
                        db=db,
                        output_directory=output_directory,
                    )

            # Si es documento 23 (Certificado de egreso PIE), usar función específica
            if document_id == 23:
                return DocumentsClass._generate_school_integration_exit_certificate_from_scratch(
                    document_id=document_id,
                    cert_data=document_data,
                    db=db,
                    output_directory=output_directory
                )
            
            # Si es documento 25 (Certificado Ley TEA - acompañamiento apoderado), usar función específica
            if document_id == 25:
                return DocumentsClass._generate_tea_certificate_from_scratch(
                    document_id=document_id,
                    cert_data=document_data,
                    db=db,
                    output_directory=output_directory
                )
            
            # Si es documento 29 (Test de Conners - Profesor Abreviado + Cuestionario de Conducta)
            if document_id == 29:
                return DocumentsClass._generate_conners_teacher_from_scratch(
                    document_id=document_id,
                    conners_data=document_data,
                    db=db,
                    output_directory=output_directory
                )
            
            # Si es documento 31 (Pauta de evaluación pedagógica - Docente de aula - 1º Básico)
            if document_id == 31:
                return DocumentsClass._generate_pedagogical_evaluation_classroom_first_grade_from_scratch(
                    document_id=document_id,
                    doc_data=document_data,
                    db=db,
                    output_directory=output_directory
                )
            
            # Si es documento 32 (Pauta de evaluación pedagógica - Docente de aula - 2º Básico)
            if document_id == 32:
                return DocumentsClass._generate_pedagogical_evaluation_classroom_second_grade_from_scratch(
                    document_id=document_id,
                    doc_data=document_data,
                    db=db,
                    output_directory=output_directory
                )
            
            # Si es documento 33 (Pauta de evaluación pedagógica - Docente de aula - 3º Básico)
            if document_id == 33:
                return DocumentsClass._generate_pedagogical_evaluation_classroom_third_grade_from_scratch(
                    document_id=document_id,
                    doc_data=document_data,
                    db=db,
                    output_directory=output_directory
                )
            
            if document_id == 34:
                return DocumentsClass._generate_pedagogical_evaluation_classroom_fourth_grade_from_scratch(
                    document_id=document_id,
                    doc_data=document_data,
                    db=db,
                    output_directory=output_directory
                )
            if document_id == 35:
                return DocumentsClass._generate_pedagogical_evaluation_classroom_fifth_grade_from_scratch(
                    document_id=document_id,
                    doc_data=document_data,
                    db=db,
                    output_directory=output_directory
                )
            if document_id == 36:
                return DocumentsClass._generate_pedagogical_evaluation_classroom_sixth_grade_from_scratch(
                    document_id=document_id,
                    doc_data=document_data,
                    db=db,
                    output_directory=output_directory
                )
            if document_id == 37:
                return DocumentsClass._generate_pedagogical_evaluation_classroom_seventh_grade_from_scratch(
                    document_id=document_id,
                    doc_data=document_data,
                    db=db,
                    output_directory=output_directory
                )
            if document_id == 38:
                return DocumentsClass._generate_pedagogical_evaluation_classroom_eighth_grade_from_scratch(
                    document_id=document_id,
                    doc_data=document_data,
                    db=db,
                    output_directory=output_directory
                )
            if document_id == 39:
                return DocumentsClass._generate_pedagogical_evaluation_classroom_first_grade_secondary_from_scratch(
                    document_id=document_id,
                    doc_data=document_data,
                    db=db,
                    output_directory=output_directory
                )
            if document_id == 40:
                return DocumentsClass._generate_pedagogical_evaluation_classroom_second_grade_secondary_from_scratch(
                    document_id=document_id,
                    doc_data=document_data,
                    db=db,
                    output_directory=output_directory
                )
            
            # Si hay template_path, usar método de sustitución
            if template_path:
                # Si no se proporcionan tag_replacements, generar automáticamente desde document_data
                if tag_replacements is None:
                    tag_replacements = DocumentsClass._generate_tag_replacements(document_data)
                
                return DocumentsClass._generate_pdf_from_template(
                    template_path=template_path,
                    tag_replacements=tag_replacements,
                    document_id=document_id,
                    document_data=document_data,
                    output_directory=output_directory
                )
            else:
                # Si no hay template, crear desde cero usando ReportLab
                return DocumentsClass._generate_pdf_from_scratch(
                    document_id=document_id,
                    document_data=document_data,
                    db=db,
                    output_directory=output_directory
                )
        except Exception as e:
            return {
                "status": "error",
                "message": f"Error generando PDF: {str(e)}",
                "filename": None,
                "file_path": None
            }
    
    @staticmethod
    def _generate_tag_replacements(document_data: Dict[str, Any]) -> Dict[str, str]:
        """
        Genera automáticamente el mapeo de etiquetas a valores desde document_data.
        Convierte todas las claves del diccionario a formato [KEY] y sus valores a string.
        """
        tag_replacements = {}
        
        # Función auxiliar para formatear fechas
        def format_date(value: Any) -> str:
            if not value:
                return ""
            try:
                if isinstance(value, str):
                    date_obj = datetime.strptime(value, "%Y-%m-%d").date()
                    return date_obj.strftime("%d/%m/%Y")
                elif hasattr(value, 'strftime'):
                    return value.strftime("%d/%m/%Y")
                return ""
            except:
                return str(value) if value else ""
        
        # Generar etiquetas para cada campo
        for key, value in document_data.items():
            if value is not None:
                # Si es una fecha, formatearla
                if 'date' in key.lower() or 'fecha' in key.lower():
                    tag_replacements[f"[{key.upper()}]"] = format_date(value)
                else:
                    tag_replacements[f"[{key.upper()}]"] = str(value) if value else ""
        
        return tag_replacements
    
    @staticmethod
    def _generate_pdf_from_template(
        template_path: str,
        tag_replacements: Dict[str, str],
        document_id: int,
        document_data: Dict[str, Any],
        output_directory: str = "files/system/students"
    ) -> Dict[str, Any]:
        """
        Genera un PDF sustituyendo datos en un template PDF existente.
        """
        try:
            template_file = Path(template_path)
            
            if not template_file.exists():
                return {
                    "status": "error",
                    "message": "Template PDF no encontrado",
                    "filename": None,
                    "file_path": None
                }
            
            # Validar que sea PDF
            if template_file.suffix.lower() != '.pdf':
                return {
                    "status": "error",
                    "message": f"El template debe ser un archivo PDF, se encontró: {template_file.suffix}",
                    "filename": None,
                    "file_path": None
                }
            
            # Crear nombre único para el archivo generado
            student_name = document_data.get("student_fullname", document_data.get("student_name", "documento")).replace(" ", "_")
            unique_filename = f"documento_{document_id}_{student_name}_{uuid.uuid4().hex[:8]}.pdf"
            output_file = Path(output_directory) / unique_filename
            
            # Asegurar que el directorio de salida existe
            output_file.parent.mkdir(parents=True, exist_ok=True)
            
            # Abrir el archivo PDF
            pdf_document = fitz.open(template_file)
            
            # Procesar cada página del PDF
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                
                # Obtener todo el texto de la página en formato dict para búsqueda más precisa
                page_dict = None
                try:
                    page_dict = page.get_text("dict")
                except:
                    pass
                
                # Reemplazar cada tag
                for tag, value in tag_replacements.items():
                    text_instances = []
                    
                    # Intentar diferentes variaciones del tag
                    tag_variations = [
                        tag,  # Tag original con corchetes
                        tag.upper(),  # Mayúsculas
                        tag.lower(),  # Minúsculas
                    ]
                    
                    # Buscar cada variación usando search_for
                    for tag_variant in tag_variations:
                        try:
                            found = page.search_for(tag_variant, flags=fitz.TEXT_DEHYPHENATE)
                            if found:
                                text_instances.extend(found)
                                break
                        except:
                            try:
                                found = page.search_for(tag_variant)
                                if found:
                                    text_instances.extend(found)
                                    break
                            except:
                                pass
                    
                    # Si aún no se encuentra, buscar usando get_text("dict")
                    if not text_instances and page_dict:
                        try:
                            for block in page_dict.get("blocks", []):
                                if "lines" in block:
                                    for line in block["lines"]:
                                        full_line_text = ""
                                        line_spans = []
                                        for span in line.get("spans", []):
                                            span_text = span.get("text", "")
                                            full_line_text += span_text
                                            line_spans.append((span, span_text))
                                        
                                        full_line_text_upper = full_line_text.upper()
                                        for tag_variant in tag_variations:
                                            tag_variant_upper = tag_variant.upper()
                                            if tag_variant_upper in full_line_text_upper:
                                                matching_spans = []
                                                for span, span_text in line_spans:
                                                    span_text_upper = span_text.upper()
                                                    if tag_variant_upper in span_text_upper:
                                                        matching_spans.append(span)
                                                
                                                if matching_spans:
                                                    x0 = min(s.get("bbox", [0, 0, 0, 0])[0] for s in matching_spans)
                                                    y0 = min(s.get("bbox", [0, 0, 0, 0])[1] for s in matching_spans)
                                                    x1 = max(s.get("bbox", [0, 0, 0, 0])[2] for s in matching_spans)
                                                    y1 = max(s.get("bbox", [0, 0, 0, 0])[3] for s in matching_spans)
                                                    rect = fitz.Rect(x0 - 2, y0 - 2, x1 + 2, y1 + 2)
                                                    text_instances.append(rect)
                                                break
                        except:
                            pass
                    
                    # Reemplazar todas las instancias encontradas
                    if text_instances:
                        seen_rects = set()
                        unique_instances = []
                        page_width = page.rect.width
                        page_height = page.rect.height
                        
                        for rect in text_instances:
                            if isinstance(rect, (list, tuple)) and len(rect) == 4:
                                rect = fitz.Rect(rect[0], rect[1], rect[2], rect[3])
                            
                            if (rect.x0 >= 0 and rect.y0 >= 0 and 
                                rect.x1 <= page_width and rect.y1 <= page_height and
                                rect.x1 > rect.x0 and rect.y1 > rect.y0):
                                
                                rect_key = (round(rect.x0, 1), round(rect.y0, 1), round(rect.x1, 1), round(rect.y1, 1))
                                if rect_key not in seen_rects:
                                    seen_rects.add(rect_key)
                                    unique_instances.append(rect)
                        
                        for rect in unique_instances:
                            if rect.x1 > rect.x0 and rect.y1 > rect.y0:
                                page.add_redact_annot(rect, fill=(1, 1, 1))
                                page.apply_redactions()
                                
                                fontsize = 10
                                try:
                                    text_rect = fitz.Rect(rect.x0, rect.y0, rect.x1, rect.y1)
                                    rc = page.insert_textbox(
                                        text_rect,
                                        value,
                                        fontsize=fontsize,
                                        color=(0, 0, 0),
                                        align=0,
                                        overlay=True
                                    )
                                    
                                    if rc < 0:
                                        page.insert_text(
                                            (rect.x0 + 2, rect.y0 + fontsize),
                                            value,
                                            fontsize=fontsize,
                                            color=(0, 0, 0),
                                            overlay=True
                                        )
                                except:
                                    try:
                                        page.insert_text(
                                            (rect.x0 + 2, rect.y0 + fontsize),
                                            value,
                                            fontsize=fontsize,
                                            color=(0, 0, 0),
                                            overlay=True
                                        )
                                    except:
                                        pass
            
            # Guardar el PDF modificado
            pdf_document.save(output_file)
            pdf_document.close()
            
            return {
                "status": "success",
                "message": "PDF generado exitosamente desde template",
                "filename": unique_filename,
                "file_path": str(output_file)
            }
            
        except Exception as e:
            return {
                "status": "error",
                "message": f"Error generando PDF desde template: {str(e)}",
                "filename": None,
                "file_path": None
            }
    
    @staticmethod
    def _generate_pdf_from_scratch(
        document_id: int,
        document_data: Dict[str, Any],
        db: Optional[Session] = None,
        output_directory: str = "files/system/students"
    ) -> Dict[str, Any]:
        """
        Genera un PDF desde cero usando ReportLab.
        """
        try:
            if not REPORTLAB_AVAILABLE:
                return {
                    "status": "error",
                    "message": "ReportLab no está instalado. Instala con: pip install reportlab",
                    "filename": None,
                    "file_path": None
                }
            
            # Crear nombre único para el archivo generado
            student_name = document_data.get("student_fullname", document_data.get("student_name", "documento")).replace(" ", "_")
            unique_filename = f"documento_{document_id}_{student_name}_{uuid.uuid4().hex[:8]}.pdf"
            output_file = Path(output_directory) / unique_filename
            
            # Asegurar que el directorio de salida existe
            output_file.parent.mkdir(parents=True, exist_ok=True)
            
            # Crear el documento PDF
            doc = SimpleDocTemplate(
                str(output_file),
                pagesize=A4,
                rightMargin=2*cm,
                leftMargin=2*cm,
                topMargin=2*cm,
                bottomMargin=2*cm
            )
            
            elements = []
            styles = getSampleStyleSheet()
            
            # Estilos personalizados
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                textColor=colors.HexColor('#000000'),
                spaceAfter=20,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            
            subtitle_style = ParagraphStyle(
                'Subtitle',
                parent=styles['Normal'],
                fontSize=12,
                textColor=colors.HexColor('#000000'),
                spaceAfter=10,
                fontName='Helvetica-Bold'
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#000000'),
                alignment=TA_JUSTIFY,
                leading=14
            )
            
            # Función auxiliar para obtener valor
            def get_value(key: str, default: str = "") -> str:
                value = document_data.get(key)
                if value is None:
                    return default
                return str(value) if value else default
            
            # Función auxiliar para formatear fechas
            def format_date(date_str: Optional[str]) -> str:
                if not date_str:
                    return ""
                try:
                    if isinstance(date_str, str):
                        date_obj = datetime.strptime(date_str, "%Y-%m-%d").date()
                        return date_obj.strftime("%d/%m/%Y")
                    return ""
                except:
                    return str(date_str) if date_str else ""
            
            # Título genérico
            document_title = document_data.get("document_title", f"DOCUMENTO {document_id}")
            elements.append(Paragraph(document_title, title_style))
            elements.append(Spacer(1, 0.3*inch))
            
            # Crear tabla con todos los datos del documento
            table_data = []
            for key, value in document_data.items():
                if value and key not in ['document_title']:
                    # Formatear la clave para mostrar
                    display_key = key.replace('_', ' ').title()
                    display_value = format_date(value) if 'date' in key.lower() or 'fecha' in key.lower() else str(value)
                    table_data.append([display_key + ":", display_value])
            
            if table_data:
                data_table = Table(table_data, colWidths=[5*cm, 11*cm])
                data_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
                    ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                elements.append(data_table)
            
            # Construir PDF
            doc.build(elements)
            
            return {
                "status": "success",
                "message": "PDF generado exitosamente desde cero",
                "filename": unique_filename,
                "file_path": str(output_file)
            }
            
        except ImportError:
            return {
                "status": "error",
                "message": "ReportLab no está instalado. Instala con: pip install reportlab",
                "filename": None,
                "file_path": None
            }
        except Exception as e:
            return {
                "status": "error",
                "message": f"Error generando PDF desde cero: {str(e)}",
                "filename": None,
                "file_path": None
            }

    @staticmethod
    def parent_authorization(
        original_file_path: str, 
        student_fullname: str,
        parent_relation: str = "",
        student_course_school: str = "",
        parent_fullname: str = "",
        parent_rut: str = "",
        city: str = "",
        day: str = "",
        month: str = "",
        year: str = "",  # Año (últimos 2 dígitos)
        parent_signature: str = "",
        yes_marker: str = "●",  # Punto negro para [Y]
        no_marker: str = "",  # En blanco para [N]
        output_directory: str = "files/system/students"
    ):
        """
        Procesa un documento PDF de autorización de padres, reemplazando los tags con los valores correspondientes.
        
        Args:
            original_file_path: Ruta del archivo PDF original
            student_fullname: Nombre completo del estudiante
            parent_relation: Relación con el estudiante (ej: "Padre", "Madre", "Tutor")
            student_course_school: Curso y establecimiento (ej: "3ro A - Escuela San Juan")
            parent_fullname: Nombre completo del padre/madre/tutor
            output_directory: Directorio donde se guardará el archivo procesado
        
        Returns:
            dict: Diccionario con status, message, filename y file_path
        """
        try:
            original_file = Path(original_file_path)
            
            if not original_file.exists():
                return {
                    "status": "error",
                    "message": "Documento original no encontrado",
                    "filename": None,
                    "file_path": None
                }
            
            # Crear una copia del documento con un nombre único
            unique_filename = f"autorizacion_{student_fullname.replace(' ', '_')}_{uuid.uuid4().hex[:8]}.pdf"
            output_file = Path(output_directory) / unique_filename
            
            # Asegurar que el directorio de salida existe
            output_file.parent.mkdir(parents=True, exist_ok=True)
            
            # Abrir el archivo PDF
            pdf_document = fitz.open(original_file)
            
            # Mapeo de tags a valores (si está vacío, usar string vacío en lugar de saltar)
            tag_replacements = {
                "[STUDENT_FULLNAME]": student_fullname or "",
                "[PARENT_RELATION]": parent_relation or "",
                "[STUDENT_COURSE_SCHOOL_LOCATION]": student_course_school or "",
                "[PARENT_FULLNAME]": parent_fullname or "",
                "[PARENT_RUT]": parent_rut or "",
                "[PARENT_SIGNATURE]": "",  # Firma del padre/madre (se deja en blanco)
                "[CITY]": city or "",
                "[DAY]": day or "",
                "[MONTH]": month or "",
                "[YEAR]": year or "",  # Año (últimos 2 dígitos)
                "[Y]": yes_marker or "●",  # Punto negro para [Y]
                "[N]": no_marker or ""  # En blanco para [N]
            }
            
            found_and_modified = False
            
            # Procesar cada página del PDF
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                
                # Obtener todo el texto de la página en formato dict para búsqueda más precisa
                page_dict = None
                try:
                    page_dict = page.get_text("dict")
                except:
                    pass
                
                # Reemplazar cada tag
                for tag, value in tag_replacements.items():
                    # Siempre reemplazar, incluso si el valor está vacío (dejará el campo vacío)
                    text_instances = []
                    
                    # Intentar diferentes variaciones del tag
                    # Para [Y] y [N] solo buscar con corchetes para evitar reemplazos incorrectos
                    if tag == "[Y]" or tag == "[N]":
                        tag_variations = [
                            tag,  # Tag original con corchetes
                            tag.upper(),  # Mayúsculas
                            tag.lower(),  # Minúsculas
                        ]
                    else:
                        tag_variations = [
                            tag,  # Tag original con corchetes
                            tag.replace("[", "").replace("]", ""),  # Sin corchetes
                            tag.upper(),  # Mayúsculas
                            tag.lower(),  # Minúsculas
                            tag.replace("[", "").replace("]", "").upper(),  # Sin corchetes, mayúsculas
                            tag.replace("[", "").replace("]", "").lower(),  # Sin corchetes, minúsculas
                        ]
                    
                    # Buscar cada variación usando search_for
                    for tag_variant in tag_variations:
                        try:
                            # Buscar con flags de case insensitive y dehyphenate
                            found = page.search_for(tag_variant, flags=fitz.TEXT_DEHYPHENATE)
                            if found:
                                text_instances.extend(found)
                                break  # Si se encuentra, no buscar más variaciones con search_for
                        except:
                            try:
                                # Si falla con flags, intentar sin flags
                                found = page.search_for(tag_variant)
                                if found:
                                    text_instances.extend(found)
                                    break
                            except:
                                pass
                    
                    # Si aún no se encuentra, buscar usando get_text("dict") para posición exacta
                    # Este método es más preciso cuando el tag está dividido en múltiples partes
                    if not text_instances and page_dict:
                        try:
                            # Buscar en bloques de texto
                            for block in page_dict.get("blocks", []):
                                if "lines" in block:
                                    for line in block["lines"]:
                                        full_line_text = ""
                                        line_spans = []
                                        for span in line.get("spans", []):
                                            span_text = span.get("text", "")
                                            full_line_text += span_text
                                            line_spans.append((span, span_text))
                                        
                                        # Buscar tag en el texto completo de la línea (case insensitive)
                                        full_line_text_upper = full_line_text.upper()
                                        for tag_variant in tag_variations:
                                            tag_variant_upper = tag_variant.upper()
                                            # Buscar con case insensitive
                                            if tag_variant_upper in full_line_text_upper:
                                                # Encontrar qué span(s) contiene(n) el tag
                                                # Si el tag está en múltiples spans, calcular el rectángulo combinado
                                                matching_spans = []
                                                for span, span_text in line_spans:
                                                    span_text_upper = span_text.upper()
                                                    if tag_variant_upper in span_text_upper:
                                                        matching_spans.append(span)
                                                
                                                if matching_spans:
                                                    # Calcular el rectángulo que cubre todos los spans que contienen el tag
                                                    x0 = min(s.get("bbox", [0, 0, 0, 0])[0] for s in matching_spans)
                                                    y0 = min(s.get("bbox", [0, 0, 0, 0])[1] for s in matching_spans)
                                                    x1 = max(s.get("bbox", [0, 0, 0, 0])[2] for s in matching_spans)
                                                    y1 = max(s.get("bbox", [0, 0, 0, 0])[3] for s in matching_spans)
                                                    rect = fitz.Rect(x0 - 2, y0 - 2, x1 + 2, y1 + 2)
                                                    text_instances.append(rect)
                                                break
                                        
                                        # Solo buscar sin espacios si el tag es lo suficientemente largo para evitar falsos positivos
                                        # No buscar sin espacios para tags cortos como [Y] o [N]
                                        if len(tag.replace("[", "").replace("]", "")) > 3:
                                            full_line_clean = full_line_text.replace(" ", "").replace("_", "").replace("-", "").upper()
                                            for tag_variant in tag_variations:
                                                tag_clean = tag_variant.replace("[", "").replace("]", "").replace(" ", "").replace("_", "").replace("-", "").upper()
                                                if len(tag_clean) > 3 and tag_clean in full_line_clean:
                                                    # Encontrar spans que contengan partes del tag
                                                    matching_spans = []
                                                    for span, span_text in line_spans:
                                                        span_clean = span_text.replace(" ", "").replace("_", "").replace("-", "").upper()
                                                        if tag_clean in span_clean:
                                                            matching_spans.append(span)
                                                    
                                                    if matching_spans:
                                                        x0 = min(s.get("bbox", [0, 0, 0, 0])[0] for s in matching_spans)
                                                        y0 = min(s.get("bbox", [0, 0, 0, 0])[1] for s in matching_spans)
                                                        x1 = max(s.get("bbox", [0, 0, 0, 0])[2] for s in matching_spans)
                                                        y1 = max(s.get("bbox", [0, 0, 0, 0])[3] for s in matching_spans)
                                                        rect = fitz.Rect(x0 - 2, y0 - 2, x1 + 2, y1 + 2)
                                                        text_instances.append(rect)
                                                    break
                        except Exception as e:
                            pass
                    
                    # Reemplazar todas las instancias encontradas
                    if text_instances:
                        # Eliminar duplicados manteniendo el orden
                        seen_rects = set()
                        unique_instances = []
                        page_width = page.rect.width
                        page_height = page.rect.height
                        
                        for rect in text_instances:
                            # Validar que el rectángulo esté dentro de los límites de la página
                            if (rect.x0 >= 0 and rect.y0 >= 0 and 
                                rect.x1 <= page_width and rect.y1 <= page_height and
                                rect.x1 > rect.x0 and rect.y1 > rect.y0):
                                
                                # Validar que el rectángulo no sea demasiado grande (más del 50% de la página)
                                rect_width = rect.x1 - rect.x0
                                rect_height = rect.y1 - rect.y0
                                if rect_width < page_width * 0.5 and rect_height < page_height * 0.5:
                                    rect_key = (round(rect.x0, 2), round(rect.y0, 2), round(rect.x1, 2), round(rect.y1, 2))
                                    if rect_key not in seen_rects:
                                        seen_rects.add(rect_key)
                                        unique_instances.append(rect)
                        
                        # Aplicar redacciones primero (cubrir tags originales) solo si hay instancias válidas
                        if unique_instances:
                            for rect in unique_instances:
                                try:
                                    page.add_redact_annot(rect, fill=(1, 1, 1))  # Blanco
                                except:
                                    pass
                            
                            # Aplicar todas las redacciones de una vez
                            try:
                                page.apply_redactions()
                            except:
                                pass
                        
                        # Luego insertar los nuevos valores
                        for rect in unique_instances:
                            try:
                                # Insertar el nuevo texto en la misma posición
                                text_rect = fitz.Rect(rect.x0, rect.y0, rect.x1, rect.y1)
                                
                                # Si el valor es el punto negro [Y], usar tamaño de fuente más grande
                                fontsize = 18 if value == "●" else 10
                                
                                # Intentar insertar con textbox
                                rc = page.insert_textbox(
                                    text_rect,
                                    value,
                                    fontsize=fontsize,
                                    color=(0, 0, 0),
                                    align=0,
                                    overlay=True
                                )
                                
                                # Si no cabe o falla, usar insert_text con el tamaño apropiado
                                if rc < 0:
                                    page.insert_text(
                                        (rect.x0, rect.y1 - 2),
                                        value,
                                        fontsize=fontsize,
                                        color=(0, 0, 0),
                                        overlay=True
                                    )
                            except Exception as e:
                                # Si falla, intentar solo con insert_text en la posición
                                try:
                                    fontsize = 18 if value == "●" else 10
                                    page.insert_text(
                                        (rect.x0, rect.y1 - 2),
                                        value,
                                        fontsize=fontsize,
                                        color=(0, 0, 0),
                                        overlay=True
                                    )
                                except:
                                    pass
                    
                    found_and_modified = True
            
            # Guardar el PDF modificado
            pdf_document.save(output_file)
            pdf_document.close()
            
            return {
                "status": "success",
                "message": "Documento procesado exitosamente",
                "filename": unique_filename,
                "file_path": str(output_file)
            }
            
        except Exception as e:
            return {
                "status": "error",
                "message": f"Error procesando el documento: {str(e)}",
                "filename": None,
                "file_path": None
            }

    def get(self, id: int) -> Any:
        """
        Obtiene un documento por su ID.
        """
        try:
            document = self.db.query(DocumentModel).filter(
                DocumentModel.id == id,
                DocumentModel.deleted_date.is_(None)
            ).first()

            if document:
                return {
                    "id": document.id,
                    "document_type_id": document.document_type_id,
                    "career_type_id": document.career_type_id,
                    "document": document.document,
                    "added_date": document.added_date.strftime("%Y-%m-%d %H:%M:%S") if document.added_date else None,
                    "updated_date": document.updated_date.strftime("%Y-%m-%d %H:%M:%S") if document.updated_date else None
                }
            else:
                return {"status": "error", "message": "No se encontraron datos para el documento especificado."}

        except Exception as e:
            error_message = str(e)
            return {"status": "error", "message": error_message}

    def get_all(self, document_type_id: Optional[int] = None, career_type_id: Optional[int] = None) -> Any:
        """
        Obtiene la lista de documentos almacenados.
        Solo devuelve documentos que no tienen deleted_date (no eliminados).
        """
        try:
            query = self.db.query(DocumentModel).filter(DocumentModel.deleted_date.is_(None))

            if document_type_id is not None:
                query = query.filter(DocumentModel.document_type_id == document_type_id)
            
            # Filtrar por career_type_id solo si se proporciona y no está vacío
            if career_type_id is not None and career_type_id != 0:
                query = query.filter(DocumentModel.career_type_id == career_type_id)

            documents = query.order_by(DocumentModel.document.asc()).all()

            return [
                {
                    "id": doc.id,
                    "document_type_id": doc.document_type_id,
                    "career_type_id": doc.career_type_id,
                    "document": doc.document
                }
                for doc in documents
            ]

        except Exception as e:
            return {
                "status": "error",
                "message": str(e)
            }

    def store(self, student_id: int, document_type_id: int, file_path: str) -> Any:
        """
        Almacena un documento dependiendo del tipo.
        - Si document_type_id == 1 (birth certificate), guarda en birth_certificate_documents y folders
        - Para otros tipos, solo guarda en folders
        """
        try:
            # Buscar el document_id correspondiente al document_type_id (solo no eliminados)
            document = self.db.query(DocumentModel).filter(
                DocumentModel.document_type_id == document_type_id,
                DocumentModel.deleted_date.is_(None)
            ).first()
            
            if not document:
                return {
                    "status": "error",
                    "message": f"No se encontró documento con document_type_id {document_type_id}"
                }
            
            detail_id = None
            
            # Si es document_type_id == 1 (birth certificate), guardar primero en birth_certificate_documents
            if document_type_id == 1:
                # Crear registro en birth_certificate_documents
                new_birth_cert = BirthCertificateDocumentModel(
                    student_id=student_id,
                    birth_certificate=file_path,
                    added_date=datetime.now(),
                    updated_date=datetime.now()
                )
                self.db.add(new_birth_cert)
                self.db.flush()  # Para obtener el ID sin hacer commit completo aún
                detail_id = new_birth_cert.id
            
            # Si es health evaluation (document_id = 4), actualizar registro existente en lugar de crear nuevo
            if document.id == 4:
                # Buscar registro con file vacío (null) para este estudiante y documento
                folder_without_file = self.db.query(FolderModel).filter(
                    FolderModel.student_id == student_id,
                    FolderModel.document_id == document.id,
                    FolderModel.file.is_(None)
                ).first()
                
                if folder_without_file:
                    # Actualizar el registro existente con file vacío
                    folder_without_file.file = file_path
                    folder_without_file.updated_date = datetime.now()
                    
                    self.db.commit()
                    self.db.refresh(folder_without_file)
                    
                    return {
                        "status": "success",
                        "message": "Documento actualizado exitosamente",
                        "document_id": folder_without_file.id,
                        "version_id": folder_without_file.version_id
                    }
                else:
                    # Si no hay registro con file vacío, actualizar la versión más reciente
                    last_version = self.db.query(FolderModel).filter(
                        FolderModel.student_id == student_id,
                        FolderModel.document_id == document.id
                    ).order_by(FolderModel.version_id.desc()).first()
                    
                    if last_version:
                        # Actualizar el registro más reciente
                        last_version.file = file_path
                        last_version.updated_date = datetime.now()
                        
                        self.db.commit()
                        self.db.refresh(last_version)
                        
                        return {
                            "status": "success",
                            "message": "Documento actualizado exitosamente",
                            "document_id": last_version.id,
                            "version_id": last_version.version_id
                        }
            
            # Para otros documentos, crear nueva versión
            # Buscar la última versión para este estudiante y documento
            last_version = self.db.query(FolderModel).filter(
                FolderModel.student_id == student_id,
                FolderModel.document_id == document.id
            ).order_by(FolderModel.version_id.desc()).first()
            
            # Determinar el nuevo version_id
            if last_version:
                new_version_id = last_version.version_id + 1
            else:
                new_version_id = 1
            
            # Crear el nuevo registro en folders
            new_folder = FolderModel(
                school_id=None,
                course_id=None,
                student_id=student_id,
                document_id=document.id,
                version_id=new_version_id,
                detail_id=detail_id,  # ID de birth_certificate_documents si es tipo 1
                professional_id=0,
                file=file_path,
                period_year=None,
                added_date=datetime.now(),
                updated_date=datetime.now(),
                deleted_date=None,
            )
            
            self.db.add(new_folder)
            self.db.commit()
            self.db.refresh(new_folder)
            
            return {
                "status": "success",
                "message": "Documento creado exitosamente",
                "document_id": new_folder.id,
                "version_id": new_version_id
            }
                
        except Exception as e:
            self.db.rollback()
            return {
                "status": "error",
                "message": str(e)
            }

    @staticmethod
    def generate_health_evaluation_pdf(
        template_path: str,
        evaluation_data: Dict[str, Any],
        db: Optional[Session] = None,
        output_directory: str = "files/system/students"
    ) -> Dict[str, Any]:
        """
        Genera un PDF de evaluación de salud rellenando campos de formulario AcroForm.
        Usa pypdf para leer y rellenar los campos del formulario.
        
        Args:
            template_path: Ruta del archivo PDF template con campos AcroForm
            evaluation_data: Diccionario con los datos de la evaluación de salud
            db: Sesión de base de datos (opcional, para obtener datos relacionados)
            output_directory: Directorio donde se guardará el archivo generado
        
        Returns:
            dict: Diccionario con status, message, filename y file_path
        """
        if not PYPDF_AVAILABLE:
            return {
                "status": "error",
                "message": "pypdf no está instalado. Instálelo con: pip install pypdf",
                "filename": None,
                "file_path": None
            }
        
        try:
            template_file = Path(template_path)
            
            if not template_file.exists():
                return {
                    "status": "error",
                    "message": "Template PDF no encontrado",
                    "filename": None,
                    "file_path": None
                }
            
            # Validar que sea PDF
            if template_file.suffix.lower() != '.pdf':
                return {
                    "status": "error",
                    "message": f"El template debe ser un archivo PDF, se encontró: {template_file.suffix}",
                    "filename": None,
                    "file_path": None
                }
            
            # Crear nombre único para el archivo generado
            student_name = evaluation_data.get("full_name", "estudiante").replace(" ", "_")
            unique_filename = f"evaluacion_salud_{student_name}_{uuid.uuid4().hex[:8]}.pdf"
            output_file = Path(output_directory) / unique_filename
            
            # Asegurar que el directorio de salida existe
            output_file.parent.mkdir(parents=True, exist_ok=True)
            
            # Leer el PDF template
            reader = PdfReader(template_file)
            writer = PdfWriter()
            
            # Copiar las páginas
            for page in reader.pages:
                writer.add_page(page)
            
            # Copiar el AcroForm del reader al writer
            # Necesitamos copiar el AcroForm después de agregar las páginas
            try:
                # Primero, obtener el AcroForm del reader
                reader_root = reader.trailer.get("/Root", {})
                if "/AcroForm" in reader_root:
                    reader_acro_form = reader_root["/AcroForm"]
                    
                    # Crear el root del writer si no existe
                    if not hasattr(writer, '_root_object') or writer._root_object is None:
                        from pypdf.generic import DictionaryObject
                        writer._root_object = DictionaryObject()
                        # Copiar otras propiedades del root si existen
                        for key, value in reader_root.items():
                            if key != "/AcroForm":
                                writer._root_object[NameObject(key)] = value
                    
                    # Intentar clonar el AcroForm usando el método de pypdf
                    try:
                        cloned_acro_form = reader_acro_form.clone(writer, force_duplicate=True)
                        writer._root_object[NameObject("/AcroForm")] = cloned_acro_form
                    except Exception as clone_error:
                        # Si falla, el AcroForm puede no estar disponible en el writer
                        # pero los campos deberían estar en las páginas
                        pass
            except Exception as e:
                # Continuar de todas formas, puede que los campos estén en las páginas
                pass
            
            # Función auxiliar para formatear fechas
            def format_date(date_str: Optional[str]) -> str:
                if not date_str:
                    return ""
                try:
                    if isinstance(date_str, str):
                        date_obj = datetime.strptime(date_str, "%Y-%m-%d").date()
                        return date_obj.strftime("%d/%m/%Y")
                    return ""
                except:
                    return str(date_str) if date_str else ""
            
            # Función auxiliar para obtener valor o string vacío
            def get_value(key: str, default: str = "") -> str:
                value = evaluation_data.get(key)
                if value is None:
                    return default
                return str(value) if value else default
            
            # Calcular edad en años y meses desde fecha de nacimiento
            def age_years_months(born_date_val, reference_date_val=None):
                if not born_date_val:
                    return ""
                try:
                    from datetime import date
                    if isinstance(born_date_val, str):
                        born = datetime.strptime(born_date_val.strip()[:10], "%Y-%m-%d").date()
                    else:
                        born = born_date_val
                    if reference_date_val:
                        ref = (
                            datetime.strptime(str(reference_date_val).strip()[:10], "%Y-%m-%d").date()
                            if isinstance(reference_date_val, str)
                            else reference_date_val
                        )
                    else:
                        ref = date.today()
                    if born > ref:
                        return ""
                    total_months = (ref.year - born.year) * 12 + (ref.month - born.month)
                    if ref.day < born.day:
                        total_months -= 1
                    years = total_months // 12
                    months = total_months % 12
                    if months == 0:
                        return f"{years} años"
                    return f"{years} años {months} meses"
                except Exception:
                    return ""
            
            # Fecha de referencia para la edad (evaluación o hoy)
            ref_date = evaluation_data.get("evaluation_date") or None
            age_str = age_years_months(evaluation_data.get("born_date"), ref_date)
            if not age_str and evaluation_data.get("age") is not None:
                age_str = str(evaluation_data.get("age"))  # fallback solo años
            
            # Inspeccionar los campos del formulario en el PDF
            form_fields = {}
            try:
                # Método 1: Obtener campos de texto usando get_form_text_fields()
                text_fields = reader.get_form_text_fields()
                if text_fields:
                    form_fields.update(text_fields)
                
                # Método 2: Obtener todos los campos accediendo directamente al AcroForm
                root = reader.trailer.get("/Root", {})
                if "/AcroForm" in root:
                    acro_form = root["/AcroForm"]
                    if "/Fields" in acro_form:
                        fields = acro_form["/Fields"]
                        
                        def extract_field_names(field_list):
                            for field_ref in field_list:
                                field = field_ref.get_object()
                                if "/Kids" in field:
                                    extract_field_names(field["/Kids"])
                                if "/T" in field:
                                    field_name = field["/T"]
                                    if field_name:
                                        form_fields[field_name] = None
                        
                        extract_field_names(fields)
                
                # Método 3: También buscar en las anotaciones de cada página
                for page_num, page in enumerate(reader.pages):
                    if "/Annots" in page:
                        for annot in page["/Annots"]:
                            annot_obj = annot.get_object()
                            if "/FT" in annot_obj and "/T" in annot_obj:
                                field_name = annot_obj.get("/T")
                                if field_name:
                                    form_fields[field_name] = None
            except Exception as e:
                pass
            
            # Obtener género (F o M)
            gender_marker = ""
            gender_name = evaluation_data.get("gender_name", "")
            if gender_name:
                gender_upper = gender_name.upper()
                if "FEMENINO" in gender_upper or "F" in gender_upper:
                    gender_marker = "F"
                elif "MASCULINO" in gender_upper or "M" in gender_upper:
                    gender_marker = "M"
            
            # Obtener valores para checkboxes basado en consultation_reason_id
            # Si consultation_reason_id == 1, marcar EGS (Examen General de Salud)
            # Si consultation_reason_id != 1, marcar DDD (Diagnóstico de Discapacidad o Déficit)
            consultation_reason_id = evaluation_data.get("consultation_reason_id")
            is_egs = (consultation_reason_id == 1)
            is_ddd = (consultation_reason_id != 1 and consultation_reason_id is not None)
            
            # EGS = Examen General de Salud (checkbox)
            # DDD = Diagnóstico de Discapacidad o Déficit (checkbox)
            
            # Procedencia: procedence_id 1=SP, 2=P, 3=E, 4=O (checkboxes)
            procedence_id = evaluation_data.get("procedence_id")
            is_sp = (procedence_id == 1)
            is_p = (procedence_id == 2)   # Particular
            is_e = (procedence_id == 3)   # Escuela
            is_o = (procedence_id == 4)   # Otro
            # OE = texto "otro": solo si O está marcado (procedence_id==4) usar procedence_other; si no, en blanco
            oe_value = get_value("procedence_other", "") if is_o else ""
            
            # Preparar datos para los campos del formulario
            # Mapeo según los nombres reales de los campos encontrados en el PDF
            field_mapping = {
                # Identificación del estudiante
                "Nombres y Apellidos": get_value("full_name", ""),
                "RUN": get_value("identification_number", ""),
                "Nacionalidad": evaluation_data.get("nationality_name", ""),
                "Fecha de Nacimiento": format_date(evaluation_data.get("born_date")),
                "BD": format_date(evaluation_data.get("born_date")),  # Campo real del formulario PDF
                "Edad": age_str,
                "Age": age_str,  # Campo real del formulario PDF (edad en años y meses desde fecha nacimiento)
                # Lenguas
                "OL": get_value("native_language", ""),  # Lengua de origen / nativa
                "HL": get_value("language_usually_used", ""),  # Lengua habitual
                # Género (botones de radio) - usar strings para update_page_form_field_values
                "F": "/Yes" if gender_marker == "F" else "/Off",
                "M": "/Yes" if gender_marker == "M" else "/Off",
                # Motivo de consulta
                "MOTIVO DE CONSULTA": get_value("consultation_reason_detail", ""),
                "DD": get_value("consultation_reason_detail", ""),  # Campo real del formulario PDF
                # Identificación del profesional
                "Nombres y Apellidos_2": evaluation_data.get("professional_fullname", ""),
                "Rut": get_value("professional_identification_number", ""),
                "PRN": get_value("professional_registration_number", ""),  # Nº registro profesional
                "Especialidad": evaluation_data.get("professional_specialty_name", get_value("professional_specialty", "")),
                "Procedencia Salud pública Particular Escuela Otro": get_value("procedence_other", ""),
                # Procedencia checkboxes: 1=SP, 2=P, 3=E, 4=O
                "SP": "/Yes" if is_sp else "/Off",
                "E": "/Yes" if is_e else "/Off",
                "P": "/Yes" if is_p else "/Off",
                "O": "/Yes" if is_o else "/Off",
                # OE = texto "otro": solo con valor si O está marcado (procedence_id==4)
                "OE": oe_value,
                "FonoEmail contacto": get_value("professional_contact", ""),
                # Fechas
                "Fecha evaluación": format_date(get_value("evaluation_date")),
                "Fecha reevaluación": format_date(get_value("reevaluation_date")),
                # Examen del estado de salud
                "EXÁMEN DEL ESTADO DE SALUD GENERAL DEL ESTUDIANTE Presenciaausencia de patologías o dificultades de salud que incidan en o expliquen sus necesidades educativas especiales por ejemplo bronquitis crónica problemas sensoriales etcRow1": get_value("general_assessment", ""),
                # Diagnóstico
                "DIAGNÓSTICO  Presencia de un trastorno déficit o discapacidad": get_value("diagnosis", ""),
                "Considere especificaciones del grado y etiología del déficit y pronóstico": get_value("diagnosis", ""),
                # Indicaciones
                "INDICACIONES": get_value("indications", ""),
                "Señale tratamiento médico necesidades de interconsulta exámenes o ayudas técnicas cuando sea el caso u otras recomendaciones u observaciones relevantes para ella estudiante en función de su diagnóstico": get_value("indications", ""),
                # Diagnóstico de discapacidad (campo de texto)
                "DIAGNÓSTICO DE DISCAPACIDAD O DÉFICIT señale cual": get_value("diagnosis", ""),
                # Checkboxes - Basado en consultation_reason_id
                # Si consultation_reason_id == 1: marcar EGS, desmarcar DDD
                # Si consultation_reason_id != 1: marcar DDD, desmarcar EGS
                "EXAMEN GENERAL DE SALUD": "/Yes" if is_egs else "/Off",
                "EGS": "/Yes" if is_egs else "/Off",
                "DDD": "/Yes" if is_ddd else "/Off",
            }
            
            # Construir el diccionario final de datos a rellenar
            # Solo incluir campos que realmente existen en el PDF
            form_data = {}
            
            # Mapeo normalizado (minúsculas, sin espacios) para coincidir con nombres del PDF
            field_mapping_normalized = {k.strip().lower(): k for k in field_mapping.keys()}
            
            for field_name in form_fields.keys():
                value = None
                # 1) Coincidencia exacta
                if field_name in field_mapping:
                    value = field_mapping[field_name]
                # 2) Coincidencia sin distinguir mayúsculas/minúsculas (BD, DD, Age, etc.)
                elif field_name.strip().lower() in field_mapping_normalized:
                    canonical_key = field_mapping_normalized[field_name.strip().lower()]
                    value = field_mapping[canonical_key]
                
                if value is not None:
                    # Checkboxes y radio: siempre agregar
                    if value == "/Off" or value == "/Yes":
                        form_data[field_name] = value
                    # DD y OE: siempre agregar para que el campo aparezca (vacío o con valor)
                    elif field_name.strip().lower() in ("dd", "oe"):
                        form_data[field_name] = value if isinstance(value, str) else (value or "")
                    elif isinstance(value, str) and value.strip():
                        # Resto de campos de texto: solo si tiene contenido
                        form_data[field_name] = value
                else:
                    # Intentar mapeo por palabras clave si no hay coincidencia exacta
                    field_lower = field_name.strip().lower()
                    if "nombre" in field_lower and "apellido" in field_lower and "2" not in field_name:
                        # Es el campo de nombres del estudiante
                        value = get_value("full_name", "")
                        if value:
                            form_data[field_name] = value
                    elif "run" in field_lower and field_name.strip().upper() == "RUN":
                        value = get_value("identification_number", "")
                        if value:
                            form_data[field_name] = value
                    elif "nacionalidad" in field_lower:
                        value = evaluation_data.get("nationality_name", "")
                        if value:
                            form_data[field_name] = value
                    elif field_lower == "dd" or field_lower.endswith(".dd") or ("." in field_lower and field_lower.split(".")[-1].strip() == "dd"):
                        # Campo DD = detalle / motivo de consulta (siempre agregar, vacío o con valor)
                        form_data[field_name] = get_value("consultation_reason_detail", "")
                    elif field_lower == "bd" or field_lower.endswith(".bd") or ("." in field_lower and field_lower.split(".")[-1].strip() == "bd"):
                        value = format_date(evaluation_data.get("born_date"))
                        if value:
                            form_data[field_name] = value
                    elif field_lower == "age" or field_lower.endswith(".age") or ("." in field_lower and field_lower.split(".")[-1].strip() == "age"):
                        if age_str:
                            form_data[field_name] = age_str
                    elif field_name == "F" or field_name == "M":
                        # Botones de radio para género - usar strings
                        if field_name == "F":
                            form_data[field_name] = "/Yes" if gender_marker == "F" else "/Off"
                        elif field_name == "M":
                            form_data[field_name] = "/Yes" if gender_marker == "M" else "/Off"
            
            # Rellenar los campos del formulario y hacerlos de solo lectura
            if form_data:
                # Separar campos de texto, radio buttons y checkboxes
                text_fields = {}
                radio_fields = {}
                checkbox_fields = {}
                for field_name, value in form_data.items():
                    if field_name in ["F", "M"]:
                        radio_fields[field_name] = value
                    elif isinstance(value, str) and value in ["/Yes", "/Off"]:
                        # Es un checkbox (botón que puede estar marcado o no)
                        checkbox_fields[field_name] = value
                    else:
                        # Campos de texto: con valor; DD y OE siempre (para vacío cuando corresponda)
                        fn_lower = field_name.strip().lower()
                        if value and str(value).strip() or fn_lower == "dd" or fn_lower == "oe":
                            text_fields[field_name] = value if isinstance(value, str) else (value or "")
                
                # Lookup normalizado: algunos visores/PDFs usan nombres con encoding distinto
                text_fields_normalized = {k.strip().lower(): (k, v) for k, v in text_fields.items()}
                checkbox_fields_normalized = {k.strip().lower(): v for k, v in checkbox_fields.items()}
                radio_fields_normalized = {k.strip().lower(): v for k, v in radio_fields.items()}
                def get_text_value(name_str):
                    if name_str in text_fields:
                        return text_fields[name_str]
                    norm = name_str.strip().lower()
                    if norm in text_fields_normalized:
                        return text_fields_normalized[norm][1]
                    return None
                def _norm(name):
                    return (name or "").strip().lstrip("/").lower()
                def get_checkbox_value(name_str):
                    n = (name_str or "").strip().lstrip("/")
                    if n in checkbox_fields:
                        return checkbox_fields[n]
                    return checkbox_fields_normalized.get(_norm(name_str))
                def get_radio_value(name_str):
                    n = (name_str or "").strip().lstrip("/")
                    if n in radio_fields:
                        return radio_fields[n]
                    return radio_fields_normalized.get(_norm(name_str))
                
                # Pedir al visor que regenere apariencia desde /V (algunos PDFs no muestran valor si no)
                try:
                    if hasattr(writer, '_root_object') and writer._root_object and "/AcroForm" in writer._root_object:
                        af = writer._root_object["/AcroForm"]
                        af[NameObject("/NeedAppearances")] = BooleanObject(True)
                except Exception:
                    pass
                
                # Rellenar todos los campos directamente desde las anotaciones de las páginas
                # Este método es más confiable que usar el AcroForm del root
                try:
                    for page_num in range(len(writer.pages)):
                        page = writer.pages[page_num]
                        if "/Annots" in page:
                            annotations = page["/Annots"]
                            if annotations:
                                for annot_ref in annotations:
                                    try:
                                        annot = annot_ref.get_object()
                                        if "/T" in annot and "/FT" in annot:
                                            # Normalizar a str: pypdf puede devolver NameObject/TextStringObject
                                            field_name = annot["/T"]
                                            field_name_str = str(field_name).strip() if field_name else ""
                                            field_type = annot["/FT"]
                                            field_type_str = str(field_type).strip() if field_type else ""
                                            
                                            # Rellenar campos de texto (DD y OE siempre; resto solo si tiene contenido)
                                            if field_type_str == "/Tx":
                                                value = get_text_value(field_name_str)
                                                fn_lower = field_name_str.strip().lower()
                                                is_dd_or_oe = fn_lower == "dd" or fn_lower == "oe"
                                                if value is not None and (is_dd_or_oe or (value and str(value).strip())):
                                                    val_str = str(value).strip() if value else ""
                                                    annot[NameObject("/V")] = TextStringObject(val_str)
                                                    # Algunos visores usan el padre: actualizar también si tiene /Parent
                                                    if "/Parent" in annot and (is_dd_or_oe or (value and str(value).strip())):
                                                        try:
                                                            parent = annot["/Parent"].get_object()
                                                            parent[NameObject("/V")] = TextStringObject(val_str)
                                                        except Exception:
                                                            pass
                                            
                                            # Rellenar radio buttons (género F/M) y checkboxes (procedencia SP,E,P,O; EGS, DDD)
                                            if field_type_str == "/Btn":
                                                radio_val = get_radio_value(field_name_str)
                                                if radio_val is not None:
                                                    if radio_val == "/Yes":
                                                        annot[NameObject("/V")] = NameObject("/Yes")
                                                        annot[NameObject("/AS")] = NameObject("/Yes")
                                                    else:
                                                        annot[NameObject("/V")] = NameObject("/Off")
                                                        annot[NameObject("/AS")] = NameObject("/Off")
                                                else:
                                                    cb_val = get_checkbox_value(field_name_str)
                                                    if cb_val is not None:
                                                        if cb_val == "/Yes":
                                                            annot[NameObject("/V")] = NameObject("/Yes")
                                                            annot[NameObject("/AS")] = NameObject("/Yes")
                                                        else:
                                                            annot[NameObject("/V")] = NameObject("/Off")
                                                            annot[NameObject("/AS")] = NameObject("/Off")
                                    except Exception as e:
                                        continue
                    
                    # Escribir también en el árbol AcroForm /Fields (algunos visores solo leen /V del campo, no del widget)
                    try:
                        writer_root = writer._root_object
                        if writer_root and "/AcroForm" in writer_root:
                            acro_form = writer_root["/AcroForm"]
                            if "/Fields" in acro_form:
                                def set_text_in_fields(field_list):
                                    if not field_list:
                                        return
                                    for field_ref in field_list:
                                        try:
                                            field = field_ref.get_object()
                                            if "/Kids" in field and field["/Kids"]:
                                                set_text_in_fields(field["/Kids"])
                                            if "/T" in field and "/FT" in field:
                                                fn = str(field["/T"]).strip() if field["/T"] else ""
                                                ft = str(field["/FT"]).strip() if field["/FT"] else ""
                                                if ft == "/Tx":
                                                    val = get_text_value(fn)
                                                    fn_low = fn.strip().lower()
                                                    is_dd_oe = fn_low == "dd" or fn_low == "oe"
                                                    if val is not None and (is_dd_oe or (val and str(val).strip())):
                                                        field[NameObject("/V")] = TextStringObject(str(val).strip() if val else "")
                                        except Exception:
                                            continue
                                set_text_in_fields(acro_form["/Fields"])
                    except Exception:
                        pass
                except Exception as e:
                    
                    # Método alternativo: intentar desde el AcroForm del root
                    try:
                        writer_root = writer._root_object
                        if writer_root and "/AcroForm" in writer_root:
                            acro_form = writer_root["/AcroForm"]
                            if "/Fields" in acro_form:
                                fields = acro_form["/Fields"]
                                
                                def process_all_fields(field_list):
                                    if not field_list:
                                        return
                                    try:
                                        for field_ref in field_list:
                                            try:
                                                field = field_ref.get_object()
                                                if "/Kids" in field and field["/Kids"]:
                                                    process_all_fields(field["/Kids"])
                                                if "/T" in field:
                                                    field_name = field["/T"]
                                                    field_name_str = str(field_name).strip() if field_name else ""
                                                    value_t = get_text_value(field_name_str)
                                                    fn_alt = field_name_str.strip().lower()
                                                    is_dd_oe_alt = fn_alt == "dd" or fn_alt == "oe"
                                                    if value_t is not None and (is_dd_oe_alt or (value_t and str(value_t).strip())):
                                                        field[NameObject("/V")] = TextStringObject(str(value_t).strip() if value_t else "")
                                                    rv = get_radio_value(field_name_str)
                                                    if rv is not None:
                                                        field[NameObject("/V")] = NameObject("/Yes" if rv == "/Yes" else "/Off")
                                                        field[NameObject("/AS")] = NameObject("/Yes" if rv == "/Yes" else "/Off")
                                                    else:
                                                        cv = get_checkbox_value(field_name_str)
                                                        if cv is not None:
                                                            field[NameObject("/V")] = NameObject("/Yes" if cv == "/Yes" else "/Off")
                                                            field[NameObject("/AS")] = NameObject("/Yes" if cv == "/Yes" else "/Off")
                                            except:
                                                continue
                                    except:
                                        pass
                                
                                process_all_fields(fields)
                    except Exception as e2:
                        pass
                
                # Manejar radio buttons directamente desde el AcroForm (si aún no se procesaron)
                if radio_fields:
                    try:
                        writer_root = writer._root_object
                        if writer_root and "/AcroForm" in writer_root:
                            acro_form = writer_root["/AcroForm"]
                            if "/Fields" in acro_form:
                                fields = acro_form["/Fields"]
                                
                                # Validar que fields no esté vacío
                                if not fields or len(fields) == 0:
                                    return
                                
                                # Función recursiva para procesar campos (incluyendo campos anidados)
                                def process_radio_fields(field_list):
                                    if not field_list:
                                        return
                                    if not isinstance(field_list, list):
                                        return
                                    try:
                                        for i, field_ref in enumerate(field_list):
                                            try:
                                                if field_ref is None:
                                                    continue
                                                
                                                field = field_ref.get_object()
                                                
                                                # Si el campo tiene subcampos (Kids), procesarlos recursivamente
                                                if "/Kids" in field and field["/Kids"]:
                                                    process_radio_fields(field["/Kids"])
                                                
                                                # Si el campo tiene nombre: radio (F/M) o checkbox (SP,E,P,O,EGS,DDD)
                                                if "/T" in field:
                                                    field_name = field["/T"]
                                                    field_name_str = (str(field_name).strip().lstrip("/") if field_name else "") or ""
                                                    rv = get_radio_value(field_name_str)
                                                    if rv is not None:
                                                        field[NameObject("/V")] = NameObject("/Yes" if rv == "/Yes" else "/Off")
                                                        field[NameObject("/AS")] = NameObject("/Yes" if rv == "/Yes" else "/Off")
                                                    else:
                                                        cv = get_checkbox_value(field_name_str)
                                                        if cv is not None:
                                                            field[NameObject("/V")] = NameObject("/Yes" if cv == "/Yes" else "/Off")
                                                            field[NameObject("/AS")] = NameObject("/Yes" if cv == "/Yes" else "/Off")
                                            except Exception as e:
                                                continue
                                    except Exception as e:
                                        pass
                                
                                # Procesar todos los campos
                                process_radio_fields(fields)
                    except Exception as e:
                        pass
                
                # Hacer los campos de solo lectura después de rellenarlos
                try:
                    # Acceder al AcroForm del writer
                    if hasattr(writer, '_root_object') and writer._root_object is not None:
                        writer_root = writer._root_object
                        if "/AcroForm" in writer_root:
                            acro_form = writer_root["/AcroForm"]
                            if "/Fields" in acro_form:
                                fields = acro_form["/Fields"]
                                
                                # Validar que fields no esté vacío y sea una lista
                                if fields and isinstance(fields, list):
                                    # Función recursiva para procesar campos (incluyendo campos anidados)
                                    def process_fields(field_list):
                                        if not field_list:
                                            return
                                        if not isinstance(field_list, list):
                                            return
                                        try:
                                            for i, field_ref in enumerate(field_list):
                                                try:
                                                    if field_ref is None:
                                                        continue
                                                    
                                                    field = field_ref.get_object()
                                                    
                                                    # Si el campo tiene subcampos (Kids), procesarlos recursivamente
                                                    if "/Kids" in field and field["/Kids"]:
                                                        process_fields(field["/Kids"])
                                                    
                                                    # Si el campo tiene nombre, hacerlo de solo lectura
                                                    if "/T" in field:
                                                        field_name = field["/T"]
                                                        field_name_str = str(field_name).strip() if field_name else ""
                                                        if field_name_str in form_data:
                                                            # Hacer el campo de solo lectura (bit 1 = ReadOnly)
                                                            if "/Ff" in field:
                                                                current_flags = int(field["/Ff"])
                                                                field[NameObject("/Ff")] = NumberObject(current_flags | 1)
                                                            else:
                                                                field[NameObject("/Ff")] = NumberObject(1)
                                                except Exception as e:
                                                    continue
                                        except Exception as e:
                                            pass
                                
                                # Procesar todos los campos
                                process_fields(fields)
                except Exception as e2:
                    pass
                    import traceback
                    traceback.print_exc()
                else:
                    pass
            
            # Guardar el PDF rellenado
            try:
                with open(output_file, "wb") as output_stream:
                    writer.write(output_stream)
            except Exception as e:
                import traceback
                traceback.print_exc()
                return {
                    "status": "error",
                    "message": f"Error guardando PDF: {str(e)}",
                    "filename": None,
                    "file_path": None
                }
            
            return {
                "status": "success",
                "message": "PDF generado exitosamente",
                "filename": unique_filename,
                "file_path": str(output_file)
            }
                
        except Exception as e:
            import traceback
            error_traceback = traceback.format_exc()
            return {
                "status": "error",
                "message": f"Error generando PDF: {str(e)}",
                "filename": None,
                "file_path": None
            }

    @staticmethod
    def generate_anamnesis_pdf(
        template_path: str,
        anamnesis_data: Dict[str, Any],
        db: Optional[Session] = None,
        output_directory: str = "files/system/students"
    ) -> Dict[str, Any]:
        """
        Genera un PDF de anamnesis rellenando campos de formulario AcroForm.
        Usa pypdf para leer y rellenar los campos del formulario.
        Template: anamnesis_student.pdf en files/original_student_files/
        Al final marca los campos como solo lectura (quita el formulario editable).
        """
        if not PYPDF_AVAILABLE:
            return {
                "status": "error",
                "message": "pypdf no está instalado. Instálelo con: pip install pypdf",
                "filename": None,
                "file_path": None
            }
        try:
            import json as _json
            template_file = Path(template_path)
            if not template_file.exists():
                return {"status": "error", "message": "Template PDF no encontrado", "filename": None, "file_path": None}
            if template_file.suffix.lower() != '.pdf':
                return {"status": "error", "message": "El template debe ser un archivo PDF", "filename": None, "file_path": None}

            student_name = (anamnesis_data.get("student_full_name") or "estudiante").replace(" ", "_")
            unique_filename = f"anamnesis_{student_name}_{uuid.uuid4().hex[:8]}.pdf"
            output_file = Path(output_directory) / unique_filename
            output_file.parent.mkdir(parents=True, exist_ok=True)

            reader = PdfReader(template_file)
            writer = PdfWriter()
            for page in reader.pages:
                writer.add_page(page)

            try:
                reader_root = reader.trailer.get("/Root", {})
                if "/AcroForm" in reader_root:
                    from pypdf.generic import DictionaryObject
                    if not hasattr(writer, '_root_object') or writer._root_object is None:
                        writer._root_object = DictionaryObject()
                        for key, value in reader_root.items():
                            if key != "/AcroForm":
                                writer._root_object[NameObject(key)] = value
                    try:
                        writer._root_object[NameObject("/AcroForm")] = reader_root["/AcroForm"].clone(writer, force_duplicate=True)
                    except Exception:
                        pass
            except Exception:
                pass

            def format_date(val) -> str:
                if not val:
                    return ""
                try:
                    if isinstance(val, str):
                        dt = datetime.strptime(val.strip()[:10], "%Y-%m-%d")
                        return dt.strftime("%d/%m/%Y")
                    if hasattr(val, 'strftime'):
                        return val.strftime("%d/%m/%Y")
                    return str(val)
                except Exception:
                    return str(val) if val else ""

            def to_str(val) -> str:
                if val is None:
                    return ""
                if isinstance(val, (list, dict)):
                    try:
                        return _json.dumps(val, ensure_ascii=False)
                    except Exception:
                        return str(val)
                return str(val).strip() if val else ""

            def get_val(key: str, fmt_date: bool = False) -> str:
                v = anamnesis_data.get(key)
                if v is None:
                    return ""
                return format_date(v) if fmt_date else to_str(v)

            # Obtener campos del formulario
            form_fields = {}
            try:
                if reader.get_form_text_fields():
                    form_fields.update(reader.get_form_text_fields())
                root = reader.trailer.get("/Root", {})
                if "/AcroForm" in root and "/Fields" in root["/AcroForm"]:
                    def extract_names(fl):
                        for ref in fl:
                            obj = ref.get_object()
                            if "/Kids" in obj:
                                extract_names(obj["/Kids"])
                            if "/T" in obj and obj["/T"]:
                                form_fields[obj["/T"]] = None
                    extract_names(root["/AcroForm"]["/Fields"])
                for page in reader.pages:
                    if "/Annots" in page:
                        for annot in page["/Annots"]:
                            a = annot.get_object()
                            if "/FT" in a and "/T" in a and a.get("/T"):
                                form_fields[a["/T"]] = None
            except Exception:
                pass

            # Resolver género (gender_id -> texto)
            d = anamnesis_data
            gender_name = ""
            if db and d.get("gender_id"):
                try:
                    from app.backend.db.models import GenderModel
                    g = db.query(GenderModel).filter(GenderModel.id == d.get("gender_id")).first()
                    if g and g.gender:
                        gender_name = g.gender.strip()
                except Exception:
                    pass
            if not gender_name:
                gender_name = anamnesis_data.get("gender_name") or ""

            # Pares: nombre campo PDF (reales del formulario) -> valor
            nombre_val = get_val("student_full_name")
            fecha_nac_val = format_date(d.get("born_date"))
            _addr, _ph = get_val("address"), get_val("phone")

            # Calcular años y meses de edad
            age_years = ""
            age_months = ""
            if d.get("born_date"):
                try:
                    from datetime import date as _date
                    bd = None
                    if isinstance(d.get("born_date"), str):
                        bd = datetime.strptime(str(d.get("born_date"))[:10], "%Y-%m-%d").date()
                    elif hasattr(d.get("born_date"), "strftime"):
                        bd = d.get("born_date")
                    if bd:
                        hoy = _date.today()
                        total_meses = (hoy.year - bd.year) * 12 + (hoy.month - bd.month)
                        if hoy.day < bd.day:
                            total_meses -= 1
                        age_years = str(total_meses // 12)
                        age_months = str(total_meses % 12)
                except Exception:
                    pass
            if not age_years and d.get("age"):
                age_years = str(d.get("age"))

            # Resolver país (nationality_id -> nombre)
            country_name = ""
            if db and d.get("nationality_id"):
                try:
                    from app.backend.db.models import NationalityModel
                    n = db.query(NationalityModel).filter(NationalityModel.id == d.get("nationality_id")).first()
                    if n and n.nationality:
                        country_name = n.nationality.strip()
                except Exception:
                    pass
            if not country_name:
                country_name = d.get("nationality_name") or ""

            # Sexo: F o M para checkboxes sex_w / sex_m
            is_female = "femenino" in gender_name.lower() or "f" == gender_name.strip().lower()
            is_male = "masculino" in gender_name.lower() or "m" == gender_name.strip().lower()

            # Grado dominio: ["understand","speak","read","write"] (también acepta legacy comprende,habla,lee,escribe)
            nd = d.get("native_language_domain") or []
            ud = d.get("language_used_domain") or []
            if isinstance(nd, str):
                try:
                    nd = _json.loads(nd) if nd else []
                except Exception:
                    nd = []
            if isinstance(ud, str):
                try:
                    ud = _json.loads(ud) if ud else []
                except Exception:
                    ud = []

            def _has(arr, *keys):
                """Retorna Yes si alguno de keys está en arr. keys: ("understand","comprende") para compat."""
                if not arr:
                    return "Off"
                for key in keys:
                    if isinstance(arr, list):
                        for x in arr:
                            v = x.get("value", x) if isinstance(x, dict) else x
                            if str(v).lower().strip() == str(key).lower():
                                return "Yes"
                        if key in arr:
                            return "Yes"
                    elif isinstance(arr, dict):
                        if arr.get(key) or arr.get(str(key).capitalize()):
                            return "Yes"
                return "Off"

            field_mapping = {
                "full_name": nombre_val,
                "student_full_name": nombre_val,
                "birth_day": fecha_nac_val,
                "born_date": fecha_nac_val,
                "sex_f": "Yes" if is_female else "Off",
                "sex_m": "Yes" if is_male else "Off",
                "years": age_years,
                "months": age_months,
                "country": country_name,
                "address": get_val("address"),
                "phone": get_val("phone"),
                "domicilio": get_val("address"),
                "telefono": get_val("phone"),
                "Domicilio actual": get_val("address"),
                "Teléfono": get_val("phone"),
                "domicilio_actual": get_val("address"),
                "direccion": get_val("address"),
                "Address": get_val("address"),
                "Phone": get_val("phone"),
                "Telefono": get_val("phone"),
                "celular": get_val("phone"),
                "mother_language": get_val("native_language"),
                "used_language": get_val("language_used"),
                "course": get_val("current_schooling"),
                "school": get_val("school_name"),
                "nombre": nombre_val,
                "sexo": gender_name,
                "fecha_nacimiento": fecha_nac_val,
                "ml_o_1": _has(nd, "understand", "comprende"),
                "ml_o_01": _has(nd, "understand", "comprende"),
                "comprende": _has(nd, "understand", "comprende"),
                "ml_o_2": _has(nd, "speak", "habla"),
                "ml_o 2": _has(nd, "speak", "habla"),
                "ml_o_02": _has(nd, "speak", "habla"),
                "habla": _has(nd, "speak", "habla"),
                "mlo2": _has(nd, "speak", "habla"),
                "ml_o_3": _has(nd, "read", "lee"),
                "ml_o_4": _has(nd, "write", "escribe"),
                "lengua_materna_comprende": _has(nd, "understand", "comprende"),
                "lengua_materna_habla": _has(nd, "speak", "habla"),
                "lengua_materna_lee": _has(nd, "read", "lee"),
                "lengua_materna_escribe": _has(nd, "write", "escribe"),
                "grado_dominio_comprende": _has(nd, "understand", "comprende"),
                "grado_dominio_habla": _has(nd, "speak", "habla"),
                "grado_dominio_lee": _has(nd, "read", "lee"),
                "grado_dominio_escribe": _has(nd, "write", "escribe"),
                "grado_dominio_ml_comprende": _has(nd, "understand", "comprende"),
                "grado_dominio_ml_habla": _has(nd, "speak", "habla"),
                "grado_dominio_ml_lee": _has(nd, "read", "lee"),
                "grado_dominio_ml_escribe": _has(nd, "write", "escribe"),
                "ul_o_1": _has(ud, "understand", "comprende"),
                "ul_o_2": _has(ud, "speak", "habla"),
                "ul_o_3": _has(ud, "read", "lee"),
                "ul_o_4": _has(ud, "write", "escribe"),
                "ul_comprende": _has(ud, "understand", "comprende"),
                "ul_habla": _has(ud, "speak", "habla"),
                "ul_lee": _has(ud, "read", "lee"),
                "ul_escribe": _has(ud, "write", "escribe"),
                "grado_dominio_ul_comprende": _has(ud, "understand", "comprende"),
                "grado_dominio_ul_habla": _has(ud, "speak", "habla"),
                "grado_dominio_ul_lee": _has(ud, "read", "lee"),
                "grado_dominio_ul_escribe": _has(ud, "write", "escribe"),
                "used_comprende": _has(ud, "understand", "comprende"),
                "used_habla": _has(ud, "speak", "habla"),
                "used_lee": _has(ud, "read", "lee"),
                "used_escribe": _has(ud, "write", "escribe"),
                "age": get_val("age"),
                "Edad": get_val("age"),
                "gender_id": get_val("gender_id"),
                "address": get_val("address"),
                "phone": get_val("phone"),
                "native_language": get_val("native_language"),
                "language_used": get_val("language_used"),
                "current_schooling": get_val("current_schooling"),
                "school_name": get_val("school_name"),
                "interview_reason": get_val("interview_reason"),
                "problem_definition": get_val("interview_reason"),
                "Motivo de la entrevista": get_val("interview_reason"),
                "Motivo entrevista": get_val("interview_reason"),
                "diagnosis_detail": get_val("diagnosis_detail"),
                "diagnosisdetail": get_val("diagnosis_detail"),
                "Detalle diagnóstico": get_val("diagnosis_detail"),
                "specialists": get_val("specialists"),
            }
            # Sección 5: Especialistas por área (desde specialists dict o list)
            _spec_raw = d.get("specialists")
            _spec = _spec_raw
            if isinstance(_spec_raw, str):
                try:
                    _spec = _json.loads(_spec_raw) if _spec_raw else {}
                except Exception:
                    _spec = {}
            if isinstance(_spec, list):
                _name_to_key = {
                    "pediatra": "pediatrics", "pediatría": "pediatrics", "pediatrics": "pediatrics",
                    "kinesiología": "kinesiology", "kinesiology": "kinesiology",
                    "genética": "genetic", "genetic": "genetic",
                    "fonoaudiología": "speech_therapy", "speech_therapy": "speech_therapy",
                    "neurología": "neurology", "neurology": "neurology",
                    "psicología": "psychology", "psychology": "psychology",
                    "psiquiatría": "psychiatry", "psychiatry": "psychiatry",
                    "psicopedagogía": "educational_psychology", "psychopedagogia": "educational_psychology",
                    "terapia ocupacional": "occupational_therapy", "occupational_therapy": "occupational_therapy",
                    "otro": "another_health", "another_health": "another_health",
                }
                _spec_dict = {}
                for _item in _spec:
                    if isinstance(_item, dict):
                        _sval = _item.get("status") or _item.get("detail") or _item.get("value") or _item.get("text") or ""
                        _name = (_item.get("name") or _item.get("specialty") or _item.get("speciality") or _item.get("area") or "").strip().lower()
                        _skey = _name_to_key.get(_name) or _name.replace(" ", "_").replace("á", "a").replace("í", "i").replace("ó", "o").replace("ú", "u").replace("é", "e")
                        if not _skey:
                            continue
                        if _skey in _name_to_key:
                            _skey = _name_to_key[_skey]
                        _spec_dict[_skey] = _sval
                _spec = _spec_dict
            _spec_keys = ("pediatrics", "kinesiology", "genetic", "speech_therapy", "neurology",
                         "psychology", "psychiatry", "educational_psychology", "occupational_therapy", "another_health")
            _spec_aliases = {
                "pediatrics": ("pediatrics", "pediatria", "pediatría"),
                "kinesiology": ("kinesiology", "kinesiologia", "kinesiología"),
                "genetic": ("genetic", "genetico", "genético"),
                "speech_therapy": ("speech_therapy", "fonoaudiologia", "fonoaudiología"),
                "neurology": ("neurology", "neurologia", "neurología"),
                "psychology": ("psychology", "psicologia", "psicología"),
                "psychiatry": ("psychiatry", "psiquiatria", "psiquiatría"),
                "educational_psychology": ("educational_psychology", "psychopedagogia", "psychopedagogía", "psicopedagogia"),
                "occupational_therapy": ("occupational_therapy", "terapia_ocupacional"),
                "another_health": ("another_health", "otro", "other"),
            }
            if isinstance(_spec, dict):
                for _sk in _spec_keys:
                    v = None
                    for _alias in _spec_aliases.get(_sk, (_sk,)):
                        v = _spec.get(_alias) or _spec.get(_alias.replace("_", " ").title())
                        if v is not None:
                            break
                    field_mapping[_sk] = to_str(v) if v is not None else ""
                for _sk, _aliases in _spec_aliases.items():
                    _v = field_mapping.get(_sk, "")
                    for _al in _aliases:
                        field_mapping[_al] = _v
            else:
                for _sk in _spec_keys:
                    field_mapping[_sk] = ""
                for _sk, _aliases in _spec_aliases.items():
                    for _al in _aliases:
                        field_mapping[_al] = ""
            field_mapping.update({
                "first_year_notes": get_val("first_year_notes"),
                "birth_type_id": get_val("birth_type_id"),
                "birth_reason": get_val("birth_reason"),
                "anotherbirthreason": "" if d.get("birth_type_id") == 1 else (get_val("birth_reason") or ""),
                "birth_weight": get_val("birth_weight"),
                "birth_height": get_val("birth_height"),
                "birthheight": get_val("birth_height"),
                "talla_nacimiento": get_val("birth_height"),
                "first_year_observations": get_val("first_year_observations"),
                "sm_head_control": get_val("sm_head_control"),
                "sm_sits_alone": get_val("sm_sits_alone"),
                "sm_walks_without_support": get_val("sm_walks_without_support"),
                "sm_first_words": get_val("sm_first_words"),
                "sm_first_phrases": get_val("sm_first_phrases"),
                "sm_dresses_alone": get_val("sm_dresses_alone"),
                "sm_bladder_day": get_val("sm_bladder_day"),
                "sm_bladder_night": get_val("sm_bladder_night"),
                "sm_bowel_day": get_val("sm_bowel_day"),
                "sm_bowel_night": get_val("sm_bowel_night"),
                "sm_observations_1": get_val("sm_observations_1"),
                "sm_motor_activity": get_val("sm_motor_activity"),
                "sm_muscle_tone": get_val("sm_muscle_tone"),
                "sm_lateral_dominance": get_val("sm_lateral_dominance"),
                "sm_observations_2": get_val("sm_observations_2"),
                "vision_hearing_observations": get_val("vision_hearing_observations"),
                "language_communication_method": get_val("language_communication_method"),
                "language_communication_other": get_val("language_communication_other"),
                "language_oral_loss": get_val("language_oral_loss"),
                "language_observations": get_val("language_observations"),
                "social_observations": get_val("social_observations"),
                "health_problems_treatment": get_val("health_problems_treatment"),
                "health_diet": get_val("health_diet"),
                "health_diet_other": get_val("health_diet_other"),
                "health_weight": get_val("health_weight"),
                "health_sleep_pattern": get_val("health_sleep_pattern"),
                "health_sleep_hours": get_val("health_sleep_hours"),
                "health_sleeps_alone": get_val("health_sleeps_alone"),
                "health_sleeps_specify": get_val("health_sleeps_specify"),
                "health_mood_behavior": get_val("health_mood_behavior"),
                "health_mood_other": get_val("health_mood_other"),
                "health_current_observations": get_val("health_current_observations"),
                "family_health_history": get_val("family_health_history"),
                "family_health_observations": get_val("family_health_observations"),
                "school_entry_age": get_val("school_entry_age"),
                "schools_count": get_val("schools_count"),
                "teaching_modality": get_val("teaching_modality"),
                "changes_reason": get_val("changes_reason"),
                "repeated_courses": get_val("repeated_courses"),
                "repeated_reason": get_val("repeated_reason"),
                "current_level": get_val("current_level"),
                "family_attitude": get_val("family_attitude"),
                "performance_assessment": get_val("performance_assessment"),
                "performance_reasons": get_val("performance_reasons"),
                "response_difficulties_other": get_val("response_difficulties_other"),
                "response_success_other": get_val("response_success_other"),
                "rewards_other": get_val("rewards_other"),
                "supporters_other_professionals": get_val("supporters_other_professionals"),
                "native_language_domain": get_val("native_language_domain"),
                "language_used_domain": get_val("language_used_domain"),
                "first_year_conditions": get_val("first_year_conditions"),
                "response_difficulties": get_val("response_difficulties"),
                "response_success": get_val("response_success"),
                "rewards": get_val("rewards"),
                "supporters": get_val("supporters"),
                "expectations": get_val("expectations"),
                "environment": get_val("environment"),
                "nationality_id": get_val("nationality_id"),
            })
            # Informantes, entrevistadores, miembros del hogar como texto
            def _join_informants():
                lst = d.get("informants") or []
                return "; ".join(
                    f"{x.get('name', '')} ({x.get('relationship', '')}) - {x.get('presence', '')}"
                    for x in lst if isinstance(x, dict)
                )
            def _join_interviewers():
                lst = d.get("interviewers") or []
                return "; ".join(
                    f"{x.get('role', '')}" for x in lst if isinstance(x, dict)
                )
            def _join_household():
                lst = d.get("household_members") or []
                return "; ".join(
                    f"{x.get('name', '')} ({x.get('relationship', '')}) - {x.get('age', '')} - {x.get('schooling', '')}"
                    for x in lst if isinstance(x, dict)
                )
            field_mapping["informants"] = _join_informants()
            field_mapping["Informantes"] = _join_informants()
            # Informantes individuales (sección 2): Fecha entrevista, Nombre, Relación, En presencia de
            inf_list = d.get("informants") or []
            for i in range(4):
                idx = i + 1
                inf = inf_list[i] if i < len(inf_list) and isinstance(inf_list[i], dict) else {}
                date_val = format_date(inf.get("interview_date"))
                name_val = to_str(inf.get("name"))
                rel_val = to_str(inf.get("relationship"))
                pres_val = to_str(inf.get("presence"))
                field_mapping[f"informant_interview_date_{idx}"] = date_val
                field_mapping[f"informant_interview_name_{idx}"] = name_val
                field_mapping[f"informant_relation_with_student_{idx}"] = rel_val
                field_mapping[f"informant_presence_{idx}"] = pres_val
                field_mapping[f"informant_{idx}_interview_date"] = date_val
                field_mapping[f"informant_{idx}_name"] = name_val
                field_mapping[f"informant_{idx}_relation"] = rel_val
                field_mapping[f"informant_{idx}_presence"] = pres_val
                field_mapping[f"inf{idx}_fecha"] = date_val
                field_mapping[f"inf{idx}_nombre"] = name_val
                field_mapping[f"inf{idx}_relacion"] = rel_val
                field_mapping[f"inf{idx}_presencia"] = pres_val
                field_mapping[f"informante_{idx}_fecha"] = date_val
                field_mapping[f"informante_{idx}_nombre"] = name_val
                field_mapping[f"informante_{idx}_relacion"] = rel_val
                field_mapping[f"informante_{idx}_presencia"] = pres_val
            _interviewers_str = _join_interviewers()
            field_mapping["interviewers"] = _interviewers_str
            field_mapping["Entrevistadores"] = _interviewers_str
            # Sección 3: Entrevistadores individuales (fecha, nombre, rol/cargo)
            int_list = d.get("interviewers") or []
            for i in range(4):
                idx = i + 1
                int_item = int_list[i] if i < len(int_list) and isinstance(int_list[i], dict) else {}
                int_date = format_date(int_item.get("interview_date"))
                int_name = to_str(int_item.get("name") or int_item.get("role"))
                int_role = to_str(int_item.get("role"))
                field_mapping[f"identification_interview_date_{idx}"] = int_date
                field_mapping[f"identification_interview_name_{idx}"] = int_name
                field_mapping[f"job_position_{idx}"] = int_role
            field_mapping["household_members"] = _join_household()
            field_mapping["Miembros del hogar"] = _join_household()
            # Checkboxes 1=Sí -> /Yes, otro -> /Off
            _ck = (
                "diagnosis_has", "birth_medical_assistance", "first_year_periodic_health_checkups", "first_year_vaccines",
                "sm_walking_stability", "sm_frequent_falls", "vision_interested_stimuli", "vision_irritated_eyes",
                "vision_headaches", "vision_squints", "vision_follows_movement", "vision_abnormal_movements",
                "vision_erroneous_behaviors", "vision_diagnosis", "hearing_interested_stimuli", "hearing_recognizes_voices",
                "hearing_turns_head", "hearing_ears_to_tv", "hearing_covers_ears", "hearing_earaches",
                "hearing_pronunciation_adequate", "hearing_diagnosis", "health_vaccines_up_to_date", "health_epilepsy",
                "health_heart_problems", "health_paraplegia", "health_hearing_loss", "health_vision_loss",
                "health_motor_disorder", "health_bronchorespiratory", "health_infectious_disease", "health_emotional_disorder",
                "health_behavioral_disorder", "health_other", "health_sleep_insomnia", "health_sleep_nightmares",
                "health_sleep_terrors", "health_sleep_sleepwalking", "health_sleep_good_mood", "attended_kindergarten",
                "learning_difficulty", "participation_difficulty", "disruptive_behavior", "attends_regularly",
                "attends_gladly", "family_support_homework", "friends",
                "language_exp_babbles", "language_exp_vocalizes_gestures", "language_exp_emits_words",
                "language_exp_emits_phrases", "language_exp_relates_experiences", "language_exp_clear_pronunciation",
                "language_comp_identifies_objects", "language_comp_identifies_people", "language_comp_understands_abstract",
                "language_comp_responds_coherently", "language_comp_follows_simple_instructions",
                "language_comp_follows_complex_instructions", "language_comp_follows_group_instructions",
                "language_comp_understands_stories", "social_relates_spontaneously", "social_explains_behaviors",
                "social_participates_groups", "social_prefers_individual", "social_echolalic_language",
                "social_difficulty_adapting", "social_relates_collaboratively", "social_respects_social_norms",
                "social_respects_school_norms", "social_shows_humor", "social_stereotyped_movements",
                "social_frequent_tantrums", "sm_fine_grab", "sm_fine_grip", "sm_fine_pinch", "sm_fine_draw",
                "sm_fine_write", "sm_fine_thread", "sm_cog_reacts_familiar", "sm_cog_demands_company",
                "sm_cog_smiles_babbles", "sm_cog_manipulates_explores", "sm_cog_understands_prohibitions",
                "sm_cog_poor_eye_hand",
            )
            for k in _ck:
                v = d.get(k)
                if v is not None:
                    field_mapping[k] = "Yes" if v == 1 else "Off"
            # Diagnóstico previo: pdhy (Sí), pdhn (No) - Yes para activar (igual que sex_f/sex_m)
            _dh = d.get("diagnosis_has")
            field_mapping["pdhy"] = "Yes" if _dh == 1 else "Off"
            field_mapping["pdhn"] = "Yes" if _dh == 2 else "Off"
            # Tipo de parto: birth_type_id 1=normal, 2=cesárea, 3=asistido, 4=fórceps
            _bt = d.get("birth_type_id")
            field_mapping["normal_birth"] = "Yes" if _bt == 1 else "Off"
            field_mapping["normalbirth"] = "Yes" if _bt == 1 else "Off"
            field_mapping["cesarean"] = "Yes" if _bt == 2 else "Off"
            field_mapping["assisted"] = "Yes" if _bt == 3 else "Off"
            field_mapping["forceps"] = "Yes" if _bt == 4 else "Off"
            # Asistencia médica en el parto: birth_medical_assistance 1=Sí, 2=No
            _bma = d.get("birth_medical_assistance")
            field_mapping["birth_assistant_y"] = "Yes" if _bma == 1 else "Off"
            field_mapping["birth_assistant_n"] = "Yes" if _bma == 2 else "Off"
            # Desnutrición (first_year_conditions.desnutricion) 1=Sí, 2=No
            _fyc = d.get("first_year_conditions") or {}
            if isinstance(_fyc, str):
                try:
                    _fyc = _json.loads(_fyc) if _fyc else {}
                except Exception:
                    _fyc = {}
            _des = _fyc.get("desnutricion") if isinstance(_fyc, dict) else None
            _des_y = "Yes" if _des == 1 else "Off"
            _des_n = "Yes" if _des == 2 else "Off"
            field_mapping["dy"] = _des_y
            field_mapping["dn"] = _des_n
            _obs = _fyc.get("obesidad") if isinstance(_fyc, dict) else None
            _obs_y, _obs_n = "Yes" if _obs == 1 else "Off", "Yes" if _obs == 2 else "Off"
            field_mapping["obesityy"] = _obs_y
            field_mapping["obesityn"] = _obs_n
            _hf = _fyc.get("fiebre_alta") if isinstance(_fyc, dict) else None
            _hf_y, _hf_n = "Yes" if _hf == 1 else "Off", "Yes" if _hf == 2 else "Off"
            field_mapping["hfy"] = _hf_y
            field_mapping["hfn"] = _hf_n
            _sz = _fyc.get("convulsiones") if isinstance(_fyc, dict) else None
            _sz_y, _sz_n = "Yes" if _sz == 1 else "Off", "Yes" if _sz == 2 else "Off"
            field_mapping["seizuresy"] = _sz_y
            field_mapping["seizuresn"] = _sz_n
            _hosp = _fyc.get("hospitalizaciones") if isinstance(_fyc, dict) else None
            _hosp_y, _hosp_n = "Yes" if _hosp == 1 else "Off", "Yes" if _hosp == 2 else "Off"
            field_mapping["hospitalizationsy"] = _hosp_y
            field_mapping["hospitalizationsn"] = _hosp_n
            _inj = _fyc.get("traumatismos") if isinstance(_fyc, dict) else None
            _inj_y, _inj_n = "Yes" if _inj == 1 else "Off", "Yes" if _inj == 2 else "Off"
            field_mapping["injuy"] = _inj_y
            field_mapping["injun"] = _inj_n
            _intox = _fyc.get("intoxicacion") if isinstance(_fyc, dict) else None
            _intox_y, _intox_n = "Yes" if _intox == 1 else "Off", "Yes" if _intox == 2 else "Off"
            field_mapping["iny"] = _intox_y
            field_mapping["inn"] = _intox_n
            _asma = _fyc.get("asma") if isinstance(_fyc, dict) else None
            _asma_y, _asma_n = "Yes" if _asma == 1 else "Off", "Yes" if _asma == 2 else "Off"
            field_mapping["ay"] = _asma_y
            field_mapping["an"] = _asma_n
            _resp = _fyc.get("enfermedad_respiratoria") if isinstance(_fyc, dict) else None
            _resp_y, _resp_n = "Yes" if _resp == 1 else "Off", "Yes" if _resp == 2 else "Off"
            field_mapping["riy"] = _resp_y
            field_mapping["rin"] = _resp_n
            _enc = _fyc.get("encefalitis") if isinstance(_fyc, dict) else None
            _enc_y, _enc_n = "Yes" if _enc == 1 else "Off", "Yes" if _enc == 2 else "Off"
            field_mapping["encpy"] = _enc_y
            field_mapping["encpn"] = _enc_n
            _men = _fyc.get("meningitis") if isinstance(_fyc, dict) else None
            _men_y, _men_n = "Yes" if _men == 1 else "Off", "Yes" if _men == 2 else "Off"
            field_mapping["mgy"] = _men_y
            field_mapping["mgn"] = _men_n
            _otro_txt = (_fyc.get("otras") if isinstance(_fyc, dict) else None) or d.get("first_year_conditions_other_specify")
            field_mapping["otro"] = to_str(_otro_txt) if _otro_txt else ""
            _otro = _fyc.get("otro") if isinstance(_fyc, dict) else None
            _otro_y, _otro_n = "Yes" if _otro == 1 else "Off", "Yes" if _otro == 2 else "Off"
            field_mapping["oy"] = _otro_y
            field_mapping["on"] = _otro_n

            # Construir form_data: solo incluir campos que existen en el PDF
            form_data = {}
            norm_map = {k.strip().lower(): k for k in field_mapping.keys()}
            pdf_field_names = list(form_fields.keys())
            for fn in form_fields.keys():
                val = None
                if fn in field_mapping:
                    val = field_mapping[fn]
                elif fn.strip().lower() in norm_map:
                    val = field_mapping[norm_map[fn.strip().lower()]]
                if val is None and "." in fn:
                    short = fn.split(".")[-1].strip()
                    if short in field_mapping:
                        val = field_mapping[short]
                    elif short.strip().lower() in norm_map:
                        val = field_mapping[norm_map[short.strip().lower()]]
                if val is None:
                    fl = fn.strip().lower()
                    if ("nombre" in fl and "completo" in fl) or ("nombre" in fl and "apellido" in fl and "2" not in fn) or fl == "nombre":
                        val = get_val("student_full_name")
                    elif "sexo" in fl or "genero" in fl or "género" in fl:
                        val = gender_name
                    elif ("fecha" in fl and "nacimiento" in fl) or fl == "bd" or fl == "fecha_nacimiento":
                        val = format_date(d.get("born_date"))
                    elif "domicilio" in fl or "direccion" in fl or "address" in fl or "dirección" in fl:
                        val = get_val("address")
                    elif "telefono" in fl or "teléfono" in fl or "phone" in fl or "celular" in fl:
                        val = get_val("phone")
                    elif "informant" in fl or "informante" in fl or "inf" in fl:
                        import re
                        m = re.search(r'[_\s]?(\d)[_\s]', fn) or re.search(r'(\d)\s*$', fn)
                        idx_match = int(m.group(1)) if m and m.group(1).isdigit() else 1
                        inf = (d.get("informants") or [{}])[idx_match - 1] if idx_match <= 4 else {}
                        inf = inf if isinstance(inf, dict) else {}
                        if "date" in fl or "fecha" in fl:
                            val = format_date(inf.get("interview_date"))
                        elif "name" in fl or "nombre" in fl:
                            val = to_str(inf.get("name"))
                        elif "relation" in fl or "relacion" in fl or "relación" in fl:
                            val = to_str(inf.get("relationship"))
                        elif "presence" in fl or "presencia" in fl:
                            val = to_str(inf.get("presence"))
                    elif "motivo" in fl and "entrevista" in fl:
                        val = get_val("interview_reason")
                    elif ("antecedentes" in fl and ("embarazo" in fl or "parto" in fl)) or "first_year_notes" in fl:
                        val = get_val("first_year_notes")
                    elif ("birth" in fl and "height" in fl) or ("talla" in fl and ("nacimiento" in fl or "birth" in fl or "parto" in fl)) or fl == "birth_height":
                        val = get_val("birth_height")
                    elif "edad" in fl or fl == "age":
                        val = get_val("age")
                    elif "direccion" in fl or "dirección" in fl:
                        val = get_val("address")
                    elif "telefono" in fl or "teléfono" in fl:
                        val = get_val("phone")
                    elif "lengua" in fl and "materna" in fl:
                        if "comprende" in fl or "comp" in fl:
                            val = _has(nd, "comprende")
                        elif "habla" in fl:
                            val = _has(nd, "habla")
                        elif "lee" in fl and "escribe" not in fl:
                            val = _has(nd, "lee")
                        elif "escribe" in fl:
                            val = _has(nd, "escribe")
                        else:
                            val = get_val("native_language")
                    elif ("ml_o" in fl or "ml_o" in fn) and "ul" not in fl:
                        if "1" in fl and "2" not in fl and "3" not in fl and "4" not in fl:
                            val = _has(nd, "comprende")
                        elif "2" in fl or "02" in fl:
                            val = _has(nd, "habla")
                        elif "3" in fl:
                            val = _has(nd, "lee")
                        elif "4" in fl:
                            val = _has(nd, "escribe")
                    elif ("comprende" in fl or "habla" in fl or "lee" in fl or "escribe" in fl) and ("ml" in fl or "materna" in fl):
                        if "comprende" in fl or "comp" in fl:
                            val = _has(nd, "comprende")
                        elif "habla" in fl:
                            val = _has(nd, "habla")
                        elif "lee" in fl and "escribe" not in fl:
                            val = _has(nd, "lee")
                        elif "escribe" in fl:
                            val = _has(nd, "escribe")
                    elif ("grado" in fl and "dominio" in fl) or "grado_dominio" in fl:
                        if "ul" in fl or "uso" in fl or "lengua_uso" in fl:
                            if "comprende" in fl or "comp" in fl:
                                val = _has(ud, "comprende")
                            elif "habla" in fl:
                                val = _has(ud, "habla")
                            elif "lee" in fl and "escribe" not in fl:
                                val = _has(ud, "lee")
                            elif "escribe" in fl:
                                val = _has(ud, "escribe")
                        else:
                            if "comprende" in fl or "comp" in fl:
                                val = _has(nd, "comprende")
                            elif "habla" in fl:
                                val = _has(nd, "habla")
                            elif "lee" in fl and "escribe" not in fl:
                                val = _has(nd, "lee")
                            elif "escribe" in fl:
                                val = _has(nd, "escribe")
                    elif fl in ("pediatrics", "kinesiology", "genetic", "speech_therapy", "neurology",
                               "psychology", "psychiatry", "educational_psychology", "occupational_therapy", "another_health"):
                        val = field_mapping.get(fl, "")
                    elif "pediatrics" in fl or "pediatri" in fl:
                        val = field_mapping.get("pediatrics", "")
                    elif "kinesiology" in fl or "kinesiolog" in fl:
                        val = field_mapping.get("kinesiology", "")
                    elif "genetic" in fl:
                        val = field_mapping.get("genetic", "")
                    elif "speech_therapy" in fl or "fonoaudiolog" in fl:
                        val = field_mapping.get("speech_therapy", "")
                    elif "neurology" in fl or "neurolog" in fl:
                        val = field_mapping.get("neurology", "")
                    elif "psychology" in fl and "educational" not in fl and "psychiatry" not in fl:
                        val = field_mapping.get("psychology", "")
                    elif "psychiatry" in fl or "psiquiatr" in fl:
                        val = field_mapping.get("psychiatry", "")
                    elif "educational_psychology" in fl or "psychopedagog" in fl:
                        val = field_mapping.get("educational_psychology", "")
                    elif "occupational_therapy" in fl or "terapia ocupacional" in fl:
                        val = field_mapping.get("occupational_therapy", "")
                    elif "another_health" in fl or ("otro" in fl and ("especialista" in fl or "health" in fl)):
                        val = field_mapping.get("another_health", "")
                    elif "traumat" in fl or "injuries" in fl:
                        val = field_mapping.get("injuy") if fl.endswith("y") or "_y" in fl else field_mapping.get("injun")
                    elif "intox" in fl or "poison" in fl:
                        val = field_mapping.get("iny") if fl.endswith("y") or "_y" in fl else field_mapping.get("inn")
                    elif "asma" in fl and ("y" in fl or "n" in fl):
                        val = field_mapping.get("ay") if fl.endswith("y") or "_y" in fl else field_mapping.get("an")
                    elif "encefal" in fl and ("y" in fl or "n" in fl):
                        val = field_mapping.get("encpy") if fl.endswith("y") or "_y" in fl else field_mapping.get("encpn")
                    elif "lengua" in fl and "uso" in fl:
                        if "comprende" in fl or "comp" in fl:
                            val = _has(ud, "comprende")
                        elif "habla" in fl:
                            val = _has(ud, "habla")
                        elif "lee" in fl and "escribe" not in fl:
                            val = _has(ud, "lee")
                        elif "escribe" in fl:
                            val = _has(ud, "escribe")
                        else:
                            val = get_val("language_used")
                if val is not None:
                    if val in ("/Off", "/Yes", "Off", "Yes"):
                        form_data[fn] = val
                    elif isinstance(val, str):
                        form_data[fn] = val
            _ck_vals = ("/Yes", "/Off", "Yes", "Off")
            _ck_matched = [k for k, v in form_data.items() if v in _ck_vals]
            _spec_matched = [k for k in form_data if any(x in k.lower() for x in ("pediatrics", "kinesiology", "genetic", "speech", "neurology", "psychology", "psychiatry", "psychopedagog", "occupational", "another_health"))]
            _gd_matched = [k for k in form_data if "grado" in k.lower() or "dominio" in k.lower() or "ml_o_" in k or "ul_o_" in k or "comprende" in k.lower() or "habla" in k.lower()]

            if form_data:
                text_f = {k: v for k, v in form_data.items() if v not in _ck_vals}
                cb_f = {k: v for k, v in form_data.items() if v in _ck_vals}
                text_norm = {k.strip().lower(): (k, v) for k, v in text_f.items()}
                cb_norm = {k.strip().lower(): v for k, v in cb_f.items()}

                def get_t(n):
                    if n in text_f:
                        return text_f[n]
                    nlow = n.strip().lower()
                    if nlow in text_norm:
                        return text_norm[nlow][1]
                    if "." in n:
                        short = n.split(".")[-1].strip()
                        if short in text_f:
                            return text_f[short]
                        if short.lower() in text_norm:
                            return text_norm[short.lower()][1]
                    for k, v in text_f.items():
                        if k.lower() in nlow or nlow in k.lower():
                            return v
                    return None

                def get_c(n):
                    v = cb_f.get(n) or cb_norm.get((n or "").strip().lower())
                    if v is not None:
                        return v
                    short = n.split(".")[-1].strip() if "." in n else n
                    v = cb_f.get(short) or cb_norm.get(short.lower())
                    if v is not None:
                        return v
                    nnorm = (n or "").strip().lower().replace(" ", "_").replace("-", "_")
                    v = cb_f.get(nnorm) or cb_norm.get(nnorm)
                    if v is not None:
                        return v
                    nlow = (n or "").strip().lower()
                    for k in cb_f:
                        if nlow in k.lower() or k.lower() in nlow or (short and short.lower() in k.lower()):
                            return cb_f[k]
                        knorm = k.lower().replace(" ", "_").replace("-", "_")
                        if nnorm in knorm or knorm in nnorm:
                            return cb_f[k]
                    return None

                try:
                    if hasattr(writer, '_root_object') and writer._root_object and "/AcroForm" in writer._root_object:
                        writer._root_object["/AcroForm"][NameObject("/NeedAppearances")] = BooleanObject(True)
                except Exception:
                    pass

                for page_num in range(len(writer.pages)):
                    page = writer.pages[page_num]
                    if "/Annots" not in page:
                        continue
                    for annot_ref in page["/Annots"]:
                        try:
                            a = annot_ref.get_object()
                            fn = str(a.get("/T", "") or "").strip()
                            ft = str(a.get("/FT", "") or "").strip()
                            if not fn and "/Parent" in a:
                                try:
                                    parent = a["/Parent"].get_object()
                                    fn = str(parent.get("/T", "") or "").strip()
                                    ft = ft or str(parent.get("/FT", "") or "").strip()
                                except Exception:
                                    pass
                            if not fn or not ft:
                                continue
                            if ft == "/Tx":
                                v = get_t(fn)
                                if v is not None:
                                    val_str = TextStringObject(str(v))
                                    a[NameObject("/V")] = val_str
                                    if "/Parent" in a:
                                        try:
                                            parent = a["/Parent"].get_object()
                                            parent[NameObject("/V")] = val_str
                                        except Exception:
                                            pass
                            elif ft == "/Btn":
                                cv = get_c(fn)
                                if cv is not None:
                                    v_btn = NameObject("/Yes" if cv in ("/Yes", "Yes") else "/Off")
                                    a[NameObject("/V")] = v_btn
                                    a[NameObject("/AS")] = v_btn
                                    if "/Parent" in a:
                                        try:
                                            parent = a["/Parent"].get_object()
                                            parent[NameObject("/V")] = v_btn
                                            parent[NameObject("/AS")] = v_btn
                                        except Exception:
                                            pass
                            if fn in form_data:
                                a[NameObject("/Ff")] = NumberObject(int(a.get("/Ff", 0)) | 1)
                        except Exception:
                            continue

                # Actualizar también en AcroForm /Fields (algunos visores leen de ahí)
                try:
                    if hasattr(writer, '_root_object') and writer._root_object and "/AcroForm" in writer._root_object:
                        acro = writer._root_object["/AcroForm"]
                        if "/Fields" in acro and acro["/Fields"]:
                            def set_text_in_fields(fl):
                                if not fl:
                                    return
                                for ref in fl:
                                    try:
                                        obj = ref.get_object()
                                        if "/Kids" in obj and obj["/Kids"]:
                                            set_text_in_fields(obj["/Kids"])
                                        if "/T" in obj and "/FT" in obj:
                                            fn = str(obj["/T"]).strip() if obj["/T"] else ""
                                            ft = str(obj["/FT"]).strip() if obj["/FT"] else ""
                                            if ft == "/Tx":
                                                v = get_t(fn)
                                                if v is not None:
                                                    obj[NameObject("/V")] = TextStringObject(str(v))
                                            elif ft == "/Btn":
                                                cv = get_c(fn)
                                                if cv is not None:
                                                    v_btn = NameObject("/Yes" if cv in ("/Yes", "Yes") else "/Off")
                                                    obj[NameObject("/V")] = v_btn
                                                    obj[NameObject("/AS")] = v_btn
                                    except Exception:
                                        continue
                            set_text_in_fields(acro["/Fields"])
                except Exception:
                    pass

                # Marcar campos rellenados como solo lectura (quitar formulario editable)
                try:
                    if hasattr(writer, '_root_object') and writer._root_object and "/AcroForm" in writer._root_object:
                        acro = writer._root_object["/AcroForm"]
                        if "/Fields" in acro and acro["/Fields"] and isinstance(acro["/Fields"], list):
                            def process_readonly(fl):
                                if not fl:
                                    return
                                for ref in fl:
                                    try:
                                        if ref is None:
                                            continue
                                        obj = ref.get_object()
                                        if "/Kids" in obj and obj["/Kids"]:
                                            process_readonly(obj["/Kids"])
                                        if "/T" in obj:
                                            fn = str(obj["/T"]).strip() if obj["/T"] else ""
                                            if fn in form_data:
                                                if "/Ff" in obj:
                                                    obj[NameObject("/Ff")] = NumberObject(int(obj["/Ff"]) | 1)
                                                else:
                                                    obj[NameObject("/Ff")] = NumberObject(1)
                                    except Exception:
                                        continue
                            process_readonly(acro["/Fields"])
                except Exception:
                    pass

            with open(output_file, "wb") as f:
                writer.write(f)

            return {
                "status": "success",
                "message": "PDF de anamnesis generado exitosamente",
                "filename": unique_filename,
                "file_path": str(output_file)
            }
        except Exception as e:
            import traceback
            traceback.print_exc()
            return {
                "status": "error",
                "message": f"Error generando PDF de anamnesis: {str(e)}",
                "filename": None,
                "file_path": None
            }

    @staticmethod
    def generate_progress_status_pdf(
        document_id: int,
        progress_status_data: Dict[str, Any],
        db: Optional[Session] = None,
        template_path: Optional[str] = None,
        output_directory: str = "files/system/students"
    ) -> Dict[str, Any]:
        """
        Genera un PDF de estado de avance (documento 18).
        Dependiendo del document_id:
        - Si hay template_path: Sustituye datos en el template PDF
        - Si no hay template_path o document_id específico: Crea el PDF desde cero
        
        Args:
            document_id: ID del documento (18 para estado de avance)
            progress_status_data: Diccionario con los datos del estado de avance
            db: Sesión de base de datos (opcional)
            template_path: Ruta del archivo PDF template (opcional)
            output_directory: Directorio donde se guardará el archivo generado
        
        Returns:
            dict: Diccionario con status, message, filename y file_path
        """
        # El método generate_progress_status_pdf ya existe más abajo
        # Este código duplicado debe eliminarse
        pass

    @staticmethod
    def generate_progress_status_pdf(
        document_id: int,
        progress_status_data: Dict[str, Any],
        db: Optional[Session] = None,
        template_path: Optional[str] = None,
        output_directory: str = "files/system/students"
    ) -> Dict[str, Any]:
        """
        Genera un PDF de estado de avance (documento 18).
        Dependiendo del document_id:
        - Si hay template_path: Sustituye datos en el template PDF
        - Si no hay template_path o document_id específico: Crea el PDF desde cero
        
        Args:
            document_id: ID del documento (18 para estado de avance)
            progress_status_data: Diccionario con los datos del estado de avance
            db: Sesión de base de datos (opcional)
            template_path: Ruta del archivo PDF template (opcional)
            output_directory: Directorio donde se guardará el archivo generado
        
        Returns:
            dict: Diccionario con status, message, filename y file_path
        """
        # Este método debe delegar a _generate_pdf_from_template o _generate_pdf_from_scratch
        # según corresponda. El código duplicado de tag replacement fue eliminado.
        if template_path:
            return DocumentsClass._generate_pdf_from_template(
                template_path=template_path,
                document_data=progress_status_data,
                output_directory=output_directory
            )
        else:
            return DocumentsClass._generate_progress_status_from_scratch(
                progress_status_data=progress_status_data,
                db=db,
                output_directory=output_directory
            )

    @staticmethod
    def generate_progress_status_pdf(
        document_id: int,
        progress_status_data: Dict[str, Any],
        db: Optional[Session] = None,
        template_path: Optional[str] = None,
        output_directory: str = "files/system/students"
    ) -> Dict[str, Any]:
        """
        Genera un PDF de estado de avance (documento 18).
        Dependiendo del document_id:
        - Si hay template_path: Sustituye datos en el template PDF
        - Si no hay template_path o document_id específico: Crea el PDF desde cero
        
        Args:
            document_id: ID del documento (18 para estado de avance)
            progress_status_data: Diccionario con los datos del estado de avance
            db: Sesión de base de datos (opcional)
            template_path: Ruta del archivo PDF template (opcional)
            output_directory: Directorio donde se guardará el archivo generado
        
        Returns:
            dict: Diccionario con status, message, filename y file_path
        """
        # Este método debe delegar a _generate_pdf_from_template o _generate_pdf_from_scratch
        # según corresponda. El código duplicado de tag replacement fue eliminado.
        if template_path:
            return DocumentsClass._generate_pdf_from_template(
                template_path=template_path,
                document_data=progress_status_data,
                output_directory=output_directory
            )
        else:
            return DocumentsClass._generate_progress_status_from_scratch(
                progress_status_data=progress_status_data,
                db=db,
                output_directory=output_directory
            )

    @staticmethod
    def generate_progress_status_pdf(
        document_id: int,
        progress_status_data: Dict[str, Any],
        db: Optional[Session] = None,
        template_path: Optional[str] = None,
        output_directory: str = "files/system/students"
    ) -> Dict[str, Any]:
        """
        Genera un PDF de estado de avance (documento 18).
        Dependiendo del document_id:
        - Si hay template_path: Sustituye datos en el template PDF
        - Si no hay template_path o document_id específico: Crea el PDF desde cero
        
        Args:
            document_id: ID del documento (determina el método de generación)
            progress_status_data: Diccionario con los datos del estado de avance
            db: Sesión de base de datos (opcional, para obtener datos relacionados)
            template_path: Ruta del archivo PDF template (opcional, si es None se crea desde cero)
            output_directory: Directorio donde se guardará el archivo generado
        
        Returns:
            dict: Diccionario con status, message, filename y file_path
        """
        try:
            # Si hay template_path, usar método de sustitución
            if template_path:
                return DocumentsClass._generate_progress_status_from_template(
                    template_path=template_path,
                    progress_status_data=progress_status_data,
                    output_directory=output_directory
                )
            else:
                # Si no hay template, crear desde cero usando ReportLab
                return DocumentsClass._generate_progress_status_from_scratch(
                    document_id=document_id,
                    progress_status_data=progress_status_data,
                    db=db,
                    output_directory=output_directory
                )
        except Exception as e:
            return {
                "status": "error",
                "message": f"Error generando PDF: {str(e)}",
                "filename": None,
                "file_path": None
            }
    
    @staticmethod
    def _generate_progress_status_from_template(
        template_path: str,
        progress_status_data: Dict[str, Any],
        output_directory: str = "files/system/students"
    ) -> Dict[str, Any]:
        """
        Genera un PDF de estado de avance sustituyendo datos en un template PDF existente.
        """
        try:
            template_file = Path(template_path)
            
            if not template_file.exists():
                return {
                    "status": "error",
                    "message": "Template PDF no encontrado",
                    "filename": None,
                    "file_path": None
                }
            
            # Validar que sea PDF
            if template_file.suffix.lower() != '.pdf':
                return {
                    "status": "error",
                    "message": f"El template debe ser un archivo PDF, se encontró: {template_file.suffix}",
                    "filename": None,
                    "file_path": None
                }
            
            # Crear nombre único para el archivo generado
            student_name = progress_status_data.get("student_fullname", "estudiante").replace(" ", "_")
            unique_filename = f"estado_avance_{student_name}_{uuid.uuid4().hex[:8]}.pdf"
            output_file = Path(output_directory) / unique_filename
            
            # Asegurar que el directorio de salida existe
            output_file.parent.mkdir(parents=True, exist_ok=True)
            
            # Abrir el archivo PDF
            pdf_document = fitz.open(template_file)
            
            # Función auxiliar para formatear fechas
            def format_date(date_str: Optional[str]) -> str:
                if not date_str:
                    return ""
                try:
                    if isinstance(date_str, str):
                        date_obj = datetime.strptime(date_str, "%Y-%m-%d").date()
                        return date_obj.strftime("%d/%m/%Y")
                    return ""
                except:
                    return str(date_str) if date_str else ""
            
            # Función auxiliar para obtener valor o string vacío
            def get_value(key: str, default: str = "") -> str:
                value = progress_status_data.get(key)
                if value is None:
                    return default
                return str(value) if value else default
            
            # Mapeo de etiquetas a valores
            tag_replacements = {
                "[STUDENT_FULLNAME]": get_value("student_fullname", ""),
                "[STUDENT_NAME]": get_value("student_name", ""),
                "[STUDENT_LASTNAME]": get_value("student_lastname", ""),
                "[PROGRESS_DATE]": format_date(get_value("progress_date", "")),
                "[PEDAGOGICAL_LANGUAGE]": get_value("pedagogical_language", ""),
                "[PEDAGOGICAL_MATHEMATICS]": get_value("pedagogical_mathematics", ""),
                "[PSYCHOPEDAGOGICAL]": get_value("psychopedagogical", ""),
                "[SPEECH_THERAPY]": get_value("speech_therapy", ""),
                "[PSYCHOLOGICAL]": get_value("psychological", ""),
                "[KINESIOLOGY]": get_value("kinesiology", ""),
                "[OCCUPATIONAL_THERAPY]": get_value("occupational_therapy", ""),
                "[DEAF_CO_EDUCATOR]": get_value("deaf_co_educator", ""),
                "[SYNTHESIS_COMMENTS]": get_value("synthesis_comments", ""),
                "[SUGGESTIONS_FAMILY]": get_value("suggestions_family", ""),
                "[SUGGESTIONS_ESTABLISHMENT]": get_value("suggestions_establishment", ""),
                "[GUARDIAN_NAME]": get_value("guardian_name", ""),
                "[GUARDIAN_LASTNAME]": get_value("guardian_lastname", ""),
                "[RESPONSIBLE_PROFESSIONALS]": get_value("responsible_professionals", ""),
            }
            
            # Procesar cada página del PDF
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                
                # Obtener todo el texto de la página en formato dict para búsqueda más precisa
                page_dict = None
                try:
                    page_dict = page.get_text("dict")
                except:
                    pass
                
                # Reemplazar cada tag
                for tag, value in tag_replacements.items():
                    # Siempre reemplazar, incluso si el valor está vacío (dejará el campo vacío)
                    text_instances = []
                    
                    # Intentar diferentes variaciones del tag
                    tag_variations = [
                        tag,  # Tag original con corchetes
                        tag.upper(),  # Mayúsculas
                        tag.lower(),  # Minúsculas
                    ]
                    
                    # Buscar cada variación usando search_for
                    for tag_variant in tag_variations:
                        try:
                            # Buscar con flags de case insensitive y dehyphenate
                            found = page.search_for(tag_variant, flags=fitz.TEXT_DEHYPHENATE)
                            if found:
                                text_instances.extend(found)
                                break  # Si se encuentra, no buscar más variaciones con search_for
                        except:
                            try:
                                # Si falla con flags, intentar sin flags
                                found = page.search_for(tag_variant)
                                if found:
                                    text_instances.extend(found)
                                    break
                            except:
                                pass
                    
                    # Si aún no se encuentra, buscar usando get_text("dict") para posición exacta
                    if not text_instances and page_dict:
                        try:
                            # Buscar en bloques de texto
                            for block in page_dict.get("blocks", []):
                                if "lines" in block:
                                    for line in block["lines"]:
                                        full_line_text = ""
                                        line_spans = []
                                        for span in line.get("spans", []):
                                            span_text = span.get("text", "")
                                            full_line_text += span_text
                                            line_spans.append((span, span_text))
                                        
                                        # Buscar tag en el texto completo de la línea (case insensitive)
                                        full_line_text_upper = full_line_text.upper()
                                        for tag_variant in tag_variations:
                                            tag_variant_upper = tag_variant.upper()
                                            # Buscar con case insensitive
                                            if tag_variant_upper in full_line_text_upper:
                                                # Encontrar qué span(s) contiene(n) el tag
                                                matching_spans = []
                                                for span, span_text in line_spans:
                                                    span_text_upper = span_text.upper()
                                                    if tag_variant_upper in span_text_upper:
                                                        matching_spans.append(span)
                                                
                                                if matching_spans:
                                                    # Calcular el rectángulo que cubre todos los spans que contienen el tag
                                                    x0 = min(s.get("bbox", [0, 0, 0, 0])[0] for s in matching_spans)
                                                    y0 = min(s.get("bbox", [0, 0, 0, 0])[1] for s in matching_spans)
                                                    x1 = max(s.get("bbox", [0, 0, 0, 0])[2] for s in matching_spans)
                                                    y1 = max(s.get("bbox", [0, 0, 0, 0])[3] for s in matching_spans)
                                                    rect = fitz.Rect(x0 - 2, y0 - 2, x1 + 2, y1 + 2)
                                                    text_instances.append(rect)
                                                break
                        except Exception as e:
                            pass
                    
                    # Reemplazar todas las instancias encontradas
                    if text_instances:
                        # Eliminar duplicados manteniendo el orden
                        seen_rects = set()
                        unique_instances = []
                        page_width = page.rect.width
                        page_height = page.rect.height
                        
                        for rect in text_instances:
                            # Convertir a Rect si es una tupla
                            if isinstance(rect, (list, tuple)) and len(rect) == 4:
                                rect = fitz.Rect(rect[0], rect[1], rect[2], rect[3])
                            
                            # Validar que el rectángulo esté dentro de los límites de la página
                            if (rect.x0 >= 0 and rect.y0 >= 0 and 
                                rect.x1 <= page_width and rect.y1 <= page_height and
                                rect.x1 > rect.x0 and rect.y1 > rect.y0):
                                
                                rect_key = (round(rect.x0, 1), round(rect.y0, 1), round(rect.x1, 1), round(rect.y1, 1))
                                if rect_key not in seen_rects:
                                    seen_rects.add(rect_key)
                                    unique_instances.append(rect)
                        
                        # Reemplazar cada instancia única
                        for rect in unique_instances:
                            # Validar el rectángulo antes de aplicar redacción
                            if rect.x1 > rect.x0 and rect.y1 > rect.y0:
                                # Redactar el texto original (cubrir con blanco)
                                page.add_redact_annot(rect, fill=(1, 1, 1))
                                page.apply_redactions()
                                
                                # Insertar el nuevo texto
                                fontsize = 10
                                
                                try:
                                    # Intentar insertar el texto en un textbox
                                    text_rect = fitz.Rect(rect.x0, rect.y0, rect.x1, rect.y1)
                                    rc = page.insert_textbox(
                                        text_rect,
                                        value,
                                        fontsize=fontsize,
                                        color=(0, 0, 0),
                                        align=0,
                                        overlay=True
                                    )
                                    
                                    # Si el texto no cabe, usar insert_text como fallback
                                    if rc < 0:
                                        page.insert_text(
                                            (rect.x0 + 2, rect.y0 + fontsize),
                                            value,
                                            fontsize=fontsize,
                                            color=(0, 0, 0),
                                            overlay=True
                                        )
                                except Exception as e:
                                    # Si falla, intentar método alternativo
                                    try:
                                        page.insert_text(
                                            (rect.x0 + 2, rect.y0 + fontsize),
                                            value,
                                            fontsize=fontsize,
                                            color=(0, 0, 0),
                                            overlay=True
                                        )
                                    except:
                                        pass
            
            # Guardar el PDF modificado
            pdf_document.save(output_file)
            pdf_document.close()
            
            return {
                "status": "success",
                "message": "PDF generado exitosamente desde template",
                "filename": unique_filename,
                "file_path": str(output_file)
            }
            
        except Exception as e:
            return {
                "status": "error",
                "message": f"Error generando PDF desde template: {str(e)}",
                "filename": None,
                "file_path": None
            }
    
    @staticmethod
    def _generate_progress_status_from_scratch(
        document_id: int,
        progress_status_data: Dict[str, Any],
        db: Optional[Session] = None,
        output_directory: str = "files/system/students"
    ) -> Dict[str, Any]:
        """
        Genera un PDF de estado de avance desde cero usando ReportLab.
        """
        try:
            if not REPORTLAB_AVAILABLE:
                return {
                    "status": "error",
                    "message": "ReportLab no está instalado. Instala con: pip install reportlab",
                    "filename": None,
                    "file_path": None
                }
            
            # Crear nombre único para el archivo generado
            student_name = progress_status_data.get("student_fullname", "estudiante").replace(" ", "_")
            unique_filename = f"estado_avance_{student_name}_{uuid.uuid4().hex[:8]}.pdf"
            output_file = Path(output_directory) / unique_filename
            
            # Asegurar que el directorio de salida existe
            output_file.parent.mkdir(parents=True, exist_ok=True)
            
            # Crear el documento PDF con márgenes iguales a ambos lados para centrar el contenido
            doc = SimpleDocTemplate(
                str(output_file),
                pagesize=A4,
                rightMargin=2.5*cm,  # Márgenes iguales a ambos lados
                leftMargin=2.5*cm,   # Márgenes iguales a ambos lados
                topMargin=3*cm,
                bottomMargin=2.5*cm
            )
            
            elements = []
            styles = getSampleStyleSheet()
            
            # Estilos personalizados para documento gubernamental (estilo formulario)
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                textColor=colors.HexColor('#000000'),
                spaceAfter=25,
                spaceBefore=15,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            
            section_style = ParagraphStyle(
                'SectionStyle',
                parent=styles['Normal'],
                fontSize=12,
                textColor=colors.HexColor('#000000'),
                spaceAfter=12,
                spaceBefore=15,
                fontName='Helvetica-Bold',
                alignment=TA_LEFT,
                backColor=colors.HexColor('#E8E8E8'),
                borderPadding=6
            )
            
            subtitle_style = ParagraphStyle(
                'Subtitle',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#000000'),
                spaceAfter=6,
                spaceBefore=10,
                fontName='Helvetica-Bold'
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#000000'),
                alignment=TA_LEFT,
                leading=12,
                spaceAfter=4
            )
            
            label_style = ParagraphStyle(
                'LabelStyle',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#000000'),
                alignment=TA_LEFT,
                fontName='Helvetica-Bold',
                backColor=colors.HexColor('#F5F5F5'),
                borderPadding=4
            )
            
            # Función auxiliar para obtener valor
            def get_value(key: str, default: str = "") -> str:
                value = progress_status_data.get(key)
                if value is None:
                    return default
                return str(value) if value else default
            
            # Función auxiliar para formatear fechas
            def format_date(date_str: Optional[str]) -> str:
                if not date_str:
                    return ""
                try:
                    if isinstance(date_str, str):
                        date_obj = datetime.strptime(date_str, "%Y-%m-%d").date()
                        return date_obj.strftime("%d/%m/%Y")
                    return ""
                except:
                    return str(date_str) if date_str else ""
            
            # Título principal
            elements.append(Paragraph("ESTADO DE AVANCE", title_style))
            elements.append(Spacer(1, 0.4*inch))
            
            # I. IDENTIFICACIÓN DEL ESTUDIANTE
            elements.append(Paragraph("I. IDENTIFICACIÓN DEL ESTUDIANTE", section_style))
            elements.append(Spacer(1, 0.2*inch))
            
            student_fullname = get_value("student_fullname", "")
            student_rut = get_value("student_rut", "")
            student_age = get_value("student_age", "")
            progress_date = format_date(get_value("progress_date", ""))
            school_name = get_value("school_name", "")
            course_name = get_value("course_name", "")
            nee_name = get_value("nee_name", "")
            guardian_name = get_value("guardian_name", "")
            guardian_lastname = get_value("guardian_lastname", "")
            guardian_relationship = get_value("guardian_relationship", "")
            guardian_fullname = f"{guardian_name} {guardian_lastname}".strip() if (guardian_name or guardian_lastname) else ""
            
            # Información del estudiante en formato vertical (etiqueta arriba, valor abajo)
            # Estructura: 2 columnas de ancho, cada campo tiene etiqueta en una fila y valor en otra
            identification_data = []
            spans = []  # Lista para almacenar los SPANs

            # Par 1: RUT y Nombre (lado a lado)
            # Fila 1: Etiquetas
            if student_rut or student_fullname:
                row1_labels = [
                    Paragraph("<b>RUT:</b>", label_style) if student_rut else "",
                    Paragraph("<b>Nombre:</b>", label_style) if student_fullname else ""
                ]
                identification_data.append(row1_labels)
                
                # Fila 2: Valores
                row1_values = [
                    Paragraph(student_rut, normal_style) if student_rut else "",
                    Paragraph(student_fullname, normal_style) if student_fullname else ""
                ]
                identification_data.append(row1_values)

            # Par 2: Edad y Fecha Estado Avance
            if student_age or progress_date:
                row2_labels = [
                    Paragraph("<b>Edad:</b>", label_style) if student_age else "",
                    Paragraph("<b>Fecha Estado Avance:</b>", label_style) if progress_date else ""
                ]
                identification_data.append(row2_labels)
                
                row2_values = [
                    Paragraph(student_age, normal_style) if student_age else "",
                    Paragraph(progress_date, normal_style) if progress_date else ""
                ]
                identification_data.append(row2_values)

            # Par 3: Educacional y Curso
            if school_name or course_name:
                row3_labels = [
                    Paragraph("<b>Educacional:</b>", label_style) if school_name else "",
                    Paragraph("<b>Curso:</b>", label_style) if course_name else ""
                ]
                identification_data.append(row3_labels)
                
                row3_values = [
                    Paragraph(school_name, normal_style) if school_name else "",
                    Paragraph(course_name, normal_style) if course_name else ""
                ]
                identification_data.append(row3_values)

            # NEE: Ocupa 2 filas y 2 columnas (todo el ancho)
            if nee_name:
                # Fila NEE-1: Etiqueta (ocupa 2 columnas)
                identification_data.append([
                    Paragraph("<b>Necesidad Educativa Especial (NEE):</b>", label_style),
                    ""  # Columna vacía que se unirá
                ])
                nee_label_row = len(identification_data) - 1
                spans.append(('SPAN', (0, nee_label_row), (1, nee_label_row)))
                
                # Fila NEE-2: Valor (ocupa 2 columnas)
                identification_data.append([
                    Paragraph(nee_name, normal_style),
                    ""  # Columna vacía que se unirá
                ])
                nee_value_row = len(identification_data) - 1
                spans.append(('SPAN', (0, nee_value_row), (1, nee_value_row)))

            # Par 4: Relación con Estudiante y Apoderado/Guardián
            if guardian_relationship or guardian_fullname:
                row4_labels = [
                    Paragraph("<b>Relación con Estudiante:</b>", label_style) if guardian_relationship else "",
                    Paragraph("<b>Apoderado/Guardián:</b>", label_style) if guardian_fullname else ""
                ]
                identification_data.append(row4_labels)
                
                row4_values = [
                    Paragraph(guardian_relationship, normal_style) if guardian_relationship else "",
                    Paragraph(guardian_fullname, normal_style) if guardian_fullname else ""
                ]
                identification_data.append(row4_values)
            
            # Crear tabla con formato vertical (2 columnas de ancho)
            if identification_data:
                identification_table = Table(identification_data, colWidths=[7.5*cm, 7.5*cm])
                table_style = [
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 10),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
                ]
                identification_table.setStyle(TableStyle(table_style))
                elements.append(identification_table)
            
            elements.append(Spacer(1, 0.3*inch))
            
            # II. PROFESIONALES RESPONSABLES
            responsible_professionals = get_value("responsible_professionals_names", get_value("responsible_professionals", ""))
            if responsible_professionals:
                elements.append(Paragraph("II. PROFESIONALES RESPONSABLES", section_style))
                elements.append(Spacer(1, 0.15*inch))
                elements.append(Paragraph(responsible_professionals, normal_style))
                elements.append(Spacer(1, 0.3*inch))
            
            # III. ESTADO DE AVANCE POR ÁREA
            elements.append(Paragraph("III. ESTADO DE AVANCE POR ÁREA", section_style))
            elements.append(Spacer(1, 0.2*inch))
            
            areas = [
                ("Área Pedagógica - Lenguaje", get_value("pedagogical_language", "")),
                ("Área Pedagógica - Matemática", get_value("pedagogical_mathematics", "")),
                ("Área Psicopedagógica", get_value("psychopedagogical", "")),
                ("Área Fonoaudiológica", get_value("speech_therapy", "")),
                ("Área Psicológica", get_value("psychological", "")),
                ("Área Kinesiológica", get_value("kinesiology", "")),
                ("Área Terapia Ocupacional", get_value("occupational_therapy", "")),
                ("Área Co-educador Sordo/a", get_value("deaf_co_educator", ""))
            ]
            
            areas_with_content = [(name, content) for name, content in areas if content]
            
            if areas_with_content:
                for idx, (area_name, area_content) in enumerate(areas_with_content):
                    elements.append(Paragraph(f"<b>{area_name}</b>", subtitle_style))
                    elements.append(Paragraph(area_content, normal_style))
                    if idx < len(areas_with_content) - 1:
                        elements.append(Spacer(1, 0.2*inch))
            else:
                elements.append(Paragraph("No se registra información en esta sección.", normal_style))
            
            elements.append(Spacer(1, 0.3*inch))
            
            # IV. SÍNTESIS, COMENTARIOS U OBSERVACIONES
            synthesis = get_value("synthesis_comments", "")
            if synthesis:
                elements.append(Paragraph("IV. SÍNTESIS, COMENTARIOS U OBSERVACIONES", section_style))
                elements.append(Spacer(1, 0.15*inch))
                elements.append(Paragraph(synthesis, normal_style))
                elements.append(Spacer(1, 0.3*inch))
            
            # V. SUGERENCIAS
            suggestions_family = get_value("suggestions_family", "")
            suggestions_establishment = get_value("suggestions_establishment", "")
            
            if suggestions_family or suggestions_establishment:
                elements.append(Paragraph("V. SUGERENCIAS", section_style))
                elements.append(Spacer(1, 0.15*inch))
                
                if suggestions_family:
                    elements.append(Paragraph("<b>a) Sugerencias a la familia:</b>", subtitle_style))
                    elements.append(Paragraph(suggestions_family, normal_style))
                    elements.append(Spacer(1, 0.2*inch))
                
                if suggestions_establishment:
                    elements.append(Paragraph("<b>b) Sugerencias al establecimiento educacional:</b>", subtitle_style))
                    elements.append(Paragraph(suggestions_establishment, normal_style))
                    elements.append(Spacer(1, 0.2*inch))
            
            # Construir PDF (sin espacio final excesivo que cause página en blanco)
            doc.build(elements)
            
            return {
                "status": "success",
                "message": "PDF generado exitosamente desde cero",
                "filename": unique_filename,
                "file_path": str(output_file)
            }
            
        except ImportError:
            return {
                "status": "error",
                "message": "ReportLab no está instalado. Instala con: pip install reportlab",
                "filename": None,
                "file_path": None
            }
        except Exception as e:
            return {
                "status": "error",
                "message": f"Error generando PDF desde cero: {str(e)}",
                "filename": None,
                "file_path": None
            }

    @staticmethod
    def _generate_progress_status_individual_support_from_scratch(
        document_id: int,
        ps_data: Dict[str, Any],
        db: Optional[Session] = None,
        output_directory: str = "files/system/students"
    ) -> Dict[str, Any]:
        """
        Genera un PDF de Estado de Avance PAI (documento 19) desde cero usando ReportLab.
        Datos desde progress_status_individual_support.
        """
        try:
            if not REPORTLAB_AVAILABLE:
                return {
                    "status": "error",
                    "message": "ReportLab no está instalado. Instala con: pip install reportlab",
                    "filename": None,
                    "file_path": None
                }

            student_name = ps_data.get("student_full_name", "estudiante").replace(" ", "_")
            unique_filename = f"estado_avance_pai_{student_name}_{uuid.uuid4().hex[:8]}.pdf"
            output_file = Path(output_directory) / unique_filename
            output_file.parent.mkdir(parents=True, exist_ok=True)

            # Márgenes ajustados para mejor distribución
            doc = SimpleDocTemplate(
                str(output_file),
                pagesize=A4,
                rightMargin=2*cm,
                leftMargin=2*cm,
                topMargin=2*cm,
                bottomMargin=2*cm
            )
            elements = []
            styles = getSampleStyleSheet()

            title_style = ParagraphStyle(
                'CertTitle19',
                parent=styles['Heading1'],
                fontSize=14,
                textColor=colors.HexColor('#000000'),
                spaceAfter=18,
                spaceBefore=8,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            section_style = ParagraphStyle(
                'CertSection19',
                parent=styles['Normal'],
                fontSize=11,
                textColor=colors.HexColor('#000000'),
                spaceAfter=8,
                spaceBefore=12,
                fontName='Helvetica-Bold',
                alignment=TA_LEFT,
                backColor=colors.HexColor('#E8E8E8'),
                borderPadding=8
            )
            subtitle_style = ParagraphStyle(
                'CertSubtitle19',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#000000'),
                spaceAfter=4,
                spaceBefore=8,
                fontName='Helvetica-Bold'
            )
            normal_style = ParagraphStyle(
                'CertNormal19',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#000000'),
                alignment=TA_LEFT,
                leading=13,
                spaceAfter=2
            )
            label_style = ParagraphStyle(
                'CertLabel19',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#000000'),
                alignment=TA_LEFT,
                fontName='Helvetica-Bold',
                backColor=colors.HexColor('#F0F0F0'),
                borderPadding=5
            )

            def get_value(key: str, default: str = "") -> str:
                v = ps_data.get(key)
                if v is None:
                    return default
                return str(v).strip() if v else default

            def format_date(date_str: Optional[str]) -> str:
                if not date_str:
                    return ""
                try:
                    if isinstance(date_str, str):
                        date_obj = datetime.strptime(date_str, "%Y-%m-%d").date()
                        return date_obj.strftime("%d/%m/%Y")
                    return ""
                except Exception:
                    return str(date_str) if date_str else ""

            elements.append(Paragraph("ESTADO DE AVANCE PAI", title_style))
            elements.append(Spacer(1, 0.4*inch))

            # I. IDENTIFICACIÓN DEL ESTUDIANTE (solo esta sección por ahora)
            elements.append(Paragraph("I. IDENTIFICACIÓN DEL ESTUDIANTE", section_style))
            elements.append(Spacer(1, 0.2*inch))

            student_fullname = get_value("student_full_name", "")
            student_rut = get_value("student_identification_number", "")
            student_age = get_value("student_age", "")
            student_born_date = format_date(get_value("student_born_date", ""))
            progress_date = format_date(get_value("progress_date", ""))
            period_label = get_value("period_label", "")
            school_name = get_value("student_school", "")
            course_name = get_value("course_name", "")
            nee_name = get_value("nee_name", "")
            guardian_fullname = get_value("guardian_fullname", "")
            guardian_rut = get_value("guardian_rut", "")
            guardian_relationship = get_value("guardian_relationship", "")

            # Formato: una fila por campo, [Etiqueta | Valor]. Columnas amplias para evitar corte de texto.
            identification_data = []
            spans = []
            col_label = 5*cm
            col_value = 11*cm

            def _row(label: str, value: str):
                identification_data.append([
                    Paragraph(f"<b>{label}</b>", label_style),
                    Paragraph(DocumentsClass._escape_html_for_paragraph(value or "—"), normal_style)
                ])

            _row("RUT:", student_rut)
            _row("Nombre:", student_fullname)
            _row("Fecha de Nacimiento:", student_born_date)
            _row("Edad:", str(student_age) if student_age else "")
            _row("Fecha Estado Avance:", progress_date)
            _row("Período:", period_label)
            _row("Establecimiento Educacional:", school_name)
            _row("Curso:", course_name)
            # NEE: ocupa todo el ancho
            identification_data.append([
                Paragraph("<b>Necesidad Educativa Especial (NEE):</b>", label_style),
                Paragraph(DocumentsClass._escape_html_for_paragraph(nee_name or "—"), normal_style)
            ])
            _row("Relación con Estudiante:", guardian_relationship)
            _row("Apoderado/Guardián:", guardian_fullname)
            _row("RUT Apoderado:", guardian_rut)

            if identification_data:
                identification_table = Table(identification_data, colWidths=[col_label, col_value])
                table_style = [
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 10),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
                ]
                if spans:
                    table_style.extend(spans)
                identification_table.setStyle(TableStyle(table_style))
                elements.append(identification_table)

            elements.append(Spacer(1, 0.3*inch))

            # II. PROFESIONALES RESPONSABLES
            elements.append(Paragraph("II. PROFESIONALES RESPONSABLES", section_style))
            elements.append(Spacer(1, 0.15*inch))
            responsible = get_value("responsible_professionals_names", "")
            if responsible:
                elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(responsible), normal_style))
            else:
                elements.append(Paragraph("No se registra información en esta sección.", normal_style))
            elements.append(Spacer(1, 0.3*inch))

            # III. OBJETIVOS PAI Y NIVEL DE AVANCE
            elements.append(Paragraph("III. OBJETIVOS PAI Y NIVEL DE AVANCE", section_style))
            elements.append(Spacer(1, 0.2*inch))
            pai_objectives = ps_data.get("pai_objectives")
            obj_list = []
            if pai_objectives:
                obj_list = pai_objectives if isinstance(pai_objectives, list) else []
                if isinstance(pai_objectives, str):
                    try:
                        import json
                        obj_list = json.loads(pai_objectives)
                    except Exception:
                        obj_list = []
            obj_list_valid = [o for o in obj_list if isinstance(o, dict) and (o.get("description") or o.get("progress_level"))]
            if obj_list_valid:
                for idx, obj in enumerate(obj_list_valid):
                    num = obj.get("number") or (idx + 1)
                    desc = obj.get("description") or ""
                    level = obj.get("progress_level") or ""
                    elements.append(Paragraph(f"<b>Objetivo {num}:</b>", subtitle_style))
                    if desc:
                        elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(desc), normal_style))
                    if level:
                        elements.append(Paragraph(f"<b>Nivel de avance:</b> {DocumentsClass._escape_html_for_paragraph(level)}", normal_style))
                    if idx < len(obj_list_valid) - 1:
                        elements.append(Spacer(1, 0.15*inch))
            else:
                elements.append(Paragraph("No se registra información en esta sección.", normal_style))
            elements.append(Spacer(1, 0.3*inch))

            # IV. OBSERVACIONES PAI
            elements.append(Paragraph("IV. OBSERVACIONES PAI", section_style))
            elements.append(Spacer(1, 0.15*inch))
            pai_observations = get_value("pai_observations", "")
            if pai_observations:
                elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(pai_observations), normal_style))
            else:
                elements.append(Paragraph("No se registra información en esta sección.", normal_style))
            elements.append(Spacer(1, 0.3*inch))

            # V. SUGERENCIAS
            elements.append(Paragraph("V. SUGERENCIAS", section_style))
            elements.append(Spacer(1, 0.15*inch))
            suggestions_family = get_value("suggestions_family", "")
            suggestions_establishment = get_value("suggestions_establishment", "")
            if suggestions_family or suggestions_establishment:
                if suggestions_family:
                    elements.append(Paragraph("<b>a) Sugerencias a la familia:</b>", subtitle_style))
                    elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(suggestions_family), normal_style))
                    elements.append(Spacer(1, 0.2*inch))
                if suggestions_establishment:
                    elements.append(Paragraph("<b>b) Sugerencias al establecimiento educacional:</b>", subtitle_style))
                    elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(suggestions_establishment), normal_style))
                    elements.append(Spacer(1, 0.2*inch))
            else:
                elements.append(Paragraph("No se registra información en esta sección.", normal_style))

            doc.build(elements)
            return {
                "status": "success",
                "message": "PDF generado exitosamente desde cero",
                "filename": unique_filename,
                "file_path": str(output_file)
            }
        except ImportError:
            return {
                "status": "error",
                "message": "ReportLab no está instalado. Instala con: pip install reportlab",
                "filename": None,
                "file_path": None
            }
        except Exception as e:
            return {
                "status": "error",
                "message": f"Error generando PDF Estado de Avance PAI: {str(e)}",
                "filename": None,
                "file_path": None
            }

    @staticmethod
    def _generate_cesp_from_scratch(
        document_id: int,
        cesp_data: Dict[str, Any],
        db: Optional[Session] = None,
        output_directory: str = "files/system/students"
    ) -> Dict[str, Any]:
        """
        Genera un PDF del documento 20 (CESP - Plan de Acompañamiento Emocional y Conductual / PAEC) desde cero usando ReportLab.
        """
        try:
            if not REPORTLAB_AVAILABLE:
                return {
                    "status": "error",
                    "message": "ReportLab no está instalado. Instala con: pip install reportlab",
                    "filename": None,
                    "file_path": None
                }
            student_name_safe = (cesp_data.get("student_fullname") or "estudiante").replace(" ", "_")
            unique_filename = f"cesp_paec_{student_name_safe}_{uuid.uuid4().hex[:8]}.pdf"
            output_file = Path(output_directory) / unique_filename
            output_file.parent.mkdir(parents=True, exist_ok=True)
            doc = SimpleDocTemplate(
                str(output_file),
                pagesize=A4,
                rightMargin=2.5*cm,
                leftMargin=2.5*cm,
                topMargin=2.5*cm,
                bottomMargin=2.5*cm
            )
            elements = []
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                textColor=colors.HexColor('#000000'),
                spaceAfter=25,
                spaceBefore=15,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            section_style = ParagraphStyle(
                'SectionStyle',
                parent=styles['Normal'],
                fontSize=12,
                textColor=colors.HexColor('#000000'),
                spaceAfter=12,
                spaceBefore=15,
                fontName='Helvetica-Bold',
                alignment=TA_LEFT,
                backColor=colors.HexColor('#E8E8E8'),
                borderPadding=6
            )
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=11,
                textColor=colors.HexColor('#000000'),
                alignment=TA_LEFT,
                leading=13,
                spaceAfter=4
            )
            label_style = ParagraphStyle(
                'LabelStyle',
                parent=styles['Normal'],
                fontSize=11,
                textColor=colors.HexColor('#000000'),
                alignment=TA_LEFT,
                fontName='Helvetica-Bold',
                backColor=colors.HexColor('#F5F5F5'),
                borderPadding=4
            )
            def get_value(key: str, default: str = "") -> str:
                v = cesp_data.get(key)
                if v is None:
                    return default
                return str(v).strip() if v else default
            def hr_line():
                t = Table([[""]], colWidths=[16*cm], rowHeights=[0.02*inch])
                t.setStyle(TableStyle([('LINEBELOW', (0, 0), (-1, 0), 0.5, colors.grey)]))
                return t
            def format_date(date_str: Optional[str]) -> str:
                if not date_str:
                    return ""
                s = str(date_str).strip()[:10]
                if not s:
                    return ""
                for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%Y/%m/%d"):
                    try:
                        date_obj = datetime.strptime(s, fmt).date()
                        return date_obj.strftime("%d/%m/%Y")
                    except Exception:
                        continue
                return s
            elements.append(Paragraph("Plan de Acompañamiento Emocional y Conductual", title_style))
            elements.append(Spacer(1, 0.3*inch))
            elements.append(Paragraph("I. Datos personales o generales", section_style))
            elements.append(Spacer(1, 0.15*inch))
            student_fullname = get_value("student_fullname")
            student_rut = get_value("student_rut")
            student_born_date = format_date(get_value("student_born_date"))
            student_age = get_value("student_age")
            student_nee = get_value("student_nee")
            student_school = get_value("student_school")
            student_course = get_value("student_course")
            elaboration_date = format_date(get_value("elaboration_date"))
            period_label = get_value("period_type_id") == "2" and "Semestral" or "Anual"
            def si_no(val: str) -> str:
                if not val: return val
                v = (val or "").strip().lower()
                if v == "yes": return "Si"
                if v == "no": return "No"
                return val.strip()
            pharmacological = si_no(get_value("pharmacological_treatment"))
            external_specialists = si_no(get_value("external_specialists"))
            def cell_para(text: str, bold: bool = False):
                content = (text or "—").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                tag = "b" if bold else ""
                return Paragraph(f"<{tag}>{content}</{tag}>" if tag else content, normal_style)
            gen_data = [
                [cell_para("RUT", True), cell_para(student_rut)],
                [cell_para("Nombre", True), cell_para(student_fullname)],
                [cell_para("Fecha de nacimiento", True), cell_para(student_born_date)],
                [cell_para("Edad", True), cell_para(student_age)],
                [cell_para("NEE", True), cell_para(student_nee)],
                [cell_para("Establecimiento", True), cell_para(student_school)],
                [cell_para("Curso", True), cell_para(student_course)],
                [cell_para("Fecha de elaboración", True), cell_para(elaboration_date)],
                [cell_para("Tipo (Período)", True), cell_para(period_label)],
                [cell_para("¿Tratamiento farmacológico?", True), cell_para(pharmacological)],
                [cell_para("¿Especialistas externos?", True), cell_para(external_specialists)],
            ]
            gen_table = Table(gen_data, colWidths=[7*cm, 9*cm])
            hr_style = [('LINEBELOW', (0, i), (-1, i), 0.5, colors.grey) for i in range(len(gen_data))]
            gen_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 11),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
                ('TOPPADDING', (0, 0), (-1, -1), 10),
                ('LEFTPADDING', (0, 0), (-1, -1), 8),
                ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ] + hr_style))
            elements.append(gen_table)
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("II. Identificación del apoderado informante y del profesional participante", section_style))
            elements.append(Spacer(1, 0.15*inch))
            guardians = cesp_data.get("guardians") or []
            if guardians:
                guard_headers = [Paragraph("<b>RUT</b>", label_style), Paragraph("<b>Nombre</b>", label_style), Paragraph("<b>Teléfono</b>", label_style), Paragraph("<b>Correo</b>", label_style)]
                guard_rows = [guard_headers]
                for g in guardians:
                    if isinstance(g, dict):
                        guard_rows.append([
                            (g.get("identification_number") or "").strip() or "—",
                            (g.get("name") or "").strip() or "—",
                            (g.get("phone") or "").strip() or "—",
                            (g.get("email") or "").strip() or "—",
                        ])
                if len(guard_rows) > 1:
                    gtable = Table(guard_rows, colWidths=[3*cm, 5.5*cm, 3*cm, 4.5*cm])
                    gtable.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e0e0e0')),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, -1), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                        ('TOPPADDING', (0, 0), (-1, -1), 8),
                        ('LEFTPADDING', (0, 0), (-1, -1), 6),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ]))
                    elements.append(gtable)
                    elements.append(Spacer(1, 0.15*inch))
            participant_name = cesp_data.get("participant_professional_name") or ""
            participant_role = (cesp_data.get("participant_professional") or {})
            if isinstance(participant_role, dict):
                participant_role = participant_role.get("professional_role") or ""
            else:
                participant_role = ""
            if participant_name or participant_role:
                part_headers = [Paragraph("<b>Nombre</b>", label_style), Paragraph("<b>Cargo / Rol</b>", label_style)]
                part_data = [[participant_name or "—", participant_role or "—"]]
                part_table = Table([part_headers] + part_data, colWidths=[8*cm, 8*cm])
                part_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e0e0e0')),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ]))
                elements.append(KeepTogether([part_table, Spacer(1, 0.2*inch)]))
            elements.append(Paragraph("III. Identificación del equipo de apoyo", section_style))
            elements.append(Spacer(1, 0.15*inch))
            support_list = cesp_data.get("support_team_members") or []
            if support_list:
                sup_headers = [Paragraph("<b>Nombre</b>", label_style), Paragraph("<b>Cargo / Rol</b>", label_style), Paragraph("<b>Roles de apoyo</b>", label_style), Paragraph("<b>Contacto</b>", label_style)]
                sup_rows = [sup_headers]
                for s in support_list:
                    if isinstance(s, dict):
                        contact = " | ".join(filter(None, [s.get("phone"), s.get("email")])) or "—"
                        sup_rows.append([
                            s.get("professional_name") or "—",
                            s.get("professional_role") or "—",
                            (s.get("support_roles") or "—")[:80] + ("..." if len((s.get("support_roles") or "")) > 80 else ""),
                            contact,
                        ])
                if len(sup_rows) > 1:
                    stable = Table(sup_rows, colWidths=[4*cm, 3.5*cm, 4*cm, 4.5*cm])
                    stable.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e0e0e0')),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, -1), 9),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                        ('TOPPADDING', (0, 0), (-1, -1), 5),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ]))
                    elements.append(stable)
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("IV. Perfil del estudiante", section_style))
            elements.append(Spacer(1, 0.1*inch))
            profile_fields = [
                ("Interacción", "profile_interaction"),
                ("Involucramiento", "profile_involvement"),
                ("Repertorio conductual", "profile_behavior_repertoire"),
                ("Habilidades", "profile_skills"),
                ("Desafíos", "profile_challenges"),
                ("Necesidades de apoyo", "profile_support_needs"),
                ("Intereses", "profile_interests"),
            ]
            for label, key in profile_fields:
                val = get_value(key)
                elements.append(Paragraph(f"<b>{label}</b>", normal_style))
                elements.append(hr_line())
                elements.append(Paragraph((val[:2000] + ("..." if len(val) > 2000 else "")) if val else "—", normal_style))
                elements.append(Spacer(1, 0.08*inch))
            elements.append(Paragraph("V. Estresores y prevención", section_style))
            elements.append(Spacer(1, 0.1*inch))
            for label, key in [("Estresores y desencadenantes", "stressors_triggers"), ("Medidas de prevención", "prevention_measures")]:
                val = get_value(key)
                elements.append(Paragraph(f"<b>{label}</b>", normal_style))
                elements.append(hr_line())
                elements.append(Paragraph((val[:2000] + ("..." if len(val) > 2000 else "")) if val else "—", normal_style))
                elements.append(Spacer(1, 0.08*inch))
            elements.append(Paragraph("VI. Sugerencias", section_style))
            elements.append(Spacer(1, 0.1*inch))
            val_sug = get_value("suggestions_special")
            elements.append(Paragraph("<b>Sugerencias</b>", normal_style))
            elements.append(hr_line())
            elements.append(Paragraph((val_sug[:2000] + ("..." if len(val_sug) > 2000 else "")) if val_sug else "—", normal_style))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("VII. Estrategias por fases", section_style))
            elements.append(Spacer(1, 0.1*inch))
            for phase in [1, 2, 3, 4]:
                m = get_value(f"strategies_phase{phase}_manifestations")
                s = get_value(f"strategies_phase{phase}_strategies")
                elements.append(Paragraph(f"<b>Fase {phase}</b>", normal_style))
                elements.append(hr_line())
                elements.append(Paragraph("<b>Manifestaciones</b>", normal_style))
                elements.append(hr_line())
                elements.append(Spacer(1, 0.06*inch))
                elements.append(Paragraph((m[:1500] + "..." if len(m) > 1500 else m) if m else "—", normal_style))
                elements.append(Spacer(1, 0.08*inch))
                elements.append(Paragraph("<b>Estrategias</b>", normal_style))
                elements.append(hr_line())
                elements.append(Spacer(1, 0.06*inch))
                elements.append(Paragraph((s[:1500] + "..." if len(s) > 1500 else s) if s else "—", normal_style))
                elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("VIII. Registro DEC", section_style))
            elements.append(Spacer(1, 0.1*inch))
            dec_records = cesp_data.get("dec_records") or []
            dec_type_labels = {1: "Desregulación emocional/conductual", 2: "Otra incidencia", 3: "Otro"}
            if dec_records:
                for idx, rec in enumerate(dec_records):
                    if not isinstance(rec, dict):
                        continue
                    rec_date = format_date(rec.get("incident_date") or "")
                    rec_time = rec.get("incident_time") or ""
                    rec_type_id = rec.get("action_incident_type_id")
                    rec_type = dec_type_labels.get(rec_type_id, str(rec_type_id) if rec_type_id else "—")
                    rec_title = (rec.get("title") or "").strip() or "—"
                    elements.append(Paragraph(f"<b>Registro {idx + 1}</b>", normal_style))
                    elements.append(hr_line())
                    elements.append(Paragraph(f"<b>Fecha</b>", normal_style))
                    elements.append(hr_line())
                    elements.append(Spacer(1, 0.04*inch))
                    elements.append(Paragraph(f"{rec_date} {rec_time}".strip() or "—", normal_style))
                    elements.append(Paragraph(f"<b>Tipo</b>", normal_style))
                    elements.append(hr_line())
                    elements.append(Spacer(1, 0.04*inch))
                    elements.append(Paragraph(rec_type, normal_style))
                    elements.append(Paragraph(f"<b>Título</b>", normal_style))
                    elements.append(hr_line())
                    elements.append(Spacer(1, 0.04*inch))
                    elements.append(Paragraph(rec_title, normal_style))
                    for label, key in [("Antecedentes", "background"), ("Conducta", "conduct"), ("Consecuencias", "consequences"), ("Recomendaciones", "recommendations")]:
                        val = (rec.get(key) or "").strip() if rec.get(key) else ""
                        elements.append(Paragraph(f"<b>{label}</b>", normal_style))
                        elements.append(hr_line())
                        elements.append(Spacer(1, 0.04*inch))
                        elements.append(Paragraph((val[:1500] + "..." if len(val) > 1500 else val) if val else "—", normal_style))
                    elements.append(Spacer(1, 0.15*inch))
            else:
                elements.append(Paragraph("No se registran incidentes DEC para este estudiante.", normal_style))
            doc.build(elements)
            return {
                "status": "success",
                "message": "PDF CESP generado correctamente",
                "filename": unique_filename,
                "file_path": str(output_file)
            }
        except Exception as e:
            return {
                "status": "error",
                "message": f"Error generando PDF CESP: {str(e)}",
                "filename": None,
                "file_path": None
            }

    @staticmethod
    def _generate_individual_support_plan_from_scratch(
        document_id: int,
        isp_data: Dict[str, Any],
        db: Optional[Session] = None,
        output_directory: str = "files/system/students"
    ) -> Dict[str, Any]:
        """
        Genera un PDF de Plan de Apoyo Individual (PAI) desde cero usando ReportLab.
        """
        try:
            if not REPORTLAB_AVAILABLE:
                return {
                    "status": "error",
                    "message": "ReportLab no está instalado. Instala con: pip install reportlab",
                    "filename": None,
                    "file_path": None
                }
            
            # Crear nombre único para el archivo generado
            student_name = isp_data.get("student_full_name", "estudiante").replace(" ", "_")
            unique_filename = f"plan_apoyo_individual_{student_name}_{uuid.uuid4().hex[:8]}.pdf"
            output_file = Path(output_directory) / unique_filename
            
            # Asegurar que el directorio de salida existe
            output_file.parent.mkdir(parents=True, exist_ok=True)
            
            # Crear el documento PDF con márgenes iguales a ambos lados para centrar el contenido
            doc = SimpleDocTemplate(
                str(output_file),
                pagesize=A4,
                rightMargin=2.5*cm,
                leftMargin=2.5*cm,
                topMargin=3*cm,
                bottomMargin=2.5*cm
            )
            
            elements = []
            styles = getSampleStyleSheet()
            
            # Estilos personalizados para documento gubernamental
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                textColor=colors.HexColor('#000000'),
                spaceAfter=25,
                spaceBefore=15,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            
            section_style = ParagraphStyle(
                'SectionStyle',
                parent=styles['Normal'],
                fontSize=12,
                textColor=colors.HexColor('#000000'),
                spaceAfter=12,
                spaceBefore=15,
                fontName='Helvetica-Bold',
                alignment=TA_LEFT,
                backColor=colors.HexColor('#E8E8E8'),
                borderPadding=6
            )
            
            subtitle_style = ParagraphStyle(
                'Subtitle',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#000000'),
                spaceAfter=6,
                spaceBefore=10,
                fontName='Helvetica-Bold'
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#000000'),
                alignment=TA_LEFT,
                leading=12,
                spaceAfter=4
            )
            
            label_style = ParagraphStyle(
                'LabelStyle',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#000000'),
                alignment=TA_LEFT,
                fontName='Helvetica-Bold',
                backColor=colors.HexColor('#F5F5F5'),
                borderPadding=4
            )
            
            # Función auxiliar para obtener valor
            def get_value(key: str, default: str = "") -> str:
                value = isp_data.get(key)
                if value is None:
                    return default
                return str(value) if value else default
            
            # Función auxiliar para formatear fechas
            def format_date(date_str: Optional[str]) -> str:
                if not date_str:
                    return ""
                try:
                    if isinstance(date_str, str):
                        date_obj = datetime.strptime(date_str, "%Y-%m-%d").date()
                        return date_obj.strftime("%d/%m/%Y")
                    return ""
                except:
                    return str(date_str) if date_str else ""
            
            # Título principal
            elements.append(Paragraph("PLAN DE APOYO INDIVIDUAL (PAI)", title_style))
            elements.append(Spacer(1, 0.4*inch))
            
            # I. IDENTIFICACIÓN DEL ESTUDIANTE
            elements.append(Paragraph("I. IDENTIFICACIÓN DEL ESTUDIANTE", section_style))
            elements.append(Spacer(1, 0.2*inch))
            
            student_fullname = get_value("student_full_name", "")
            student_rut = get_value("student_identification_number", "")
            student_age = get_value("student_age", "")
            student_born_date = format_date(get_value("student_born_date", ""))
            elaboration_date = format_date(get_value("elaboration_date", ""))
            school_name = get_value("student_school", "")
            course_name = get_value("course_name", "")
            nee_name = get_value("nee_name", "")
            
            # Información del estudiante en formato vertical (2 columnas)
            identification_data = []
            spans = []
            
            # Par 1: RUT y Nombre
            if student_rut or student_fullname:
                row1_labels = [
                    Paragraph("<b>RUT:</b>", label_style) if student_rut else "",
                    Paragraph("<b>Nombre:</b>", label_style) if student_fullname else ""
                ]
                identification_data.append(row1_labels)
                
                row1_values = [
                    Paragraph(student_rut, normal_style) if student_rut else "",
                    Paragraph(student_fullname, normal_style) if student_fullname else ""
                ]
                identification_data.append(row1_values)
            
            # Par 2: Fecha de Nacimiento y Edad
            if student_born_date or student_age:
                row2_labels = [
                    Paragraph("<b>Fecha de Nacimiento:</b>", label_style) if student_born_date else "",
                    Paragraph("<b>Edad:</b>", label_style) if student_age else ""
                ]
                identification_data.append(row2_labels)
                
                row2_values = [
                    Paragraph(student_born_date, normal_style) if student_born_date else "",
                    Paragraph(student_age, normal_style) if student_age else ""
                ]
                identification_data.append(row2_values)
            
            # Par 3: Establecimiento Educacional y Curso
            if school_name or course_name:
                row3_labels = [
                    Paragraph("<b>Establecimiento Educacional:</b>", label_style) if school_name else "",
                    Paragraph("<b>Curso:</b>", label_style) if course_name else ""
                ]
                identification_data.append(row3_labels)
                
                row3_values = [
                    Paragraph(school_name, normal_style) if school_name else "",
                    Paragraph(course_name, normal_style) if course_name else ""
                ]
                identification_data.append(row3_values)
            
            # NEE: Ocupa 2 columnas
            if nee_name:
                identification_data.append([
                    Paragraph("<b>Necesidad Educativa Especial (NEE):</b>", label_style),
                    ""
                ])
                nee_label_row = len(identification_data) - 1
                spans.append(('SPAN', (0, nee_label_row), (1, nee_label_row)))
                
                identification_data.append([
                    Paragraph(nee_name, normal_style),
                    ""
                ])
                nee_value_row = len(identification_data) - 1
                spans.append(('SPAN', (0, nee_value_row), (1, nee_value_row)))
            
            # Fecha de Elaboración
            if elaboration_date:
                row4_labels = [
                    Paragraph("<b>Fecha de Elaboración:</b>", label_style),
                    ""
                ]
                identification_data.append(row4_labels)
                row4_label_row = len(identification_data) - 1
                spans.append(('SPAN', (0, row4_label_row), (1, row4_label_row)))
                
                row4_values = [
                    Paragraph(elaboration_date, normal_style),
                    ""
                ]
                identification_data.append(row4_values)
                row4_value_row = len(identification_data) - 1
                spans.append(('SPAN', (0, row4_value_row), (1, row4_value_row)))
            
            # Crear tabla de identificación
            if identification_data:
                identification_table = Table(identification_data, colWidths=[7.5*cm, 7.5*cm])
                table_style = [
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 8),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
                ]
                table_style.extend(spans)
                identification_table.setStyle(TableStyle(table_style))
                elements.append(identification_table)
            
            elements.append(Spacer(1, 0.3*inch))
            
            # II. PROFESIONALES ASOCIADOS
            professionals = isp_data.get("professionals", [])
            if professionals:
                elements.append(Paragraph("II. PROFESIONALES ASOCIADOS", section_style))
                elements.append(Spacer(1, 0.2*inch))
                
                # Crear tabla con profesionales, 1 columna, misma estructura que identificación
                professionals_data = []
                professionals_spans = []
                
                for idx, prof in enumerate(professionals):
                    prof_name = prof.get("professional_name", "")
                    prof_career = prof.get("career_type_name", "")
                    prof_reg = prof.get("registration_number", "")
                    prof_days = prof.get("days_hours", "")
                    prof_from = format_date(prof.get("from_date", ""))
                    prof_to = format_date(prof.get("to_date", ""))
                    prof_period = f"{prof_from} - {prof_to}" if (prof_from and prof_to) else (prof_from or prof_to or "")
                    prof_modality = prof.get("support_modality", "")
                    
                    # Nombre del Profesional (ocupa 1 columna completa)
                    if prof_name:
                        professionals_data.append([
                            Paragraph(f"<b>Profesional {idx + 1}:</b>", label_style),
                            ""
                        ])
                        prof_name_label_row = len(professionals_data) - 1
                        professionals_spans.append(('SPAN', (0, prof_name_label_row), (1, prof_name_label_row)))
                        
                        professionals_data.append([
                            Paragraph(prof_name, normal_style),
                            ""
                        ])
                        prof_name_value_row = len(professionals_data) - 1
                        professionals_spans.append(('SPAN', (0, prof_name_value_row), (1, prof_name_value_row)))
                    
                    # Especialidad (ocupa 2 columnas - ancho completo)
                    if prof_career:
                        professionals_data.append([
                            Paragraph("<b>Especialidad:</b>", label_style),
                            ""
                        ])
                        career_label_row = len(professionals_data) - 1
                        professionals_spans.append(('SPAN', (0, career_label_row), (1, career_label_row)))
                        
                        professionals_data.append([
                            Paragraph(prof_career, normal_style),
                            ""
                        ])
                        career_value_row = len(professionals_data) - 1
                        professionals_spans.append(('SPAN', (0, career_value_row), (1, career_value_row)))
                    
                    # N° Registro (ocupa 2 columnas - ancho completo)
                    if prof_reg:
                        professionals_data.append([
                            Paragraph("<b>N° Registro:</b>", label_style),
                            ""
                        ])
                        reg_label_row = len(professionals_data) - 1
                        professionals_spans.append(('SPAN', (0, reg_label_row), (1, reg_label_row)))
                        
                        professionals_data.append([
                            Paragraph(prof_reg, normal_style),
                            ""
                        ])
                        reg_value_row = len(professionals_data) - 1
                        professionals_spans.append(('SPAN', (0, reg_value_row), (1, reg_value_row)))
                    
                    # Días y Horarios (ocupa 2 columnas - ancho completo)
                    if prof_days:
                        professionals_data.append([
                            Paragraph("<b>Días y Horarios:</b>", label_style),
                            ""
                        ])
                        days_label_row = len(professionals_data) - 1
                        professionals_spans.append(('SPAN', (0, days_label_row), (1, days_label_row)))
                        
                        professionals_data.append([
                            Paragraph(prof_days, normal_style),
                            ""
                        ])
                        days_value_row = len(professionals_data) - 1
                        professionals_spans.append(('SPAN', (0, days_value_row), (1, days_value_row)))
                    
                    # Período (ocupa 2 columnas - ancho completo)
                    if prof_period:
                        professionals_data.append([
                            Paragraph("<b>Período:</b>", label_style),
                            ""
                        ])
                        period_label_row = len(professionals_data) - 1
                        professionals_spans.append(('SPAN', (0, period_label_row), (1, period_label_row)))
                        
                        professionals_data.append([
                            Paragraph(prof_period, normal_style),
                            ""
                        ])
                        period_value_row = len(professionals_data) - 1
                        professionals_spans.append(('SPAN', (0, period_value_row), (1, period_value_row)))
                    
                    # Modalidad (ocupa 2 columnas)
                    if prof_modality:
                        professionals_data.append([
                            Paragraph("<b>Modalidad de Apoyo:</b>", label_style),
                            ""
                        ])
                        modality_label_row = len(professionals_data) - 1
                        professionals_spans.append(('SPAN', (0, modality_label_row), (1, modality_label_row)))
                        
                        professionals_data.append([
                            Paragraph(prof_modality, normal_style),
                            ""
                        ])
                        modality_value_row = len(professionals_data) - 1
                        professionals_spans.append(('SPAN', (0, modality_value_row), (1, modality_value_row)))
                    
                    # Agregar espacio entre profesionales
                    if idx < len(professionals) - 1:
                        professionals_data.append(["", ""])
                
                # Crear tabla con 2 columnas (pero cada profesional usa ambas cuando es necesario)
                if professionals_data:
                    professionals_table = Table(professionals_data, colWidths=[7.5*cm, 7.5*cm])
                    professionals_table_style = [
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('LEFTPADDING', (0, 0), (-1, -1), 8),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                        ('TOPPADDING', (0, 0), (-1, -1), 8),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
                    ]
                    professionals_table_style.extend(professionals_spans)
                    professionals_table.setStyle(TableStyle(professionals_table_style))
                    elements.append(professionals_table)
                
                elements.append(Spacer(1, 0.3*inch))
            
            # III. FORTALEZAS DEL ESTUDIANTE
            social_affective = get_value("social_affective_strengths", "")
            cognitive = get_value("cognitive_strengths", "")
            curricular = get_value("curricular_strengths", "")
            family = get_value("family_strengths", "")
            
            if social_affective or cognitive or curricular or family:
                elements.append(Paragraph("III. FORTALEZAS DEL/LA ESTUDIANTE", section_style))
                elements.append(Spacer(1, 0.2*inch))
                
                if social_affective:
                    elements.append(Paragraph("<b>a) Fortalezas Socioafectivas:</b>", subtitle_style))
                    elements.append(Paragraph(social_affective, normal_style))
                    elements.append(Spacer(1, 0.15*inch))
                
                if cognitive:
                    elements.append(Paragraph("<b>b) Fortalezas Cognitivas:</b>", subtitle_style))
                    elements.append(Paragraph(cognitive, normal_style))
                    elements.append(Spacer(1, 0.15*inch))
                
                if curricular:
                    elements.append(Paragraph("<b>c) Fortalezas Curriculares:</b>", subtitle_style))
                    elements.append(Paragraph(curricular, normal_style))
                    elements.append(Spacer(1, 0.15*inch))
                
                if family:
                    elements.append(Paragraph("<b>d) Fortalezas Familiares:</b>", subtitle_style))
                    elements.append(Paragraph(family, normal_style))
                    elements.append(Spacer(1, 0.15*inch))
                
                elements.append(Spacer(1, 0.3*inch))
            
            # IV. PROPUESTA DE INTERVENCIÓN
            elements.append(Paragraph("IV. PROPUESTA DE INTERVENCIÓN", section_style))
            elements.append(Spacer(1, 0.2*inch))
            
            # Áreas de intervención
            intervention_areas = [
                ("Educación Diferencial", get_value("intervention_ed_diferencial", ""), get_value("intervention_ed_diferencial_strategies", "")),
                ("Psicopedagogía", get_value("intervention_psicopedagogia", ""), get_value("intervention_psicopedagogia_strategies", "")),
                ("Fonoaudiología", get_value("intervention_fonoaudiologia", ""), get_value("intervention_fonoaudiologia_strategies", "")),
                ("Psicología", get_value("intervention_psicologia", ""), get_value("intervention_psicologia_strategies", "")),
                ("Terapia Ocupacional", get_value("intervention_terapia_ocupacional", ""), get_value("intervention_terapia_ocupacional_strategies", "")),
                ("Kinesiología", get_value("intervention_kinesiologia", ""), get_value("intervention_kinesiologia_strategies", "")),
                ("Co-educador Sordo", get_value("intervention_coeducador_sordo", ""), get_value("intervention_coeducador_sordo_strategies", "")),
                ("Intérprete Lengua de Señas", get_value("intervention_int_lengua_senas", ""), get_value("intervention_int_lengua_senas_strategies", ""))
            ]
            
            areas_with_content = [(name, objectives, strategies) for name, objectives, strategies in intervention_areas if objectives or strategies]
            
            if areas_with_content:
                for idx, (area_name, objectives, strategies) in enumerate(areas_with_content):
                    elements.append(Paragraph(f"<b>{area_name}</b>", subtitle_style))
                    if objectives:
                        elements.append(Paragraph("<b>Objetivos:</b>", subtitle_style))
                        elements.append(Paragraph(objectives, normal_style))
                        elements.append(Spacer(1, 0.1*inch))
                    if strategies:
                        elements.append(Paragraph("<b>Estrategias:</b>", subtitle_style))
                        elements.append(Paragraph(strategies, normal_style))
                    if idx < len(areas_with_content) - 1:
                        elements.append(Spacer(1, 0.2*inch))
            else:
                elements.append(Paragraph("No se registra información en esta sección.", normal_style))
            
            elements.append(Spacer(1, 0.3*inch))
            
            # V. SEGUIMIENTO DEL PAI
            follow_up = get_value("follow_up_pai", "")
            if follow_up:
                elements.append(Paragraph("IV. SEGUIMIENTO DEL PAI", section_style))
                elements.append(Spacer(1, 0.15*inch))
                elements.append(Paragraph(follow_up, normal_style))
            
            # Construir PDF
            doc.build(elements)
            
            return {
                "status": "success",
                "message": "PDF generado exitosamente desde cero",
                "filename": unique_filename,
                "file_path": str(output_file)
            }
            
        except ImportError:
            return {
                "status": "error",
                "message": "ReportLab no está instalado. Instala con: pip install reportlab",
                "filename": None,
                "file_path": None
            }
        except Exception as e:
            return {
                "status": "error",
                "message": f"Error generando PDF desde cero: {str(e)}",
                "filename": None,
                "file_path": None
            }

    @staticmethod
    def _escape_html_for_paragraph(s: Optional[str]) -> str:
        """Escapa caracteres HTML para usar en ReportLab Paragraph."""
        if s is None or not isinstance(s, str):
            return ""
        return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    @staticmethod
    def _generate_fonoaudiological_report_from_scratch(
        document_id: int,
        report_data: Dict[str, Any],
        db: Optional[Session] = None,
        output_directory: str = "files/system/students"
    ) -> Dict[str, Any]:
        """
        Genera un PDF de Informe Fonoaudiológico (Documento 8) desde cero usando ReportLab.
        Mismo estilo que PAI y Estado de Avance.
        """
        try:
            if not REPORTLAB_AVAILABLE:
                return {
                    "status": "error",
                    "message": "ReportLab no está instalado. Instala con: pip install reportlab",
                    "filename": None,
                    "file_path": None
                }

            student_name = report_data.get("student_full_name", "estudiante").replace(" ", "_")
            unique_filename = f"informe_fonoaudologico_{student_name}_{uuid.uuid4().hex[:8]}.pdf"
            output_file = Path(output_directory) / unique_filename
            output_file.parent.mkdir(parents=True, exist_ok=True)

            doc = SimpleDocTemplate(
                str(output_file),
                pagesize=A4,
                rightMargin=2.5*cm,
                leftMargin=2.5*cm,
                topMargin=3*cm,
                bottomMargin=2.5*cm
            )
            elements = []
            styles = getSampleStyleSheet()

            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                textColor=colors.HexColor('#000000'),
                spaceAfter=25,
                spaceBefore=15,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            section_style = ParagraphStyle(
                'SectionStyle',
                parent=styles['Normal'],
                fontSize=12,
                textColor=colors.HexColor('#000000'),
                spaceAfter=12,
                spaceBefore=15,
                fontName='Helvetica-Bold',
                alignment=TA_LEFT,
                backColor=colors.HexColor('#E8E8E8'),
                borderPadding=6
            )
            subtitle_style = ParagraphStyle(
                'Subtitle',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#000000'),
                spaceAfter=6,
                spaceBefore=10,
                fontName='Helvetica-Bold'
            )
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#000000'),
                alignment=TA_LEFT,
                leading=12,
                spaceAfter=4
            )
            label_style = ParagraphStyle(
                'LabelStyle',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#000000'),
                alignment=TA_LEFT,
                fontName='Helvetica-Bold',
                backColor=colors.HexColor('#F5F5F5'),
                borderPadding=4
            )

            def get_value(key: str, default: str = "") -> str:
                v = report_data.get(key)
                if v is None:
                    return default
                return str(v).strip() if v else default

            def format_date(date_val) -> str:
                if not date_val:
                    return ""
                try:
                    if isinstance(date_val, str):
                        dt = datetime.strptime(date_val, "%Y-%m-%d").date()
                        return dt.strftime("%d-%m-%Y")
                    return ""
                except Exception:
                    return str(date_val) if date_val else ""

            def add_section(title: str):
                elements.append(Paragraph(title, section_style))
                elements.append(Spacer(1, 0.2*inch))

            def add_text_block(label: str, text: str):
                if not text:
                    return
                elements.append(Paragraph(f"<b>{label}</b>", subtitle_style))
                elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(text), normal_style))
                elements.append(Spacer(1, 0.15*inch))

            # Título
            elements.append(Paragraph("Informe Fonoaudiológico (informal)", title_style))
            elements.append(Spacer(1, 0.4*inch))

            # I. IDENTIFICACIÓN DEL/LA ESTUDIANTE
            # Orden: Nombre, RUT, Fecha nacimiento, Establecimiento ID, Curso, Profesional/es responsable/s, Fecha Informe, Tipo
            add_section("I. Identificación del/la estudiante")
            student_fullname = get_value("student_full_name", "")
            student_rut = get_value("student_identification_number", "")
            student_born = format_date(get_value("student_born_date", ""))
            establishment = get_value("establishment_id", "")
            course_name = get_value("course_name", "")
            prof_names = get_value("responsible_professionals_names", "")
            report_date = format_date(get_value("report_date", ""))
            type_id = report_data.get("type_id")
            type_label = "Ingreso" if type_id == 1 else ("Reevaluación" if type_id == 2 else get_value("type_label", ""))

            id_data = []
            id_spans = []

            def _add_row(lbl: str, val: str):
                if not val:
                    return
                r = len(id_data)
                id_data.append([Paragraph(f"<b>{lbl}</b>", label_style), ""])
                id_spans.append(('SPAN', (0, r), (1, r)))
                id_data.append([Paragraph(DocumentsClass._escape_html_for_paragraph(val), normal_style), ""])
                id_spans.append(('SPAN', (0, r + 1), (1, r + 1)))

            _add_row("Nombre", student_fullname)
            _add_row("RUT", student_rut)
            _add_row("Fecha de nacimiento", student_born)
            _add_row("Establecimiento ID", establishment)
            _add_row("Curso", course_name)
            _add_row("Profesional/es responsable/s", prof_names)
            _add_row("Fecha de Informe", report_date)
            _add_row("Tipo", type_label)

            if id_data:
                tbl = Table(id_data, colWidths=[7.5*cm, 7.5*cm])
                tbl.setStyle(TableStyle([
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 8),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
                ] + id_spans))
                elements.append(tbl)
            elements.append(Spacer(1, 0.3*inch))

            # Salto de página antes de la sección II
            elements.append(PageBreak())

            # II. DESARROLLO DEL INFORME
            add_section("II. Desarrollo del Informe")
            add_text_block("Motivo de la evaluación", get_value("reason_evaluation", ""))
            add_text_block("Instrumentos de evaluación utilizado", get_value("evaluation_instruments", ""))
            add_text_block("Antecedentes relevantes", get_value("relevant_background", ""))
            add_text_block("Conductas observadas durante la evaluación", get_value("behaviors_observed", ""))

            # III. ANÁLISIS DE ÁREAS
            add_section("III. Análisis de áreas")
            add_text_block("Estructuras y funciones orofaciales y auditivas", get_value("orofacial_auditory", ""))
            add_text_block("Nivel fonológico", get_value("phonological_level", ""))
            add_text_block("Nivel morfosintáctico", get_value("morphosyntactic_level", ""))
            add_text_block("Nivel semántico", get_value("semantic_level", ""))
            add_text_block("Nivel pragmático", get_value("pragmatic_level", ""))
            add_text_block("Observaciones adicionales", get_value("additional_observations", ""))

            # IV. SÍNTESIS DIAGNÓSTICA Y/O CONCLUSIONES
            add_section("IV. Síntesis diagnóstica y/o conclusiones")
            diag = get_value("diagnostic_synthesis", "")
            if diag:
                add_text_block("Síntesis diagnóstica y/o conclusiones", diag)

            # V. SUGERENCIAS
            sug_fam = get_value("suggestions_family", "")
            sug_est = get_value("suggestions_establishment", "")
            if sug_fam or sug_est:
                add_section("V. Sugerencias")
                if sug_fam:
                    add_text_block("Sugerencias a la familia", sug_fam)
                if sug_est:
                    add_text_block("Sugerencias al establecimiento", sug_est)

            doc.build(elements)
            return {
                "status": "success",
                "message": "PDF generado exitosamente",
                "filename": unique_filename,
                "file_path": str(output_file)
            }
        except ImportError:
            return {
                "status": "error",
                "message": "ReportLab no está instalado. Instala con: pip install reportlab",
                "filename": None,
                "file_path": None
            }
        except Exception as e:
            return {
                "status": "error",
                "message": f"Error generando informe fonoaudiológico: {str(e)}",
                "filename": None,
                "file_path": None
            }

    @staticmethod
    def _generate_idtel_report_from_scratch(
        document_id: int,
        report_data: Dict[str, Any],
        db: Optional[Session] = None,
        output_directory: str = "files/system/students"
    ) -> Dict[str, Any]:
        """
        Genera un PDF del Informe fonoaudiológico IDTEL (Documento 9) desde cero usando ReportLab.
        """
        try:
            if not REPORTLAB_AVAILABLE:
                return {
                    "status": "error",
                    "message": "ReportLab no está instalado. Instala con: pip install reportlab",
                    "filename": None,
                    "file_path": None
                }
            student_name = report_data.get("student_full_name", "estudiante").replace(" ", "_")
            unique_filename = f"informe_idtel_{student_name}_{uuid.uuid4().hex[:8]}.pdf"
            output_file = Path(output_directory) / unique_filename
            output_file.parent.mkdir(parents=True, exist_ok=True)
            doc = SimpleDocTemplate(
                str(output_file),
                pagesize=A4,
                rightMargin=2.5*cm,
                leftMargin=2.5*cm,
                topMargin=3*cm,
                bottomMargin=2.5*cm
            )
            elements = []
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                textColor=colors.HexColor('#000000'),
                spaceAfter=25,
                spaceBefore=15,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            section_style = ParagraphStyle(
                'SectionStyle',
                parent=styles['Normal'],
                fontSize=12,
                textColor=colors.HexColor('#000000'),
                spaceAfter=12,
                spaceBefore=15,
                fontName='Helvetica-Bold',
                alignment=TA_LEFT,
                backColor=colors.HexColor('#E8E8E8'),
                borderPadding=6
            )
            subtitle_style = ParagraphStyle(
                'Subtitle',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#000000'),
                spaceAfter=6,
                spaceBefore=10,
                fontName='Helvetica-Bold'
            )
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#000000'),
                alignment=TA_LEFT,
                leading=12,
                spaceAfter=4
            )
            label_style = ParagraphStyle(
                'LabelStyle',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#000000'),
                alignment=TA_LEFT,
                fontName='Helvetica-Bold',
                backColor=colors.HexColor('#F5F5F5'),
                borderPadding=4
            )

            def get_value(key: str, default: str = "") -> str:
                v = report_data.get(key)
                if v is None:
                    return default
                return str(v).strip() if v else default

            def format_date(date_val) -> str:
                if not date_val:
                    return ""
                try:
                    s = str(date_val).strip()[:10]
                    if not s:
                        return ""
                    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
                        try:
                            dt = datetime.strptime(s, fmt).date()
                            return dt.strftime("%d/%m/%Y")
                        except Exception:
                            continue
                    return s
                except Exception:
                    return str(date_val) if date_val else ""

            def add_section(title: str):
                elements.append(Paragraph(title, section_style))
                elements.append(Spacer(1, 0.2*inch))

            def add_text_block(label: str, text: str):
                if not text:
                    return
                elements.append(Paragraph(f"<b>{label}</b>", subtitle_style))
                elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(text), normal_style))
                elements.append(Spacer(1, 0.15*inch))

            elements.append(Paragraph("Informe fonoaudiológico IDTEL", title_style))
            elements.append(Spacer(1, 0.4*inch))

            add_section("I. Identificación del/la estudiante")
            id_data = []
            id_spans = []
            def _add_row(lbl: str, val: str):
                if val is None:
                    val = ""
                val = str(val).strip()
                r = len(id_data)
                id_data.append([Paragraph(f"<b>{lbl}</b>", label_style), ""])
                id_spans.append(('SPAN', (0, r), (1, r)))
                id_data.append([Paragraph(DocumentsClass._escape_html_for_paragraph(val or "—"), normal_style), ""])
                id_spans.append(('SPAN', (0, r + 1), (1, r + 1)))

            age_range = report_data.get("idtel_age_range") or "6_7"
            age_range_label = "8 a 9 años 11 meses" if age_range == "8_9" else "6 a 7 años 11 meses"
            _add_row("Nombre", get_value("student_full_name"))
            _add_row("RUT", get_value("student_identification_number"))
            _add_row("Fecha de nacimiento", format_date(get_value("student_born_date")))
            _add_row("Edad", get_value("student_age"))
            _add_row("Establecimiento", get_value("establishment_id"))
            _add_row("Curso", get_value("course_name"))
            _add_row("Profesional/es responsable/s", get_value("responsible_professionals_names"))
            _add_row("Fecha de informe", format_date(get_value("report_date")))
            _add_row("Rango etario IDTEL", age_range_label)
            if id_data:
                tbl = Table(id_data, colWidths=[7.5*cm, 7.5*cm])
                tbl.setStyle(TableStyle([
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 8),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
                ] + id_spans))
                elements.append(tbl)
            elements.append(Spacer(1, 0.3*inch))

            add_section("II. Motivo y antecedentes")
            add_text_block("Motivo de la evaluación", get_value("reason_evaluation"))
            add_text_block("Antecedentes relevantes", get_value("relevant_background"))
            add_text_block("Instrumentos de evaluación utilizados", get_value("evaluation_instruments"))
            add_text_block("Conductas observadas durante la evaluación", get_value("behaviors_observed"))

            add_section("III. Puntajes IDTEL")
            score_rows = [
                ("Fonológico", report_data.get("score_phonological")),
                ("Morfosintáctico", report_data.get("score_morphosyntactic")),
                ("Semántico", report_data.get("score_semantic")),
                ("Pragmático", report_data.get("score_pragmatic")),
                ("Total IDTEL", report_data.get("score_total_idtel")),
            ]
            score_table_data = [[Paragraph("<b>Área</b>", label_style), Paragraph("<b>Puntaje</b>", label_style)]]
            for area, val in score_rows:
                if val is not None and str(val).strip() != "":
                    score_table_data.append([area, str(val)])
            if len(score_table_data) > 1:
                st = Table(score_table_data, colWidths=[5*cm, 4*cm])
                st.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e0e0e0')),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('PADDING', (0, 0), (-1, -1), 6),
                ]))
                elements.append(st)
            elements.append(Spacer(1, 0.2*inch))

            add_section("IV. Análisis por niveles")
            add_text_block("Nivel fonológico", get_value("phonological_level"))
            add_text_block("Nivel morfosintáctico", get_value("morphosyntactic_level"))
            add_text_block("Nivel semántico", get_value("semantic_level"))
            add_text_block("Nivel pragmático", get_value("pragmatic_level"))

            add_section("V. Síntesis diagnóstica y/o conclusiones")
            diag = get_value("diagnostic_synthesis")
            if diag:
                add_text_block("Síntesis diagnóstica y/o conclusiones", diag)

            sug_fam = get_value("suggestions_family")
            sug_est = get_value("suggestions_establishment")
            if sug_fam or sug_est:
                add_section("VI. Sugerencias")
                if sug_fam:
                    add_text_block("Sugerencias a la familia", sug_fam)
                if sug_est:
                    add_text_block("Sugerencias al establecimiento", sug_est)

            doc.build(elements)
            return {
                "status": "success",
                "message": "PDF informe IDTEL generado correctamente",
                "filename": unique_filename,
                "file_path": str(output_file)
            }
        except Exception as e:
            return {
                "status": "error",
                "message": f"Error generando informe IDTEL: {str(e)}",
                "filename": None,
                "file_path": None
            }

    @staticmethod
    def _generate_psychomotor_evaluation_report_from_scratch(
        document_id: int,
        report_data: Dict[str, Any],
        db: Optional[Session] = None,
        output_directory: str = "files/system/students"
    ) -> Dict[str, Any]:
        """
        Genera PDF del Informe de evaluación psicomotriz con ReportLab.
        """
        try:
            if not REPORTLAB_AVAILABLE:
                return {
                    "status": "error",
                    "message": "ReportLab no está instalado. Instala con: pip install reportlab",
                    "filename": None,
                    "file_path": None,
                }

            def _suggestions_text(v) -> str:
                if v is None:
                    return ""
                if isinstance(v, list):
                    lines = [str(x).strip() for x in v if str(x).strip()]
                    return "\n".join(f"• {t}" for t in lines) if lines else ""
                return str(v).strip()

            student_name = str(report_data.get("student_full_name", "estudiante")).replace(" ", "_")
            unique_filename = f"informe_psicomotriz_{student_name}_{uuid.uuid4().hex[:8]}.pdf"
            output_file = Path(output_directory) / unique_filename
            output_file.parent.mkdir(parents=True, exist_ok=True)
            doc = SimpleDocTemplate(
                str(output_file),
                pagesize=A4,
                rightMargin=2.5 * cm,
                leftMargin=2.5 * cm,
                topMargin=3 * cm,
                bottomMargin=2.5 * cm,
            )
            elements = []
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                "PsicoEvalTitle",
                parent=styles["Heading1"],
                fontSize=16,
                textColor=colors.HexColor("#000000"),
                spaceAfter=25,
                spaceBefore=15,
                alignment=TA_CENTER,
                fontName="Helvetica-Bold",
            )
            section_style = ParagraphStyle(
                "PsicoEvalSection",
                parent=styles["Normal"],
                fontSize=12,
                textColor=colors.HexColor("#000000"),
                spaceAfter=12,
                spaceBefore=15,
                fontName="Helvetica-Bold",
                alignment=TA_LEFT,
                backColor=colors.HexColor("#E8E8E8"),
                borderPadding=6,
            )
            subtitle_style = ParagraphStyle(
                "PsicoEvalSubtitle",
                parent=styles["Normal"],
                fontSize=10,
                textColor=colors.HexColor("#000000"),
                spaceAfter=6,
                spaceBefore=10,
                fontName="Helvetica-Bold",
            )
            normal_style = ParagraphStyle(
                "PsicoEvalNormal",
                parent=styles["Normal"],
                fontSize=10,
                textColor=colors.HexColor("#000000"),
                alignment=TA_LEFT,
                leading=12,
                spaceAfter=4,
            )
            label_style = ParagraphStyle(
                "PsicoEvalLabel",
                parent=styles["Normal"],
                fontSize=10,
                textColor=colors.HexColor("#000000"),
                alignment=TA_LEFT,
                fontName="Helvetica-Bold",
                backColor=colors.HexColor("#F5F5F5"),
                borderPadding=4,
            )

            def get_value(key: str, default: str = "") -> str:
                v = report_data.get(key)
                if v is None:
                    return default
                return str(v).strip() if v else default

            def format_date(date_val) -> str:
                if not date_val:
                    return ""
                try:
                    s = str(date_val).strip()[:10]
                    if not s:
                        return ""
                    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
                        try:
                            dt = datetime.strptime(s, fmt).date()
                            return dt.strftime("%d/%m/%Y")
                        except Exception:
                            continue
                    return s
                except Exception:
                    return str(date_val) if date_val else ""

            def add_section(title: str):
                elements.append(Paragraph(title, section_style))
                elements.append(Spacer(1, 0.2 * inch))

            def add_text_block(label: str, text: str):
                if not text:
                    return
                elements.append(Paragraph(f"<b>{label}</b>", subtitle_style))
                elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(text), normal_style))
                elements.append(Spacer(1, 0.15 * inch))

            elements.append(Paragraph("Informe de evaluación psicomotriz", title_style))
            elements.append(Spacer(1, 0.4 * inch))

            add_section("I. Identificación del/la estudiante")
            id_data = []
            id_spans = []

            def _add_row(lbl: str, val: str):
                if val is None:
                    val = ""
                val = str(val).strip()
                r = len(id_data)
                id_data.append([Paragraph(f"<b>{lbl}</b>", label_style), ""])
                id_spans.append(("SPAN", (0, r), (1, r)))
                id_data.append([Paragraph(DocumentsClass._escape_html_for_paragraph(val or "—"), normal_style), ""])
                id_spans.append(("SPAN", (0, r + 1), (1, r + 1)))

            _add_row("Nombre", get_value("student_full_name"))
            _add_row("RUT", get_value("student_identification_number"))
            _add_row("Fecha de nacimiento", format_date(get_value("student_born_date")))
            _add_row("Edad", get_value("student_age"))
            _add_row("Establecimiento", get_value("establishment_id"))
            _add_row("Curso", get_value("course_name"))
            _add_row("Profesional/es responsable/s", get_value("responsible_professionals_names"))
            _add_row("Fecha de informe", format_date(get_value("report_date")))
            if id_data:
                tbl = Table(id_data, colWidths=[7.5 * cm, 7.5 * cm])
                tbl.setStyle(
                    TableStyle(
                        [
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("LEFTPADDING", (0, 0), (-1, -1), 8),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                            ("TOPPADDING", (0, 0), (-1, -1), 8),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
                        ]
                        + id_spans
                    )
                )
                elements.append(tbl)
            elements.append(Spacer(1, 0.3 * inch))

            add_section("II. Desarrollo del informe")
            add_text_block("Motivo de la evaluación", get_value("reason_evaluation"))
            add_text_block("Instrumentos de evaluación utilizados", get_value("evaluation_instruments"))
            add_text_block("Antecedentes relevantes", get_value("relevant_background"))
            add_text_block("Conductas observadas durante la evaluación", get_value("behaviors_observed"))

            add_section("III. Análisis de áreas")
            add_text_block("Análisis general", get_value("area_analysis_general"))
            add_text_block("Desempeño motor grueso", get_value("area_gross_motor"))
            add_text_block("Integración motora y sensorial", get_value("area_motor_sensory_integration"))
            add_text_block("Habilidades motoras finas relacionadas", get_value("area_fine_motor_skills"))
            add_text_block("Desarrollo postural", get_value("area_postural_development"))
            add_text_block("Adaptaciones físicas al entorno escolar", get_value("area_physical_adaptations"))
            add_text_block("Desempeño funcional en actividades escolares", get_value("area_functional_school"))
            add_text_block("Desarrollo del esquema corporal", get_value("area_body_schema"))

            add_section("IV. Síntesis diagnóstica y/o conclusiones")
            add_text_block("Síntesis", get_value("diagnostic_synthesis"))

            sug_fam = _suggestions_text(report_data.get("suggestions_family"))
            sug_est = _suggestions_text(report_data.get("suggestions_establishment"))
            if sug_fam or sug_est:
                add_section("V. Sugerencias")
                if sug_fam:
                    add_text_block("Sugerencias a la familia", sug_fam)
                if sug_est:
                    add_text_block("Sugerencias al establecimiento", sug_est)

            doc.build(elements)
            return {
                "status": "success",
                "message": "PDF informe psicomotriz generado correctamente",
                "filename": unique_filename,
                "file_path": str(output_file),
            }
        except Exception as e:
            return {
                "status": "error",
                "message": f"Error generando informe psicomotriz: {str(e)}",
                "filename": None,
                "file_path": None,
            }

    @staticmethod
    def _generate_school_integration_exit_certificate_from_scratch(
        document_id: int,
        cert_data: Dict[str, Any],
        db: Optional[Session] = None,
        output_directory: str = "files/system/students"
    ) -> Dict[str, Any]:
        """
        Genera el PDF del Certificado de egreso PIE (Documento 23) desde cero.
        Formato: título, declaración con profesional/estudiante/nee, apoderado, firma del profesional.
        """
        try:
            if not REPORTLAB_AVAILABLE:
                return {
                    "status": "error",
                    "message": "ReportLab no está instalado. Instala con: pip install reportlab",
                    "filename": None,
                    "file_path": None
                }
            def g(k: str, d: str = "") -> str:
                v = cert_data.get(k)
                if v is None:
                    return d
                return str(v).strip() if v else d

            student_name = g("student_full_name", "estudiante").replace(" ", "_")
            unique_filename = f"certificado_egreso_pie_{student_name}_{uuid.uuid4().hex[:8]}.pdf"
            output_file = Path(output_directory) / unique_filename
            output_file.parent.mkdir(parents=True, exist_ok=True)

            doc = SimpleDocTemplate(
                str(output_file),
                pagesize=A4,
                rightMargin=2.5*cm,
                leftMargin=2.5*cm,
                topMargin=3*cm,
                bottomMargin=2.5*cm
            )
            elements = []
            styles = getSampleStyleSheet()

            title_style = ParagraphStyle(
                'CertTitle',
                parent=styles['Heading1'],
                fontSize=20,
                textColor=colors.HexColor('#000000'),
                spaceAfter=28,
                spaceBefore=18,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            normal_style = ParagraphStyle(
                'CertNormal',
                parent=styles['Normal'],
                fontSize=12,
                textColor=colors.HexColor('#000000'),
                alignment=TA_JUSTIFY,
                leading=17,
                spaceAfter=10
            )
            center_style = ParagraphStyle(
                'CertCenter',
                parent=styles['Normal'],
                fontSize=12,
                textColor=colors.HexColor('#000000'),
                alignment=TA_CENTER,
                leading=14,
                spaceAfter=4
            )

            prof_name = DocumentsClass._escape_html_for_paragraph(g("professional_fullname", ""))
            prof_rut = DocumentsClass._escape_html_for_paragraph(g("professional_rut", ""))
            prof_role = DocumentsClass._escape_html_for_paragraph(g("professional_role", ""))
            st_name = DocumentsClass._escape_html_for_paragraph(g("student_full_name", ""))
            st_rut = DocumentsClass._escape_html_for_paragraph(g("student_rut", ""))
            course = DocumentsClass._escape_html_for_paragraph(g("course_name", "Sin curso"))
            establishment = DocumentsClass._escape_html_for_paragraph(g("establishment_name", ""))
            nee = DocumentsClass._escape_html_for_paragraph(g("nee_name", ""))
            guardian_name = DocumentsClass._escape_html_for_paragraph(g("guardian_fullname", ""))
            guardian_rut = DocumentsClass._escape_html_for_paragraph(g("guardian_rut", ""))

            if not prof_name:
                prof_name = "________________"
            if not prof_rut:
                prof_rut = "________________"
            if not prof_role:
                prof_role = "________________"
            if not st_name:
                st_name = "________________"
            if not st_rut:
                st_rut = "________________"
            if not establishment:
                establishment = "________________"
            if not nee:
                nee = "________________"
            if not guardian_name:
                guardian_name = "________________"
            if not guardian_rut:
                guardian_rut = "________________"

            elements.append(Paragraph("Certificado de egreso PIE", title_style))
            elements.append(Spacer(1, 0.5*inch))

            decl = (
                f'Por medio de este documento, el/la profesional <b>{prof_name}</b>, RUT <b>{prof_rut}</b>, '
                f'con cargo de <b>{prof_role}</b>, certifica que, luego del proceso de revaluación diagnóstica '
                f'integral realizado según lo indicado por el Decreto 170, el/la estudiante <b>{st_name}</b> '
                f'RUT <b>{st_rut}</b> quien cursa <b>{course}</b> en el establecimiento <b>{establishment}</b>, '
                f'ha superado su Necesidad Educativa Especial (NEE) asociada a <b>{nee}</b> por lo que se '
                f'considera su egreso del Programa de Integración Escolar PIE.'
            )
            elements.append(Paragraph(decl, normal_style))
            elements.append(Spacer(1, 0.4*inch))

            ack = (
                f'En tanto, el apoderado {guardian_name}, RUT {guardian_rut} '
                f'queda en total conocimiento de esta información.'
            )
            elements.append(Paragraph(ack, normal_style))
            elements.append(Spacer(1, 0.6*inch))

            elements.append(Paragraph("_________________________", center_style))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph(prof_name, center_style))
            elements.append(Paragraph(prof_rut, center_style))
            elements.append(Paragraph(prof_role, center_style))

            doc.build(elements)
            return {
                "status": "success",
                "message": "PDF generado exitosamente",
                "filename": unique_filename,
                "file_path": str(output_file)
            }
        except ImportError:
            return {
                "status": "error",
                "message": "ReportLab no está instalado. Instala con: pip install reportlab",
                "filename": None,
                "file_path": None
            }
        except Exception as e:
            return {
                "status": "error",
                "message": f"Error generando certificado de egreso PIE: {str(e)}",
                "filename": None,
                "file_path": None
            }

    @staticmethod
    def _generate_tea_certificate_from_scratch(
        document_id: int,
        cert_data: Dict[str, Any],
        db: Optional[Session] = None,
        output_directory: str = "files/system/students"
    ) -> Dict[str, Any]:
        """
        Genera el PDF del Certificado Ley TEA (Documento 25) desde cero.
        Certificado para respaldar la salida del apoderado del trabajo (Ley N°21.545 / artículo 66 quinquies Código del Trabajo).
        """
        try:
            if not REPORTLAB_AVAILABLE:
                return {
                    "status": "error",
                    "message": "ReportLab no está instalado. Instala con: pip install reportlab",
                    "filename": None,
                    "file_path": None
                }
            def g(k: str, d: str = "") -> str:
                v = cert_data.get(k)
                if v is None:
                    return d
                return str(v).strip() if v else d

            student_name = g("student_full_name", "estudiante").replace(" ", "_")
            unique_filename = f"certificado_ley_tea_{student_name}_{uuid.uuid4().hex[:8]}.pdf"
            output_file = Path(output_directory) / unique_filename
            output_file.parent.mkdir(parents=True, exist_ok=True)

            doc = SimpleDocTemplate(
                str(output_file),
                pagesize=A4,
                rightMargin=2.5*cm,
                leftMargin=2.5*cm,
                topMargin=2.5*cm,
                bottomMargin=2.5*cm
            )
            elements = []
            styles = getSampleStyleSheet()

            title_style = ParagraphStyle(
                'TeaCertTitle',
                parent=styles['Heading1'],
                fontSize=16,
                textColor=colors.HexColor('#000000'),
                spaceAfter=20,
                spaceBefore=12,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            normal_style = ParagraphStyle(
                'TeaCertNormal',
                parent=styles['Normal'],
                fontSize=11,
                textColor=colors.HexColor('#000000'),
                alignment=TA_JUSTIFY,
                leading=16,
                spaceAfter=12
            )

            establishment = DocumentsClass._escape_html_for_paragraph(g("establishment_name", ""))
            guardian_name = DocumentsClass._escape_html_for_paragraph(g("guardian_fullname", ""))
            guardian_rut = DocumentsClass._escape_html_for_paragraph(g("guardian_rut", ""))
            st_name = DocumentsClass._escape_html_for_paragraph(g("student_full_name", ""))
            course = DocumentsClass._escape_html_for_paragraph(g("course_name", ""))
            nee_name = DocumentsClass._escape_html_for_paragraph(g("nee_name", ""))
            attendance_date = DocumentsClass._escape_html_for_paragraph(g("attendance_date", ""))
            attendance_time = DocumentsClass._escape_html_for_paragraph(g("attendance_time", ""))

            if not establishment:
                establishment = "________________"
            if not guardian_name:
                guardian_name = "________________"
            if not guardian_rut:
                guardian_rut = "________________"
            if not st_name:
                st_name = "________________"
            if not course:
                course = "________________"
            if not nee_name:
                nee_name = "________________"
            if not attendance_date:
                attendance_date = "________________"
            if not attendance_time:
                attendance_time = "________________"

            elements.append(Paragraph("Certificado de asistencia del apoderado", title_style))
            elements.append(Spacer(1, 0.4*inch))

            p1 = (
                f'Por medio del presente certificado, y en el marco de la Ley N°21.545, conocida como Ley TEA, '
                f'el establecimiento educacional <b>{establishment}</b> informa que <b>{guardian_name}</b>, '
                f'RUT <b>{guardian_rut}</b>, apoderado/a de <b>{st_name}</b>, matriculado/a en <b>{course}</b>, '
                f'quien posee diagnóstico de <b>{nee_name}</b>.'
            )
            elements.append(Paragraph(p1, normal_style))
            elements.append(Spacer(1, 0.2*inch))

            p2 = (
                f'Se confirma que el/la apoderado/a antes mencionado/a asistió al establecimiento el día '
                f'<b>{attendance_date}</b>, en el horario de <b>{attendance_time}</b>, en el contexto de apoyo al/la '
                f'estudiante o de participación en coordinación con el equipo educativo, de acuerdo al Protocolo '
                f'de Desregulación Emocional y Conductual y al artículo 66 quinquies del Código del Trabajo.'
            )
            elements.append(Paragraph(p2, normal_style))
            elements.append(Spacer(1, 0.2*inch))

            p3 = (
                'El presente certificado tiene por objeto respaldar y justificar su salida desde el lugar de trabajo ante el empleador.'
            )
            elements.append(Paragraph(p3, normal_style))
            elements.append(Spacer(1, 0.2*inch))

            p4 = (
                '"Los padres, madres o cuidadores de personas con trastorno del espectro autista tendrán derecho a '
                'ausentarse de sus labores para acompañar a la persona de la cual son responsables en las atenciones '
                'que requiera, sin que ello pueda significar una disminución de sus remuneraciones ni el no pago de '
                'las mismas, conforme al artículo 66 quinquies del Código del Trabajo."'
            )
            elements.append(Paragraph(p4, normal_style))
            elements.append(Spacer(1, 0.2*inch))

            p5 = (
                'Se extiende el presente documento a solicitud del/la interesado/a, para los fines que estime pertinentes.'
            )
            elements.append(Paragraph(p5, normal_style))

            doc.build(elements)
            return {
                "status": "success",
                "message": "PDF generado exitosamente",
                "filename": unique_filename,
                "file_path": str(output_file)
            }
        except ImportError:
            return {
                "status": "error",
                "message": "ReportLab no está instalado. Instala con: pip install reportlab",
                "filename": None,
                "file_path": None
            }
        except Exception as e:
            return {
                "status": "error",
                "message": f"Error generando certificado Ley TEA: {str(e)}",
                "filename": None,
                "file_path": None
            }

    @staticmethod
    def _generate_pedagogical_evaluation_classroom_first_grade_from_scratch(
        document_id: int,
        doc_data: Dict[str, Any],
        db: Optional[Session] = None,
        output_directory: str = "files/system/students"
    ) -> Dict[str, Any]:
        """
        Genera el PDF del documento 31 (Pauta de evaluación pedagógica - Docente de aula - 1º Básico)
        desde cero, con todos los datos y respuestas en español.
        """
        try:
            if not REPORTLAB_AVAILABLE:
                return {
                    "status": "error",
                    "message": "ReportLab no está instalado. Instala con: pip install reportlab",
                    "filename": None,
                    "file_path": None
                }

            def get_value(key: str, default: str = "") -> str:
                v = doc_data.get(key)
                if v is None:
                    return default
                return str(v).strip() if v else default

            def format_date(date_val) -> str:
                if not date_val:
                    return ""
                try:
                    if isinstance(date_val, str) and len(date_val) >= 10:
                        d = datetime.strptime(date_val[:10], "%Y-%m-%d").date()
                        return d.strftime("%d/%m/%Y")
                    if hasattr(date_val, "strftime"):
                        return date_val.strftime("%d/%m/%Y")
                    return str(date_val)[:10]
                except Exception:
                    return str(date_val)[:10] if date_val else ""

            student_name = get_value("student_full_name", "estudiante").replace(" ", "_")
            unique_filename = f"pauta_eval_pedagogica_1basico_{student_name}_{uuid.uuid4().hex[:8]}.pdf"
            output_file = Path(output_directory) / unique_filename
            output_file.parent.mkdir(parents=True, exist_ok=True)

            doc = SimpleDocTemplate(
                str(output_file),
                pagesize=A4,
                rightMargin=2*cm,
                leftMargin=2*cm,
                topMargin=2*cm,
                bottomMargin=2*cm
            )
            elements = []
            styles = getSampleStyleSheet()

            title_style = ParagraphStyle(
                'Doc31Title',
                parent=styles['Heading1'],
                fontSize=14,
                textColor=colors.HexColor('#000000'),
                spaceAfter=14,
                spaceBefore=8,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            section_style = ParagraphStyle(
                'Doc31Section',
                parent=styles['Normal'],
                fontSize=11,
                textColor=colors.HexColor('#000000'),
                spaceAfter=6,
                spaceBefore=10,
                fontName='Helvetica-Bold',
                backColor=colors.HexColor('#E8E8E8'),
                borderPadding=6
            )
            subtitle_style = ParagraphStyle(
                'Doc31Subtitle',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#000000'),
                spaceAfter=4,
                spaceBefore=6,
                fontName='Helvetica-Bold'
            )
            normal_style = ParagraphStyle(
                'Doc31Normal',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#000000'),
                alignment=TA_LEFT,
                leading=12,
                spaceAfter=2
            )
            legend_style = ParagraphStyle(
                'Doc31Legend',
                parent=styles['Normal'],
                fontSize=9,
                textColor=colors.HexColor('#333333'),
                alignment=TA_LEFT,
                spaceAfter=6,
                spaceBefore=2
            )
            header_cell_style = ParagraphStyle(
                'Doc31HeaderCell',
                parent=styles['Normal'],
                fontSize=9,
                textColor=colors.white,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            cell_center_style = ParagraphStyle(
                'Doc31CellCenter',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#000000'),
                alignment=TA_CENTER
            )
            # Leyenda de respuestas: S, G, O, P/V, N, N/O
            DOC31_RESPONSE_COLS = ["S", "G", "O", "P/V", "N", "N/O"]
            DOC31_LEGEND_TEXT = "S: Siempre   G: Generalmente   O: Ocasionalmente   P/V: Pocas veces   N: Nunca   N/O: No observado"
            # Indicadores del primer cuadro (III. Actitud) - 12 ítems
            DOC31_ACTITUD_INDICADORES = [
                "Manifiesta curiosidad e interés por aprender",
                "Demuestra disposición para compartir ideas, experiencias y opiniones con los demás",
                "Presta atención a las actividades y dinámicas de la clase",
                "Adopta una actitud positiva hacia sí mismo y hacia el aprendizaje",
                "Inicia, persiste y finaliza tareas de forma autónoma",
                "Realiza sus trabajos de manera ordenada y sistemática",
                "Aborda y resuelve problemas de manera flexible y creativa durante las actividades",
                "Muestra entusiasmo y motivación frente a nuevos desafíos",
                "Se adapta con flexibilidad a cambios, imprevistos o nuevas rutinas de trabajo",
                "Participa activamente en actividades grupales y colaborativas",
                "Es receptivo a la retroalimentación y la aplica para mejorar",
                "Busca apoyo o aclara dudas cuando enfrenta dificultades",
            ]
            # Indicadores del segundo cuadro (IV. Lenguaje) - 13 ítems
            DOC31_LENGUAJE_INDICADORES = [
                "Identifica los sonidos que componen las palabras (conciencia fonológica)",
                "Reconoce que las palabras son unidades de significado separadas por espacios",
                "Asocia sonido con letra (principio alfabético)",
                "Lee palabras de uso frecuente en forma global",
                "Comprende textos breves leídos por un adulto",
                "Responde preguntas literales sobre el texto",
                "Produce escritura con letras convencionales",
                "Escribe su nombre y palabras conocidas",
                "Sigue la direccionalidad de la escritura (izquierda a derecha)",
                "Diferencia entre letras y números",
                "Participa en conversaciones respetando turnos",
                "Expresa ideas o experiencias en forma oral",
                "Comprende y sigue instrucciones orales",
            ]
            # Indicadores del tercer cuadro (V. Matemática) - 11 ítems
            DOC31_MATEMATICA_INDICADORES = [
                "Lee números del 0 al 20 y los representa de forma concreta, pictórica y simbólica",
                "Compara y ordena números del 0 al 20 de menor a mayor y viceversa",
                "Compone y descompone números del 0 al 20 de manera aditiva",
                "Determina las unidades y decenas en números del 0 al 20",
                "Demuestra comprensión de la adición y la sustracción de números del 0 al 20, resolviendo sumas y restas progresivas",
                "Reconoce, crea y continúa patrones repetitivos (sonidos, figuras, ritmos) y patrones numéricos hasta el 20, en secuencias crecientes y decrecientes",
                "Describe la posición de objetos y personas en relación a sí mismos y a otros, utilizando un lenguaje común (por ejemplo, derecha, izquierda)",
                "Identifica y relaciona figuras 3D y 2D en el entorno, utilizando material concreto",
                "Identifica y compara la longitud de objetos, utilizando términos como largo y corto",
                "Recolecta y registra datos para responder preguntas estadísticas sobre sí mismo y su entorno, utilizando bloques, tablas de conteo y pictogramas",
                "Construye, lee e interpreta pictogramas",
            ]

            def normalize_response_doc31(val):
                if val is None or (isinstance(val, str) and not val.strip()):
                    return None
                v = str(val).strip().upper().replace(" ", "")
                if v in ("S", "SIEMPRE"):
                    return "S"
                if v in ("G", "GENERALMENTE"):
                    return "G"
                if v in ("O", "OCASIONALMENTE"):
                    return "O"
                if v in ("P/V", "PV", "POCASVECES", "P"):
                    return "P/V"
                if v in ("N/O", "NO", "NOOBSERVADO"):
                    return "N/O"
                if v in ("N", "NUNCA"):
                    return "N"
                # Por si viene solo una letra
                if len(v) == 1 and v in "SGOPN":
                    return "P/V" if v == "P" else v
                return None

            def build_doc31_indicator_table(header_label, item_prefix, num_items, get_val_fn, indicator_labels=None):
                """Construye tabla con leyenda, encabezado (Indicadores + S, G, O, P/V, N, N/O) y filas con check.
                Si indicator_labels está definido y tiene al menos num_items, se usa como texto del indicador."""
                rows = []
                rows.append([Paragraph(DocumentsClass._escape_html_for_paragraph("Indicadores"), header_cell_style)] + [Paragraph(c, header_cell_style) for c in DOC31_RESPONSE_COLS])
                for i in range(1, num_items + 1):
                    val = get_val_fn(i)
                    col_key = normalize_response_doc31(val)
                    if indicator_labels and len(indicator_labels) >= i:
                        label_text = indicator_labels[i - 1]
                    else:
                        label_text = f"{item_prefix} {i}"
                    cells = [Paragraph(DocumentsClass._escape_html_for_paragraph(label_text), normal_style)]
                    for c in DOC31_RESPONSE_COLS:
                        cells.append(Paragraph("✓" if col_key == c else "", cell_center_style))
                    rows.append(cells)
                col_ind = 8*cm
                col_opt = (16*cm - col_ind) / 6
                col_widths = [col_ind] + [col_opt] * 6
                tbl = Table(rows, colWidths=col_widths)
                tbl.setStyle(TableStyle([
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                    ('ALIGN', (1, 0), (-1, -1), 'CENTER'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ('TOPPADDING', (0, 0), (-1, -1), 5),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#505050')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ]))
                return tbl

            # Título
            elements.append(Paragraph("Pauta de evaluación pedagógica - Docente de aula - 1º Básico", title_style))
            elements.append(Spacer(1, 0.25*inch))

            # I. Datos del estudiante e informe
            elements.append(Paragraph("I. DATOS DEL ESTUDIANTE E INFORME", section_style))
            elements.append(Spacer(1, 0.15*inch))
            col_label = 5*cm
            col_value = 11*cm
            id_data = []
            id_data.append([Paragraph("<b>Nombre completo:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("student_full_name", "—")), normal_style)])
            id_data.append([Paragraph("<b>RUT:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("student_identification_number", "—")), normal_style)])
            id_data.append([Paragraph("<b>Fecha de nacimiento:</b>", subtitle_style), Paragraph(format_date(doc_data.get("student_born_date")) or "—", normal_style)])
            id_data.append([Paragraph("<b>Edad:</b>", subtitle_style), Paragraph(get_value("student_age", "—"), normal_style)])
            id_data.append([Paragraph("<b>Establecimiento:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("establishment_id", "—")), normal_style)])
            id_data.append([Paragraph("<b>Curso:</b>", subtitle_style), Paragraph(get_value("course", "—"), normal_style)])
            id_data.append([Paragraph("<b>Fecha del informe:</b>", subtitle_style), Paragraph(format_date(doc_data.get("report_date")) or "—", normal_style)])
            id_data.append([Paragraph("<b>Repeticiones:</b>", subtitle_style), Paragraph(get_value("repetitions", "—"), normal_style)])
            report_type_val = (get_value("report_type", "") or "").strip()
            if report_type_val.lower() == "anual":
                report_type_val = "Anual"
            elif report_type_val.lower() == "semestral":
                report_type_val = "Semestral"
            id_data.append([Paragraph("<b>Tipo de informe:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(report_type_val or "—"), normal_style)])
            tbl_id = Table(id_data, colWidths=[col_label, col_value])
            tbl_id.setStyle(TableStyle([
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('LEFTPADDING', (0, 0), (-1, -1), 8),
                ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
            ]))
            elements.append(tbl_id)
            elements.append(Spacer(1, 0.25*inch))

            # II. Fortalezas de la situación escolar (school_situation_strengths puede ser JSON)
            elements.append(Paragraph("II. FORTALEZAS DE LA SITUACIÓN ESCOLAR", section_style))
            elements.append(Spacer(1, 0.1*inch))
            strengths_raw = doc_data.get("school_situation_strengths")
            strength_labels_es = {
                "strength_creative": "Creativo/a",
                "strength_autonomous": "Autónomo/a",
                "strength_motivated": "Motivado/a",
                "strength_tolerant_frustration": "Tolerante a la frustración",
                "strength_respects_rules": "Respeta las normas",
                "strength_follows_instructions": "Sigue instrucciones",
                "strength_peers_relationship": "Buena relación con pares",
                "strength_attends_gladly": "Asiste con gusto",
                "strength_family_communication": "Comunicación con la familia",
                "strength_persistent": "Persistente",
                "strength_self_confident": "Seguro/a de sí mismo/a",
                "strength_participates": "Participa",
                "strength_pays_attention": "Presta atención",
                "strength_interested_learning": "Interesado/a en el aprendizaje",
                "strength_completes_tasks": "Completa tareas",
                "strength_adults_relationship": "Buena relación con adultos",
                "strength_attends_regularly": "Asiste regularmente",
                "strength_family_committed": "Familia comprometida",
            }
            strengths_text = ""
            if strengths_raw is not None and strengths_raw != "":
                if isinstance(strengths_raw, dict):
                    strengths_list = [strength_labels_es.get(k, k) for k, v in strengths_raw.items() if v is True]
                    strengths_text = ", ".join(strengths_list) if strengths_list else "Ninguna marcada."
                elif isinstance(strengths_raw, str):
                    try:
                        import json
                        parsed = json.loads(strengths_raw)
                        if isinstance(parsed, dict):
                            strengths_list = [strength_labels_es.get(k, k) for k, v in parsed.items() if v is True]
                            strengths_text = ", ".join(strengths_list) if strengths_list else "Ninguna marcada."
                        else:
                            strengths_text = str(strengths_raw).strip() or "—"
                    except Exception:
                        strengths_text = str(strengths_raw).strip() or "—"
                else:
                    strengths_text = str(strengths_raw).strip() or "—"
            if not strengths_text:
                strengths_text = "—"
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(strengths_text), normal_style))
            elements.append(Spacer(1, 0.15*inch))
            elements.append(Paragraph("Observaciones generales", subtitle_style))
            elements.append(Spacer(1, 0.1*inch))
            obs = get_value("observations", "")
            if obs:
                elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(obs), normal_style))
            else:
                elements.append(Paragraph("—", normal_style))
            elements.append(Spacer(1, 0.25*inch))

            # III. Actitud (ítems 1-12: tabla con preguntas y check)
            elements.append(Paragraph("III. ACTITUD", section_style))
            elements.append(Spacer(1, 0.08*inch))
            elements.append(Paragraph(DOC31_LEGEND_TEXT, legend_style))
            elements.append(Spacer(1, 0.06*inch))
            att_table = build_doc31_indicator_table(
                "Actitud", "Actitud", 12, lambda i: get_value(f"attitude_{i}", ""),
                indicator_labels=DOC31_ACTITUD_INDICADORES
            )
            elements.append(att_table)
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("Observaciones actitud:", subtitle_style))
            obs_att = get_value("observations_attitude", "")
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(obs_att) if obs_att else "—", normal_style))
            elements.append(Spacer(1, 0.25*inch))

            # IV. Lenguaje (ítems 1-13: tabla con preguntas y check)
            elements.append(Paragraph("IV. LENGUAJE", section_style))
            elements.append(Spacer(1, 0.08*inch))
            elements.append(Paragraph(DOC31_LEGEND_TEXT, legend_style))
            elements.append(Spacer(1, 0.06*inch))
            lang_table = build_doc31_indicator_table(
                "Lenguaje", "Lenguaje", 13, lambda i: get_value(f"language_{i}", ""),
                indicator_labels=DOC31_LENGUAJE_INDICADORES
            )
            elements.append(lang_table)
            elements.append(Spacer(1, 0.12*inch))
            # Tipo de lectura, Nivel de comprensión, Nivel de escritura (pueden ser varios, separados por coma)
            reading_type_labels = {
                "en_desarrollo": "En desarrollo",
                "silabica": "Silábica",
                "palabra_a_palabra": "Palabra a palabra",
                "por_unidades_cortas": "Por unidades cortas",
                "fluida": "Fluida",
            }
            comprehension_labels = {
                "en_desarrollo": "En desarrollo",
                "literal": "Literal",
                "inferencial": "Inferencial",
                "critica": "Crítica",
            }
            writing_labels = {
                "en_desarrollo": "En desarrollo",
                "copia": "Copia",
                "dictado": "Dictado",
                "espontanea": "Espontánea",
            }

            def format_multi_value(raw: str, labels_dict: dict) -> str:
                if not raw or not str(raw).strip():
                    return "—"
                parts = [p.strip().lower().replace(" ", "_") for p in str(raw).split(",") if p and p.strip()]
                if not parts:
                    return "—"
                display = [labels_dict.get(p) or p for p in parts]
                return ", ".join(display)

            rt_display = format_multi_value(get_value("reading_type", ""), reading_type_labels)
            cl_display = format_multi_value(get_value("comprehension_level", ""), comprehension_labels)
            wl_display = format_multi_value(get_value("writing_level", ""), writing_labels)
            elements.append(Paragraph("<b>Tipo de lectura:</b> " + DocumentsClass._escape_html_for_paragraph(rt_display), normal_style))
            elements.append(Paragraph("<b>Nivel de comprensión:</b> " + DocumentsClass._escape_html_for_paragraph(cl_display), normal_style))
            elements.append(Paragraph("<b>Nivel de escritura:</b> " + DocumentsClass._escape_html_for_paragraph(wl_display), normal_style))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("Observaciones lenguaje:", subtitle_style))
            obs_lang = get_value("observations_language", "")
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(obs_lang) if obs_lang else "—", normal_style))
            elements.append(Spacer(1, 0.25*inch))

            # V. Matemática (ítems 1-11: tabla con preguntas y check)
            elements.append(Paragraph("V. MATEMÁTICA", section_style))
            elements.append(Spacer(1, 0.08*inch))
            elements.append(Paragraph(DOC31_LEGEND_TEXT, legend_style))
            elements.append(Spacer(1, 0.06*inch))
            math_table = build_doc31_indicator_table(
                "Matemática", "Matemática", 11, lambda i: get_value(f"mathematics_{i}", ""),
                indicator_labels=DOC31_MATEMATICA_INDICADORES
            )
            elements.append(math_table)
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("Observaciones matemática:", subtitle_style))
            obs_math = get_value("observations_mathematics", "")
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(obs_math) if obs_math else "—", normal_style))

            doc.build(elements)
            return {
                "status": "success",
                "message": "PDF generado exitosamente desde cero",
                "filename": unique_filename,
                "file_path": str(output_file)
            }
        except ImportError:
            return {
                "status": "error",
                "message": "ReportLab no está instalado. Instala con: pip install reportlab",
                "filename": None,
                "file_path": None
            }
        except Exception as e:
            return {
                "status": "error",
                "message": f"Error generando PDF documento 31: {str(e)}",
                "filename": None,
                "file_path": None
            }

    @staticmethod
    def _generate_pedagogical_evaluation_classroom_second_grade_from_scratch(
        document_id: int,
        doc_data: Dict[str, Any],
        db: Optional[Session] = None,
        output_directory: str = "files/system/students"
    ) -> Dict[str, Any]:
        """
        Genera el PDF del documento 32 (Pauta de evaluación pedagógica - Docente de aula - 2º Básico)
        desde cero. 15 indicadores de lenguaje, 12 de matemática.
        """
        try:
            if not REPORTLAB_AVAILABLE:
                return {
                    "status": "error",
                    "message": "ReportLab no está instalado. Instala con: pip install reportlab",
                    "filename": None,
                    "file_path": None
                }

            def get_value(key: str, default: str = "") -> str:
                v = doc_data.get(key)
                if v is None:
                    return default
                return str(v).strip() if v else default

            def format_date(date_val) -> str:
                if not date_val:
                    return ""
                try:
                    if isinstance(date_val, str) and len(date_val) >= 10:
                        d = datetime.strptime(date_val[:10], "%Y-%m-%d").date()
                        return d.strftime("%d/%m/%Y")
                    if hasattr(date_val, "strftime"):
                        return date_val.strftime("%d/%m/%Y")
                    return str(date_val)[:10]
                except Exception:
                    return str(date_val)[:10] if date_val else ""

            student_name = get_value("student_full_name", "estudiante").replace(" ", "_")
            unique_filename = f"pauta_eval_pedagogica_2basico_{student_name}_{uuid.uuid4().hex[:8]}.pdf"
            output_file = Path(output_directory) / unique_filename
            output_file.parent.mkdir(parents=True, exist_ok=True)

            doc = SimpleDocTemplate(
                str(output_file),
                pagesize=A4,
                rightMargin=2*cm,
                leftMargin=2*cm,
                topMargin=2*cm,
                bottomMargin=2*cm
            )
            elements = []
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'Doc32Title', parent=styles['Heading1'], fontSize=14, textColor=colors.HexColor('#000000'),
                spaceAfter=14, spaceBefore=8, alignment=TA_CENTER, fontName='Helvetica-Bold'
            )
            section_style = ParagraphStyle(
                'Doc32Section', parent=styles['Normal'], fontSize=11, textColor=colors.HexColor('#000000'),
                spaceAfter=6, spaceBefore=10, fontName='Helvetica-Bold',
                backColor=colors.HexColor('#E8E8E8'), borderPadding=6
            )
            subtitle_style = ParagraphStyle(
                'Doc32Subtitle', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#000000'),
                spaceAfter=4, spaceBefore=6, fontName='Helvetica-Bold'
            )
            normal_style = ParagraphStyle(
                'Doc32Normal', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#000000'),
                alignment=TA_LEFT, leading=12, spaceAfter=2
            )
            legend_style = ParagraphStyle(
                'Doc32Legend', parent=styles['Normal'], fontSize=9, textColor=colors.HexColor('#333333'),
                alignment=TA_LEFT, spaceAfter=6, spaceBefore=2
            )
            header_cell_style = ParagraphStyle(
                'Doc32HeaderCell', parent=styles['Normal'], fontSize=9, textColor=colors.white,
                alignment=TA_CENTER, fontName='Helvetica-Bold'
            )
            cell_center_style = ParagraphStyle(
                'Doc32CellCenter', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#000000'), alignment=TA_CENTER
            )
            DOC32_RESPONSE_COLS = ["S", "G", "O", "P/V", "N", "N/O"]
            DOC32_LEGEND_TEXT = "S: Siempre   G: Generalmente   O: Ocasionalmente   P/V: Pocas veces   N: Nunca   N/O: No observado"
            DOC32_ACTITUD_INDICADORES = [
                "Manifiesta curiosidad e interés por aprender",
                "Demuestra disposición para compartir ideas, experiencias y opiniones con los demás",
                "Presta atención a las actividades y dinámicas de la clase",
                "Adopta una actitud positiva hacia sí mismo y hacia el aprendizaje",
                "Inicia, persiste y finaliza tareas de forma autónoma",
                "Realiza sus trabajos de manera ordenada y sistemática",
                "Aborda y resuelve problemas de manera flexible y creativa durante las actividades",
                "Muestra entusiasmo y motivación frente a nuevos desafíos",
                "Se adapta con flexibilidad a cambios, imprevistos o nuevas rutinas de trabajo",
                "Participa activamente en actividades grupales y colaborativas",
                "Es receptivo a la retroalimentación y la aplica para mejorar",
                "Busca apoyo o aclara dudas cuando enfrenta dificultades",
            ]
            DOC32_LENGUAJE_INDICADORES = [
                "Lee textos con palabras que incluyen hiatos, diptongos, grupos consonánticos y combinaciones complejas",
                "Lee en voz alta con precisión y fluidez progresiva, autocorrigiéndose cuando es necesario",
                "Lee y comprende textos no literarios como cartas, notas e instrucciones",
                "Comprende narraciones extrayendo información explícita e implícita, recreando la secuencia de acciones, describiendo personajes, ambiente, entre otros",
                "Comprende y extrae información de explicaciones, relatos o instrucciones verbales",
                "Escribe oraciones simples con coherencia y respetando las normas gramaticales básicas",
                "Escribe de manera creativa, expresando sus ideas y pensamientos",
                "Redacta narraciones incluyendo inicio, desarrollo y desenlace claros y coherentes",
                "Planifica sus escritos a partir de observaciones y discusiones sobre temas de interés personal",
                "Escribe, revisa y edita sus textos, planificando ideas, corrigiendo errores y mejorando la redacción",
                "Escribe de forma legible y separa las palabras correctamente",
                "Participa activamente en conversaciones grupales, expresando ideas, mostrando interés y respetando turnos de habla",
                "Interactúa de acuerdo con las convenciones sociales adecuándose a distintos contextos",
                "Expone ideas de forma coherente y articulada sobre temas de interés, utilizando un vocabulario variado y pronunciación adecuada",
                "Trabaja en equipo y desempeña diferentes roles para desarrollar su lenguaje y autoestima"
            ]
            DOC32_MATEMATICA_INDICADORES = [
                "Lee números del 0 al 100 y los representa de manera concreta, pictórica y simbólica",
                "Compara y ordena números del 0 al 100, de menor a mayor y viceversa",
                "Compone y descompone números del 0 al 100 de manera aditiva",
                "Identifica las unidades y decenas en números del 0 al 100, representando las cantidades de acuerdo a su valor posicional",
                "Demuestra comprensión de la adición y sustracción en el ámbito de números del 0 al 100",
                "Demuestra comprensión de la multiplicación",
                "Crea, representa y continúa patrones numéricos",
                "Describe, compara y construye figuras 2D",
                "Describe, compara y construye figuras 3D",
                "Determina la longitud de objetos utilizando unidades de medida estandarizadas y no estandarizadas",
                "Recolecta y registra datos para responder preguntas estadísticas sobre juegos con monedas y dados, utilizando bloques, tablas de conteo y pictogramas",
                "Construye, lee e interpreta pictogramas con escala y gráficos de barra simples"
            ]

            def normalize_response(val):
                if val is None or (isinstance(val, str) and not val.strip()):
                    return None
                v = str(val).strip().upper().replace(" ", "")
                if v in ("S", "SIEMPRE"): return "S"
                if v in ("G", "GENERALMENTE"): return "G"
                if v in ("O", "OCASIONALMENTE"): return "O"
                if v in ("P/V", "PV", "POCASVECES", "P"): return "P/V"
                if v in ("N/O", "NO", "NOOBSERVADO"): return "N/O"
                if v in ("N", "NUNCA"): return "N"
                if len(v) == 1 and v in "SGOPN": return "P/V" if v == "P" else v
                return None

            def build_table(header_label, num_items, get_val_fn, indicator_labels):
                rows = []
                rows.append([Paragraph(DocumentsClass._escape_html_for_paragraph("Indicadores"), header_cell_style)] + [Paragraph(c, header_cell_style) for c in DOC32_RESPONSE_COLS])
                for i in range(1, num_items + 1):
                    val = get_val_fn(i)
                    col_key = normalize_response(val)
                    label_text = indicator_labels[i - 1] if indicator_labels and len(indicator_labels) >= i else f"{header_label} {i}"
                    cells = [Paragraph(DocumentsClass._escape_html_for_paragraph(label_text), normal_style)]
                    for c in DOC32_RESPONSE_COLS:
                        cells.append(Paragraph("✓" if col_key == c else "", cell_center_style))
                    rows.append(cells)
                col_ind = 8*cm
                col_opt = (16*cm - col_ind) / 6
                col_widths = [col_ind] + [col_opt] * 6
                tbl = Table(rows, colWidths=col_widths)
                tbl.setStyle(TableStyle([
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'), ('ALIGN', (0, 0), (0, -1), 'LEFT'), ('ALIGN', (1, 0), (-1, -1), 'CENTER'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6), ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ('TOPPADDING', (0, 0), (-1, -1), 5), ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#505050')), ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ]))
                return tbl

            def format_multi_value(raw: str, labels_dict: dict) -> str:
                if not raw or not str(raw).strip():
                    return "—"
                parts = [p.strip().lower().replace(" ", "_") for p in str(raw).split(",") if p and p.strip()]
                if not parts:
                    return "—"
                display = [labels_dict.get(p) or p for p in parts]
                return ", ".join(display)
            reading_type_labels = {"en_desarrollo": "En desarrollo", "silabica": "Silábica", "palabra_a_palabra": "Palabra a palabra", "por_unidades_cortas": "Por unidades cortas", "fluida": "Fluida"}
            comprehension_labels = {"en_desarrollo": "En desarrollo", "literal": "Literal", "inferencial": "Inferencial", "critica": "Crítica"}
            writing_labels = {"en_desarrollo": "En desarrollo", "copia": "Copia", "dictado": "Dictado", "espontanea": "Espontánea"}

            elements.append(Paragraph("Pauta de evaluación pedagógica - Docente de aula - 2º Básico", title_style))
            elements.append(Spacer(1, 0.25*inch))

            # I. Datos del estudiante e informe
            elements.append(Paragraph("I. DATOS DEL ESTUDIANTE E INFORME", section_style))
            elements.append(Spacer(1, 0.15*inch))
            col_label, col_value = 5*cm, 11*cm
            id_data = [
                [Paragraph("<b>Nombre completo:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("student_full_name", "—")), normal_style)],
                [Paragraph("<b>RUT:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("student_identification_number", "—")), normal_style)],
                [Paragraph("<b>Fecha de nacimiento:</b>", subtitle_style), Paragraph(format_date(doc_data.get("student_born_date")) or "—", normal_style)],
                [Paragraph("<b>Edad:</b>", subtitle_style), Paragraph(get_value("student_age", "—"), normal_style)],
                [Paragraph("<b>Establecimiento:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("establishment_id", "—")), normal_style)],
                [Paragraph("<b>Curso:</b>", subtitle_style), Paragraph(get_value("course", "—"), normal_style)],
                [Paragraph("<b>Fecha del informe:</b>", subtitle_style), Paragraph(format_date(doc_data.get("report_date")) or "—", normal_style)],
                [Paragraph("<b>Repeticiones:</b>", subtitle_style), Paragraph(get_value("repetitions", "—"), normal_style)],
            ]
            report_type_val = (get_value("report_type", "") or "").strip()
            report_type_val = "Anual" if report_type_val.lower() == "anual" else ("Semestral" if report_type_val.lower() == "semestral" else report_type_val or "—")
            id_data.append([Paragraph("<b>Tipo de informe:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(report_type_val), normal_style)])
            tbl_id = Table(id_data, colWidths=[col_label, col_value])
            tbl_id.setStyle(TableStyle([
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('LEFTPADDING', (0, 0), (-1, -1), 8),
                ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
            ]))
            elements.append(tbl_id)
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("II. FORTALEZAS DE LA SITUACIÓN ESCOLAR", section_style))
            elements.append(Spacer(1, 0.1*inch))
            strengths_raw = doc_data.get("school_situation_strengths")
            strength_labels_es = {"strength_creative": "Creativo/a", "strength_autonomous": "Autónomo/a", "strength_motivated": "Motivado/a", "strength_tolerant_frustration": "Tolerante a la frustración", "strength_respects_rules": "Respeta las normas", "strength_follows_instructions": "Sigue instrucciones", "strength_peers_relationship": "Buena relación con pares", "strength_attends_gladly": "Asiste con gusto", "strength_family_communication": "Comunicación con la familia", "strength_persistent": "Persistente", "strength_self_confident": "Seguro/a de sí mismo/a", "strength_participates": "Participa", "strength_pays_attention": "Presta atención", "strength_interested_learning": "Interesado/a en el aprendizaje", "strength_completes_tasks": "Completa tareas", "strength_adults_relationship": "Buena relación con adultos", "strength_attends_regularly": "Asiste regularmente", "strength_family_committed": "Familia comprometida"}
            strengths_text = "—"
            if strengths_raw not in (None, ""):
                if isinstance(strengths_raw, dict):
                    strengths_list = [strength_labels_es.get(k, k) for k, v in strengths_raw.items() if v is True]
                    strengths_text = ", ".join(strengths_list) if strengths_list else "Ninguna marcada."
                elif isinstance(strengths_raw, str):
                    try:
                        parsed = __import__("json").loads(strengths_raw)
                        if isinstance(parsed, dict):
                            strengths_list = [strength_labels_es.get(k, k) for k, v in parsed.items() if v is True]
                            strengths_text = ", ".join(strengths_list) if strengths_list else "Ninguna marcada."
                        else:
                            strengths_text = str(strengths_raw).strip() or "—"
                    except Exception:
                        strengths_text = str(strengths_raw).strip() or "—"
                else:
                    strengths_text = str(strengths_raw).strip() or "—"
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(strengths_text), normal_style))
            elements.append(Spacer(1, 0.15*inch))
            elements.append(Paragraph("Observaciones generales", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations", "")) if get_value("observations", "") else "—", normal_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("III. ACTITUD", section_style))
            elements.append(Spacer(1, 0.08*inch))
            elements.append(Paragraph(DOC32_LEGEND_TEXT, legend_style))
            elements.append(Spacer(1, 0.06*inch))
            elements.append(build_table("Actitud", 12, lambda i: get_value(f"attitude_{i}", ""), DOC32_ACTITUD_INDICADORES))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("Observaciones actitud:", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations_attitude", "")) if get_value("observations_attitude", "") else "—", normal_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("IV. LENGUAJE", section_style))
            elements.append(Spacer(1, 0.08*inch))
            elements.append(Paragraph(DOC32_LEGEND_TEXT, legend_style))
            elements.append(Spacer(1, 0.06*inch))
            elements.append(build_table("Lenguaje", 15, lambda i: get_value(f"language_{i}", ""), DOC32_LENGUAJE_INDICADORES))
            elements.append(Spacer(1, 0.12*inch))
            rt_display = format_multi_value(get_value("reading_type", ""), reading_type_labels)
            cl_display = format_multi_value(get_value("comprehension_level", ""), comprehension_labels)
            wl_display = format_multi_value(get_value("writing_level", ""), writing_labels)
            elements.append(Paragraph("<b>Tipo de lectura:</b> " + DocumentsClass._escape_html_for_paragraph(rt_display), normal_style))
            elements.append(Paragraph("<b>Nivel de comprensión:</b> " + DocumentsClass._escape_html_for_paragraph(cl_display), normal_style))
            elements.append(Paragraph("<b>Nivel de escritura:</b> " + DocumentsClass._escape_html_for_paragraph(wl_display), normal_style))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("Observaciones lenguaje:", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations_language", "")) if get_value("observations_language", "") else "—", normal_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("V. MATEMÁTICA", section_style))
            elements.append(Spacer(1, 0.08*inch))
            elements.append(Paragraph(DOC32_LEGEND_TEXT, legend_style))
            elements.append(Spacer(1, 0.06*inch))
            elements.append(build_table("Matemática", 12, lambda i: get_value(f"mathematics_{i}", ""), DOC32_MATEMATICA_INDICADORES))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("Observaciones matemática:", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations_mathematics", "")) if get_value("observations_mathematics", "") else "—", normal_style))
            doc.build(elements)
            return {"status": "success", "message": "PDF generado exitosamente desde cero", "filename": unique_filename, "file_path": str(output_file)}
        except ImportError:
            return {"status": "error", "message": "ReportLab no está instalado.", "filename": None, "file_path": None}
        except Exception as e:
            return {"status": "error", "message": f"Error generando PDF documento 32: {str(e)}", "filename": None, "file_path": None}

    @staticmethod
    def _generate_pedagogical_evaluation_classroom_third_grade_from_scratch(
        document_id: int,
        doc_data: Dict[str, Any],
        db: Optional[Session] = None,
        output_directory: str = "files/system/students"
    ) -> Dict[str, Any]:
        """
        Genera el PDF del documento 33 (Pauta de evaluación pedagógica - Docente de aula - 3º Básico)
        desde cero. 15 indicadores de lenguaje, 12 de matemática.
        """
        try:
            if not REPORTLAB_AVAILABLE:
                return {
                    "status": "error",
                    "message": "ReportLab no está instalado. Instala con: pip install reportlab",
                    "filename": None,
                    "file_path": None
                }

            def get_value(key: str, default: str = "") -> str:
                v = doc_data.get(key)
                if v is None:
                    return default
                return str(v).strip() if v else default

            def format_date(date_val) -> str:
                if not date_val:
                    return ""
                try:
                    if isinstance(date_val, str) and len(date_val) >= 10:
                        d = datetime.strptime(date_val[:10], "%Y-%m-%d").date()
                        return d.strftime("%d/%m/%Y")
                    if hasattr(date_val, "strftime"):
                        return date_val.strftime("%d/%m/%Y")
                    return str(date_val)[:10]
                except Exception:
                    return str(date_val)[:10] if date_val else ""

            student_name = get_value("student_full_name", "estudiante").replace(" ", "_")
            unique_filename = f"pauta_eval_pedagogica_3basico_{student_name}_{uuid.uuid4().hex[:8]}.pdf"
            output_file = Path(output_directory) / unique_filename
            output_file.parent.mkdir(parents=True, exist_ok=True)

            doc = SimpleDocTemplate(
                str(output_file),
                pagesize=A4,
                rightMargin=2*cm,
                leftMargin=2*cm,
                topMargin=2*cm,
                bottomMargin=2*cm
            )
            elements = []
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'Doc33Title', parent=styles['Heading1'], fontSize=14, textColor=colors.HexColor('#000000'),
                spaceAfter=14, spaceBefore=8, alignment=TA_CENTER, fontName='Helvetica-Bold'
            )
            section_style = ParagraphStyle(
                'Doc33Section', parent=styles['Normal'], fontSize=11, textColor=colors.HexColor('#000000'),
                spaceAfter=6, spaceBefore=10, fontName='Helvetica-Bold',
                backColor=colors.HexColor('#E8E8E8'), borderPadding=6
            )
            subtitle_style = ParagraphStyle(
                'Doc33Subtitle', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#000000'),
                spaceAfter=4, spaceBefore=6, fontName='Helvetica-Bold'
            )
            normal_style = ParagraphStyle(
                'Doc33Normal', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#000000'),
                alignment=TA_LEFT, leading=12, spaceAfter=2
            )
            legend_style = ParagraphStyle(
                'Doc33Legend', parent=styles['Normal'], fontSize=9, textColor=colors.HexColor('#333333'),
                alignment=TA_LEFT, spaceAfter=6, spaceBefore=2
            )
            header_cell_style = ParagraphStyle(
                'Doc33HeaderCell', parent=styles['Normal'], fontSize=9, textColor=colors.white,
                alignment=TA_CENTER, fontName='Helvetica-Bold'
            )
            cell_center_style = ParagraphStyle(
                'Doc33CellCenter', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#000000'), alignment=TA_CENTER
            )
            DOC33_RESPONSE_COLS = ["S", "G", "O", "P/V", "N", "N/O"]
            DOC33_LEGEND_TEXT = "S: Siempre   G: Generalmente   O: Ocasionalmente   P/V: Pocas veces   N: Nunca   N/O: No observado"
            DOC33_ACTITUD_INDICADORES = [
                "Manifiesta curiosidad e interés por aprender",
                "Demuestra disposición para compartir ideas, experiencias y opiniones con los demás",
                "Presta atención a las actividades y dinámicas de la clase",
                "Adopta una actitud positiva hacia sí mismo y hacia el aprendizaje",
                "Inicia, persiste y finaliza tareas de forma autónoma",
                "Realiza sus trabajos de manera ordenada y sistemática",
                "Aborda y resuelve problemas de manera flexible y creativa durante las actividades",
                "Muestra entusiasmo y motivación frente a nuevos desafíos",
                "Se adapta con flexibilidad a cambios, imprevistos o nuevas rutinas de trabajo",
                "Participa activamente en actividades grupales y colaborativas",
                "Es receptivo a la retroalimentación y la aplica para mejorar",
                "Busca apoyo o aclara dudas cuando enfrenta dificultades",
            ]
            DOC33_LENGUAJE_INDICADORES = [
                "Lee en voz alta con fluidez, respetando puntuación y pronunciación",
                "Aplica estrategias de comprensión lectora (subrayar, formular preguntas, recapitular)",
                "Comprende textos literarios y no literarios adecuados al nivel",
                "Emite opiniones sobre personajes y describe ambientes en textos narrativos",
                "Comprende e interpreta lenguaje figurado en poemas y cuentos",
                "Escribe textos breves con coherencia y estructura adecuada",
                "Planifica sus escritos y revisa ortografía y redacción",
                "Redacta narraciones con inicio, desarrollo y desenlace coherentes",
                "Participa en conversaciones respetando turnos y aportando ideas",
                "Expone ideas de forma coherente con vocabulario variado",
                "Utiliza convenciones sociales adecuadas en distintos contextos",
                "Lee y comprende instrucciones, cartas y textos informativos breves",
                "Identifica ideas principales e información explícita e implícita en textos",
                "Escribe de forma legible con letra clara y separación correcta de palabras",
                "Trabaja en equipo y desempeña diferentes roles en actividades de lenguaje",
            ]
            DOC33_MATEMATICA_INDICADORES = [
                "Lee números hasta 1.000 y los representa de forma concreta, pictórica y simbólica",
                "Compara y ordena números hasta 1.000, utilizando la recta numérica o la tabla posicional",
                "Identifica y describe unidades, decenas y centenas en números hasta 1.000, representando las cantidades de acuerdo a su valor posicional",
                "Demuestra comprensión de la adición y sustracción de números hasta 1.000",
                "Demuestra una comprensión progresiva de las tablas de multiplicar hasta 10",
                "Demuestra comprensión de la división, en el contexto de las tablas de multiplicar hasta 10x10",
                "Resuelve problemas rutinarios en contextos cotidianos, involucrando las cuatro operaciones básicas (no combinadas)",
                "Demuestra comprensión de las fracciones de uso común: 1/4, 1/3, 1/2, 2/3, 3/4",
                "Genera, describe y registra patrones numéricos utilizando diversas estrategias",
                "Demuestra comprensión de la relación entre figuras 2D y 3D",
                "Demuestra comprensión del perímetro de una figura, tanto regular como irregular",
                "Demuestra comprensión de la medición del peso en gramos (g) y kilogramos (kg)",
                "Construye, lee e interpreta pictogramas y gráficos de barra simple, basados en información recolectada o proporcionada",
            ]

            def normalize_response(val):
                if val is None or (isinstance(val, str) and not val.strip()):
                    return None
                v = str(val).strip().upper().replace(" ", "")
                if v in ("S", "SIEMPRE"): return "S"
                if v in ("G", "GENERALMENTE"): return "G"
                if v in ("O", "OCASIONALMENTE"): return "O"
                if v in ("P/V", "PV", "POCASVECES", "P"): return "P/V"
                if v in ("N/O", "NO", "NOOBSERVADO"): return "N/O"
                if v in ("N", "NUNCA"): return "N"
                if len(v) == 1 and v in "SGOPN": return "P/V" if v == "P" else v
                return None

            def build_table(header_label, num_items, get_val_fn, indicator_labels):
                rows = []
                rows.append([Paragraph(DocumentsClass._escape_html_for_paragraph("Indicadores"), header_cell_style)] + [Paragraph(c, header_cell_style) for c in DOC33_RESPONSE_COLS])
                for i in range(1, num_items + 1):
                    val = get_val_fn(i)
                    col_key = normalize_response(val)
                    label_text = indicator_labels[i - 1] if indicator_labels and len(indicator_labels) >= i else f"{header_label} {i}"
                    cells = [Paragraph(DocumentsClass._escape_html_for_paragraph(label_text), normal_style)]
                    for c in DOC33_RESPONSE_COLS:
                        cells.append(Paragraph("✓" if col_key == c else "", cell_center_style))
                    rows.append(cells)
                col_ind = 8*cm
                col_opt = (16*cm - col_ind) / 6
                col_widths = [col_ind] + [col_opt] * 6
                tbl = Table(rows, colWidths=col_widths)
                tbl.setStyle(TableStyle([
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'), ('ALIGN', (0, 0), (0, -1), 'LEFT'), ('ALIGN', (1, 0), (-1, -1), 'CENTER'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6), ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ('TOPPADDING', (0, 0), (-1, -1), 5), ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#505050')), ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ]))
                return tbl

            def format_multi_value(raw: str, labels_dict: dict) -> str:
                if not raw or not str(raw).strip():
                    return "—"
                parts = [p.strip().lower().replace(" ", "_") for p in str(raw).split(",") if p and p.strip()]
                if not parts:
                    return "—"
                display = [labels_dict.get(p) or p for p in parts]
                return ", ".join(display)
            reading_type_labels = {"en_desarrollo": "En desarrollo", "silabica": "Silábica", "palabra_a_palabra": "Palabra a palabra", "por_unidades_cortas": "Por unidades cortas", "fluida": "Fluida"}
            comprehension_labels = {"en_desarrollo": "En desarrollo", "literal": "Literal", "inferencial": "Inferencial", "critica": "Crítica"}
            writing_labels = {"en_desarrollo": "En desarrollo", "copia": "Copia", "dictado": "Dictado", "espontanea": "Espontánea"}

            elements.append(Paragraph("Pauta de evaluación pedagógica - Docente de aula - 3º Básico", title_style))
            elements.append(Spacer(1, 0.25*inch))

            elements.append(Paragraph("I. DATOS DEL ESTUDIANTE E INFORME", section_style))
            elements.append(Spacer(1, 0.15*inch))
            col_label, col_value = 5*cm, 11*cm
            id_data = [
                [Paragraph("<b>Nombre completo:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("student_full_name", "—")), normal_style)],
                [Paragraph("<b>RUT:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("student_identification_number", "—")), normal_style)],
                [Paragraph("<b>Fecha de nacimiento:</b>", subtitle_style), Paragraph(format_date(doc_data.get("student_born_date")) or "—", normal_style)],
                [Paragraph("<b>Edad:</b>", subtitle_style), Paragraph(get_value("student_age", "—"), normal_style)],
                [Paragraph("<b>Establecimiento:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("establishment_id", "—")), normal_style)],
                [Paragraph("<b>Curso:</b>", subtitle_style), Paragraph(get_value("course", "—"), normal_style)],
                [Paragraph("<b>Fecha del informe:</b>", subtitle_style), Paragraph(format_date(doc_data.get("report_date")) or "—", normal_style)],
                [Paragraph("<b>Repeticiones:</b>", subtitle_style), Paragraph(get_value("repetitions", "—"), normal_style)],
            ]
            report_type_val = (get_value("report_type", "") or "").strip()
            report_type_val = "Anual" if report_type_val.lower() == "anual" else ("Semestral" if report_type_val.lower() == "semestral" else report_type_val or "—")
            id_data.append([Paragraph("<b>Tipo de informe:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(report_type_val), normal_style)])
            tbl_id = Table(id_data, colWidths=[col_label, col_value])
            tbl_id.setStyle(TableStyle([
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('LEFTPADDING', (0, 0), (-1, -1), 8),
                ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
            ]))
            elements.append(tbl_id)
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("II. FORTALEZAS DE LA SITUACIÓN ESCOLAR", section_style))
            elements.append(Spacer(1, 0.1*inch))
            strengths_raw = doc_data.get("school_situation_strengths")
            strength_labels_es = {"strength_creative": "Creativo/a", "strength_autonomous": "Autónomo/a", "strength_motivated": "Motivado/a", "strength_tolerant_frustration": "Tolerante a la frustración", "strength_respects_rules": "Respeta las normas", "strength_follows_instructions": "Sigue instrucciones", "strength_peers_relationship": "Buena relación con pares", "strength_attends_gladly": "Asiste con gusto", "strength_family_communication": "Comunicación con la familia", "strength_persistent": "Persistente", "strength_self_confident": "Seguro/a de sí mismo/a", "strength_participates": "Participa", "strength_pays_attention": "Presta atención", "strength_interested_learning": "Interesado/a en el aprendizaje", "strength_completes_tasks": "Completa tareas", "strength_adults_relationship": "Buena relación con adultos", "strength_attends_regularly": "Asiste regularmente", "strength_family_committed": "Familia comprometida"}
            strengths_text = "—"
            if strengths_raw not in (None, ""):
                if isinstance(strengths_raw, dict):
                    strengths_list = [strength_labels_es.get(k, k) for k, v in strengths_raw.items() if v is True]
                    strengths_text = ", ".join(strengths_list) if strengths_list else "Ninguna marcada."
                elif isinstance(strengths_raw, str):
                    try:
                        parsed = __import__("json").loads(strengths_raw)
                        if isinstance(parsed, dict):
                            strengths_list = [strength_labels_es.get(k, k) for k, v in parsed.items() if v is True]
                            strengths_text = ", ".join(strengths_list) if strengths_list else "Ninguna marcada."
                        else:
                            strengths_text = str(strengths_raw).strip() or "—"
                    except Exception:
                        strengths_text = str(strengths_raw).strip() or "—"
                else:
                    strengths_text = str(strengths_raw).strip() or "—"
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(strengths_text), normal_style))
            elements.append(Spacer(1, 0.15*inch))
            elements.append(Paragraph("Observaciones generales", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations", "")) if get_value("observations", "") else "—", normal_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("III. ACTITUD", section_style))
            elements.append(Spacer(1, 0.08*inch))
            elements.append(Paragraph(DOC33_LEGEND_TEXT, legend_style))
            elements.append(Spacer(1, 0.06*inch))
            elements.append(build_table("Actitud", 12, lambda i: get_value(f"attitude_{i}", ""), DOC33_ACTITUD_INDICADORES))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("Observaciones actitud:", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations_attitude", "")) if get_value("observations_attitude", "") else "—", normal_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("IV. LENGUAJE", section_style))
            elements.append(Spacer(1, 0.08*inch))
            elements.append(Paragraph(DOC33_LEGEND_TEXT, legend_style))
            elements.append(Spacer(1, 0.06*inch))
            elements.append(build_table("Lenguaje", 15, lambda i: get_value(f"language_{i}", ""), DOC33_LENGUAJE_INDICADORES))
            elements.append(Spacer(1, 0.12*inch))
            rt_display = format_multi_value(get_value("reading_type", ""), reading_type_labels)
            cl_display = format_multi_value(get_value("comprehension_level", ""), comprehension_labels)
            wl_display = format_multi_value(get_value("writing_level", ""), writing_labels)
            elements.append(Paragraph("<b>Tipo de lectura:</b> " + DocumentsClass._escape_html_for_paragraph(rt_display), normal_style))
            elements.append(Paragraph("<b>Nivel de comprensión:</b> " + DocumentsClass._escape_html_for_paragraph(cl_display), normal_style))
            elements.append(Paragraph("<b>Nivel de escritura:</b> " + DocumentsClass._escape_html_for_paragraph(wl_display), normal_style))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("Observaciones lenguaje:", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations_language", "")) if get_value("observations_language", "") else "—", normal_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("V. MATEMÁTICA", section_style))
            elements.append(Spacer(1, 0.08*inch))
            elements.append(Paragraph(DOC33_LEGEND_TEXT, legend_style))
            elements.append(Spacer(1, 0.06*inch))
            elements.append(build_table("Matemática", 13, lambda i: get_value(f"mathematics_{i}", ""), DOC33_MATEMATICA_INDICADORES))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("Observaciones matemática:", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations_mathematics", "")) if get_value("observations_mathematics", "") else "—", normal_style))
            doc.build(elements)
            return {"status": "success", "message": "PDF generado exitosamente desde cero", "filename": unique_filename, "file_path": str(output_file)}
        except ImportError:
            return {"status": "error", "message": "ReportLab no está instalado.", "filename": None, "file_path": None}
        except Exception as e:
            return {"status": "error", "message": f"Error generando PDF documento 33: {str(e)}", "filename": None, "file_path": None}

    @staticmethod
    def _generate_pedagogical_evaluation_classroom_fourth_grade_from_scratch(
        document_id: int,
        doc_data: Dict[str, Any],
        db: Optional[Session] = None,
        output_directory: str = "files/system/students"
    ) -> Dict[str, Any]:
        """Genera el PDF del documento 34 (Pauta de evaluación pedagógica - Docente de aula - 4º Básico). 15 lenguaje, 13 matemática."""
        try:
            if not REPORTLAB_AVAILABLE:
                return {"status": "error", "message": "ReportLab no esta instalado.", "filename": None, "file_path": None}

            def get_value(key: str, default: str = "") -> str:
                v = doc_data.get(key)
                if v is None:
                    return default
                return str(v).strip() if v else default

            def format_date(date_val) -> str:
                if not date_val:
                    return ""
                try:
                    if isinstance(date_val, str) and len(date_val) >= 10:
                        d = datetime.strptime(date_val[:10], "%Y-%m-%d").date()
                        return d.strftime("%d/%m/%Y")
                    if hasattr(date_val, "strftime"):
                        return date_val.strftime("%d/%m/%Y")
                    return str(date_val)[:10]
                except Exception:
                    return str(date_val)[:10] if date_val else ""

            student_name = get_value("student_full_name", "estudiante").replace(" ", "_")
            unique_filename = f"pauta_eval_pedagogica_4basico_{student_name}_{uuid.uuid4().hex[:8]}.pdf"
            output_file = Path(output_directory) / unique_filename
            output_file.parent.mkdir(parents=True, exist_ok=True)

            doc = SimpleDocTemplate(str(output_file), pagesize=A4, rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
            elements = []
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle('Doc34Title', parent=styles['Heading1'], fontSize=14, textColor=colors.HexColor('#000000'), spaceAfter=14, spaceBefore=8, alignment=TA_CENTER, fontName='Helvetica-Bold')
            section_style = ParagraphStyle('Doc34Section', parent=styles['Normal'], fontSize=11, textColor=colors.HexColor('#000000'), spaceAfter=6, spaceBefore=10, fontName='Helvetica-Bold', backColor=colors.HexColor('#E8E8E8'), borderPadding=6)
            subtitle_style = ParagraphStyle('Doc34Subtitle', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#000000'), spaceAfter=4, spaceBefore=6, fontName='Helvetica-Bold')
            normal_style = ParagraphStyle('Doc34Normal', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#000000'), alignment=TA_LEFT, leading=12, spaceAfter=2)
            legend_style = ParagraphStyle('Doc34Legend', parent=styles['Normal'], fontSize=9, textColor=colors.HexColor('#333333'), alignment=TA_LEFT, spaceAfter=6, spaceBefore=2)
            header_cell_style = ParagraphStyle('Doc34HeaderCell', parent=styles['Normal'], fontSize=9, textColor=colors.white, alignment=TA_CENTER, fontName='Helvetica-Bold')
            cell_center_style = ParagraphStyle('Doc34CellCenter', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#000000'), alignment=TA_CENTER)
            DOC34_RESPONSE_COLS = ["S", "G", "O", "P/V", "N", "N/O"]
            DOC34_LEGEND_TEXT = "S: Siempre   G: Generalmente   O: Ocasionalmente   P/V: Pocas veces   N: Nunca   N/O: No observado"
            DOC34_ACTITUD_INDICADORES = [
                "Manifiesta curiosidad e interes por aprender", "Demuestra disposicion para compartir ideas, experiencias y opiniones con los demas",
                "Presta atencion a las actividades y dinamicas de la clase", "Adopta una actitud positiva hacia si mismo y hacia el aprendizaje",
                "Inicia, persiste y finaliza tareas de forma autonoma", "Realiza sus trabajos de manera ordenada y sistematica",
                "Aborda y resuelve problemas de manera flexible y creativa durante las actividades", "Muestra entusiasmo y motivacion frente a nuevos desafios",
                "Se adapta con flexibilidad a cambios, imprevistos o nuevas rutinas de trabajo", "Participa activamente en actividades grupales y colaborativas",
                "Es receptivo a la retroalimentacion y la aplica para mejorar", "Busca apoyo o aclara dudas cuando enfrenta dificultades",
            ]
            DOC34_LENGUAJE_INDICADORES = [
                "Lee en voz alta con fluidez, pronunciacion, puntuacion, entonacion y velocidad adecuadas",
                "Lee y comprende narraciones extrayendo informacion, identificando problemas y soluciones, comparando personajes y emitiendo opiniones",
                "Lee en forma autonoma textos no literarios, extrayendo informacion explicita e implicita y utilizando organizadores textuales",
                "Comprende e interpreta el lenguaje figurado en poemas",
                "Aplica estrategias de comprension lectora (relacionar con experiencia previa, releer, visualizar, formular preguntas, subrayar informacion clave)",
                "Aplica estrategias para determinar el significado de palabras nuevas (contexto, claves del texto, diccionario)",
                "Comprende, interactua y extrae informacion relevante de textos orales (instrucciones, explicaciones, relatos, noticias)",
                "Busca y organiza informacion de diversas fuentes (libros, internet, diarios, revistas, enciclopedias)",
                "Escribe textos de forma autonoma y clara, organizando las ideas con logica, conectores y vocabulario variado",
                "Planifica sus escritos (proposito, audiencia, generacion de ideas)",
                "Revisa y edita sus textos (corregir errores, mejorar la redaccion segun el proposito)",
                "Aplica normas de ortografia y gramatica",
                "Escribe con letra clara y legible",
                "Participa activamente en conversaciones grupales (claridad, respeto de turnos, interes, empatia)",
                "Expone ideas de forma coherente en presentaciones orales (vocabulario variado, pronunciacion y gestos adecuados)",
                "Interactua respetando convenciones sociales y adaptando la comunicacion a distintos contextos",
                "Trabaja en equipo desempenando roles variados y promoviendo habilidades comunicativas y autoestima",
            ]
            DOC34_MATEMATICA_INDICADORES = [
                "Representa y describe numeros naturales hasta 10.000",
                "Describe y aplica estrategias de calculo mental",
                "Demuestra comprension de la adicion y sustraccion de numeros hasta 1.000",
                "Demuestra comprension de la multiplicacion de numeros de tres digitos por numeros de un digito",
                "Demuestra comprension de la division con dividendos de dos digitos y divisores de un digito",
                "Resuelve problemas rutinarios y no rutinarios, seleccionando y utilizando la operacion matemática adecuada",
                "Demuestra comprension de las fracciones",
                "Resuelve adiciones y sustracciones de fracciones con igual denominador",
                "Identifica y describe patrones numericos en tablas que involucren una operacion",
                "Resuelve ecuaciones e inecuaciones de un paso que involucren adicion y sustraccion de numeros hasta 100",
                "Traslada, rota y refleja figuras 2D",
                "Construye y compara angulos con el transportador",
                "Mide longitudes con unidades estandarizadas (metros, centimetros) y realiza conversiones entre estas unidades (m a cm y viceversa)",
                "Demuestra comprension del concepto de area de un rectangulo y de un cuadrado",
                "Demuestra comprension del concepto de volumen de un cuerpo",
                "Realiza encuestas, analiza los datos y los compara con los resultados de muestras aleatorias, usando tablas y graficos",
                "Lee e interpreta pictogramas y graficos de barra simple con escala, comunicando sus conclusiones",
            ]

            def normalize_response(val):
                if val is None or (isinstance(val, str) and not val.strip()):
                    return None
                v = str(val).strip().upper().replace(" ", "")
                if v in ("S", "SIEMPRE"): return "S"
                if v in ("G", "GENERALMENTE"): return "G"
                if v in ("O", "OCASIONALMENTE"): return "O"
                if v in ("P/V", "PV", "POCASVECES", "P"): return "P/V"
                if v in ("N/O", "NO", "NOOBSERVADO"): return "N/O"
                if v in ("N", "NUNCA"): return "N"
                if len(v) == 1 and v in "SGOPN": return "P/V" if v == "P" else v
                return None

            def build_table(header_label, num_items, get_val_fn, indicator_labels):
                rows = []
                rows.append([Paragraph(DocumentsClass._escape_html_for_paragraph("Indicadores"), header_cell_style)] + [Paragraph(c, header_cell_style) for c in DOC34_RESPONSE_COLS])
                for i in range(1, num_items + 1):
                    val = get_val_fn(i)
                    col_key = normalize_response(val)
                    label_text = indicator_labels[i - 1] if indicator_labels and len(indicator_labels) >= i else f"{header_label} {i}"
                    cells = [Paragraph(DocumentsClass._escape_html_for_paragraph(label_text), normal_style)]
                    for c in DOC34_RESPONSE_COLS:
                        cells.append(Paragraph("✓" if col_key == c else "", cell_center_style))
                    rows.append(cells)
                col_ind = 8*cm
                col_opt = (16*cm - col_ind) / 6
                col_widths = [col_ind] + [col_opt] * 6
                tbl = Table(rows, colWidths=col_widths)
                tbl.setStyle(TableStyle([
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'), ('ALIGN', (0, 0), (0, -1), 'LEFT'), ('ALIGN', (1, 0), (-1, -1), 'CENTER'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6), ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ('TOPPADDING', (0, 0), (-1, -1), 5), ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#505050')), ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ]))
                return tbl

            def format_multi_value(raw: str, labels_dict: dict) -> str:
                if not raw or not str(raw).strip():
                    return "—"
                parts = [p.strip().lower().replace(" ", "_") for p in str(raw).split(",") if p and p.strip()]
                if not parts:
                    return "—"
                display = [labels_dict.get(p) or p for p in parts]
                return ", ".join(display)
            reading_type_labels = {"en_desarrollo": "En desarrollo", "silabica": "Silabica", "palabra_a_palabra": "Palabra a palabra", "por_unidades_cortas": "Por unidades cortas", "fluida": "Fluida"}
            comprehension_labels = {"en_desarrollo": "En desarrollo", "literal": "Literal", "inferencial": "Inferencial", "critica": "Critica"}
            writing_labels = {"en_desarrollo": "En desarrollo", "copia": "Copia", "dictado": "Dictado", "espontanea": "Espontanea"}

            elements.append(Paragraph("Pauta de evaluación pedagógica - Docente de aula - 4º Básico", title_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("I. DATOS DEL ESTUDIANTE E INFORME", section_style))
            elements.append(Spacer(1, 0.15*inch))
            col_label, col_value = 5*cm, 11*cm
            id_data = [
                [Paragraph("<b>Nombre completo:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("student_full_name", "—")), normal_style)],
                [Paragraph("<b>RUT:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("student_identification_number", "—")), normal_style)],
                [Paragraph("<b>Fecha de nacimiento:</b>", subtitle_style), Paragraph(format_date(doc_data.get("student_born_date")) or "—", normal_style)],
                [Paragraph("<b>Edad:</b>", subtitle_style), Paragraph(get_value("student_age", "—"), normal_style)],
                [Paragraph("<b>Establecimiento:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("establishment_id", "—")), normal_style)],
                [Paragraph("<b>Curso:</b>", subtitle_style), Paragraph(get_value("course", "—"), normal_style)],
                [Paragraph("<b>Fecha del informe:</b>", subtitle_style), Paragraph(format_date(doc_data.get("report_date")) or "—", normal_style)],
                [Paragraph("<b>Repeticiones:</b>", subtitle_style), Paragraph(get_value("repetitions", "—"), normal_style)],
            ]
            report_type_val = (get_value("report_type", "") or "").strip()
            report_type_val = "Anual" if report_type_val.lower() == "anual" else ("Semestral" if report_type_val.lower() == "semestral" else report_type_val or "—")
            id_data.append([Paragraph("<b>Tipo de informe:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(report_type_val), normal_style)])
            tbl_id = Table(id_data, colWidths=[col_label, col_value])
            tbl_id.setStyle(TableStyle([('VALIGN', (0, 0), (-1, -1), 'TOP'), ('ALIGN', (0, 0), (-1, -1), 'LEFT'), ('LEFTPADDING', (0, 0), (-1, -1), 8), ('RIGHTPADDING', (0, 0), (-1, -1), 8), ('TOPPADDING', (0, 0), (-1, -1), 6), ('BOTTOMPADDING', (0, 0), (-1, -1), 6), ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
            ]))
            elements.append(tbl_id)
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("II. FORTALEZAS DE LA SITUACION ESCOLAR", section_style))
            elements.append(Spacer(1, 0.1*inch))
            strengths_raw = doc_data.get("school_situation_strengths")
            strength_labels_es = {"strength_creative": "Creativo/a", "strength_autonomous": "Autonomo/a", "strength_motivated": "Motivado/a", "strength_tolerant_frustration": "Tolerante a la frustracion", "strength_respects_rules": "Respeta las normas", "strength_follows_instructions": "Sigue instrucciones", "strength_peers_relationship": "Buena relacion con pares", "strength_attends_gladly": "Asiste con gusto", "strength_family_communication": "Comunicacion con la familia", "strength_persistent": "Persistente", "strength_self_confident": "Seguro/a de si mismo/a", "strength_participates": "Participa", "strength_pays_attention": "Presta atencion", "strength_interested_learning": "Interesado/a en el aprendizaje", "strength_completes_tasks": "Completa tareas", "strength_adults_relationship": "Buena relacion con adultos", "strength_attends_regularly": "Asiste regularmente", "strength_family_committed": "Familia comprometida"}
            strengths_text = "—"
            if strengths_raw not in (None, ""):
                if isinstance(strengths_raw, dict):
                    strengths_list = [strength_labels_es.get(k, k) for k, v in strengths_raw.items() if v is True]
                    strengths_text = ", ".join(strengths_list) if strengths_list else "Ninguna marcada."
                elif isinstance(strengths_raw, str):
                    try:
                        parsed = __import__("json").loads(strengths_raw)
                        if isinstance(parsed, dict):
                            strengths_list = [strength_labels_es.get(k, k) for k, v in parsed.items() if v is True]
                            strengths_text = ", ".join(strengths_list) if strengths_list else "Ninguna marcada."
                        else:
                            strengths_text = str(strengths_raw).strip() or "—"
                    except Exception:
                        strengths_text = str(strengths_raw).strip() or "—"
                else:
                    strengths_text = str(strengths_raw).strip() or "—"
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(strengths_text), normal_style))
            elements.append(Spacer(1, 0.15*inch))
            elements.append(Paragraph("Observaciones generales", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations", "")) if get_value("observations", "") else "—", normal_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("III. ACTITUD", section_style))
            elements.append(Spacer(1, 0.08*inch))
            elements.append(Paragraph(DOC34_LEGEND_TEXT, legend_style))
            elements.append(Spacer(1, 0.06*inch))
            elements.append(build_table("Actitud", 12, lambda i: get_value(f"attitude_{i}", ""), DOC34_ACTITUD_INDICADORES))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("Observaciones actitud:", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations_attitude", "")) if get_value("observations_attitude", "") else "—", normal_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("IV. LENGUAJE", section_style))
            elements.append(Spacer(1, 0.08*inch))
            elements.append(Paragraph(DOC34_LEGEND_TEXT, legend_style))
            elements.append(Spacer(1, 0.06*inch))
            elements.append(build_table("Lenguaje", 17, lambda i: get_value(f"language_{i}", ""), DOC34_LENGUAJE_INDICADORES))
            elements.append(Spacer(1, 0.12*inch))
            rt_display = format_multi_value(get_value("reading_type", ""), reading_type_labels)
            cl_display = format_multi_value(get_value("comprehension_level", ""), comprehension_labels)
            wl_display = format_multi_value(get_value("writing_level", ""), writing_labels)
            elements.append(Paragraph("<b>Tipo de lectura:</b> " + DocumentsClass._escape_html_for_paragraph(rt_display), normal_style))
            elements.append(Paragraph("<b>Nivel de comprensión:</b> " + DocumentsClass._escape_html_for_paragraph(cl_display), normal_style))
            elements.append(Paragraph("<b>Nivel de escritura:</b> " + DocumentsClass._escape_html_for_paragraph(wl_display), normal_style))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("Observaciones lenguaje:", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations_language", "")) if get_value("observations_language", "") else "—", normal_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("V. MATEMÁTICA", section_style))
            elements.append(Spacer(1, 0.08*inch))
            elements.append(Paragraph(DOC34_LEGEND_TEXT, legend_style))
            elements.append(Spacer(1, 0.06*inch))
            elements.append(build_table("Matemática", 17, lambda i: get_value(f"mathematics_{i}", ""), DOC34_MATEMATICA_INDICADORES))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("Observaciones matemática:", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations_mathematics", "")) if get_value("observations_mathematics", "") else "—", normal_style))
            doc.build(elements)
            return {"status": "success", "message": "PDF generado exitosamente desde cero", "filename": unique_filename, "file_path": str(output_file)}
        except ImportError:
            return {"status": "error", "message": "ReportLab no esta instalado.", "filename": None, "file_path": None}
        except Exception as e:
            return {"status": "error", "message": f"Error generando PDF documento 34: {str(e)}", "filename": None, "file_path": None}

    @staticmethod
    def _generate_pedagogical_evaluation_classroom_fifth_grade_from_scratch(
        document_id: int,
        doc_data: Dict[str, Any],
        db: Optional[Session] = None,
        output_directory: str = "files/system/students"
    ) -> Dict[str, Any]:
        """Genera el PDF del documento 35 (Pauta de evaluación pedagógica - Docente de aula - 5º Básico). 13 actitud, 11 lenguaje, 13 matemática."""
        try:
            if not REPORTLAB_AVAILABLE:
                return {"status": "error", "message": "ReportLab no esta instalado.", "filename": None, "file_path": None}
            def get_value(key: str, default: str = "") -> str:
                v = doc_data.get(key)
                if v is None: return default
                return str(v).strip() if v else default
            def format_date(date_val) -> str:
                if not date_val: return ""
                try:
                    if isinstance(date_val, str) and len(date_val) >= 10:
                        d = datetime.strptime(date_val[:10], "%Y-%m-%d").date()
                        return d.strftime("%d/%m/%Y")
                    if hasattr(date_val, "strftime"): return date_val.strftime("%d/%m/%Y")
                    return str(date_val)[:10]
                except Exception:
                    return str(date_val)[:10] if date_val else ""
            student_name = get_value("student_full_name", "estudiante").replace(" ", "_")
            unique_filename = f"pauta_eval_pedagogica_5basico_{student_name}_{uuid.uuid4().hex[:8]}.pdf"
            output_file = Path(output_directory) / unique_filename
            output_file.parent.mkdir(parents=True, exist_ok=True)
            doc = SimpleDocTemplate(str(output_file), pagesize=A4, rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
            elements = []
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle('Doc35Title', parent=styles['Heading1'], fontSize=14, textColor=colors.HexColor('#000000'), spaceAfter=14, spaceBefore=8, alignment=TA_CENTER, fontName='Helvetica-Bold')
            section_style = ParagraphStyle('Doc35Section', parent=styles['Normal'], fontSize=11, textColor=colors.HexColor('#000000'), spaceAfter=6, spaceBefore=10, fontName='Helvetica-Bold', backColor=colors.HexColor('#E8E8E8'), borderPadding=6)
            subtitle_style = ParagraphStyle('Doc35Subtitle', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#000000'), spaceAfter=4, spaceBefore=6, fontName='Helvetica-Bold')
            normal_style = ParagraphStyle('Doc35Normal', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#000000'), alignment=TA_LEFT, leading=12, spaceAfter=2)
            legend_style = ParagraphStyle('Doc35Legend', parent=styles['Normal'], fontSize=9, textColor=colors.HexColor('#333333'), alignment=TA_LEFT, spaceAfter=6, spaceBefore=2)
            header_cell_style = ParagraphStyle('Doc35HeaderCell', parent=styles['Normal'], fontSize=9, textColor=colors.white, alignment=TA_CENTER, fontName='Helvetica-Bold')
            cell_center_style = ParagraphStyle('Doc35CellCenter', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#000000'), alignment=TA_CENTER)
            DOC35_RESPONSE_COLS = ["S", "G", "O", "P/V", "N", "N/O"]
            DOC35_LEGEND_TEXT = "S: Siempre   G: Generalmente   O: Ocasionalmente   P/V: Pocas veces   N: Nunca   N/O: No observado"
            DOC35_ACTITUD_INDICADORES = [
                "Manifiesta curiosidad e interes por aprender",
                "Adopta una actitud positiva hacia si mismo y el aprendizaje",
                "Expresa y escucha ideas de forma respetuosa",
                "Participa activamente en actividades grupales y colaborativas",
                "Muestra disposicion para compartir ideas, experiencias y opiniones con los demas",
                "Inicia, persiste y finaliza sus tareas de forma autonoma",
                "Realiza sus trabajos de manera ordenada y metodica",
                "Aborda y resuelve problemas de manera flexible y creativa",
                "Se adapta con flexibilidad a cambios, imprevistos o nuevas rutinas de trabajo",
                "Busca apoyo o aclara dudas cuando enfrenta dificultades",
                "Es receptivo a la retroalimentacion y la aplica para mejorar",
                "Demuestra una actitud de esfuerzo y perseverancia",
                "Reflexiona sobre si mismo, sus ideas y sus intereses para comprenderse y valorarse",
            ]
            DOC35_LENGUAJE_INDICADORES = [
                "Lee de manera fluida textos variados apropiados para su edad",
                "Lee e identifica un amplio repertorio de literatura, como poemas, leyendas, fabulas y novelas",
                "Comprende y analiza aspectos relevantes de los textos leidos",
                "Evalua criticamente la informacion presente en textos de diversa procedencia",
                "Busca y selecciona la informacion mas relevante sobre un tema en diversas fuentes para llevar a cabo una investigacion",
                "Planifica sus textos, estableciendo un proposito y destinatario",
                "Escribe narraciones y otros tipos de textos de manera autonoma y creativa, con estructura clara, conectores adecuados e informacion, descripciones o dialogos pertinentes",
                "Revisa y edita sus textos para satisfacer un proposito y transmitir sus ideas con claridad",
                "Comprende, interactua y obtiene informacion relevante de textos orales",
                "Dialoga para compartir, desarrollar ideas y buscar acuerdos",
                "Se expresa de manera clara y efectiva en exposiciones orales para comunicar temas de su interes",
            ]
            DOC35_MATEMATICA_INDICADORES = [
                "Resuelve ejercicios y problemas que involucran las cuatro operaciones y sus combinaciones",
                "Comprende y aplica el concepto de razon",
                "Comprende y aplica el concepto de porcentaje",
                "Comprende y resuelve ejercicios con fracciones y numeros mixtos",
                "Comprende la multiplicacion de decimales por numeros naturales de un digito o multiplos de 10",
                "Comprende la division de decimales por numeros naturales de un digito o multiplos de 10",
                "Resuelve ejercicios y problemas que involucran adicion y sustraccion de fracciones propias, impropias, numeros mixtos o decimales hasta la milesima",
                "Resuelve ecuaciones de primer grado con una incognita",
                "Identifica y calcula los angulos que se forman entre dos rectas que se cortan",
                "Comprende y calcula el area de una superficie de cubos y paralelepipedos a partir de sus redes",
                "Calcula la superficie de cubos y paralelepipedos, expresando el resultado en cm2 y m2",
                "Calcula el volumen de cubos y paralelepipedos, expresando el resultado en cm3, m3 y mm3",
                "Lee e interpreta graficos de barra doble y circulares, y comunica sus conclusiones de manera clara",
            ]
            def normalize_response(val):
                if val is None or (isinstance(val, str) and not val.strip()): return None
                v = str(val).strip().upper().replace(" ", "")
                if v in ("S", "SIEMPRE"): return "S"
                if v in ("G", "GENERALMENTE"): return "G"
                if v in ("O", "OCASIONALMENTE"): return "O"
                if v in ("P/V", "PV", "POCASVECES", "P"): return "P/V"
                if v in ("N/O", "NO", "NOOBSERVADO"): return "N/O"
                if v in ("N", "NUNCA"): return "N"
                if len(v) == 1 and v in "SGOPN": return "P/V" if v == "P" else v
                return None
            def build_table(header_label, num_items, get_val_fn, indicator_labels):
                rows = []
                rows.append([Paragraph(DocumentsClass._escape_html_for_paragraph("Indicadores"), header_cell_style)] + [Paragraph(c, header_cell_style) for c in DOC35_RESPONSE_COLS])
                for i in range(1, num_items + 1):
                    val = get_val_fn(i)
                    col_key = normalize_response(val)
                    label_text = indicator_labels[i - 1] if indicator_labels and len(indicator_labels) >= i else f"{header_label} {i}"
                    cells = [Paragraph(DocumentsClass._escape_html_for_paragraph(label_text), normal_style)]
                    for c in DOC35_RESPONSE_COLS:
                        cells.append(Paragraph("✓" if col_key == c else "", cell_center_style))
                    rows.append(cells)
                col_ind = 8*cm
                col_opt = (16*cm - col_ind) / 6
                col_widths = [col_ind] + [col_opt] * 6
                tbl = Table(rows, colWidths=col_widths)
                tbl.setStyle(TableStyle([
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'), ('ALIGN', (0, 0), (0, -1), 'LEFT'), ('ALIGN', (1, 0), (-1, -1), 'CENTER'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6), ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ('TOPPADDING', (0, 0), (-1, -1), 5), ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#505050')), ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ]))
                return tbl
            def format_multi_value(raw: str, labels_dict: dict) -> str:
                if not raw or not str(raw).strip(): return "—"
                parts = [p.strip().lower().replace(" ", "_") for p in str(raw).split(",") if p and p.strip()]
                if not parts: return "—"
                display = [labels_dict.get(p) or p for p in parts]
                return ", ".join(display)
            reading_type_labels = {"en_desarrollo": "En desarrollo", "silabica": "Silabica", "palabra_a_palabra": "Palabra a palabra", "por_unidades_cortas": "Por unidades cortas", "fluida": "Fluida"}
            comprehension_labels = {"en_desarrollo": "En desarrollo", "literal": "Literal", "inferencial": "Inferencial", "critica": "Critica"}
            writing_labels = {"en_desarrollo": "En desarrollo", "copia": "Copia", "dictado": "Dictado", "espontanea": "Espontanea"}
            elements.append(Paragraph("Pauta de evaluación pedagógica - Docente de aula - 5º Básico", title_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("I. DATOS DEL ESTUDIANTE E INFORME", section_style))
            elements.append(Spacer(1, 0.15*inch))
            col_label, col_value = 5*cm, 11*cm
            id_data = [
                [Paragraph("<b>Nombre completo:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("student_full_name", "—")), normal_style)],
                [Paragraph("<b>RUT:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("student_identification_number", "—")), normal_style)],
                [Paragraph("<b>Fecha de nacimiento:</b>", subtitle_style), Paragraph(format_date(doc_data.get("student_born_date")) or "—", normal_style)],
                [Paragraph("<b>Edad:</b>", subtitle_style), Paragraph(get_value("student_age", "—"), normal_style)],
                [Paragraph("<b>Establecimiento:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("establishment_id", "—")), normal_style)],
                [Paragraph("<b>Curso:</b>", subtitle_style), Paragraph(get_value("course", "—"), normal_style)],
                [Paragraph("<b>Fecha del informe:</b>", subtitle_style), Paragraph(format_date(doc_data.get("report_date")) or "—", normal_style)],
                [Paragraph("<b>Repeticiones:</b>", subtitle_style), Paragraph(get_value("repetitions", "—"), normal_style)],
            ]
            report_type_val = (get_value("report_type", "") or "").strip()
            report_type_val = "Anual" if report_type_val.lower() == "anual" else ("Semestral" if report_type_val.lower() == "semestral" else report_type_val or "—")
            id_data.append([Paragraph("<b>Tipo de informe:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(report_type_val), normal_style)])
            tbl_id = Table(id_data, colWidths=[col_label, col_value])
            tbl_id.setStyle(TableStyle([('VALIGN', (0, 0), (-1, -1), 'TOP'), ('ALIGN', (0, 0), (-1, -1), 'LEFT'), ('LEFTPADDING', (0, 0), (-1, -1), 8), ('RIGHTPADDING', (0, 0), (-1, -1), 8), ('TOPPADDING', (0, 0), (-1, -1), 6), ('BOTTOMPADDING', (0, 0), (-1, -1), 6), ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC'))]))
            elements.append(tbl_id)
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("II. FORTALEZAS DE LA SITUACION ESCOLAR", section_style))
            elements.append(Spacer(1, 0.1*inch))
            strengths_raw = doc_data.get("school_situation_strengths")
            strength_labels_es = {"strength_creative": "Creativo/a", "strength_autonomous": "Autonomo/a", "strength_motivated": "Motivado/a", "strength_tolerant_frustration": "Tolerante a la frustracion", "strength_respects_rules": "Respeta las normas", "strength_follows_instructions": "Sigue instrucciones", "strength_peers_relationship": "Buena relacion con pares", "strength_attends_gladly": "Asiste con gusto", "strength_family_communication": "Comunicacion con la familia", "strength_persistent": "Persistente", "strength_self_confident": "Seguro/a de si mismo/a", "strength_participates": "Participa", "strength_pays_attention": "Presta atencion", "strength_interested_learning": "Interesado/a en el aprendizaje", "strength_completes_tasks": "Completa tareas", "strength_adults_relationship": "Buena relacion con adultos", "strength_attends_regularly": "Asiste regularmente", "strength_family_committed": "Familia comprometida"}
            strengths_text = "—"
            if strengths_raw not in (None, ""):
                if isinstance(strengths_raw, dict):
                    strengths_list = [strength_labels_es.get(k, k) for k, v in strengths_raw.items() if v is True]
                    strengths_text = ", ".join(strengths_list) if strengths_list else "Ninguna marcada."
                elif isinstance(strengths_raw, str):
                    try:
                        parsed = __import__("json").loads(strengths_raw)
                        strengths_list = [strength_labels_es.get(k, k) for k, v in (parsed if isinstance(parsed, dict) else {}).items() if v is True]
                        strengths_text = ", ".join(strengths_list) if strengths_list else "Ninguna marcada."
                    except Exception:
                        strengths_text = str(strengths_raw).strip() or "—"
                else:
                    strengths_text = str(strengths_raw).strip() or "—"
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(strengths_text), normal_style))
            elements.append(Spacer(1, 0.15*inch))
            elements.append(Paragraph("Observaciones generales", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations", "")) if get_value("observations", "") else "—", normal_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("III. ACTITUD", section_style))
            elements.append(Spacer(1, 0.08*inch))
            elements.append(Paragraph(DOC35_LEGEND_TEXT, legend_style))
            elements.append(Spacer(1, 0.06*inch))
            elements.append(build_table("Actitud", 13, lambda i: get_value(f"attitude_{i}", ""), DOC35_ACTITUD_INDICADORES))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("Observaciones actitud:", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations_attitude", "")) if get_value("observations_attitude", "") else "—", normal_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("IV. LENGUAJE", section_style))
            elements.append(Spacer(1, 0.08*inch))
            elements.append(Paragraph(DOC35_LEGEND_TEXT, legend_style))
            elements.append(Spacer(1, 0.06*inch))
            elements.append(build_table("Lenguaje", 11, lambda i: get_value(f"language_{i}", ""), DOC35_LENGUAJE_INDICADORES))
            elements.append(Spacer(1, 0.12*inch))
            rt_display = format_multi_value(get_value("reading_type", ""), reading_type_labels)
            cl_display = format_multi_value(get_value("comprehension_level", ""), comprehension_labels)
            wl_display = format_multi_value(get_value("writing_level", ""), writing_labels)
            elements.append(Paragraph("<b>Tipo de lectura:</b> " + DocumentsClass._escape_html_for_paragraph(rt_display), normal_style))
            elements.append(Paragraph("<b>Nivel de comprensión:</b> " + DocumentsClass._escape_html_for_paragraph(cl_display), normal_style))
            elements.append(Paragraph("<b>Nivel de escritura:</b> " + DocumentsClass._escape_html_for_paragraph(wl_display), normal_style))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("Observaciones lenguaje:", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations_language", "")) if get_value("observations_language", "") else "—", normal_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("V. MATEMÁTICA", section_style))
            elements.append(Spacer(1, 0.08*inch))
            elements.append(Paragraph(DOC35_LEGEND_TEXT, legend_style))
            elements.append(Spacer(1, 0.06*inch))
            elements.append(build_table("Matemática", 13, lambda i: get_value(f"mathematics_{i}", ""), DOC35_MATEMATICA_INDICADORES))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("Observaciones matemática:", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations_mathematics", "")) if get_value("observations_mathematics", "") else "—", normal_style))
            doc.build(elements)
            return {"status": "success", "message": "PDF generado exitosamente desde cero", "filename": unique_filename, "file_path": str(output_file)}
        except ImportError:
            return {"status": "error", "message": "ReportLab no esta instalado.", "filename": None, "file_path": None}
        except Exception as e:
            return {"status": "error", "message": f"Error generando PDF documento 35: {str(e)}", "filename": None, "file_path": None}

    @staticmethod
    def _generate_pedagogical_evaluation_classroom_sixth_grade_from_scratch(
        document_id: int,
        doc_data: Dict[str, Any],
        db: Optional[Session] = None,
        output_directory: str = "files/system/students"
    ) -> Dict[str, Any]:
        """Genera el PDF del documento 36 (Pauta de evaluación pedagógica - Docente de aula - 6º Básico). 13 actitud, 11 lenguaje, 13 matemática."""
        try:
            if not REPORTLAB_AVAILABLE:
                return {"status": "error", "message": "ReportLab no esta instalado.", "filename": None, "file_path": None}
            def get_value(key: str, default: str = "") -> str:
                v = doc_data.get(key)
                if v is None: return default
                return str(v).strip() if v else default
            def format_date(date_val) -> str:
                if not date_val: return ""
                try:
                    if isinstance(date_val, str) and len(date_val) >= 10:
                        d = datetime.strptime(date_val[:10], "%Y-%m-%d").date()
                        return d.strftime("%d/%m/%Y")
                    if hasattr(date_val, "strftime"): return date_val.strftime("%d/%m/%Y")
                    return str(date_val)[:10]
                except Exception:
                    return str(date_val)[:10] if date_val else ""
            student_name = get_value("student_full_name", "estudiante").replace(" ", "_")
            unique_filename = f"pauta_eval_pedagogica_6basico_{student_name}_{uuid.uuid4().hex[:8]}.pdf"
            output_file = Path(output_directory) / unique_filename
            output_file.parent.mkdir(parents=True, exist_ok=True)
            doc = SimpleDocTemplate(str(output_file), pagesize=A4, rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
            elements = []
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle('Doc36Title', parent=styles['Heading1'], fontSize=14, textColor=colors.HexColor('#000000'), spaceAfter=14, spaceBefore=8, alignment=TA_CENTER, fontName='Helvetica-Bold')
            section_style = ParagraphStyle('Doc36Section', parent=styles['Normal'], fontSize=11, textColor=colors.HexColor('#000000'), spaceAfter=6, spaceBefore=10, fontName='Helvetica-Bold', backColor=colors.HexColor('#E8E8E8'), borderPadding=6)
            subtitle_style = ParagraphStyle('Doc36Subtitle', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#000000'), spaceAfter=4, spaceBefore=6, fontName='Helvetica-Bold')
            normal_style = ParagraphStyle('Doc36Normal', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#000000'), alignment=TA_LEFT, leading=12, spaceAfter=2)
            legend_style = ParagraphStyle('Doc36Legend', parent=styles['Normal'], fontSize=9, textColor=colors.HexColor('#333333'), alignment=TA_LEFT, spaceAfter=6, spaceBefore=2)
            header_cell_style = ParagraphStyle('Doc36HeaderCell', parent=styles['Normal'], fontSize=9, textColor=colors.white, alignment=TA_CENTER, fontName='Helvetica-Bold')
            cell_center_style = ParagraphStyle('Doc36CellCenter', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#000000'), alignment=TA_CENTER)
            DOC36_RESPONSE_COLS = ["S", "G", "O", "P/V", "N", "N/O"]
            DOC36_LEGEND_TEXT = "S: Siempre   G: Generalmente   O: Ocasionalmente   P/V: Pocas veces   N: Nunca   N/O: No observado"
            DOC36_ACTITUD_INDICADORES = [
                "Manifiesta curiosidad e interes por aprender",
                "Adopta una actitud positiva hacia si mismo y el aprendizaje",
                "Expresa y escucha ideas de forma respetuosa",
                "Participa activamente en actividades grupales y colaborativas",
                "Muestra disposicion para compartir ideas, experiencias y opiniones con los demas",
                "Inicia, persiste y finaliza sus tareas de forma autonoma",
                "Realiza sus trabajos de manera ordenada y metodica",
                "Aborda y resuelve problemas de manera flexible y creativa",
                "Se adapta con flexibilidad a cambios, imprevistos o nuevas rutinas de trabajo",
                "Busca apoyo o aclara dudas cuando enfrenta dificultades",
                "Es receptivo a la retroalimentacion y la aplica para mejorar",
                "Demuestra una actitud de esfuerzo y perseverancia",
                "Reflexiona sobre si mismo, sus ideas y sus intereses para comprenderse y valorarse",
            ]
            DOC36_LENGUAJE_INDICADORES = [
                "Lee de manera fluida textos variados apropiados para su edad",
                "Lee e identifica un amplio repertorio de literatura, como poemas, leyendas, fabulas y novelas",
                "Comprende y analiza aspectos relevantes de los textos leidos",
                "Evalua criticamente la informacion presente en textos de diversa procedencia",
                "Busca y selecciona la informacion mas relevante sobre un tema en diversas fuentes para llevar a cabo una investigacion",
                "Planifica sus textos, estableciendo un proposito y destinatario",
                "Escribe narraciones y otros tipos de textos de manera autonoma y creativa, con estructura clara, conectores adecuados e informacion, descripciones o dialogos pertinentes",
                "Revisa y edita sus textos para satisfacer un proposito y transmitir sus ideas con claridad",
                "Comprende, interactua y obtiene informacion relevante de textos orales",
                "Dialoga para compartir, desarrollar ideas y buscar acuerdos",
                "Se expresa de manera clara y efectiva en exposiciones orales para comunicar temas de su interes",
            ]
            DOC36_MATEMATICA_INDICADORES = [
                "Resuelve ejercicios y problemas que involucran las cuatro operaciones y sus combinaciones",
                "Comprende y aplica el concepto de razon",
                "Comprende y aplica el concepto de porcentaje",
                "Comprende y resuelve ejercicios con fracciones y numeros mixtos",
                "Comprende la multiplicacion de decimales por numeros naturales de un digito o multiplos de 10",
                "Comprende la division de decimales por numeros naturales de un digito o multiplos de 10",
                "Resuelve ejercicios y problemas que involucran adicion y sustraccion de fracciones propias, impropias, numeros mixtos o decimales hasta la milesima",
                "Resuelve ecuaciones de primer grado con una incognita",
                "Identifica y calcula los angulos que se forman entre dos rectas que se cortan",
                "Comprende y calcula el area de una superficie de cubos y paralelepipedos a partir de sus redes",
                "Calcula la superficie de cubos y paralelepipedos, expresando el resultado en cm2 y m2",
                "Calcula el volumen de cubos y paralelepipedos, expresando el resultado en cm3, m3 y mm3",
                "Lee e interpreta graficos de barra doble y circulares, y comunica sus conclusiones de manera clara",
            ]
            def normalize_response(val):
                if val is None or (isinstance(val, str) and not val.strip()): return None
                v = str(val).strip().upper().replace(" ", "")
                if v in ("S", "SIEMPRE"): return "S"
                if v in ("G", "GENERALMENTE"): return "G"
                if v in ("O", "OCASIONALMENTE"): return "O"
                if v in ("P/V", "PV", "POCASVECES", "P"): return "P/V"
                if v in ("N/O", "NO", "NOOBSERVADO"): return "N/O"
                if v in ("N", "NUNCA"): return "N"
                if len(v) == 1 and v in "SGOPN": return "P/V" if v == "P" else v
                return None
            def build_table(header_label, num_items, get_val_fn, indicator_labels):
                rows = []
                rows.append([Paragraph(DocumentsClass._escape_html_for_paragraph("Indicadores"), header_cell_style)] + [Paragraph(c, header_cell_style) for c in DOC36_RESPONSE_COLS])
                for i in range(1, num_items + 1):
                    val = get_val_fn(i)
                    col_key = normalize_response(val)
                    label_text = indicator_labels[i - 1] if indicator_labels and len(indicator_labels) >= i else f"{header_label} {i}"
                    cells = [Paragraph(DocumentsClass._escape_html_for_paragraph(label_text), normal_style)]
                    for c in DOC36_RESPONSE_COLS:
                        cells.append(Paragraph("✓" if col_key == c else "", cell_center_style))
                    rows.append(cells)
                col_ind = 8*cm
                col_opt = (16*cm - col_ind) / 6
                col_widths = [col_ind] + [col_opt] * 6
                tbl = Table(rows, colWidths=col_widths)
                tbl.setStyle(TableStyle([
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'), ('ALIGN', (0, 0), (0, -1), 'LEFT'), ('ALIGN', (1, 0), (-1, -1), 'CENTER'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6), ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ('TOPPADDING', (0, 0), (-1, -1), 5), ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#505050')), ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ]))
                return tbl
            def format_multi_value(raw: str, labels_dict: dict) -> str:
                if not raw or not str(raw).strip(): return "—"
                parts = [p.strip().lower().replace(" ", "_") for p in str(raw).split(",") if p and p.strip()]
                if not parts: return "—"
                display = [labels_dict.get(p) or p for p in parts]
                return ", ".join(display)
            reading_type_labels = {"en_desarrollo": "En desarrollo", "silabica": "Silabica", "palabra_a_palabra": "Palabra a palabra", "por_unidades_cortas": "Por unidades cortas", "fluida": "Fluida"}
            comprehension_labels = {"en_desarrollo": "En desarrollo", "literal": "Literal", "inferencial": "Inferencial", "critica": "Critica"}
            writing_labels = {"en_desarrollo": "En desarrollo", "copia": "Copia", "dictado": "Dictado", "espontanea": "Espontanea"}
            elements.append(Paragraph("Pauta de evaluación pedagógica - Docente de aula - 6º Básico", title_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("I. DATOS DEL ESTUDIANTE E INFORME", section_style))
            elements.append(Spacer(1, 0.15*inch))
            col_label, col_value = 5*cm, 11*cm
            id_data = [
                [Paragraph("<b>Nombre completo:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("student_full_name", "—")), normal_style)],
                [Paragraph("<b>RUT:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("student_identification_number", "—")), normal_style)],
                [Paragraph("<b>Fecha de nacimiento:</b>", subtitle_style), Paragraph(format_date(doc_data.get("student_born_date")) or "—", normal_style)],
                [Paragraph("<b>Edad:</b>", subtitle_style), Paragraph(get_value("student_age", "—"), normal_style)],
                [Paragraph("<b>Establecimiento:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("establishment_id", "—")), normal_style)],
                [Paragraph("<b>Curso:</b>", subtitle_style), Paragraph(get_value("course", "—"), normal_style)],
                [Paragraph("<b>Fecha del informe:</b>", subtitle_style), Paragraph(format_date(doc_data.get("report_date")) or "—", normal_style)],
                [Paragraph("<b>Repeticiones:</b>", subtitle_style), Paragraph(get_value("repetitions", "—"), normal_style)],
            ]
            report_type_val = (get_value("report_type", "") or "").strip()
            report_type_val = "Anual" if report_type_val.lower() == "anual" else ("Semestral" if report_type_val.lower() == "semestral" else report_type_val or "—")
            id_data.append([Paragraph("<b>Tipo de informe:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(report_type_val), normal_style)])
            tbl_id = Table(id_data, colWidths=[col_label, col_value])
            tbl_id.setStyle(TableStyle([('VALIGN', (0, 0), (-1, -1), 'TOP'), ('ALIGN', (0, 0), (-1, -1), 'LEFT'), ('LEFTPADDING', (0, 0), (-1, -1), 8), ('RIGHTPADDING', (0, 0), (-1, -1), 8), ('TOPPADDING', (0, 0), (-1, -1), 6), ('BOTTOMPADDING', (0, 0), (-1, -1), 6), ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC'))]))
            elements.append(tbl_id)
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("II. FORTALEZAS DE LA SITUACION ESCOLAR", section_style))
            elements.append(Spacer(1, 0.1*inch))
            strengths_raw = doc_data.get("school_situation_strengths")
            strength_labels_es = {"strength_creative": "Creativo/a", "strength_autonomous": "Autonomo/a", "strength_motivated": "Motivado/a", "strength_tolerant_frustration": "Tolerante a la frustracion", "strength_respects_rules": "Respeta las normas", "strength_follows_instructions": "Sigue instrucciones", "strength_peers_relationship": "Buena relacion con pares", "strength_attends_gladly": "Asiste con gusto", "strength_family_communication": "Comunicacion con la familia", "strength_persistent": "Persistente", "strength_self_confident": "Seguro/a de si mismo/a", "strength_participates": "Participa", "strength_pays_attention": "Presta atencion", "strength_interested_learning": "Interesado/a en el aprendizaje", "strength_completes_tasks": "Completa tareas", "strength_adults_relationship": "Buena relacion con adultos", "strength_attends_regularly": "Asiste regularmente", "strength_family_committed": "Familia comprometida"}
            strengths_text = "—"
            if strengths_raw not in (None, ""):
                if isinstance(strengths_raw, dict):
                    strengths_list = [strength_labels_es.get(k, k) for k, v in strengths_raw.items() if v is True]
                    strengths_text = ", ".join(strengths_list) if strengths_list else "Ninguna marcada."
                elif isinstance(strengths_raw, str):
                    try:
                        parsed = __import__("json").loads(strengths_raw)
                        strengths_list = [strength_labels_es.get(k, k) for k, v in (parsed if isinstance(parsed, dict) else {}).items() if v is True]
                        strengths_text = ", ".join(strengths_list) if strengths_list else "Ninguna marcada."
                    except Exception:
                        strengths_text = str(strengths_raw).strip() or "—"
                else:
                    strengths_text = str(strengths_raw).strip() or "—"
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(strengths_text), normal_style))
            elements.append(Spacer(1, 0.15*inch))
            elements.append(Paragraph("Observaciones generales", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations", "")) if get_value("observations", "") else "—", normal_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("III. ACTITUD", section_style))
            elements.append(Spacer(1, 0.08*inch))
            elements.append(Paragraph(DOC36_LEGEND_TEXT, legend_style))
            elements.append(Spacer(1, 0.06*inch))
            elements.append(build_table("Actitud", 13, lambda i: get_value(f"attitude_{i}", ""), DOC36_ACTITUD_INDICADORES))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("Observaciones actitud:", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations_attitude", "")) if get_value("observations_attitude", "") else "—", normal_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("IV. LENGUAJE", section_style))
            elements.append(Spacer(1, 0.08*inch))
            elements.append(Paragraph(DOC36_LEGEND_TEXT, legend_style))
            elements.append(Spacer(1, 0.06*inch))
            elements.append(build_table("Lenguaje", 11, lambda i: get_value(f"language_{i}", ""), DOC36_LENGUAJE_INDICADORES))
            elements.append(Spacer(1, 0.12*inch))
            rt_display = format_multi_value(get_value("reading_type", ""), reading_type_labels)
            cl_display = format_multi_value(get_value("comprehension_level", ""), comprehension_labels)
            wl_display = format_multi_value(get_value("writing_level", ""), writing_labels)
            elements.append(Paragraph("<b>Tipo de lectura:</b> " + DocumentsClass._escape_html_for_paragraph(rt_display), normal_style))
            elements.append(Paragraph("<b>Nivel de comprensión:</b> " + DocumentsClass._escape_html_for_paragraph(cl_display), normal_style))
            elements.append(Paragraph("<b>Nivel de escritura:</b> " + DocumentsClass._escape_html_for_paragraph(wl_display), normal_style))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("Observaciones lenguaje:", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations_language", "")) if get_value("observations_language", "") else "—", normal_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("V. MATEMÁTICA", section_style))
            elements.append(Spacer(1, 0.08*inch))
            elements.append(Paragraph(DOC36_LEGEND_TEXT, legend_style))
            elements.append(Spacer(1, 0.06*inch))
            elements.append(build_table("Matemática", 13, lambda i: get_value(f"mathematics_{i}", ""), DOC36_MATEMATICA_INDICADORES))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("Observaciones matemática:", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations_mathematics", "")) if get_value("observations_mathematics", "") else "—", normal_style))
            doc.build(elements)
            return {"status": "success", "message": "PDF generado exitosamente desde cero", "filename": unique_filename, "file_path": str(output_file)}
        except ImportError:
            return {"status": "error", "message": "ReportLab no esta instalado.", "filename": None, "file_path": None}
        except Exception as e:
            return {"status": "error", "message": f"Error generando PDF documento 36: {str(e)}", "filename": None, "file_path": None}

    @staticmethod
    def _generate_pedagogical_evaluation_classroom_seventh_grade_from_scratch(
        document_id: int,
        doc_data: Dict[str, Any],
        db: Optional[Session] = None,
        output_directory: str = "files/system/students"
    ) -> Dict[str, Any]:
        """Genera el PDF del documento 37 (Pauta de evaluación pedagógica - Docente de aula - 7º Básico). 13 actitud, 11 lenguaje, 13 matemática."""
        try:
            if not REPORTLAB_AVAILABLE:
                return {"status": "error", "message": "ReportLab no esta instalado.", "filename": None, "file_path": None}
            def get_value(key: str, default: str = "") -> str:
                v = doc_data.get(key)
                if v is None: return default
                return str(v).strip() if v else default
            def format_date(date_val) -> str:
                if not date_val: return ""
                try:
                    if isinstance(date_val, str) and len(date_val) >= 10:
                        d = datetime.strptime(date_val[:10], "%Y-%m-%d").date()
                        return d.strftime("%d/%m/%Y")
                    if hasattr(date_val, "strftime"): return date_val.strftime("%d/%m/%Y")
                    return str(date_val)[:10]
                except Exception:
                    return str(date_val)[:10] if date_val else ""
            student_name = get_value("student_full_name", "estudiante").replace(" ", "_")
            unique_filename = f"pauta_eval_pedagogica_7basico_{student_name}_{uuid.uuid4().hex[:8]}.pdf"
            output_file = Path(output_directory) / unique_filename
            output_file.parent.mkdir(parents=True, exist_ok=True)
            doc = SimpleDocTemplate(str(output_file), pagesize=A4, rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
            elements = []
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle('Doc37Title', parent=styles['Heading1'], fontSize=14, textColor=colors.HexColor('#000000'), spaceAfter=14, spaceBefore=8, alignment=TA_CENTER, fontName='Helvetica-Bold')
            section_style = ParagraphStyle('Doc37Section', parent=styles['Normal'], fontSize=11, textColor=colors.HexColor('#000000'), spaceAfter=6, spaceBefore=10, fontName='Helvetica-Bold', backColor=colors.HexColor('#E8E8E8'), borderPadding=6)
            subtitle_style = ParagraphStyle('Doc37Subtitle', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#000000'), spaceAfter=4, spaceBefore=6, fontName='Helvetica-Bold')
            normal_style = ParagraphStyle('Doc37Normal', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#000000'), alignment=TA_LEFT, leading=12, spaceAfter=2)
            legend_style = ParagraphStyle('Doc37Legend', parent=styles['Normal'], fontSize=9, textColor=colors.HexColor('#333333'), alignment=TA_LEFT, spaceAfter=6, spaceBefore=2)
            header_cell_style = ParagraphStyle('Doc37HeaderCell', parent=styles['Normal'], fontSize=9, textColor=colors.white, alignment=TA_CENTER, fontName='Helvetica-Bold')
            cell_center_style = ParagraphStyle('Doc37CellCenter', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#000000'), alignment=TA_CENTER)
            DOC37_RESPONSE_COLS = ["S", "G", "O", "P/V", "N", "N/O"]
            DOC37_LEGEND_TEXT = "S: Siempre   G: Generalmente   O: Ocasionalmente   P/V: Pocas veces   N: Nunca   N/O: No observado"
            DOC37_ACTITUD_INDICADORES = [
                "Manifiesta curiosidad e interes por aprender",
                "Adopta una actitud positiva hacia si mismo y el aprendizaje",
                "Expresa y escucha ideas de forma respetuosa",
                "Participa activamente en actividades grupales y colaborativas",
                "Muestra disposicion para compartir ideas, experiencias y opiniones con los demas",
                "Inicia, persiste y finaliza sus tareas de forma autonoma",
                "Realiza sus trabajos de manera ordenada y metodica",
                "Aborda y resuelve problemas de manera flexible y creativa",
                "Se adapta con flexibilidad a cambios, imprevistos o nuevas rutinas de trabajo",
                "Busca apoyo o aclara dudas cuando enfrenta dificultades",
                "Es receptivo a la retroalimentacion y la aplica para mejorar",
                "Demuestra una actitud de esfuerzo y perseverancia",
                "Reflexiona sobre si mismo, sus ideas y sus intereses para comprenderse y valorarse",
            ]
            DOC37_LENGUAJE_INDICADORES = [
                "Lee de manera fluida textos variados apropiados para su edad",
                "Lee e identifica un amplio repertorio de literatura, como poemas, leyendas, fabulas y novelas",
                "Comprende y analiza aspectos relevantes de los textos leidos",
                "Evalua criticamente la informacion presente en textos de diversa procedencia",
                "Busca y selecciona la informacion mas relevante sobre un tema en diversas fuentes para llevar a cabo una investigacion",
                "Planifica sus textos, estableciendo un proposito y destinatario",
                "Escribe narraciones y otros tipos de textos de manera autonoma y creativa, con estructura clara, conectores adecuados e informacion, descripciones o dialogos pertinentes",
                "Revisa y edita sus textos para satisfacer un proposito y transmitir sus ideas con claridad",
                "Comprende, interactua y obtiene informacion relevante de textos orales",
                "Dialoga para compartir, desarrollar ideas y buscar acuerdos",
                "Se expresa de manera clara y efectiva en exposiciones orales para comunicar temas de su interes",
            ]
            DOC37_MATEMATICA_INDICADORES = [
                "Resuelve ejercicios y problemas que involucran las cuatro operaciones y sus combinaciones",
                "Comprende y aplica el concepto de razon",
                "Comprende y aplica el concepto de porcentaje",
                "Comprende y resuelve ejercicios con fracciones y numeros mixtos",
                "Comprende la multiplicacion de decimales por numeros naturales de un digito o multiplos de 10",
                "Comprende la division de decimales por numeros naturales de un digito o multiplos de 10",
                "Resuelve ejercicios y problemas que involucran adicion y sustraccion de fracciones propias, impropias, numeros mixtos o decimales hasta la milesima",
                "Resuelve ecuaciones de primer grado con una incognita",
                "Identifica y calcula los angulos que se forman entre dos rectas que se cortan",
                "Comprende y calcula el area de una superficie de cubos y paralelepipedos a partir de sus redes",
                "Calcula la superficie de cubos y paralelepipedos, expresando el resultado en cm2 y m2",
                "Calcula el volumen de cubos y paralelepipedos, expresando el resultado en cm3, m3 y mm3",
                "Lee e interpreta graficos de barra doble y circulares, y comunica sus conclusiones de manera clara",
            ]
            def normalize_response(val):
                if val is None or (isinstance(val, str) and not val.strip()): return None
                v = str(val).strip().upper().replace(" ", "")
                if v in ("S", "SIEMPRE"): return "S"
                if v in ("G", "GENERALMENTE"): return "G"
                if v in ("O", "OCASIONALMENTE"): return "O"
                if v in ("P/V", "PV", "POCASVECES", "P"): return "P/V"
                if v in ("N/O", "NO", "NOOBSERVADO"): return "N/O"
                if v in ("N", "NUNCA"): return "N"
                if len(v) == 1 and v in "SGOPN": return "P/V" if v == "P" else v
                return None
            def build_table(header_label, num_items, get_val_fn, indicator_labels):
                rows = []
                rows.append([Paragraph(DocumentsClass._escape_html_for_paragraph("Indicadores"), header_cell_style)] + [Paragraph(c, header_cell_style) for c in DOC37_RESPONSE_COLS])
                for i in range(1, num_items + 1):
                    val = get_val_fn(i)
                    col_key = normalize_response(val)
                    label_text = indicator_labels[i - 1] if indicator_labels and len(indicator_labels) >= i else f"{header_label} {i}"
                    cells = [Paragraph(DocumentsClass._escape_html_for_paragraph(label_text), normal_style)]
                    for c in DOC37_RESPONSE_COLS:
                        cells.append(Paragraph("✓" if col_key == c else "", cell_center_style))
                    rows.append(cells)
                col_ind = 8*cm
                col_opt = (16*cm - col_ind) / 6
                col_widths = [col_ind] + [col_opt] * 6
                tbl = Table(rows, colWidths=col_widths)
                tbl.setStyle(TableStyle([
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'), ('ALIGN', (0, 0), (0, -1), 'LEFT'), ('ALIGN', (1, 0), (-1, -1), 'CENTER'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6), ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ('TOPPADDING', (0, 0), (-1, -1), 5), ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#505050')), ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ]))
                return tbl
            def format_multi_value(raw: str, labels_dict: dict) -> str:
                if not raw or not str(raw).strip(): return "—"
                parts = [p.strip().lower().replace(" ", "_") for p in str(raw).split(",") if p and p.strip()]
                if not parts: return "—"
                display = [labels_dict.get(p) or p for p in parts]
                return ", ".join(display)
            reading_type_labels = {"en_desarrollo": "En desarrollo", "silabica": "Silabica", "palabra_a_palabra": "Palabra a palabra", "por_unidades_cortas": "Por unidades cortas", "fluida": "Fluida"}
            comprehension_labels = {"en_desarrollo": "En desarrollo", "literal": "Literal", "inferencial": "Inferencial", "critica": "Critica"}
            writing_labels = {"en_desarrollo": "En desarrollo", "copia": "Copia", "dictado": "Dictado", "espontanea": "Espontanea"}
            elements.append(Paragraph("Pauta de evaluación pedagógica - Docente de aula - 7º Básico", title_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("I. DATOS DEL ESTUDIANTE E INFORME", section_style))
            elements.append(Spacer(1, 0.15*inch))
            col_label, col_value = 5*cm, 11*cm
            id_data = [
                [Paragraph("<b>Nombre completo:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("student_full_name", "—")), normal_style)],
                [Paragraph("<b>RUT:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("student_identification_number", "—")), normal_style)],
                [Paragraph("<b>Fecha de nacimiento:</b>", subtitle_style), Paragraph(format_date(doc_data.get("student_born_date")) or "—", normal_style)],
                [Paragraph("<b>Edad:</b>", subtitle_style), Paragraph(get_value("student_age", "—"), normal_style)],
                [Paragraph("<b>Establecimiento:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("establishment_id", "—")), normal_style)],
                [Paragraph("<b>Curso:</b>", subtitle_style), Paragraph(get_value("course", "—"), normal_style)],
                [Paragraph("<b>Fecha del informe:</b>", subtitle_style), Paragraph(format_date(doc_data.get("report_date")) or "—", normal_style)],
                [Paragraph("<b>Repeticiones:</b>", subtitle_style), Paragraph(get_value("repetitions", "—"), normal_style)],
            ]
            report_type_val = (get_value("report_type", "") or "").strip()
            report_type_val = "Anual" if report_type_val.lower() == "anual" else ("Semestral" if report_type_val.lower() == "semestral" else report_type_val or "—")
            id_data.append([Paragraph("<b>Tipo de informe:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(report_type_val), normal_style)])
            tbl_id = Table(id_data, colWidths=[col_label, col_value])
            tbl_id.setStyle(TableStyle([('VALIGN', (0, 0), (-1, -1), 'TOP'), ('ALIGN', (0, 0), (-1, -1), 'LEFT'), ('LEFTPADDING', (0, 0), (-1, -1), 8), ('RIGHTPADDING', (0, 0), (-1, -1), 8), ('TOPPADDING', (0, 0), (-1, -1), 6), ('BOTTOMPADDING', (0, 0), (-1, -1), 6), ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC'))]))
            elements.append(tbl_id)
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("II. FORTALEZAS DE LA SITUACION ESCOLAR", section_style))
            elements.append(Spacer(1, 0.1*inch))
            strengths_raw = doc_data.get("school_situation_strengths")
            strength_labels_es = {"strength_creative": "Creativo/a", "strength_autonomous": "Autonomo/a", "strength_motivated": "Motivado/a", "strength_tolerant_frustration": "Tolerante a la frustracion", "strength_respects_rules": "Respeta las normas", "strength_follows_instructions": "Sigue instrucciones", "strength_peers_relationship": "Buena relacion con pares", "strength_attends_gladly": "Asiste con gusto", "strength_family_communication": "Comunicacion con la familia", "strength_persistent": "Persistente", "strength_self_confident": "Seguro/a de si mismo/a", "strength_participates": "Participa", "strength_pays_attention": "Presta atencion", "strength_interested_learning": "Interesado/a en el aprendizaje", "strength_completes_tasks": "Completa tareas", "strength_adults_relationship": "Buena relacion con adultos", "strength_attends_regularly": "Asiste regularmente", "strength_family_committed": "Familia comprometida"}
            strengths_text = "—"
            if strengths_raw not in (None, ""):
                if isinstance(strengths_raw, dict):
                    strengths_list = [strength_labels_es.get(k, k) for k, v in strengths_raw.items() if v is True]
                    strengths_text = ", ".join(strengths_list) if strengths_list else "Ninguna marcada."
                elif isinstance(strengths_raw, str):
                    try:
                        parsed = __import__("json").loads(strengths_raw)
                        strengths_list = [strength_labels_es.get(k, k) for k, v in (parsed if isinstance(parsed, dict) else {}).items() if v is True]
                        strengths_text = ", ".join(strengths_list) if strengths_list else "Ninguna marcada."
                    except Exception:
                        strengths_text = str(strengths_raw).strip() or "—"
                else:
                    strengths_text = str(strengths_raw).strip() or "—"
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(strengths_text), normal_style))
            elements.append(Spacer(1, 0.15*inch))
            elements.append(Paragraph("Observaciones generales", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations", "")) if get_value("observations", "") else "—", normal_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("III. ACTITUD", section_style))
            elements.append(Spacer(1, 0.08*inch))
            elements.append(Paragraph(DOC37_LEGEND_TEXT, legend_style))
            elements.append(Spacer(1, 0.06*inch))
            elements.append(build_table("Actitud", 13, lambda i: get_value(f"attitude_{i}", ""), DOC37_ACTITUD_INDICADORES))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("Observaciones actitud:", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations_attitude", "")) if get_value("observations_attitude", "") else "—", normal_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("IV. LENGUAJE", section_style))
            elements.append(Spacer(1, 0.08*inch))
            elements.append(Paragraph(DOC37_LEGEND_TEXT, legend_style))
            elements.append(Spacer(1, 0.06*inch))
            elements.append(build_table("Lenguaje", 11, lambda i: get_value(f"language_{i}", ""), DOC37_LENGUAJE_INDICADORES))
            elements.append(Spacer(1, 0.12*inch))
            rt_display = format_multi_value(get_value("reading_type", ""), reading_type_labels)
            cl_display = format_multi_value(get_value("comprehension_level", ""), comprehension_labels)
            wl_display = format_multi_value(get_value("writing_level", ""), writing_labels)
            elements.append(Paragraph("<b>Tipo de lectura:</b> " + DocumentsClass._escape_html_for_paragraph(rt_display), normal_style))
            elements.append(Paragraph("<b>Nivel de comprensión:</b> " + DocumentsClass._escape_html_for_paragraph(cl_display), normal_style))
            elements.append(Paragraph("<b>Nivel de escritura:</b> " + DocumentsClass._escape_html_for_paragraph(wl_display), normal_style))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("Observaciones lenguaje:", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations_language", "")) if get_value("observations_language", "") else "—", normal_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("V. MATEMÁTICA", section_style))
            elements.append(Spacer(1, 0.08*inch))
            elements.append(Paragraph(DOC37_LEGEND_TEXT, legend_style))
            elements.append(Spacer(1, 0.06*inch))
            elements.append(build_table("Matemática", 13, lambda i: get_value(f"mathematics_{i}", ""), DOC37_MATEMATICA_INDICADORES))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("Observaciones matemática:", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations_mathematics", "")) if get_value("observations_mathematics", "") else "—", normal_style))
            doc.build(elements)
            return {"status": "success", "message": "PDF generado exitosamente desde cero", "filename": unique_filename, "file_path": str(output_file)}
        except ImportError:
            return {"status": "error", "message": "ReportLab no esta instalado.", "filename": None, "file_path": None}
        except Exception as e:
            return {"status": "error", "message": f"Error generando PDF documento 37: {str(e)}", "filename": None, "file_path": None}

    @staticmethod
    def _generate_pedagogical_evaluation_classroom_eighth_grade_from_scratch(
        document_id: int,
        doc_data: Dict[str, Any],
        db: Optional[Session] = None,
        output_directory: str = "files/system/students"
    ) -> Dict[str, Any]:
        """Genera el PDF del documento 38 (Pauta de evaluación pedagógica - Docente de aula - 8º Básico). 13 actitud, 11 lenguaje, 13 matemática."""
        try:
            if not REPORTLAB_AVAILABLE:
                return {"status": "error", "message": "ReportLab no esta instalado.", "filename": None, "file_path": None}

            def get_value(key: str, default: str = "") -> str:
                v = doc_data.get(key)
                if v is None:
                    return default
                return str(v).strip() if v else default

            def format_date(date_val) -> str:
                if not date_val:
                    return ""
                try:
                    if isinstance(date_val, str) and len(date_val) >= 10:
                        d = datetime.strptime(date_val[:10], "%Y-%m-%d").date()
                        return d.strftime("%d/%m/%Y")
                    if hasattr(date_val, "strftime"):
                        return date_val.strftime("%d/%m/%Y")
                    return str(date_val)[:10]
                except Exception:
                    return str(date_val)[:10] if date_val else ""

            student_name = get_value("student_full_name", "estudiante").replace(" ", "_")
            unique_filename = f"pauta_eval_pedagogica_first_grade_secondary_{student_name}_{uuid.uuid4().hex[:8]}.pdf"
            output_file = Path(output_directory) / unique_filename
            output_file.parent.mkdir(parents=True, exist_ok=True)

            doc = SimpleDocTemplate(str(output_file), pagesize=A4, rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
            elements = []
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle('Doc38Title', parent=styles['Heading1'], fontSize=14, textColor=colors.HexColor('#000000'), spaceAfter=14, spaceBefore=8, alignment=TA_CENTER, fontName='Helvetica-Bold')
            section_style = ParagraphStyle('Doc38Section', parent=styles['Normal'], fontSize=11, textColor=colors.HexColor('#000000'), spaceAfter=6, spaceBefore=10, fontName='Helvetica-Bold', backColor=colors.HexColor('#E8E8E8'), borderPadding=6)
            subtitle_style = ParagraphStyle('Doc38Subtitle', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#000000'), spaceAfter=4, spaceBefore=6, fontName='Helvetica-Bold')
            normal_style = ParagraphStyle('Doc38Normal', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#000000'), alignment=TA_LEFT, leading=12, spaceAfter=2)
            legend_style = ParagraphStyle('Doc38Legend', parent=styles['Normal'], fontSize=9, textColor=colors.HexColor('#333333'), alignment=TA_LEFT, spaceAfter=6, spaceBefore=2)
            header_cell_style = ParagraphStyle('Doc38HeaderCell', parent=styles['Normal'], fontSize=9, textColor=colors.white, alignment=TA_CENTER, fontName='Helvetica-Bold')
            cell_center_style = ParagraphStyle('Doc38CellCenter', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#000000'), alignment=TA_CENTER)
            DOC38_RESPONSE_COLS = ["S", "G", "O", "P/V", "N", "N/O"]
            DOC38_LEGEND_TEXT = "S: Siempre   G: Generalmente   O: Ocasionalmente   P/V: Pocas veces   N: Nunca   N/O: No observado"
            # Doc 38 (8vo Basico): 13 actitud, 11 lenguaje, 13 matematica (pauta oficial)
            DOC38_ACTITUD_INDICADORES = [
                "Manifiesta curiosidad e interes por aprender",
                "Adopta una actitud positiva hacia si mismo y el aprendizaje",
                "Expresa y escucha ideas de forma respetuosa",
                "Participa activamente en actividades grupales y colaborativas",
                "Muestra disposicion para compartir ideas, experiencias y opiniones con los demas",
                "Inicia, persiste y finaliza sus tareas de forma autonoma",
                "Realiza sus trabajos de manera ordenada y metodica",
                "Aborda y resuelve problemas de manera flexible y creativa",
                "Se adapta con flexibilidad a cambios, imprevistos o nuevas rutinas de trabajo",
                "Busca apoyo o aclara dudas cuando enfrenta dificultades",
                "Es receptivo a la retroalimentacion y la aplica para mejorar",
                "Demuestra una actitud de esfuerzo y perseverancia",
                "Reflexiona sobre si mismo, sus ideas y sus intereses para comprenderse y valorarse",
            ]
            DOC38_LENGUAJE_INDICADORES = [
                "Lee de manera fluida textos variados apropiados para su edad",
                "Lee e identifica un amplio repertorio de literatura, como poemas, leyendas, fabulas y novelas",
                "Comprende y analiza aspectos relevantes de los textos leidos",
                "Evalua criticamente la informacion presente en textos de diversa procedencia",
                "Busca y selecciona la informacion mas relevante sobre un tema en diversas fuentes para llevar a cabo una investigacion",
                "Planifica sus textos, estableciendo un proposito y destinatario",
                "Escribe narraciones y otros tipos de textos de manera autonoma y creativa, con estructura clara, conectores adecuados e informacion, descripciones o dialogos pertinentes",
                "Revisa y edita sus textos para satisfacer un proposito y transmitir sus ideas con claridad",
                "Comprende, interactua y obtiene informacion relevante de textos orales",
                "Dialoga para compartir, desarrollar ideas y buscar acuerdos",
                "Se expresa de manera clara y efectiva en exposiciones orales para comunicar temas de su interes",
            ]
            DOC38_MATEMATICA_INDICADORES = [
                "Resuelve ejercicios y problemas que involucran las cuatro operaciones y sus combinaciones",
                "Comprende y aplica el concepto de razon",
                "Comprende y aplica el concepto de porcentaje",
                "Comprende y resuelve ejercicios con fracciones y numeros mixtos",
                "Comprende la multiplicacion de decimales por numeros naturales de un digito o multiplos de 10",
                "Comprende la division de decimales por numeros naturales de un digito o multiplos de 10",
                "Resuelve ejercicios y problemas que involucran adicion y sustraccion de fracciones propias, impropias, numeros mixtos o decimales hasta la milesima",
                "Resuelve ecuaciones de primer grado con una incognita",
                "Identifica y calcula los angulos que se forman entre dos rectas que se cortan",
                "Comprende y calcula el area de una superficie de cubos y paralelepipedos a partir de sus redes",
                "Calcula la superficie de cubos y paralelepipedos, expresando el resultado en cm2 y m2",
                "Calcula el volumen de cubos y paralelepipedos, expresando el resultado en cm3, m3 y mm3",
                "Lee e interpreta graficos de barra doble y circulares, y comunica sus conclusiones de manera clara",
            ]

            def normalize_response(val):
                if val is None or (isinstance(val, str) and not val.strip()):
                    return None
                v = str(val).strip().upper().replace(" ", "")
                if v in ("S", "SIEMPRE"): return "S"
                if v in ("G", "GENERALMENTE"): return "G"
                if v in ("O", "OCASIONALMENTE"): return "O"
                if v in ("P/V", "PV", "POCASVECES", "P"): return "P/V"
                if v in ("N/O", "NO", "NOOBSERVADO"): return "N/O"
                if v in ("N", "NUNCA"): return "N"
                if len(v) == 1 and v in "SGOPN": return "P/V" if v == "P" else v
                return None

            def build_table(header_label, num_items, get_val_fn, indicator_labels):
                rows = []
                rows.append([Paragraph(DocumentsClass._escape_html_for_paragraph("Indicadores"), header_cell_style)] + [Paragraph(c, header_cell_style) for c in DOC38_RESPONSE_COLS])
                for i in range(1, num_items + 1):
                    val = get_val_fn(i)
                    col_key = normalize_response(val)
                    label_text = indicator_labels[i - 1] if indicator_labels and len(indicator_labels) >= i else f"{header_label} {i}"
                    cells = [Paragraph(DocumentsClass._escape_html_for_paragraph(label_text), normal_style)]
                    for c in DOC38_RESPONSE_COLS:
                        cells.append(Paragraph("✓" if col_key == c else "", cell_center_style))
                    rows.append(cells)
                col_ind = 8*cm
                col_opt = (16*cm - col_ind) / 6
                col_widths = [col_ind] + [col_opt] * 6
                tbl = Table(rows, colWidths=col_widths)
                tbl.setStyle(TableStyle([
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'), ('ALIGN', (0, 0), (0, -1), 'LEFT'), ('ALIGN', (1, 0), (-1, -1), 'CENTER'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6), ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ('TOPPADDING', (0, 0), (-1, -1), 5), ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#505050')), ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ]))
                return tbl

            def format_multi_value(raw: str, labels_dict: dict) -> str:
                if not raw or not str(raw).strip():
                    return "—"
                parts = [p.strip().lower().replace(" ", "_") for p in str(raw).split(",") if p and p.strip()]
                if not parts:
                    return "—"
                display = [labels_dict.get(p) or p for p in parts]
                return ", ".join(display)
            reading_type_labels = {"en_desarrollo": "En desarrollo", "silabica": "Silabica", "palabra_a_palabra": "Palabra a palabra", "por_unidades_cortas": "Por unidades cortas", "fluida": "Fluida"}
            comprehension_labels = {"en_desarrollo": "En desarrollo", "literal": "Literal", "inferencial": "Inferencial", "critica": "Critica"}
            writing_labels = {"en_desarrollo": "En desarrollo", "copia": "Copia", "dictado": "Dictado", "espontanea": "Espontanea"}

            elements.append(Paragraph("Pauta de evaluación pedagógica - Docente de aula - 8º Básico", title_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("I. DATOS DEL ESTUDIANTE E INFORME", section_style))
            elements.append(Spacer(1, 0.15*inch))
            col_label, col_value = 5*cm, 11*cm
            id_data = [
                [Paragraph("<b>Nombre completo:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("student_full_name", "—")), normal_style)],
                [Paragraph("<b>RUT:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("student_identification_number", "—")), normal_style)],
                [Paragraph("<b>Fecha de nacimiento:</b>", subtitle_style), Paragraph(format_date(doc_data.get("student_born_date")) or "—", normal_style)],
                [Paragraph("<b>Edad:</b>", subtitle_style), Paragraph(get_value("student_age", "—"), normal_style)],
                [Paragraph("<b>Establecimiento:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("establishment_id", "—")), normal_style)],
                [Paragraph("<b>Curso:</b>", subtitle_style), Paragraph(get_value("course", "—"), normal_style)],
                [Paragraph("<b>Fecha del informe:</b>", subtitle_style), Paragraph(format_date(doc_data.get("report_date")) or "—", normal_style)],
                [Paragraph("<b>Repeticiones:</b>", subtitle_style), Paragraph(get_value("repetitions", "—"), normal_style)],
            ]
            report_type_val = (get_value("report_type", "") or "").strip()
            report_type_val = "Anual" if report_type_val.lower() == "anual" else ("Semestral" if report_type_val.lower() == "semestral" else report_type_val or "—")
            id_data.append([Paragraph("<b>Tipo de informe:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(report_type_val), normal_style)])
            tbl_id = Table(id_data, colWidths=[col_label, col_value])
            tbl_id.setStyle(TableStyle([('VALIGN', (0, 0), (-1, -1), 'TOP'), ('ALIGN', (0, 0), (-1, -1), 'LEFT'), ('LEFTPADDING', (0, 0), (-1, -1), 8), ('RIGHTPADDING', (0, 0), (-1, -1), 8), ('TOPPADDING', (0, 0), (-1, -1), 6), ('BOTTOMPADDING', (0, 0), (-1, -1), 6), ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
            ]))
            elements.append(tbl_id)
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("II. FORTALEZAS DE LA SITUACION ESCOLAR", section_style))
            elements.append(Spacer(1, 0.1*inch))
            strengths_raw = doc_data.get("school_situation_strengths")
            strength_labels_es = {"strength_creative": "Creativo/a", "strength_autonomous": "Autonomo/a", "strength_motivated": "Motivado/a", "strength_tolerant_frustration": "Tolerante a la frustracion", "strength_respects_rules": "Respeta las normas", "strength_follows_instructions": "Sigue instrucciones", "strength_peers_relationship": "Buena relacion con pares", "strength_attends_gladly": "Asiste con gusto", "strength_family_communication": "Comunicacion con la familia", "strength_persistent": "Persistente", "strength_self_confident": "Seguro/a de si mismo/a", "strength_participates": "Participa", "strength_pays_attention": "Presta atencion", "strength_interested_learning": "Interesado/a en el aprendizaje", "strength_completes_tasks": "Completa tareas", "strength_adults_relationship": "Buena relacion con adultos", "strength_attends_regularly": "Asiste regularmente", "strength_family_committed": "Familia comprometida"}
            strengths_text = "—"
            if strengths_raw not in (None, ""):
                if isinstance(strengths_raw, dict):
                    strengths_list = [strength_labels_es.get(k, k) for k, v in strengths_raw.items() if v is True]
                    strengths_text = ", ".join(strengths_list) if strengths_list else "Ninguna marcada."
                elif isinstance(strengths_raw, str):
                    try:
                        parsed = __import__("json").loads(strengths_raw)
                        if isinstance(parsed, dict):
                            strengths_list = [strength_labels_es.get(k, k) for k, v in parsed.items() if v is True]
                            strengths_text = ", ".join(strengths_list) if strengths_list else "Ninguna marcada."
                        else:
                            strengths_text = str(strengths_raw).strip() or "—"
                    except Exception:
                        strengths_text = str(strengths_raw).strip() or "—"
                else:
                    strengths_text = str(strengths_raw).strip() or "—"
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(strengths_text), normal_style))
            elements.append(Spacer(1, 0.15*inch))
            elements.append(Paragraph("Observaciones generales", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations", "")) if get_value("observations", "") else "—", normal_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("III. ACTITUD", section_style))
            elements.append(Spacer(1, 0.08*inch))
            elements.append(Paragraph(DOC38_LEGEND_TEXT, legend_style))
            elements.append(Spacer(1, 0.06*inch))
            elements.append(build_table("Actitud", 13, lambda i: get_value(f"attitude_{i}", ""), DOC38_ACTITUD_INDICADORES))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("Observaciones actitud:", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations_attitude", "")) if get_value("observations_attitude", "") else "—", normal_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("IV. LENGUA Y LITERATURA", section_style))
            elements.append(Spacer(1, 0.08*inch))
            elements.append(Paragraph(DOC38_LEGEND_TEXT, legend_style))
            elements.append(Spacer(1, 0.06*inch))
            elements.append(build_table("Lengua y literatura", 11, lambda i: get_value(f"language_{i}", ""), DOC38_LENGUAJE_INDICADORES))
            elements.append(Spacer(1, 0.12*inch))
            rt_display = format_multi_value(get_value("reading_type", ""), reading_type_labels)
            cl_display = format_multi_value(get_value("comprehension_level", ""), comprehension_labels)
            wl_display = format_multi_value(get_value("writing_level", ""), writing_labels)
            elements.append(Paragraph("<b>Tipo de lectura:</b> " + DocumentsClass._escape_html_for_paragraph(rt_display), normal_style))
            elements.append(Paragraph("<b>Nivel de comprensión:</b> " + DocumentsClass._escape_html_for_paragraph(cl_display), normal_style))
            elements.append(Paragraph("<b>Nivel de escritura:</b> " + DocumentsClass._escape_html_for_paragraph(wl_display), normal_style))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("Observaciones lenguaje:", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations_language", "")) if get_value("observations_language", "") else "—", normal_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("V. MATEMÁTICA", section_style))
            elements.append(Spacer(1, 0.08*inch))
            elements.append(Paragraph(DOC38_LEGEND_TEXT, legend_style))
            elements.append(Spacer(1, 0.06*inch))
            elements.append(build_table("Matemática", 13, lambda i: get_value(f"mathematics_{i}", ""), DOC38_MATEMATICA_INDICADORES))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("Observaciones matemática:", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations_mathematics", "")) if get_value("observations_mathematics", "") else "—", normal_style))
            doc.build(elements)
            return {"status": "success", "message": "PDF generado exitosamente desde cero", "filename": unique_filename, "file_path": str(output_file)}
        except ImportError:
            return {"status": "error", "message": "ReportLab no esta instalado.", "filename": None, "file_path": None}
        except Exception as e:
            return {"status": "error", "message": f"Error generando PDF documento 39: {str(e)}", "filename": None, "file_path": None}

    @staticmethod
    def _generate_pedagogical_evaluation_classroom_second_grade_secondary_from_scratch(
        document_id: int,
        doc_data: Dict[str, Any],
        db: Optional[Session] = None,
        output_directory: str = "files/system/students"
    ) -> Dict[str, Any]:
        """Genera el PDF del documento 40 (Pauta de evaluación pedagógica - Docente de aula - 2º Medio). 13 actitud, 12 lengua y literatura, 6 matemática."""
        try:
            if not REPORTLAB_AVAILABLE:
                return {"status": "error", "message": "ReportLab no esta instalado.", "filename": None, "file_path": None}
            def get_value(key: str, default: str = "") -> str:
                v = doc_data.get(key)
                if v is None: return default
                return str(v).strip() if v else default
            def format_date(date_val) -> str:
                if not date_val: return ""
                try:
                    if isinstance(date_val, str) and len(date_val) >= 10:
                        d = datetime.strptime(date_val[:10], "%Y-%m-%d").date()
                        return d.strftime("%d/%m/%Y")
                    if hasattr(date_val, "strftime"): return date_val.strftime("%d/%m/%Y")
                    return str(date_val)[:10]
                except Exception:
                    return str(date_val)[:10] if date_val else ""
            student_name = get_value("student_full_name", "estudiante").replace(" ", "_")
            unique_filename = f"pauta_eval_pedagogica_second_grade_secondary_{student_name}_{uuid.uuid4().hex[:8]}.pdf"
            output_file = Path(output_directory) / unique_filename
            output_file.parent.mkdir(parents=True, exist_ok=True)
            doc = SimpleDocTemplate(str(output_file), pagesize=A4, rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
            elements = []
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle('Doc39Title', parent=styles['Heading1'], fontSize=14, textColor=colors.HexColor('#000000'), spaceAfter=14, spaceBefore=8, alignment=TA_CENTER, fontName='Helvetica-Bold')
            section_style = ParagraphStyle('Doc39Section', parent=styles['Normal'], fontSize=11, textColor=colors.HexColor('#000000'), spaceAfter=6, spaceBefore=10, fontName='Helvetica-Bold', backColor=colors.HexColor('#E8E8E8'), borderPadding=6)
            subtitle_style = ParagraphStyle('Doc39Subtitle', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#000000'), spaceAfter=4, spaceBefore=6, fontName='Helvetica-Bold')
            normal_style = ParagraphStyle('Doc39Normal', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#000000'), alignment=TA_LEFT, leading=12, spaceAfter=2)
            legend_style = ParagraphStyle('Doc39Legend', parent=styles['Normal'], fontSize=9, textColor=colors.HexColor('#333333'), alignment=TA_LEFT, spaceAfter=6, spaceBefore=2)
            header_cell_style = ParagraphStyle('Doc39HeaderCell', parent=styles['Normal'], fontSize=9, textColor=colors.white, alignment=TA_CENTER, fontName='Helvetica-Bold')
            cell_center_style = ParagraphStyle('Doc39CellCenter', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#000000'), alignment=TA_CENTER)
            DOC39_RESPONSE_COLS = ["S", "G", "O", "P/V", "N", "N/O"]
            DOC39_LEGEND_TEXT = "S: Siempre   G: Generalmente   O: Ocasionalmente   P/V: Pocas veces   N: Nunca   N/O: No observado"
            DOC39_ACTITUD_INDICADORES = [
                "Manifiesta curiosidad e interes por aprender", "Adopta una actitud positiva hacia si mismo y el aprendizaje",
                "Expresa y escucha ideas de forma respetuosa", "Participa activamente en actividades grupales y colaborativas",
                "Muestra disposicion para compartir ideas, experiencias y opiniones con los demas", "Inicia, persiste y finaliza sus tareas de forma autonoma",
                "Realiza sus trabajos de manera ordenada y metodica", "Aborda y resuelve problemas de manera flexible y creativa",
                "Se adapta con flexibilidad a cambios, imprevistos o nuevas rutinas de trabajo", "Busca apoyo o aclara dudas cuando enfrenta dificultades",
                "Es receptivo a la retroalimentacion y la aplica para mejorar", "Demuestra una actitud de esfuerzo y perseverancia",
                "Reflexiona sobre si mismo, sus ideas y sus intereses para comprenderse y valorarse",
            ]
            DOC39_LENGUAJE_INDICADORES = [
                "Comprende y analiza distintos tipos de texto apropiados para su nivel", "Formula interpretaciones fundamentadas de textos leidos o escuchados",
                "Analiza y evalua textos argumentativos, como columnas de opinion, cartas y discursos",
                "Analiza y evalua criticamente textos de medios de comunicacion, como noticias, reportajes, cartas al director, textos publicitarios o publicaciones en redes sociales",
                "Se expresa creativamente mediante la escritura de textos de diversos generos", "Redacta textos con proposito persuasivo, como cartas al director, editoriales o criticas literarias",
                "Planifica, escribe, revisa y edita sus textos considerando el contexto, destinatario y proposito comunicativo", "Escribe textos aplicando correctamente las reglas ortograficas",
                "Comprende y utiliza recursos de correferencia lexica compleja",
                "Comprende, compara y analiza textos orales y audiovisuales, como exposiciones, discursos, documentales, noticias y reportajes",
                "Dialoga de manera constructiva para debatir ideas y llegar a acuerdos", "Realiza investigaciones sobre diversos temas ajustandose al proposito planteado",
            ]
            DOC39_MATEMATICA_INDICADORES = [
                "Realiza calculos y estimaciones que involucran operaciones con numeros reales",
                "Comprende las relaciones entre potencias, raices enesimas y logaritmos, aplicando sus propiedades en diferentes contextos",
                "Comprende la funcion cuadratica",
                "Comprende la funcion inversa",
                "Comprende las razones trigonometricas (seno, coseno y tangente) en triangulos rectangulos y las aplica para resolver problemas",
                "Utiliza permutaciones y combinaciones sencillas para calcular probabilidades de eventos y resolver problemas",
            ]
            def normalize_response(val):
                if val is None or (isinstance(val, str) and not val.strip()): return None
                v = str(val).strip().upper().replace(" ", "")
                if v in ("S", "SIEMPRE"): return "S"
                if v in ("G", "GENERALMENTE"): return "G"
                if v in ("O", "OCASIONALMENTE"): return "O"
                if v in ("P/V", "PV", "POCASVECES", "P"): return "P/V"
                if v in ("N/O", "NO", "NOOBSERVADO"): return "N/O"
                if v in ("N", "NUNCA"): return "N"
                if len(v) == 1 and v in "SGOPN": return "P/V" if v == "P" else v
                return None
            def build_table(header_label, num_items, get_val_fn, indicator_labels):
                rows = []
                rows.append([Paragraph(DocumentsClass._escape_html_for_paragraph("Indicadores"), header_cell_style)] + [Paragraph(c, header_cell_style) for c in DOC39_RESPONSE_COLS])
                for i in range(1, num_items + 1):
                    val = get_val_fn(i)
                    col_key = normalize_response(val)
                    label_text = indicator_labels[i - 1] if indicator_labels and len(indicator_labels) >= i else f"{header_label} {i}"
                    cells = [Paragraph(DocumentsClass._escape_html_for_paragraph(label_text), normal_style)]
                    for c in DOC39_RESPONSE_COLS:
                        cells.append(Paragraph("✓" if col_key == c else "", cell_center_style))
                    rows.append(cells)
                col_ind = 8*cm
                col_opt = (16*cm - col_ind) / 6
                col_widths = [col_ind] + [col_opt] * 6
                tbl = Table(rows, colWidths=col_widths)
                tbl.setStyle(TableStyle([
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'), ('ALIGN', (0, 0), (0, -1), 'LEFT'), ('ALIGN', (1, 0), (-1, -1), 'CENTER'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6), ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ('TOPPADDING', (0, 0), (-1, -1), 5), ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#505050')), ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ]))
                return tbl
            def format_multi_value(raw: str, labels_dict: dict) -> str:
                if not raw or not str(raw).strip(): return "—"
                parts = [p.strip().lower().replace(" ", "_") for p in str(raw).split(",") if p and p.strip()]
                if not parts: return "—"
                display = [labels_dict.get(p) or p for p in parts]
                return ", ".join(display)
            reading_type_labels = {"en_desarrollo": "En desarrollo", "silabica": "Silabica", "palabra_a_palabra": "Palabra a palabra", "por_unidades_cortas": "Por unidades cortas", "fluida": "Fluida"}
            comprehension_labels = {"en_desarrollo": "En desarrollo", "literal": "Literal", "inferencial": "Inferencial", "critica": "Critica"}
            writing_labels = {"en_desarrollo": "En desarrollo", "copia": "Copia", "dictado": "Dictado", "espontanea": "Espontanea"}
            elements.append(Paragraph("Pauta de evaluación pedagógica - Docente de aula - 2º Medio", title_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("I. DATOS DEL ESTUDIANTE E INFORME", section_style))
            elements.append(Spacer(1, 0.15*inch))
            col_label, col_value = 5*cm, 11*cm
            id_data = [
                [Paragraph("<b>Nombre completo:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("student_full_name", "—")), normal_style)],
                [Paragraph("<b>RUT:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("student_identification_number", "—")), normal_style)],
                [Paragraph("<b>Fecha de nacimiento:</b>", subtitle_style), Paragraph(format_date(doc_data.get("student_born_date")) or "—", normal_style)],
                [Paragraph("<b>Edad:</b>", subtitle_style), Paragraph(get_value("student_age", "—"), normal_style)],
                [Paragraph("<b>Establecimiento:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("establishment_id", "—")), normal_style)],
                [Paragraph("<b>Curso:</b>", subtitle_style), Paragraph(get_value("course", "—"), normal_style)],
                [Paragraph("<b>Fecha del informe:</b>", subtitle_style), Paragraph(format_date(doc_data.get("report_date")) or "—", normal_style)],
                [Paragraph("<b>Repeticiones:</b>", subtitle_style), Paragraph(get_value("repetitions", "—"), normal_style)],
            ]
            report_type_val = (get_value("report_type", "") or "").strip()
            report_type_val = "Anual" if report_type_val.lower() == "anual" else ("Semestral" if report_type_val.lower() == "semestral" else report_type_val or "—")
            id_data.append([Paragraph("<b>Tipo de informe:</b>", subtitle_style), Paragraph(DocumentsClass._escape_html_for_paragraph(report_type_val), normal_style)])
            tbl_id = Table(id_data, colWidths=[col_label, col_value])
            tbl_id.setStyle(TableStyle([('VALIGN', (0, 0), (-1, -1), 'TOP'), ('ALIGN', (0, 0), (-1, -1), 'LEFT'), ('LEFTPADDING', (0, 0), (-1, -1), 8), ('RIGHTPADDING', (0, 0), (-1, -1), 8), ('TOPPADDING', (0, 0), (-1, -1), 6), ('BOTTOMPADDING', (0, 0), (-1, -1), 6), ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC'))]))
            elements.append(tbl_id)
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("II. FORTALEZAS DE LA SITUACION ESCOLAR", section_style))
            elements.append(Spacer(1, 0.1*inch))
            strengths_raw = doc_data.get("school_situation_strengths")
            strength_labels_es = {"strength_creative": "Creativo/a", "strength_autonomous": "Autonomo/a", "strength_motivated": "Motivado/a", "strength_tolerant_frustration": "Tolerante a la frustracion", "strength_respects_rules": "Respeta las normas", "strength_follows_instructions": "Sigue instrucciones", "strength_peers_relationship": "Buena relacion con pares", "strength_attends_gladly": "Asiste con gusto", "strength_family_communication": "Comunicacion con la familia", "strength_persistent": "Persistente", "strength_self_confident": "Seguro/a de si mismo/a", "strength_participates": "Participa", "strength_pays_attention": "Presta atencion", "strength_interested_learning": "Interesado/a en el aprendizaje", "strength_completes_tasks": "Completa tareas", "strength_adults_relationship": "Buena relacion con adultos", "strength_attends_regularly": "Asiste regularmente", "strength_family_committed": "Familia comprometida"}
            strengths_text = "—"
            if strengths_raw not in (None, ""):
                if isinstance(strengths_raw, dict):
                    strengths_list = [strength_labels_es.get(k, k) for k, v in strengths_raw.items() if v is True]
                    strengths_text = ", ".join(strengths_list) if strengths_list else "Ninguna marcada."
                elif isinstance(strengths_raw, str):
                    try:
                        parsed = __import__("json").loads(strengths_raw)
                        strengths_list = [strength_labels_es.get(k, k) for k, v in (parsed if isinstance(parsed, dict) else {}).items() if v is True]
                        strengths_text = ", ".join(strengths_list) if strengths_list else "Ninguna marcada."
                    except Exception:
                        strengths_text = str(strengths_raw).strip() or "—"
                else:
                    strengths_text = str(strengths_raw).strip() or "—"
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(strengths_text), normal_style))
            elements.append(Spacer(1, 0.15*inch))
            elements.append(Paragraph("Observaciones generales", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations", "")) if get_value("observations", "") else "—", normal_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("III. ACTITUD", section_style))
            elements.append(Spacer(1, 0.08*inch))
            elements.append(Paragraph(DOC39_LEGEND_TEXT, legend_style))
            elements.append(Spacer(1, 0.06*inch))
            elements.append(build_table("Actitud", 13, lambda i: get_value(f"attitude_{i}", ""), DOC39_ACTITUD_INDICADORES))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("Observaciones actitud:", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations_attitude", "")) if get_value("observations_attitude", "") else "—", normal_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("IV. LENGUA Y LITERATURA", section_style))
            elements.append(Spacer(1, 0.08*inch))
            elements.append(Paragraph(DOC39_LEGEND_TEXT, legend_style))
            elements.append(Spacer(1, 0.06*inch))
            elements.append(build_table("Lengua y literatura", 12, lambda i: get_value(f"language_{i}", ""), DOC39_LENGUAJE_INDICADORES))
            elements.append(Spacer(1, 0.12*inch))
            rt_display = format_multi_value(get_value("reading_type", ""), reading_type_labels)
            cl_display = format_multi_value(get_value("comprehension_level", ""), comprehension_labels)
            wl_display = format_multi_value(get_value("writing_level", ""), writing_labels)
            elements.append(Paragraph("<b>Tipo de lectura:</b> " + DocumentsClass._escape_html_for_paragraph(rt_display), normal_style))
            elements.append(Paragraph("<b>Nivel de comprensión:</b> " + DocumentsClass._escape_html_for_paragraph(cl_display), normal_style))
            elements.append(Paragraph("<b>Nivel de escritura:</b> " + DocumentsClass._escape_html_for_paragraph(wl_display), normal_style))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("Observaciones lenguaje:", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations_language", "")) if get_value("observations_language", "") else "—", normal_style))
            elements.append(Spacer(1, 0.25*inch))
            elements.append(Paragraph("V. MATEMÁTICA", section_style))
            elements.append(Spacer(1, 0.08*inch))
            elements.append(Paragraph(DOC39_LEGEND_TEXT, legend_style))
            elements.append(Spacer(1, 0.06*inch))
            elements.append(build_table("Matemática", 6, lambda i: get_value(f"mathematics_{i}", ""), DOC39_MATEMATICA_INDICADORES))
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("Observaciones matemática:", subtitle_style))
            elements.append(Paragraph(DocumentsClass._escape_html_for_paragraph(get_value("observations_mathematics", "")) if get_value("observations_mathematics", "") else "—", normal_style))
            doc.build(elements)
            return {"status": "success", "message": "PDF generado exitosamente desde cero", "filename": unique_filename, "file_path": str(output_file)}
        except ImportError:
            return {"status": "error", "message": "ReportLab no esta instalado.", "filename": None, "file_path": None}
        except Exception as e:
            return {"status": "error", "message": f"Error generando PDF documento 40: {str(e)}", "filename": None, "file_path": None}

    @staticmethod
    def _generate_conners_teacher_from_scratch(
        document_id: int,
        conners_data: Dict[str, Any],
        db: Optional[Session] = None,
        output_directory: str = "files/system/students"
    ) -> Dict[str, Any]:
        """
        Genera el PDF del Test de Conners (Documento 29) desde cero.
        Dos partes con cuadrados: 1) Escala Conners 10 ítems (0-3), 2) Cuestionario de conducta 18 ítems (N/P/B/M).
        """
        # Textos de los 10 ítems de la Escala Conners (II)
        CONNERS_SCALE_ITEMS = [
            "Inquieto o hiperactivo",
            "Excitable, impulsivo",
            "Molesta a otros niños",
            "No termina lo que empieza—poca capacidad de atención",
            "Se mueve constantemente",
            "Desatento, se distrae con facilidad",
            "Exige que se cumplan sus demandas de inmediato—se frustra con facilidad",
            "Llora a menudo y con facilidad",
            "Los cambios de humor son rápidos y drásticos",
            "Estallidos de mal genio, conducta explosiva e impredecible",
        ]
        # Textos de los 18 ítems del Cuestionario de Conducta (III)
        CONNERS_CONDUCT_ITEMS = [
            "Emite sonidos molestos en situaciones inapropiadas.",
            "Habla en exceso.",
            "Interrumpe o se entromete en conversaciones o juegos.",
            "No espera su turno cuando debe hacer fila o en juegos.",
            "Responde inesperadamente o antes de que se haya concluido una pregunta.",
            "Tiene dificultades para permanecer sentado cuando se le pide.",
            "Se levanta en situaciones en que debería permanecer sentado.",
            "Corre o salta en situaciones inapropiadas.",
            "Tiene dificultades para jugar o dedicarse tranquilamente a actividades de ocio.",
            "Está en marcha o actúa como si tuviera un motor.",
            "Pierde cosas necesarias para tareas o actividades.",
            "Se distrae con estímulos externos.",
            "Muestra olvido en las actividades cotidianas.",
            "No sigue instrucciones y no termina tareas.",
            "No se lleva bien con la mayoría de sus compañeros/as.",
            "Evita o se muestra reacio a tareas que requieren esfuerzo mental sostenido.",
            "Comete frecuentes y numerosos errores, por olvidar o desatender detalles.",
            "Tiene dificultades de aprendizaje escolar.",
        ]
        try:
            if not REPORTLAB_AVAILABLE:
                return {
                    "status": "error",
                    "message": "ReportLab no está instalado. Instala con: pip install reportlab",
                    "filename": None,
                    "file_path": None
                }

            def g(k: str, d: str = ""):
                v = conners_data.get(k)
                if v is None:
                    return d
                return str(v).strip() if v else d

            student_name = g("student_full_name", "estudiante").replace(" ", "_")
            unique_filename = f"conners_teacher_{student_name}_{uuid.uuid4().hex[:8]}.pdf"
            output_file = Path(output_directory) / unique_filename
            output_file.parent.mkdir(parents=True, exist_ok=True)

            # Mismos márgenes y criterio de diseño que otros PDFs (cert. egreso PIE, cert. TEA)
            doc = SimpleDocTemplate(
                str(output_file),
                pagesize=A4,
                rightMargin=2.5*cm,
                leftMargin=2.5*cm,
                topMargin=2.5*cm,
                bottomMargin=2.5*cm
            )
            elements = []
            styles = getSampleStyleSheet()

            # Título: mismo criterio que certificados (centrado, negrita)
            title_style = ParagraphStyle(
                'ConnersTitle',
                parent=styles['Heading1'],
                fontSize=16,
                textColor=colors.HexColor('#000000'),
                spaceAfter=20,
                spaceBefore=12,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            section_style = ParagraphStyle(
                'ConnersSection',
                parent=styles['Heading2'],
                fontSize=12,
                textColor=colors.HexColor('#000000'),
                spaceAfter=8,
                spaceBefore=12,
                fontName='Helvetica-Bold'
            )
            normal_style = ParagraphStyle(
                'ConnersNormal',
                parent=styles['Normal'],
                fontSize=11,
                textColor=colors.HexColor('#000000'),
                leading=16,
                spaceAfter=8
            )
            cell_style_center = ParagraphStyle(
                'ConnersCell',
                parent=styles['Normal'],
                fontSize=10,
                alignment=TA_CENTER,
                textColor=colors.HexColor('#000000')
            )
            cell_style_left = ParagraphStyle(
                'ConnersCellLeft',
                parent=styles['Normal'],
                fontSize=9,
                alignment=TA_LEFT,
                textColor=colors.HexColor('#000000'),
                leftIndent=2,
                rightIndent=2,
            )

            # Encabezado
            elements.append(Paragraph("Test de Conners", title_style))
            elements.append(Spacer(1, 0.2*inch))

            student_full = DocumentsClass._escape_html_for_paragraph(g("student_full_name", "—"))
            student_rut = DocumentsClass._escape_html_for_paragraph(g("student_rut", "—"))
            eval_date_raw = g("evaluation_date", "")
            eval_date = "—"
            if eval_date_raw:
                try:
                    s = str(eval_date_raw).strip()[:10]
                    if len(s) == 10 and s[4] == "-" and s[7] == "-":
                        d = datetime.strptime(s, "%Y-%m-%d").date()
                        eval_date = d.strftime("%d/%m/%Y")
                    else:
                        eval_date = eval_date_raw
                except Exception:
                    eval_date = eval_date_raw if eval_date_raw else "—"
            evaluator = DocumentsClass._escape_html_for_paragraph(g("evaluator_name", "—"))
            eval_type = g("evaluation_type", "ingreso")
            eval_type_label = "Reevaluación" if eval_type == "reevaluacion" else "Ingreso"
            total_score = conners_data.get("total_score")
            total_score_str = str(total_score) if total_score is not None else "—"
            comments = DocumentsClass._escape_html_for_paragraph(g("comments_observations", ""))

            born_date_str = DocumentsClass._escape_html_for_paragraph(g("student_born_date", "—"))
            # Edad a la fecha de evaluación (años y meses)
            age_str = "—"
            born_raw = g("student_born_date", "")
            ref_date = None
            if eval_date_raw:
                try:
                    s = str(eval_date_raw).strip()[:10]
                    if len(s) == 10 and s[4] == "-" and s[7] == "-":
                        ref_date = datetime.strptime(s, "%Y-%m-%d").date()
                except Exception:
                    pass
            if ref_date is None:
                ref_date = date.today()
            if born_raw:
                try:
                    # Puede venir DD/MM/YYYY o YYYY-MM-DD
                    s = str(born_raw).strip()
                    if len(s) >= 10 and s[2] == "/" and s[5] == "/":
                        birth_d = datetime.strptime(s[:10], "%d/%m/%Y").date()
                    elif len(s) >= 10 and s[4] == "-" and s[7] == "-":
                        birth_d = datetime.strptime(s[:10], "%Y-%m-%d").date()
                    else:
                        birth_d = None
                    if birth_d:
                        years = ref_date.year - birth_d.year
                        if (ref_date.month, ref_date.day) < (birth_d.month, birth_d.day):
                            years -= 1
                        # Meses desde el último cumpleaños (siempre mostrar año y mes)
                        if (ref_date.month, ref_date.day) >= (birth_d.month, birth_d.day):
                            months = (ref_date.month - birth_d.month) + (0 if ref_date.day >= birth_d.day else -1)
                        else:
                            months = (12 - birth_d.month + ref_date.month - 1) + (0 if ref_date.day >= birth_d.day else -1)
                        if months < 0:
                            months += 12
                            years -= 1
                        if years < 0:
                            age_str = "—"
                        else:
                            age_str = f"{years} año{'s' if years != 1 else ''} {months} mes{'es' if months != 1 else ''}"
                except Exception:
                    pass
            age_str = DocumentsClass._escape_html_for_paragraph(age_str)
            info_data = [
                [Paragraph("<b>RUT:</b>", normal_style), Paragraph(student_rut, normal_style)],
                [Paragraph("<b>Estudiante:</b>", normal_style), Paragraph(student_full, normal_style)],
                [Paragraph("<b>Fecha de&nbsp;nacimiento:</b>", normal_style), Paragraph(born_date_str, normal_style)],
                [Paragraph("<b>Edad:</b>", normal_style), Paragraph(age_str, normal_style)],
                [Paragraph("<b>Fecha evaluación:</b>", normal_style), Paragraph(eval_date, normal_style)],
                [Paragraph("<b>Evaluador(a):</b>", normal_style), Paragraph(evaluator, normal_style)],
                [Paragraph("<b>Tipo:</b>", normal_style), Paragraph(eval_type_label, normal_style)],
            ]
            info_table = Table(info_data, colWidths=[5.5*cm, 10.5*cm])
            info_table.setStyle(TableStyle([
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
            ]))
            elements.append(info_table)
            elements.append(Spacer(1, 0.25*inch))

            # --- II. Escala Conners (10 ítems) - cuadrados 0,1,2,3 ---
            elements.append(Paragraph("II. Escala Conners (10 ítems)", section_style))
            elements.append(Paragraph("0 = Nada,  1 = Un poco,  2 = Bastante,  3 = Mucho", normal_style))
            elements.append(Spacer(1, 0.1*inch))

            scores_list = conners_data.get("scores") or []
            scores_by_item = {int(s.get("item_index", 0)): int(s.get("score", 0)) for s in scores_list if s.get("item_index") is not None}

            # Ancho útil A4 menos márgenes (21 - 5 = 16 cm); columna ítem ancha para el texto
            table_width = 16*cm
            col_item = 9*cm
            col_opt = (table_width - col_item) / 4
            col_widths = [col_item, col_opt, col_opt, col_opt, col_opt]
            header_cells_1 = [
                Paragraph(DocumentsClass._escape_html_for_paragraph("Ítem"), cell_style_center),
                Paragraph("Nada<br/>0", cell_style_center),
                Paragraph("Un poco<br/>1", cell_style_center),
                Paragraph("Bastante<br/>2", cell_style_center),
                Paragraph("Mucho<br/>3", cell_style_center),
            ]
            table_data_1 = [header_cells_1]
            for i in range(1, 11):
                sc = scores_by_item.get(i, -1)
                desc = CONNERS_SCALE_ITEMS[i - 1] if i <= len(CONNERS_SCALE_ITEMS) else ""
                item_text = f"{i}. {desc}" if desc else str(i)
                row = [Paragraph(DocumentsClass._escape_html_for_paragraph(item_text), cell_style_left)]
                for opt in range(4):
                    cell_text = "X" if sc == opt else ""
                    row.append(Paragraph(cell_text, cell_style_center))
                table_data_1.append(row)

            tbl1 = Table(table_data_1, colWidths=col_widths)
            tbl1.setStyle(TableStyle([
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                ('ALIGN', (1, 0), (-1, -1), 'CENTER'),
                ('LEFTPADDING', (0, 0), (-1, -1), 4),
                ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#333333')),
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E0E0E0')),
            ]))
            elements.append(tbl1)
            elements.append(Spacer(1, 0.12*inch))
            elements.append(Paragraph(f"<b>Puntaje total (0-30):</b> {total_score_str}", normal_style))
            elements.append(Spacer(1, 0.25*inch))

            # --- III. Cuestionario de conducta (18 ítems) - en la segunda hoja ---
            elements.append(PageBreak())
            elements.append(Paragraph("III. Cuestionario de Conducta (18 ítems)", section_style))
            elements.append(Paragraph("N = Nada,  P = Poco,  B = Bastante,  M = Mucho", normal_style))
            elements.append(Spacer(1, 0.1*inch))

            conduct_list = conners_data.get("conduct_responses") or []
            conduct_by_item = {}
            for r in conduct_list:
                idx = r.get("item_index")
                if idx is not None:
                    val = (r.get("response") or "").strip().lower()[:1]
                    if val in ("n", "p", "b", "m"):
                        conduct_by_item[int(idx)] = val

            col_widths2 = [col_item, col_opt, col_opt, col_opt, col_opt]
            header_cells_2 = [
                Paragraph(DocumentsClass._escape_html_for_paragraph("Ítem"), cell_style_center),
                Paragraph("Nada<br/>N", cell_style_center),
                Paragraph("Poco<br/>P", cell_style_center),
                Paragraph("Bastante<br/>B", cell_style_center),
                Paragraph("Mucho<br/>M", cell_style_center),
            ]
            table_data_2 = [header_cells_2]
            for i in range(1, 19):
                resp = conduct_by_item.get(i, "")
                desc = CONNERS_CONDUCT_ITEMS[i - 1] if i <= len(CONNERS_CONDUCT_ITEMS) else ""
                item_text = f"{i}. {desc}" if desc else str(i)
                row = [Paragraph(DocumentsClass._escape_html_for_paragraph(item_text), cell_style_left)]
                for opt in ["n", "p", "b", "m"]:
                    cell_text = "X" if resp == opt else ""
                    row.append(Paragraph(cell_text, cell_style_center))
                table_data_2.append(row)

            tbl2 = Table(table_data_2, colWidths=col_widths2)
            tbl2.setStyle(TableStyle([
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                ('ALIGN', (1, 0), (-1, -1), 'CENTER'),
                ('LEFTPADDING', (0, 0), (-1, -1), 4),
                ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#333333')),
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E0E0E0')),
            ]))
            elements.append(tbl2)
            elements.append(Spacer(1, 0.3*inch))

            # --- IV. Comentarios y observaciones ---
            elements.append(Paragraph("IV. Comentarios y observaciones", section_style))
            elements.append(Paragraph(comments if comments else "—", normal_style))

            doc.build(elements)
            return {
                "status": "success",
                "message": "PDF Conners generado exitosamente",
                "filename": unique_filename,
                "file_path": str(output_file)
            }
        except ImportError:
            return {
                "status": "error",
                "message": "ReportLab no está instalado. Instala con: pip install reportlab",
                "filename": None,
                "file_path": None
            }
        except Exception as e:
            return {
                "status": "error",
                "message": f"Error generando PDF Conners: {str(e)}",
                "filename": None,
                "file_path": None
            }

    @staticmethod
    def generate_health_evaluation_document(
        template_path: str,
        evaluation_data: Dict[str, Any],
        output_directory: str = "files/original_student_files"
    ) -> Dict[str, Any]:
        """
        Genera un documento Word de evaluación de salud a partir de un template.
        Reemplaza los campos del template con los datos de la evaluación.
        
        Los campos pueden identificarse de varias formas:
        1. Campos de formulario (form fields) con nombres específicos
        2. Placeholders de texto como [FULL_NAME], {full_name}, etc.
        3. Texto que identifica la posición seguido de un campo
        
        Args:
            template_path: Ruta del archivo Word template
            evaluation_data: Diccionario con los datos de la evaluación de salud
            output_directory: Directorio donde se guardará el archivo generado
        
        Returns:
            dict: Diccionario con status, message, filename y file_path
        """
        converted_path = None  # Variable para rastrear archivos temporales
        try:
            template_file = Path(template_path)
            
            if not template_file.exists():
                return {
                    "status": "error",
                    "message": "Template no encontrado",
                    "filename": None,
                    "file_path": None
                }
            
            # Validar y convertir si es necesario
            file_extension = template_file.suffix.lower()
            
            # Si es .doc, intentar convertirlo a .docx temporalmente
            if file_extension == '.doc':
                converted_path = None
                temp_dir = tempfile.gettempdir()
                temp_docx = os.path.join(temp_dir, f"temp_{uuid.uuid4().hex[:8]}.docx")
                
                # Detectar el sistema operativo
                is_windows = platform.system() == "Windows"
                
                # Intentar convertir usando LibreOffice (funciona en Windows y Linux)
                try:
                    # En Linux, el comando puede ser 'libreoffice' o 'soffice'
                    # En Windows, generalmente es 'soffice.exe' o la ruta completa
                    libreoffice_cmd = None
                    
                    if is_windows:
                        # Intentar diferentes rutas comunes en Windows
                        # Buscar en el PATH primero, luego en ubicaciones comunes
                        possible_paths = [
                            "soffice.exe",  # En PATH
                            "soffice",      # En PATH sin extensión
                            r"C:\Program Files\LibreOffice\program\soffice.exe",
                            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                            r"C:\Program Files\LibreOffice 7\program\soffice.exe",
                            r"C:\Program Files\LibreOffice 6\program\soffice.exe",
                            r"C:\Program Files\LibreOffice 5\program\soffice.exe",
                        ]
                        # También buscar en AppData (instalaciones portables)
                        appdata = os.getenv('LOCALAPPDATA', '')
                        if appdata:
                            possible_paths.append(os.path.join(appdata, r"Programs\LibreOffice\program\soffice.exe"))
                    else:
                        # En Linux, generalmente está en el PATH
                        possible_paths = ["libreoffice", "soffice", "/usr/bin/libreoffice", "/usr/bin/soffice"]
                    
                    # Buscar el comando disponible
                    for cmd in possible_paths:
                        try:
                            result = subprocess.run(
                                [cmd, "--version"],
                                capture_output=True,
                                timeout=5
                            )
                            if result.returncode == 0:
                                libreoffice_cmd = cmd
                                break
                        except (FileNotFoundError, subprocess.TimeoutExpired):
                            continue
                    
                    if not libreoffice_cmd:
                        # Si no se encuentra LibreOffice, intentar win32com solo en Windows
                        if is_windows:
                            try:
                                if not WIN32_AVAILABLE:
                                    raise ImportError("win32com no disponible")
                                
                                # Abrir Word y convertir
                                word = win32com.client.Dispatch("Word.Application")
                                word.Visible = False
                                doc = word.Documents.Open(str(template_file.absolute()))
                                doc.SaveAs2(temp_docx, FileFormat=16)  # 16 = wdFormatXMLDocument (.docx)
                                doc.Close()
                                word.Quit()
                                
                                template_file = Path(temp_docx)
                                converted_path = temp_docx
                            except ImportError:
                                # pywin32 no está instalado
                                if is_windows:
                                    return {
                                        "status": "error",
                                        "message": "No se encontró LibreOffice ni pywin32 para usar Microsoft Word. Opciones: 1) Instala LibreOffice (recomendado), 2) Instala pywin32: pip install pywin32 (requiere Word instalado), 3) Convierte el archivo manualmente a .docx.",
                                        "filename": None,
                                        "file_path": None
                                    }
                                else:
                                    return {
                                        "status": "error",
                                        "message": "LibreOffice no está instalado. Para convertir archivos .doc a .docx en Ubuntu/Linux, instala LibreOffice: sudo apt-get update && sudo apt-get install -y libreoffice",
                                        "filename": None,
                                        "file_path": None
                                    }
                            except Exception as e:
                                error_str = str(e)
                                # Detectar si el error es porque Word no está instalado
                                if "Word.Application" in error_str or "COM" in error_str or "0x80040154" in error_str:
                                    return {
                                        "status": "error",
                                        "message": "Microsoft Word no está instalado o no está disponible. Opciones: 1) Instala LibreOffice (recomendado, gratuito): https://www.libreoffice.org/download/, 2) Instala Microsoft Word, 3) Convierte el archivo FU_EVALUACION-DE_SALUD_2024-3.doc manualmente a .docx usando Word o LibreOffice.",
                                        "filename": None,
                                        "file_path": None
                                    }
                                else:
                                    return {
                                        "status": "error",
                                        "message": f"Error al convertir .doc a .docx usando Word: {error_str}. Asegúrate de tener Microsoft Word instalado o instala LibreOffice (gratis): https://www.libreoffice.org/download/",
                                        "filename": None,
                                        "file_path": None
                                    }
                        else:
                            return {
                                "status": "error",
                                "message": "LibreOffice no está instalado. Para convertir archivos .doc a .docx en Ubuntu, instala LibreOffice: sudo apt-get update && sudo apt-get install -y libreoffice",
                                "filename": None,
                                "file_path": None
                            }
                    else:
                        # Convertir usando LibreOffice
                        result = subprocess.run([
                            libreoffice_cmd, "--headless", "--convert-to", "docx",
                            "--outdir", temp_dir,
                            str(template_file.absolute())
                        ], capture_output=True, timeout=30)
                        
                        if result.returncode == 0:
                            # LibreOffice crea el archivo con el mismo nombre pero .docx
                            converted_docx = Path(temp_dir) / f"{template_file.stem}.docx"
                            if converted_docx.exists():
                                converted_docx.rename(temp_docx)
                                template_file = Path(temp_docx)
                                converted_path = temp_docx
                            else:
                                return {
                                    "status": "error",
                                    "message": "No se pudo crear el archivo convertido. Verifica los permisos del directorio temporal.",
                                    "filename": None,
                                    "file_path": None
                                }
                        else:
                            error_msg = result.stderr.decode('utf-8', errors='ignore') if result.stderr else 'Error desconocido'
                            return {
                                "status": "error",
                                "message": f"Error al convertir .doc a .docx con LibreOffice: {error_msg}",
                                "filename": None,
                                "file_path": None
                            }
                            
                except subprocess.TimeoutExpired:
                    return {
                        "status": "error",
                        "message": "Timeout al convertir el archivo .doc. El proceso tardó demasiado.",
                        "filename": None,
                        "file_path": None
                    }
                except Exception as e:
                    return {
                        "status": "error",
                        "message": f"Error inesperado al convertir .doc a .docx: {str(e)}",
                        "filename": None,
                        "file_path": None
                    }
            
            elif file_extension != '.docx':
                return {
                    "status": "error",
                    "message": f"Formato de archivo no soportado: {file_extension}. Solo se soporta .docx o .doc",
                    "filename": None,
                    "file_path": None
                }
            
            # Crear nombre único para el archivo generado (PDF en lugar de DOCX)
            student_name = evaluation_data.get("full_name", "estudiante").replace(" ", "_")
            unique_filename_docx = f"evaluacion_salud_{student_name}_{uuid.uuid4().hex[:8]}.docx"
            unique_filename = f"evaluacion_salud_{student_name}_{uuid.uuid4().hex[:8]}.pdf"
            output_file_docx = Path(output_directory) / unique_filename_docx
            output_file = Path(output_directory) / unique_filename
            
            # Asegurar que el directorio de salida existe
            output_file.parent.mkdir(parents=True, exist_ok=True)
            
            # Abrir el documento template
            try:
                doc = Document(template_file)
            except Exception as e:
                return {
                    "status": "error",
                    "message": f"Error al abrir el template. Asegúrate de que es un archivo .docx válido: {str(e)}",
                    "filename": None,
                    "file_path": None
                }
            
            # Función auxiliar para formatear fechas
            def format_date(date_str: Optional[str]) -> str:
                if not date_str:
                    return ""
                try:
                    if isinstance(date_str, str):
                        date_obj = datetime.strptime(date_str, "%Y-%m-%d").date()
                        return date_obj.strftime("%d/%m/%Y")
                    return ""
                except:
                    return str(date_str) if date_str else ""
            
            # Función auxiliar para obtener valor o string vacío
            def get_value(key: str, default: str = "") -> str:
                value = evaluation_data.get(key)
                if value is None:
                    return default
                return str(value) if value else default
            
            # Mapeo de campos: diferentes formas de identificar cada campo
            # Cada campo puede tener múltiples identificadores
            field_mapping = {
                "full_name": {
                    "value": get_value("full_name"),
                    "placeholders": ["[FULL_NAME]", "{full_name}", "[NOMBRE_COMPLETO]", "{nombre_completo}", "FULL_NAME", "full_name"]
                },
                "identification_number": {
                    "value": get_value("identification_number"),
                    "placeholders": ["[IDENTIFICATION_NUMBER]", "{identification_number}", "[RUT]", "{rut}", "[IDENTIFICACION]", "{identificacion}"]
                },
                "born_date": {
                    "value": format_date(get_value("born_date")),
                    "placeholders": ["[BORN_DATE]", "{born_date}", "[FECHA_NACIMIENTO]", "{fecha_nacimiento}", "[FECHA_NAC]", "{fecha_nac}"]
                },
                "age": {
                    "value": get_value("age"),
                    "placeholders": ["[AGE]", "{age}", "[EDAD]", "{edad}"]
                },
                "native_language": {
                    "value": get_value("native_language"),
                    "placeholders": ["[NATIVE_LANGUAGE]", "{native_language}", "[LENGUA_NATIVA]", "{lengua_nativa}", "[LENGUA_FAMILIA]", "{lengua_familia}"]
                },
                "language_usually_used": {
                    "value": get_value("language_usually_used"),
                    "placeholders": ["[LANGUAGE_USUALLY_USED]", "{language_usually_used}", "[LENGUA_HABITUAL]", "{lengua_habitual}", "[LENGUA_USO]", "{lengua_uso}"]
                },
                "consultation_reason_detail": {
                    "value": get_value("consultation_reason_detail"),
                    "placeholders": ["[CONSULTATION_REASON_DETAIL]", "{consultation_reason_detail}", "[MOTIVO_CONSULTA]", "{motivo_consulta}", "[DETALLE_MOTIVO]", "{detalle_motivo}"]
                },
                "professional_identification_number": {
                    "value": get_value("professional_identification_number"),
                    "placeholders": ["[PROFESSIONAL_IDENTIFICATION_NUMBER]", "{professional_identification_number}", "[RUT_PROFESIONAL]", "{rut_profesional}", "[RUT_MEDICO]", "{rut_medico}"]
                },
                "professional_registration_number": {
                    "value": get_value("professional_registration_number"),
                    "placeholders": ["[PROFESSIONAL_REGISTRATION_NUMBER]", "{professional_registration_number}", "[REGISTRO_PROFESIONAL]", "{registro_profesional}", "[N_REGISTRO]", "{n_registro}"]
                },
                "professional_specialty": {
                    "value": get_value("professional_specialty"),
                    "placeholders": ["[PROFESSIONAL_SPECIALTY]", "{professional_specialty}", "[ESPECIALIDAD]", "{especialidad}", "[ESPECIALIDAD_PROFESIONAL]", "{especialidad_profesional}"]
                },
                "procedence_other": {
                    "value": get_value("procedence_other"),
                    "placeholders": ["[PROCEDENCE_OTHER]", "{procedence_other}", "[PROCEDENCIA_OTRO]", "{procedencia_otro}", "[OTRO_PROCEDENCIA]", "{otro_procedencia}"]
                },
                "professional_contact": {
                    "value": get_value("professional_contact"),
                    "placeholders": ["[PROFESSIONAL_CONTACT]", "{professional_contact}", "[CONTACTO_PROFESIONAL]", "{contacto_profesional}", "[FONO_EMAIL]", "{fono_email}"]
                },
                "evaluation_date": {
                    "value": format_date(get_value("evaluation_date")),
                    "placeholders": ["[EVALUATION_DATE]", "{evaluation_date}", "[FECHA_EVALUACION]", "{fecha_evaluacion}", "[FECHA_EVAL]", "{fecha_eval}"]
                },
                "reevaluation_date": {
                    "value": format_date(get_value("reevaluation_date")),
                    "placeholders": ["[REEVALUATION_DATE]", "{reevaluation_date}", "[FECHA_REEVALUACION]", "{fecha_reevaluacion}", "[FECHA_REVAL]", "{fecha_reval}"]
                },
                "general_assessment": {
                    "value": get_value("general_assessment"),
                    "placeholders": ["[GENERAL_ASSESSMENT]", "{general_assessment}", "[VALORACION_GENERAL]", "{valoracion_general}", "[VALORACION]", "{valoracion}"]
                },
                "diagnosis": {
                    "value": get_value("diagnosis"),
                    "placeholders": ["[DIAGNOSIS]", "{diagnosis}", "[DIAGNOSTICO]", "{diagnostico}"]
                },
                "indications": {
                    "value": get_value("indications"),
                    "placeholders": ["[INDICATIONS]", "{indications}", "[INDICACIONES]", "{indicaciones}"]
                }
            }
            
            # Función para reemplazar todos los placeholders en un texto
            def replace_placeholders(text: str) -> str:
                result = text
                for field_data in field_mapping.values():
                    for placeholder in field_data["placeholders"]:
                        result = result.replace(placeholder, field_data["value"])
                return result
            
            # Reemplazar texto en párrafos
            for paragraph in doc.paragraphs:
                original_text = paragraph.text
                new_text = replace_placeholders(original_text)
                if original_text != new_text:
                    # Limpiar el párrafo y agregar el nuevo texto
                    paragraph.clear()
                    paragraph.add_run(new_text)
            
            # Reemplazar texto en tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            original_text = paragraph.text
                            new_text = replace_placeholders(original_text)
                            if original_text != new_text:
                                # Limpiar el párrafo y agregar el nuevo texto
                                paragraph.clear()
                                paragraph.add_run(new_text)
            
            # Intentar reemplazar campos de formulario (form fields)
            # Esto requiere acceso al XML del documento
            try:
                if not DOCX_EXTRA_AVAILABLE:
                    raise ImportError("docx.oxml no disponible")
                
                # Buscar campos de formulario en el documento
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        # Buscar campos de formulario por nombre
                        for field_name, field_data in field_mapping.items():
                            # Los campos de formulario pueden tener diferentes nombres
                            # Intentar buscar por el nombre del campo o variaciones
                            pass  # python-docx tiene limitaciones con campos de formulario
            except:
                # Si no se pueden procesar campos de formulario, continuar con placeholders
                pass
            
            # Guardar el documento Word generado temporalmente
            doc.save(str(output_file_docx))
            
            # Limpiar archivo temporal si se convirtió de .doc
            if converted_path and Path(converted_path).exists():
                try:
                    os.unlink(converted_path)
                except:
                    pass
            
            # Convertir el Word a PDF usando LibreOffice
            
            is_windows = platform.system() == "Windows"
            libreoffice_cmd = None
            
            # Buscar LibreOffice
            if is_windows:
                possible_paths = [
                    "soffice.exe",
                    "soffice",
                    r"C:\Program Files\LibreOffice\program\soffice.exe",
                    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                    r"C:\Program Files\LibreOffice 7\program\soffice.exe",
                    r"C:\Program Files\LibreOffice 6\program\soffice.exe",
                ]
                appdata = os.getenv('LOCALAPPDATA', '')
                if appdata:
                    possible_paths.append(os.path.join(appdata, r"Programs\LibreOffice\program\soffice.exe"))
            else:
                possible_paths = ["libreoffice", "soffice", "/usr/bin/libreoffice", "/usr/bin/soffice"]
            
            # Buscar el comando disponible
            for cmd in possible_paths:
                try:
                    result = subprocess.run(
                        [cmd, "--version"],
                        capture_output=True,
                        timeout=5
                    )
                    if result.returncode == 0:
                        libreoffice_cmd = cmd
                        break
                except (FileNotFoundError, subprocess.TimeoutExpired):
                    continue
            
            if not libreoffice_cmd:
                # Si no hay LibreOffice, intentar usar Word para convertir a PDF (solo Windows)
                if is_windows:
                    try:
                        if not WIN32_AVAILABLE:
                            raise ImportError("win32com no disponible")
                        
                        # Abrir Word y convertir a PDF
                        word = win32com.client.Dispatch("Word.Application")
                        word.Visible = False
                        doc = word.Documents.Open(str(output_file_docx.absolute()))
                        # 17 = wdFormatPDF
                        doc.SaveAs2(str(output_file.absolute()), FileFormat=17)
                        doc.Close()
                        word.Quit()
                        
                        # Eliminar el archivo Word temporal
                        try:
                            if output_file_docx.exists():
                                os.unlink(output_file_docx)
                        except:
                            pass
                        
                        if output_file.exists():
                            return {
                                "status": "success",
                                "message": "PDF generado exitosamente usando Microsoft Word",
                                "filename": unique_filename,
                                "file_path": str(output_file)
                            }
                        else:
                            return {
                                "status": "error",
                                "message": "No se pudo crear el archivo PDF con Word. Verifica los permisos del directorio.",
                                "filename": None,
                                "file_path": None
                            }
                    except ImportError:
                        # pywin32 no está disponible
                        try:
                            if output_file_docx.exists():
                                os.unlink(output_file_docx)
                        except:
                            pass
                        return {
                            "status": "error",
                            "message": "LibreOffice no está instalado y pywin32 no está disponible. Se requiere LibreOffice o Microsoft Word (con pywin32) para convertir Word a PDF. Instala LibreOffice (recomendado) o pywin32: pip install pywin32",
                            "filename": None,
                            "file_path": None
                        }
                    except Exception as e:
                        # Error al usar Word
                        error_str = str(e)
                        try:
                            if output_file_docx.exists():
                                os.unlink(output_file_docx)
                        except:
                            pass
                        
                        # Detectar si el error es porque Word no está instalado
                        if "Word.Application" in error_str or "COM" in error_str or "0x80040154" in error_str or "Invalid class string" in error_str:
                            return {
                                "status": "error",
                                "message": "Microsoft Word no está instalado. Para convertir Word a PDF necesitas: 1) Instalar LibreOffice (gratis, recomendado): https://www.libreoffice.org/download/, 2) Instalar Microsoft Word, 3) O usar un servicio de conversión externo.",
                                "filename": None,
                                "file_path": None
                            }
                        else:
                            return {
                                "status": "error",
                                "message": f"Error al convertir Word a PDF usando Microsoft Word: {error_str}. Asegúrate de tener Word instalado o instala LibreOffice (gratis): https://www.libreoffice.org/download/",
                                "filename": None,
                                "file_path": None
                            }
                else:
                    # Linux sin LibreOffice
                    try:
                        if output_file_docx.exists():
                            os.unlink(output_file_docx)
                    except:
                        pass
                    return {
                        "status": "error",
                        "message": "LibreOffice no está instalado. Para convertir Word a PDF en Ubuntu/Linux, instala LibreOffice: sudo apt-get update && sudo apt-get install -y libreoffice",
                        "filename": None,
                        "file_path": None
                    }
            
            # Convertir Word a PDF
            try:
                result = subprocess.run([
                    libreoffice_cmd, "--headless", "--convert-to", "pdf",
                    "--outdir", str(output_file.parent),
                    str(output_file_docx.absolute())
                ], capture_output=True, timeout=30)
                
                if result.returncode == 0:
                    # LibreOffice crea el PDF con el mismo nombre base
                    pdf_file = output_file_docx.with_suffix('.pdf')
                    if pdf_file.exists():
                        # Mover/renombrar al nombre final deseado
                        pdf_file.rename(output_file)
                        
                        # Eliminar el archivo Word temporal
                        try:
                            if output_file_docx.exists():
                                os.unlink(output_file_docx)
                        except:
                            pass
                        
                        return {
                            "status": "success",
                            "message": "PDF generado exitosamente",
                            "filename": unique_filename,
                            "file_path": str(output_file)
                        }
                    else:
                        return {
                            "status": "error",
                            "message": "No se pudo crear el archivo PDF. Verifica los permisos del directorio.",
                            "filename": None,
                            "file_path": None
                        }
                else:
                    error_msg = result.stderr.decode('utf-8', errors='ignore') if result.stderr else 'Error desconocido'
                    # Limpiar archivo Word temporal
                    try:
                        if output_file_docx.exists():
                            os.unlink(output_file_docx)
                    except:
                        pass
                    return {
                        "status": "error",
                        "message": f"Error al convertir Word a PDF: {error_msg}",
                        "filename": None,
                        "file_path": None
                    }
            except subprocess.TimeoutExpired:
                # Limpiar archivo Word temporal
                try:
                    if output_file_docx.exists():
                        os.unlink(output_file_docx)
                except:
                    pass
                return {
                    "status": "error",
                    "message": "Timeout al convertir Word a PDF. El proceso tardó demasiado.",
                    "filename": None,
                    "file_path": None
                }
            except Exception as e:
                # Limpiar archivo Word temporal
                try:
                    if output_file_docx.exists():
                        os.unlink(output_file_docx)
                except:
                    pass
                return {
                    "status": "error",
                    "message": f"Error inesperado al convertir Word a PDF: {str(e)}",
                    "filename": None,
                    "file_path": None
                }
            
        except Exception as e:
            # Limpiar archivo temporal si existe
            if converted_path and Path(converted_path).exists():
                try:
                    os.unlink(converted_path)
                except:
                    pass
            
            return {
                "status": "error",
                "message": f"Error generando documento: {str(e)}",
                "filename": None,
                "file_path": None
            }

    @staticmethod
    def inspect_pdf_form_fields(template_path: str) -> Dict[str, Any]:
        """
        Inspecciona un PDF para identificar qué campos de formulario contiene.
        Útil para entender cómo están estructurados los campos en el template PDF.
        
        Args:
            template_path: Ruta del archivo PDF template
        
        Returns:
            dict: Información sobre los campos de formulario encontrados
        """
        try:
            template_file = Path(template_path)
            
            if not template_file.exists():
                return {
                    "status": "error",
                    "message": "Template PDF no encontrado",
                    "fields": []
                }
            
            pdf_document = fitz.open(template_file)
            
            form_fields = []
            text_placeholders = set()
            
            # Buscar campos de formulario en todas las páginas
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                
                # Obtener widgets (campos de formulario)
                widgets = page.widgets()
                for widget in widgets:
                    field_info = {
                        "field_name": widget.field_name,
                        "field_type": widget.field_type_string,
                        "field_value": widget.field_value if hasattr(widget, 'field_value') else None,
                        "rect": {
                            "x0": widget.rect.x0,
                            "y0": widget.rect.y0,
                            "x1": widget.rect.x1,
                            "y1": widget.rect.y1
                        },
                        "page": page_num + 1
                    }
                    form_fields.append(field_info)
                
                # También buscar placeholders de texto
                text = page.get_text()
                placeholders_brackets = re.findall(r'\[([^\]]+)\]', text)
                placeholders_braces = re.findall(r'\{([^\}]+)\}', text)
                
                for p in placeholders_brackets:
                    text_placeholders.add(f"[{p}]")
                for p in placeholders_braces:
                    text_placeholders.add(f"{{{p}}}")
            
            pdf_document.close()
            
            return {
                "status": "success",
                "message": "PDF inspeccionado exitosamente",
                "form_fields": form_fields,
                "text_placeholders": list(text_placeholders),
                "total_form_fields": len(form_fields)
            }
            
        except Exception as e:
            return {
                "status": "error",
                "message": f"Error inspeccionando PDF: {str(e)}",
                "fields": []
            }

    @staticmethod
    def inspect_template_fields(template_path: str) -> Dict[str, Any]:
        """
        Inspecciona un template Word para identificar qué campos contiene.
        Útil para entender cómo están estructurados los campos en el template.
        
        Args:
            template_path: Ruta del archivo Word template
        
        Returns:
            dict: Información sobre los campos encontrados en el template
        """
        try:
            template_file = Path(template_path)
            
            if not template_file.exists():
                return {
                    "status": "error",
                    "message": "Template no encontrado",
                    "fields": []
                }
            
            doc = Document(template_file)
            
            found_placeholders = set()
            found_text = []
            
            # Buscar placeholders en párrafos
            for paragraph in doc.paragraphs:
                text = paragraph.text
                # Buscar placeholders con formato [CAMPO] o {campo}
                placeholders_brackets = re.findall(r'\[([^\]]+)\]', text)
                placeholders_braces = re.findall(r'\{([^\}]+)\}', text)
                
                for p in placeholders_brackets:
                    found_placeholders.add(f"[{p}]")
                for p in placeholders_braces:
                    found_placeholders.add(f"{{{p}}}")
                
                if text.strip():
                    found_text.append(text[:100])  # Primeros 100 caracteres
            
            # Buscar en tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text = paragraph.text
                            placeholders_brackets = re.findall(r'\[([^\]]+)\]', text)
                            placeholders_braces = re.findall(r'\{([^\}]+)\}', text)
                            
                            for p in placeholders_brackets:
                                found_placeholders.add(f"[{p}]")
                            for p in placeholders_braces:
                                found_placeholders.add(f"{{{p}}}")
            
            # Intentar encontrar campos de formulario (form fields)
            form_fields = []
            try:
                # python-docx tiene limitaciones para acceder a campos de formulario
                # Necesitaríamos usar lxml directamente para esto
                pass
            except:
                pass
            
            return {
                "status": "success",
                "message": "Template inspeccionado exitosamente",
                "placeholders_found": list(found_placeholders),
                "form_fields_found": form_fields,
                "sample_text": found_text[:10]  # Primeros 10 textos como muestra
            }
            
        except Exception as e:
            return {
                "status": "error",
                "message": f"Error inspeccionando template: {str(e)}",
                "fields": []
            }

    @staticmethod
    def fill_docx_form(
        template_path: str,
        replacements: Dict[str, str],
        output_path: str,
        remove_literal_strings: Optional[List[str]] = None,
    ) -> Dict[str, Any]:
        """
        Rellena un formulario DOCX reemplazando etiquetas/placeholders.
        Soporta: {etiqueta}, [etiqueta], <<etiqueta>>, {{etiqueta}}, CONTENT CONTROLS por tag.
        Si remove_literal_strings está definido, se eliminan esas cadenas del texto (ej. placeholder "Haz clic o pulse aquí...").
        """
        try:
            doc = Document(template_path)

            def replace_in_text(text: str) -> str:
                result = text
                for tag, value in replacements.items():
                    val = str(value) if value is not None else ""
                    result = result.replace(f"{{{tag}}}", val)
                    result = result.replace(f"[{tag}]", val)
                    result = result.replace(f"<<{tag}>>", val)
                    result = result.replace(f"{{{{{tag}}}}}", val)
                return result

            def process_paragraph(para):
                full_text = para.text
                new_text = replace_in_text(full_text)
                if full_text != new_text:
                    para.clear()
                    para.add_run(new_text)

            for paragraph in doc.paragraphs:
                process_paragraph(paragraph)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            process_paragraph(para)

            for section in doc.sections:
                for header_footer in (section.header, section.first_page_header, section.footer, section.first_page_footer):
                    if header_footer is not None:
                        for para in header_footer.paragraphs:
                            process_paragraph(para)
                        if hasattr(header_footer, "tables"):
                            for tbl in header_footer.tables:
                                for row in tbl.rows:
                                    for cell in row.cells:
                                        for para in cell.paragraphs:
                                            process_paragraph(para)

            if DOCX_EXTRA_AVAILABLE:
                try:
                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement
                    W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"
                    W14_CHECKED_ATTR = f"{{{W14_NS}}}val"
                    CHK_SYMBOL_CHECKED = "\u2611"   # ☑
                    CHK_SYMBOL_UNCHECKED = "\u2610"  # ☐

                    def _set_checkbox_checked(sdt, sdtPr, checked: bool):
                        """Activa o desactiva un checkbox content control (w14:checkbox)."""
                        checkbox = None
                        for child in sdtPr:
                            if child.tag == f"{{{W14_NS}}}checkbox" or child.tag.endswith("}checkbox"):
                                checkbox = child
                                break
                        if checkbox is None:
                            return False
                        check_el = None
                        for child in checkbox:
                            if child.tag == f"{{{W14_NS}}}checked" or child.tag.endswith("}checked"):
                                check_el = child
                                break
                        val_str = "1" if checked else "0"
                        if check_el is not None:
                            check_el.set(W14_CHECKED_ATTR, val_str)
                        else:
                            from docx.oxml import parse_xml
                            check_el = parse_xml(
                                f'<w14:checked xmlns:w14="{W14_NS}" w14:val="{val_str}"/>'
                            )
                            checkbox.append(check_el)
                        # Actualizar también el texto visible en sdtContent (Word lo usa para mostrar)
                        symbol = CHK_SYMBOL_CHECKED if checked else CHK_SYMBOL_UNCHECKED
                        sdtContent = sdt.find(qn("w:sdtContent"))
                        if sdtContent is not None:
                            wt_list = list(sdtContent.iter(qn("w:t")))
                            if wt_list:
                                wt_list[0].text = symbol
                            else:
                                for p in sdtContent.iter(qn("w:p")):
                                    w_r = p.find(qn("w:r"))
                                    if w_r is not None:
                                        wt = w_r.find(qn("w:t"))
                                        if wt is not None:
                                            wt.text = symbol
                                        break
                        return True

                    def process_sdt_body(parent):
                        for sdt in parent.iter(qn("w:sdt")):
                            sdtPr = sdt.find(qn("w:sdtPr"))
                            if sdtPr is None:
                                continue
                            tag_el = sdtPr.find(qn("w:tag"))
                            if tag_el is None:
                                continue
                            tag_val = (tag_el.get(qn("w:val")) or "").strip()
                            if not tag_val:
                                continue
                            # Coincidencia insensible a mayúsculas para que Document_1 = document_1
                            rkey = next((k for k in replacements if k.strip().lower() == tag_val.strip().lower()), None)
                            if rkey is None:
                                continue
                            val = str(replacements.get(rkey, "") or "").strip()
                            is_checked = bool(val and val not in ("0", "false", "no", "off"))

                            # Si es checkbox content control (w14:checkbox), activar/desactivar
                            if _set_checkbox_checked(sdt, sdtPr, is_checked):
                                continue

                            # Campos de texto: si valor vacío, desenvolver el control para que no salga formulario editable
                            sdtContent = sdt.find(qn("w:sdtContent"))
                            if sdtContent is not None and not val:
                                for wt in sdtContent.iter(qn("w:t")):
                                    wt.text = ""
                                parent = sdt.getparent()
                                if parent is not None:
                                    idx = parent.index(sdt)
                                    children = list(sdtContent)
                                    parent.remove(sdt)
                                    for i, child in enumerate(children):
                                        parent.insert(idx + i, child)
                                continue

                            # Reemplazo de texto normal cuando hay valor
                            if sdtContent is not None:
                                wt_list = list(sdtContent.iter(qn("w:t")))
                                if wt_list:
                                    for i, wt in enumerate(wt_list):
                                        wt.text = val if i == 0 else ""
                                else:
                                    w_r = sdtContent.find(qn("w:r"))
                                    if w_r is not None:
                                        wt_el = w_r.find(qn("w:t"))
                                        if wt_el is not None:
                                            wt_el.text = val
                                        else:
                                            new_t = OxmlElement("w:t")
                                            if val:
                                                new_t.text = val
                                            new_t.set(qn("xml:space"), "preserve")
                                            w_r.append(new_t)
                                    else:
                                        for p in sdtContent.iter(qn("w:p")):
                                            w_r = p.find(qn("w:r"))
                                            if w_r is not None:
                                                wt_el = w_r.find(qn("w:t"))
                                                if wt_el is not None:
                                                    wt_el.text = val
                                                else:
                                                    new_t = OxmlElement("w:t")
                                                    if val:
                                                        new_t.text = val
                                                    new_t.set(qn("xml:space"), "preserve")
                                                    w_r.append(new_t)
                                                break
                    process_sdt_body(doc.element.body)
                    for section in doc.sections:
                        for hf in (section.header, section.footer):
                            if hf is not None and hasattr(hf, "_element"):
                                process_sdt_body(hf._element)
                except Exception:
                    pass

            # Eliminar cadenas literales (ej. placeholder) cuando Procedencia no es Otro
            if remove_literal_strings:
                def _strip_literal(para):
                    text = para.text
                    new_text = text
                    for s in remove_literal_strings:
                        if s:
                            new_text = new_text.replace(s, "")
                    if new_text != text:
                        para.clear()
                        if new_text:
                            para.add_run(new_text)
                for para in doc.paragraphs:
                    _strip_literal(para)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                _strip_literal(para)
                for section in doc.sections:
                    for hf in (section.header, section.first_page_header, section.footer, section.first_page_footer):
                        if hf is not None:
                            for para in hf.paragraphs:
                                _strip_literal(para)
                            if hasattr(hf, "tables"):
                                for tbl in hf.tables:
                                    for row in tbl.rows:
                                        for cell in row.cells:
                                            for para in cell.paragraphs:
                                                _strip_literal(para)

            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_path)
            return {
                "status": "success",
                "message": "Documento DOCX rellenado correctamente",
                "file_path": output_path,
                "filename": Path(output_path).name,
            }
        except Exception as e:
            return {
                "status": "error",
                "message": str(e),
                "file_path": None,
                "filename": None,
            }
