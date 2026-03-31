"""
Clonación OOXML del bloque de registro por asignatura (cuadro observaciones + tabla rarpf_*).
El párrafo largo «Registro de acciones realizadas por el profesor…» no se duplica; solo una vez en la plantilla.
"""

from __future__ import annotations

import logging
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


def _empty_paragraph_spacing_before(twips_before: int = 480) -> OxmlElement:
    """Párrafo vacío con espacio superior (twips = 1/20 pt; 480 ≈ 24 pt) para separar bloques clonados."""
    p = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(twips_before))
    p_pr.append(spacing)
    p.append(p_pr)
    return p

_HEADER_SNIPPET = "Registro de acciones realizadas por el profesor"


def _get_sdt_tag_val(sdt) -> str:
    sdt_pr = sdt.find(qn("w:sdtPr"))
    if sdt_pr is None:
        return ""
    tag_el = sdt_pr.find(qn("w:tag"))
    if tag_el is None:
        return ""
    return (tag_el.get(qn("w:val")) or "").strip()


def _set_sdt_tag_and_alias(sdt, new_val: str) -> None:
    sdt_pr = sdt.find(qn("w:sdtPr"))
    if sdt_pr is None:
        return
    tag_el = sdt_pr.find(qn("w:tag"))
    if tag_el is not None:
        tag_el.set(qn("w:val"), new_val)
    alias_el = sdt_pr.find(qn("w:alias"))
    if alias_el is not None:
        alias_el.set(qn("w:val"), new_val)


def _strip_sdt_id(sdt) -> None:
    pr = sdt.find(qn("w:sdtPr"))
    if pr is None:
        return
    for el in list(pr):
        if el.tag == qn("w:id") or (el.tag.endswith("}id") and "id" in el.tag.lower()):
            pr.remove(el)


def _map_tag_for_block_clone(old: str, block_index: int, rows_per: int) -> str | None:
    """
    Prototipo = bloque 0 (rarpo_1, rarpf_1..). Clon k (1..N-1) -> rarpo_{k+1}, rarpf_{k*rows + i}.
    """
    if block_index < 1:
        return None
    ol = (old or "").strip()
    if not ol:
        return None
    low = ol.lower()
    if low.startswith("rarpo_"):
        return f"rarpo_{block_index + 1}"
    if low.startswith("rarpas_"):
        return f"rarpas_{block_index + 1}"
    for prefix in ("rarpf_", "rarphp_", "rarpad_", "rarpnfd_"):
        if low.startswith(prefix):
            try:
                n = int(ol.split("_")[-1])
            except ValueError:
                return None
            return f"{prefix}{block_index * rows_per + n}"
    return None


def _retag_sdts_in_element(root, block_index: int, rows_per: int) -> None:
    for sdt in root.iter(qn("w:sdt")):
        old = _get_sdt_tag_val(sdt)
        new = _map_tag_for_block_clone(old, block_index, rows_per)
        if new:
            _set_sdt_tag_and_alias(sdt, new)
            _strip_sdt_id(sdt)


def _find_clone_fragment_range(children: list) -> tuple[int | None, int | None]:
    """
    Rango inclusivo [start, end] de lo que se CLONA por cada asignatura extra.

    No incluye el párrafo largo «Registro de acciones realizadas por el profesor…»
    (b)/c)); solo lo que va después hasta la tabla con rarpf_1 (p. ej. párrafo vacío + tabla).
    """
    tbl_idx = None
    for i, child in enumerate(children):
        if child.tag != qn("w:tbl"):
            continue
        for sdt in child.iter(qn("w:sdt")):
            if _get_sdt_tag_val(sdt) == "rarpf_1":
                tbl_idx = i
                break
        if tbl_idx is not None:
            break
    if tbl_idx is None:
        return None, None
    header_idx = None
    for j in range(tbl_idx - 1, -1, -1):
        if children[j].tag != qn("w:p"):
            break
        full = "".join((t.text or "") for t in children[j].iter(qn("w:t")))
        if _HEADER_SNIPPET in full:
            header_idx = j
            break
    if header_idx is not None:
        start = header_idx + 1
    else:
        start = max(0, tbl_idx - 2)
    end = tbl_idx
    if start > end:
        start = end
    return start, end


def clone_register_book_section_b_blocks(
    docx_path: str | Path,
    n_blocks: int,
    rows_per_register: int = 11,
) -> bool:
    """
    Duplica el bloque completo (texto b) + tabla con controles rarpo/rarp*) N veces.

    El prototipo en plantilla debe incluir la tabla donde está `rarpf_1` y el párrafo
    previo con «Registro de acciones realizadas por el profesor…».

    :param n_blocks: Número de bloques finales deseados (p. ej. una asignatura por bloque).
    """
    path = Path(docx_path)
    if n_blocks <= 1:
        return True
    try:
        doc = Document(str(path))
    except Exception as e:
        logger.warning("clone_section_b: no se abrió %s: %s", path, e)
        return False

    body = doc.element.body
    children = list(body)
    start, end = _find_clone_fragment_range(children)
    if start is None or end is None:
        logger.warning("clone_section_b: no se encontró tabla con rarpf_1 / encabezado II b)")
        return False

    fragment = [children[i] for i in range(start, end + 1)]
    insert_after = children[end]
    parent = insert_after.getparent()
    if parent is None:
        return False

    idx = parent.index(insert_after) + 1
    for copy_i in range(1, n_blocks):
        new_els = [deepcopy(el) for el in fragment]
        for el in new_els:
            _retag_sdts_in_element(el, copy_i, rows_per_register)
        for j, el in enumerate(new_els):
            parent.insert(idx + j, el)
        insert_after = new_els[-1]
        idx = parent.index(insert_after) + 1

    try:
        doc.save(str(path))
    except Exception as e:
        logger.warning("clone_section_b: save falló %s: %s", path, e)
        return False
    return True


def clone_rarpo_observation_blocks_if_needed(docx_path: str | Path, n_blocks: int) -> bool:
    """Compatibilidad: usar clone_register_book_section_b_blocks."""
    return clone_register_book_section_b_blocks(docx_path, n_blocks, rows_per_register=11)


def _map_pai_table_clone_tag(old: str, block_index: int) -> str | None:
    """Prototipo bloque 0: paih_1..5, paie_1, paio_1, aer_1, paifi_*, paift_*. Clon block_index>=1."""
    if block_index < 1:
        return None
    ol = (old or "").strip()
    if not ol:
        return None
    low = ol.lower()
    for prefix in ("paih_", "paifi_", "paift_"):
        if low.startswith(prefix):
            try:
                n = int(ol.split("_")[-1])
            except ValueError:
                return None
            return f"{prefix}{block_index * 5 + n}"
    if low == "paie_1":
        return f"paie_{block_index + 1}"
    if low == "paio_1":
        return f"paio_{block_index + 1}"
    if low == "aer_1":
        return f"aer_{block_index + 1}"
    return None


def _retag_pai_sdts_in_element(root, block_index: int) -> None:
    for sdt in root.iter(qn("w:sdt")):
        old = _get_sdt_tag_val(sdt)
        new = _map_pai_table_clone_tag(old, block_index)
        if new:
            _set_sdt_tag_and_alias(sdt, new)
            _strip_sdt_id(sdt)


def _find_pai_table_index(children: list) -> int | None:
    for i, child in enumerate(children):
        if child.tag != qn("w:tbl"):
            continue
        for sdt in child.iter(qn("w:sdt")):
            if _get_sdt_tag_val(sdt) == "paih_1":
                return i
    return None


def clone_pai_support_tables(docx_path: str | Path, n_blocks: int) -> bool:
    """
    Duplica solo la tabla del PAI (controles paih_1..5, paie_1, …), sin el párrafo
    «4. Plan de Apoyo Individual…».
    """
    path = Path(docx_path)
    if n_blocks <= 1:
        return True
    try:
        doc = Document(str(path))
    except Exception as e:
        logger.warning("clone_pai: no se abrió %s: %s", path, e)
        return False

    body = doc.element.body
    children = list(body)
    tbl_idx = _find_pai_table_index(children)
    if tbl_idx is None:
        logger.warning("clone_pai: no hay tabla con tag paih_1")
        return False

    fragment = [children[tbl_idx]]
    insert_after = children[tbl_idx]
    parent = insert_after.getparent()
    if parent is None:
        return False

    idx = parent.index(insert_after) + 1
    for copy_i in range(1, n_blocks):
        new_tbl = deepcopy(fragment[0])
        _retag_pai_sdts_in_element(new_tbl, copy_i)
        parent.insert(idx, new_tbl)
        insert_after = new_tbl
        idx = parent.index(insert_after) + 1

    try:
        doc.save(str(path))
    except Exception as e:
        logger.warning("clone_pai: save falló %s: %s", path, e)
        return False
    return True


def _map_raeg_table_clone_tag(old: str, block_index: int) -> str | None:
    """
    Prototipo bloque 1: raegee_1, raegeop_1_1, raegef_1_1, raegel_1_1, raegear_1_1, raegep_1_1 (profesional), …
    Clon k (>=1) → raegef_2_1, …  (raegeop antes que raegep en la tupla de prefijos).
    """
    if block_index < 1:
        return None
    ol = (old or "").strip()
    if not ol:
        return None
    new_b = block_index + 1
    low = ol.lower()
    prefixes = ("raegear", "raegee", "raegef", "raegel", "raegeoa", "raegeop", "raegep")
    for prefix in prefixes:
        pfx = prefix + "_"
        if not low.startswith(pfx):
            continue
        rest = ol[len(prefix) + 1 :]
        parts = rest.split("_")
        if len(parts) == 1 and parts[0] == "1":
            return f"{prefix}_{new_b}"
        if len(parts) >= 2 and parts[0] == "1" and parts[1].isdigit():
            tail = "_".join(parts[1:])
            return f"{prefix}_{new_b}_{tail}"
    return None


def _retag_raeg_sdts_in_element(root, block_index: int) -> None:
    for sdt in root.iter(qn("w:sdt")):
        old = _get_sdt_tag_val(sdt)
        new = _map_raeg_table_clone_tag(old, block_index)
        if new:
            _set_sdt_tag_and_alias(sdt, new)
            _strip_sdt_id(sdt)


def _find_raeg_table_index(children: list) -> int | None:
    """Plantilla: filas de datos usan raegef_1_1 / raegear_1_1 (no raegef_1)."""
    for i, child in enumerate(children):
        if child.tag != qn("w:tbl"):
            continue
        for sdt in child.iter(qn("w:sdt")):
            t = _get_sdt_tag_val(sdt)
            if t in ("raegear_1", "raegef_1", "raegear_1_1", "raegef_1_1", "raegep_1_1"):
                return i
    return None


def _find_raeg_fragment_start_index(children: list, tbl_idx: int) -> int:
    """
    Inicio del bloque «registro de apoyos» en el cuerpo: w:p o w:tbl con raegee_1 / raegeoa_1 / raegeop_1_1
    *antes* de la tabla de filas (raegef_1_1). Si solo se clona la tabla de datos, raegee_2 no existe.
    """
    header_tags = frozenset({"raegee_1", "raegeoa_1", "raegeop_1_1"})
    hits: list[int] = []
    for j in range(0, tbl_idx):
        ch = children[j]
        for sdt in ch.iter(qn("w:sdt")):
            if _get_sdt_tag_val(sdt) in header_tags:
                hits.append(j)
                break
    return min(hits) if hits else tbl_idx


def _find_tr_containing_sdt_tag(root, needle: str):
    for tbl in root.iter(qn("w:tbl")):
        for tr in tbl.iter(qn("w:tr")):
            for sdt in tr.iter(qn("w:sdt")):
                if _get_sdt_tag_val(sdt) == needle:
                    return tr
    return None


def _element_has_sdt_tag(root, needle: str) -> bool:
    for sdt in root.iter(qn("w:sdt")):
        if _get_sdt_tag_val(sdt) == needle:
            return True
    return False


def _find_activity_fragment_range(
    children: list,
    start_tag: str,
    end_tag: str,
) -> tuple[int | None, int | None]:
    """Rango de elementos del cuerpo entre el primer SDT start_tag y el último end_tag (inclusive)."""
    start = end = None
    for i, ch in enumerate(children):
        if _element_has_sdt_tag(ch, start_tag):
            start = i
            break
    if start is None:
        return None, None
    for i in range(start, len(children)):
        if _element_has_sdt_tag(children[i], end_tag):
            end = i
            break
    if end is None:
        return None, None
    return start, end


def _find_rafc_fragment_range(children: list) -> tuple[int | None, int | None]:
    """
    IV «Trabajo con la familia…»: párrafo Fecha (rafcf_1_1) + tablas participantes
    y campos a)–d) hasta rafcnir_1_1 (inclusive en el cuerpo del documento).
    """
    return _find_activity_fragment_range(children, "rafcf_1_1", "rafcnir_1_1")


def _find_tcee_fragment_range(children: list) -> tuple[int | None, int | None]:
    """IV «Trabajo con la comunidad…»: tceef_1_1 … tceer_1_1."""
    return _find_activity_fragment_range(children, "tceef_1_1", "tceer_1_1")


def _find_ar_fragment_range(children: list) -> tuple[int | None, int | None]:
    """V Acta de reuniones (section other): arf_1_1 … arc_1_1."""
    return _find_activity_fragment_range(children, "arf_1_1", "arc_1_1")


def _map_rafc_block_clone_tag(old: str, block_index: int) -> str | None:
    """Prototipo rafcf_1_1, …, rafcnir_1_1. Clon k>=1 → rafcf_2_1, …"""
    if block_index < 1:
        return None
    ol = (old or "").strip()
    if not ol:
        return None
    parts = ol.rsplit("_", 2)
    if len(parts) != 3 or parts[1] != "1" or parts[2] != "1":
        return None
    prefix = parts[0]
    new_b = block_index + 1
    return f"{prefix}_{new_b}_1"


def _retag_rafc_sdts_in_element(root, block_index: int) -> None:
    for sdt in root.iter(qn("w:sdt")):
        old = _get_sdt_tag_val(sdt)
        new = _map_rafc_block_clone_tag(old, block_index)
        if new:
            _set_sdt_tag_and_alias(sdt, new)
            _strip_sdt_id(sdt)


def _clone_iv_activity_blocks(
    docx_path: str | Path,
    n_blocks: int,
    find_fragment,
    log_prefix: str,
) -> bool:
    path = Path(docx_path)
    if n_blocks <= 1:
        return True
    try:
        doc = Document(str(path))
    except Exception as e:
        logger.warning("%s: no se abrió %s: %s", log_prefix, path, e)
        return False

    body = doc.element.body
    children = list(body)
    start_idx, end_idx = find_fragment(children)
    if start_idx is None or end_idx is None:
        logger.warning("%s: no se encontró el fragmento de plantilla", log_prefix)
        return False

    fragment = [children[i] for i in range(start_idx, end_idx + 1)]
    insert_after = children[end_idx]
    parent = insert_after.getparent()
    if parent is None:
        return False

    idx = parent.index(insert_after) + 1
    for copy_i in range(1, n_blocks):
        parent.insert(idx, _empty_paragraph_spacing_before())
        idx += 1
        new_els = [deepcopy(el) for el in fragment]
        for el in new_els:
            _retag_rafc_sdts_in_element(el, copy_i)
        for j, el in enumerate(new_els):
            parent.insert(idx + j, el)
        insert_after = new_els[-1]
        idx = parent.index(insert_after) + 1

    try:
        doc.save(str(path))
    except Exception as e:
        logger.warning("%s: save falló %s: %s", log_prefix, path, e)
        return False
    return True


def clone_course_activity_record_blocks(docx_path: str | Path, n_blocks: int) -> bool:
    """
    IV familia: duplica bloque (rafcf_*_1 … rafcnir_*_1).
    """
    return _clone_iv_activity_blocks(docx_path, n_blocks, _find_rafc_fragment_range, "clone_rafc")


def clone_course_community_activity_blocks(docx_path: str | Path, n_blocks: int) -> bool:
    """
    IV comunidad: duplica bloque (tceef_*_1 … tceer_*_1).
    """
    return _clone_iv_activity_blocks(docx_path, n_blocks, _find_tcee_fragment_range, "clone_tcee")


def clone_course_acta_reunion_blocks(docx_path: str | Path, n_blocks: int) -> bool:
    """
    V Acta de reuniones: duplica bloque (arf_*_1 … arc_*_1), datos section=other.
    """
    return _clone_iv_activity_blocks(docx_path, n_blocks, _find_ar_fragment_range, "clone_ar")


def _retag_raeg_duplicate_row_tag(old: str, block_1based: int, from_row: int, to_row: int) -> str | None:
    """raegef_1_1 → raegef_1_2; raegef_2_1 → raegef_2_2."""
    if not old:
        return None
    suffix = f"_{block_1based}_{from_row}"
    if old.endswith(suffix):
        return old[: -len(suffix)] + f"_{block_1based}_{to_row}"
    return None


def expand_raeg_intervention_rows(docx_path: str | Path, block_1based: int, n_rows: int) -> bool:
    """
    Duplica la fila de datos (tags raegef_*/raegel_*/raegear_*/raegep_* …) dentro de la tabla del bloque,
    para varias intervenciones: raegef_1_2, raegef_1_3, … o raegef_2_2, …
    """
    path = Path(docx_path)
    if n_rows <= 1:
        return True
    try:
        doc = Document(str(path))
    except Exception as e:
        logger.warning("expand_raeg: no se abrió %s: %s", path, e)
        return False

    needle = "raegef_1_1" if block_1based == 1 else f"raegef_{block_1based}_1"
    body = doc.element.body
    tr_el = _find_tr_containing_sdt_tag(body, needle)
    if tr_el is None:
        logger.warning("expand_raeg: fila no encontrada para tag %s", needle)
        return False

    tbl = tr_el.getparent()
    if tbl is None or tbl.tag != qn("w:tbl"):
        logger.warning("expand_raeg: padre de fila no es w:tbl")
        return False

    idx = tbl.index(tr_el)
    insert_at = idx + 1
    for r in range(2, n_rows + 1):
        new_tr = deepcopy(tr_el)
        for sdt in new_tr.iter(qn("w:sdt")):
            old = _get_sdt_tag_val(sdt)
            new_val = _retag_raeg_duplicate_row_tag(old, block_1based, 1, r)
            if new_val:
                _set_sdt_tag_and_alias(sdt, new_val)
                _strip_sdt_id(sdt)
        tbl.insert(insert_at, new_tr)
        insert_at += 1

    try:
        doc.save(str(path))
    except Exception as e:
        logger.warning("expand_raeg: save falló %s: %s", path, e)
        return False
    return True


def expand_participant_table_rows(
    docx_path: str | Path,
    block_1based: int,
    n_rows: int,
    name_tag_prefix: str,
    log_prefix: str,
) -> bool:
    """
    Tabla de participantes: una fila por asistente. name_tag_prefix = 'rafcne' | 'tceen'
    (columnas del mismo tr: rafcnia/rafcnit o tceeap/tceete).
    """
    path = Path(docx_path)
    if n_rows <= 1:
        return True
    try:
        doc = Document(str(path))
    except Exception as e:
        logger.warning("%s: no se abrió %s: %s", log_prefix, path, e)
        return False

    needle = f"{name_tag_prefix}_{block_1based}_1"
    body = doc.element.body
    tr_el = _find_tr_containing_sdt_tag(body, needle)
    if tr_el is None:
        logger.warning("%s: fila no encontrada para tag %s", log_prefix, needle)
        return False

    tbl = tr_el.getparent()
    if tbl is None or tbl.tag != qn("w:tbl"):
        logger.warning("%s: padre de fila no es w:tbl", log_prefix)
        return False

    idx = tbl.index(tr_el)
    insert_at = idx + 1
    for r in range(2, n_rows + 1):
        new_tr = deepcopy(tr_el)
        for sdt in new_tr.iter(qn("w:sdt")):
            old = _get_sdt_tag_val(sdt)
            new_val = _retag_raeg_duplicate_row_tag(old, block_1based, 1, r)
            if new_val:
                _set_sdt_tag_and_alias(sdt, new_val)
                _strip_sdt_id(sdt)
        tbl.insert(insert_at, new_tr)
        insert_at += 1

    try:
        doc.save(str(path))
    except Exception as e:
        logger.warning("%s: save falló %s: %s", log_prefix, path, e)
        return False
    return True


def expand_rafc_participant_rows(docx_path: str | Path, block_1based: int, n_rows: int) -> bool:
    """IV familia: rafcne_*, rafcnia_*, rafcnit_*."""
    return expand_participant_table_rows(docx_path, block_1based, n_rows, "rafcne", "expand_rafc_part")


def expand_tcee_participant_rows(docx_path: str | Path, block_1based: int, n_rows: int) -> bool:
    """IV comunidad: tceen_*, tceeap_*, tceete_*."""
    return expand_participant_table_rows(docx_path, block_1based, n_rows, "tceen", "expand_tcee_part")


def expand_arn_participant_rows(docx_path: str | Path, block_1based: int, n_rows: int) -> bool:
    """V acta: arn_*, arpa_*, arr_*, art_*."""
    return expand_participant_table_rows(docx_path, block_1based, n_rows, "arn", "expand_arn_part")


def _retag_rla_duplicate_row_tag(old: str, block_1based: int, from_row: int, to_row: int) -> str | None:
    """rlane_1_1 → rlane_1_2; rllr_1_1 → rllr_1_2; rlcs_1_1 → rlcs_1_2."""
    if not old:
        return None
    suffix = f"_{block_1based}_{from_row}"
    if old.endswith(suffix):
        return old[: -len(suffix)] + f"_{block_1based}_{to_row}"
    return None


def _ensure_paragraph_spacing_before(p_el, twips_before: int) -> None:
    """Espacio antes del párrafo (equivalente visual a margin-top sobre la fila de tabla)."""
    p_pr = p_el.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        p_el.insert(0, p_pr)
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:before"), str(twips_before))


def apply_learning_achievement_period_top_spacing(
    docx_path: str | Path,
    row_indices_1based: list[int],
    twips_before: int = 480,
) -> bool:
    """
    Añade espacio superior a la primera fila de cada nuevo período (sin filas vacías).
    row_indices_1based: índices de fila del tag rlane_1_k (p. ej. primera fila del 2do y 3er período).
    """
    if not row_indices_1based:
        return True
    path = Path(docx_path)
    try:
        doc = Document(str(path))
    except Exception as e:
        logger.warning("rla_spacing: no se abrió %s: %s", path, e)
        return False

    body = doc.element.body
    for r in row_indices_1based:
        needle = f"rlane_1_{r}"
        tr_el = _find_tr_containing_sdt_tag(body, needle)
        if tr_el is None:
            logger.warning("rla_spacing: fila no encontrada para %s", needle)
            continue
        first_p = None
        for tc in tr_el.findall(qn("w:tc")):
            for p in tc.iter(qn("w:p")):
                first_p = p
                break
            if first_p is not None:
                break
        if first_p is not None:
            _ensure_paragraph_spacing_before(first_p, twips_before)

    try:
        doc.save(str(path))
    except Exception as e:
        logger.warning("rla_spacing: save falló %s: %s", path, e)
        return False
    return True


def expand_learning_achievement_rows(docx_path: str | Path, block_1based: int, n_rows: int) -> bool:
    """
    Sección «3. Registro de logros de aprendizaje»: duplica la fila de la tabla
    (tags rlane_*, rllr_*, rlcs_*) para varias filas de datos: rlane_1_2, rllr_1_2, …
    """
    path = Path(docx_path)
    if n_rows <= 1:
        return True
    try:
        doc = Document(str(path))
    except Exception as e:
        logger.warning("expand_rla: no se abrió %s: %s", path, e)
        return False

    needle = "rlane_1_1" if block_1based == 1 else f"rlane_{block_1based}_1"
    body = doc.element.body
    tr_el = _find_tr_containing_sdt_tag(body, needle)
    if tr_el is None:
        logger.warning("expand_rla: fila no encontrada para tag %s", needle)
        return False

    tbl = tr_el.getparent()
    if tbl is None or tbl.tag != qn("w:tbl"):
        logger.warning("expand_rla: padre de fila no es w:tbl")
        return False

    idx = tbl.index(tr_el)
    insert_at = idx + 1
    for r in range(2, n_rows + 1):
        new_tr = deepcopy(tr_el)
        for sdt in new_tr.iter(qn("w:sdt")):
            old = _get_sdt_tag_val(sdt)
            new_val = _retag_rla_duplicate_row_tag(old, block_1based, 1, r)
            if new_val:
                _set_sdt_tag_and_alias(sdt, new_val)
                _strip_sdt_id(sdt)
        tbl.insert(insert_at, new_tr)
        insert_at += 1

    try:
        doc.save(str(path))
    except Exception as e:
        logger.warning("expand_rla: save falló %s: %s", path, e)
        return False
    return True


def clone_course_record_support_tables(docx_path: str | Path, n_blocks: int) -> bool:
    """
    Sección «2. Registro de apoyos…»: duplica el bloque completo (párrafos con nombre/objetivos + tabla),
    no solo la tabla, para que existan raegee_2, raegeop_2_1, etc.
    """
    path = Path(docx_path)
    if n_blocks <= 1:
        return True
    try:
        doc = Document(str(path))
    except Exception as e:
        logger.warning("clone_raeg: no se abrió %s: %s", path, e)
        return False

    body = doc.element.body
    children = list(body)
    tbl_idx = _find_raeg_table_index(children)
    if tbl_idx is None:
        logger.warning("clone_raeg: no hay tabla con tag raegear_1 / raegef_1")
        return False

    start_idx = _find_raeg_fragment_start_index(children, tbl_idx)
    fragment = [children[i] for i in range(start_idx, tbl_idx + 1)]
    insert_after = children[tbl_idx]
    parent = insert_after.getparent()
    if parent is None:
        return False

    idx = parent.index(insert_after) + 1
    for copy_i in range(1, n_blocks):
        parent.insert(idx, _empty_paragraph_spacing_before())
        idx += 1
        new_els = [deepcopy(el) for el in fragment]
        for el in new_els:
            _retag_raeg_sdts_in_element(el, copy_i)
        for j, el in enumerate(new_els):
            parent.insert(idx + j, el)
        insert_after = new_els[-1]
        idx = parent.index(insert_after) + 1

    try:
        doc.save(str(path))
    except Exception as e:
        logger.warning("clone_raeg: save falló %s: %s", path, e)
        return False
    return True
