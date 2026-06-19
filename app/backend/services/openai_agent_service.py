"""Integración OpenAI: archivos, Responses API y code interpreter."""

from __future__ import annotations

import logging
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any

from sqlalchemy.orm import Session

from app.backend.core.config import settings
from app.backend.db.models import AgentFileModel, AgentModel
from app.backend.utils.agent_document_index import role_instructions_to_text, strip_html
from app.backend.utils.agent_familia_template import (
    build_familia_form_rules,
    build_familia_narrative_enrichment_rules,
    build_form_template_supremacy_block,
    build_redaction_min_paragraphs_rules,
    filter_familia_template_file_names,
    is_familia_form_file,
    is_familia_tabla_file,
    pick_familia_base_template,
    resolve_familia_template_from_rows,
)
from app.backend.utils.agent_student_lookup import format_student_context_block
from app.backend.utils.agent_files import agent_dir

logger = logging.getLogger(__name__)

CONTAINER_TTL = timedelta(minutes=18)


class OpenAIUploadError(RuntimeError):
    """Fallo al subir un archivo a OpenAI."""


def _utcnow() -> datetime:
    return datetime.utcnow()


def is_container_expired(updated_at: datetime | None) -> bool:
    if not updated_at:
        return True
    return _utcnow() - updated_at > CONTAINER_TTL


def get_openai_client():
    if not settings.openai_api_key:
        raise RuntimeError("OPENAI_API_KEY no configurada")
    from openai import OpenAI

    return OpenAI(api_key=settings.openai_api_key)


def upload_bytes_to_openai(content: bytes, display_name: str) -> str:
    from io import BytesIO

    client = get_openai_client()
    handle = BytesIO(content)
    uploaded = client.files.create(file=(display_name, handle), purpose="assistants")
    return uploaded.id


def require_openai_file_upload(content: bytes, display_name: str) -> str:
    if not settings.openai_api_key:
        raise OpenAIUploadError("OPENAI_API_KEY no configurada")
    try:
        leaf_name = display_name.replace("\\", "/").split("/")[-1] or display_name
        return upload_bytes_to_openai(content, leaf_name)
    except OpenAIUploadError:
        raise
    except Exception as exc:
        raise OpenAIUploadError(str(exc)) from exc


def upload_local_file_to_openai(file_path: Path, display_name: str | None = None) -> str:
    client = get_openai_client()
    upload_name = display_name or file_path.name
    with file_path.open("rb") as handle:
        uploaded = client.files.create(file=(upload_name, handle), purpose="assistants")
    return uploaded.id


def delete_openai_file(file_id: str | None) -> None:
    if not file_id or not settings.openai_api_key:
        return
    try:
        client = get_openai_client()
        client.files.delete(file_id)
    except Exception as exc:
        logger.warning("No se pudo eliminar archivo OpenAI %s: %s", file_id, exc)


def attach_file_to_container(container_id: str, file_id: str) -> None:
    client = get_openai_client()
    client.containers.files.create(container_id, file_id=file_id)


def ensure_openai_files_in_container(agent: AgentModel, file_ids: list[str]) -> None:
    """Adjunta archivos al contenedor reutilizado (p. ej. base PIE360 recién generada)."""
    if not file_ids or not agent.openai_container_id:
        return
    if is_container_expired(agent.openai_container_updated_at):
        return
    for file_id in file_ids:
        try:
            attach_file_to_container(agent.openai_container_id, file_id)
        except Exception as exc:
            logger.warning(
                "No se pudo adjuntar %s al contenedor %s: %s",
                file_id,
                agent.openai_container_id,
                exc,
            )


def clear_agent_container(agent: AgentModel) -> None:
    agent.openai_container_id = None
    agent.openai_container_updated_at = None


def ensure_openai_file_for_row(
    db: Session,
    agent: AgentModel,
    row: AgentFileModel,
    disk_path: Path,
) -> str | None:
    if row.openai_file_id:
        if agent.openai_container_id and not is_container_expired(agent.openai_container_updated_at):
            try:
                attach_file_to_container(agent.openai_container_id, row.openai_file_id)
            except Exception as exc:
                logger.warning("Contenedor expirado o inválido, se creará uno nuevo: %s", exc)
                clear_agent_container(agent)
        return row.openai_file_id

    if not settings.openai_api_key:
        raise OpenAIUploadError("OPENAI_API_KEY no configurada")

    try:
        file_id = upload_local_file_to_openai(disk_path, row.display_name)
        row.openai_file_id = file_id
        row.openai_upload_error = None
        db.flush()

        if agent.openai_container_id and not is_container_expired(agent.openai_container_updated_at):
            try:
                attach_file_to_container(agent.openai_container_id, file_id)
            except Exception as exc:
                logger.warning("No se pudo adjuntar archivo al contenedor existente: %s", exc)
                clear_agent_container(agent)
        return file_id
    except Exception as exc:
        row.openai_upload_error = str(exc)[:500]
        logger.exception("Error subiendo archivo a OpenAI: %s", row.id)
        return None


def sync_agent_openai_files(
    db: Session,
    agent_id: str,
    *,
    only_rows: list[AgentFileModel] | None = None,
) -> list[str]:
    agent = db.query(AgentModel).filter(AgentModel.id == agent_id).first()
    if not agent:
        return []

    if only_rows is not None:
        rows = only_rows
    else:
        rows = (
            db.query(AgentFileModel)
            .filter(AgentFileModel.agent_id == agent_id)
            .order_by(AgentFileModel.uploaded_at.asc())
            .all()
        )
    file_ids: list[str] = []
    for row in rows:
        disk_path = agent_dir(agent_id) / row.id
        if not disk_path.is_file():
            continue
        openai_id = ensure_openai_file_for_row(db, agent, row, disk_path)
        if openai_id:
            file_ids.append(openai_id)
    return file_ids


def _extract_container_id(response: Any, agent: AgentModel | None = None) -> str | None:
    from app.backend.services.agent_response_files_service import _extract_container_id_from_response

    container_id = _extract_container_id_from_response(response)
    if container_id:
        return container_id
    if agent and agent.openai_container_id and not is_container_expired(agent.openai_container_updated_at):
        return agent.openai_container_id
    return None


def _agent_uploaded_file_names(db: Session, agent_id: str) -> list[str]:
    rows = (
        db.query(AgentFileModel.display_name)
        .filter(AgentFileModel.agent_id == agent_id)
        .order_by(AgentFileModel.uploaded_at.asc())
        .all()
    )
    return [row[0] for row in rows if row[0]]


def _build_familia_tabla_rules() -> str:
    return (
        "- Para informe a la familia (solo si NO hay plantilla tipo formulario en archivos): "
        "abre FORMATO INFORME DE FAMILIA.docx como plantilla base y solo completa celdas.\n"
        "- TIPOGRAFÍA DE PLANTILLA (OBLIGATORIO): copia fuente, tamaño y alineación del run "
        "existente en cada párrafo (típicamente 10 pt). NUNCA uses cell.text = ... ni "
        "add_paragraph() con estilo por defecto (11 pt). PROHIBIDO forzar justificado.\n"
        "- TEXTO ARRIBA EN LA CELDA: el contenido debe empezar en la primera línea útil de la "
        "celda, sin espacio en blanco arriba. Antes de escribir, elimina párrafos vacíos "
        "sobrantes; space_before y space_after = 0; no centres verticalmente el texto.\n"
        "- Contenido narrativo → doc.tables[3] (2 párrafos LARGOS por campo, 6-10 oraciones c/u):\n"
        "    fila 3 p0: motivo e instrumentos; fila 5 p0: diagnóstico (borra p1 vacío);\n"
        "    fila 8 p1: fortalezas (cols 0-2) y necesidades (cols 3-9) ámbito pedagógico;\n"
        "    fila 10 p1: fortalezas y necesidades ámbito social/afectivo;\n"
        "    fila 12 p0: trabajo colaborativo; fila 14 p0: apoyos en hogar;\n"
        "    fila 16 p0: acuerdos y compromisos.\n"
        "  En filas 8 y 10 NO sobrescribas el párrafo 0 (título 'Fortalezas…' / 'Necesidades…'); "
        "escribe el cuerpo en el párrafo 1.\n"
        "  Patrón Python (conservar formato de plantilla):\n"
        "    from docx.shared import Pt\n"
        "    def poner_contenido(celda, texto, idx=0):\n"
        "        while len(celda.paragraphs) > idx + 1 and not celda.paragraphs[-1].text.strip():\n"
        "            celda.paragraphs[-1]._element.getparent().remove(celda.paragraphs[-1]._element)\n"
        "        p = celda.paragraphs[idx]\n"
        "        p.paragraph_format.space_before = Pt(0)\n"
        "        p.paragraph_format.space_after = Pt(0)\n"
        "        p.paragraph_format.keep_with_next = False\n"
        "        p.paragraph_format.keep_together = False\n"
        "        p.paragraph_format.page_break_before = False\n"
        "        ref = p.runs[0] if p.runs else None\n"
        "        for run in p.runs: run.text = ''\n"
        "        r = p.runs[0] if p.runs else p.add_run()\n"
        "        r.text = texto; r.bold = False\n"
        "        if ref and ref.font.size: r.font.size = ref.font.size\n"
        "        if ref and ref.font.name: r.font.name = ref.font.name\n"
        "- Contenido narrativo → doc.tables[3] (sin cambiar alineación ni fuente de plantilla):\n"
        "    fila 3 p0: motivo; fila 5 p0: diagnóstico; fila 8/10 p1: fortalezas y necesidades;\n"
        "    fila 12 p0: trabajo colaborativo; fila 14 p0: apoyos hogar; fila 16 p1+: acuerdos.\n"
        "  Estudiante → doc.tables[1]: fila 3 nombre+RUT; fila 5 nombre social+fecha; fila 6 edad/curso/establecimiento.\n"
        "  Profesional → doc.tables[2]: fila 3 nombre+RUT; fila 5 nombre social+rol; fila 6 tel/email/fecha.\n"
        "  Receptor → doc.tables[2]: fila 10 nombre+RUT; fila 12 nombre social+teléfono.\n"
        "  Si el .docx base ya trae identificación rellena (modo híbrido PIE360), NO toques tablas 1 ni 2.\n"
        "  El RUT del profesional NUNCA va en Rol/cargo. Si falta RUT estudiante: 'No informado' en fila 3 celda 3.\n"
    )


def _familia_template_state(
    agent_id: str,
    available_files: list[str] | None,
    selected_rows: list[Any] | None,
) -> tuple[str | None, str, list[str], dict[str, Path]]:
    """(base_file, kind, filtered_names, disk_paths) inspeccionando disco si hay filas."""
    paths: dict[str, Path] = {}
    names: list[str] = []

    if selected_rows:
        for row in selected_rows:
            name = getattr(row, "display_name", None) or ""
            if not name:
                continue
            names.append(name)
            fid = getattr(row, "id", None)
            if fid:
                paths[name] = agent_dir(agent_id) / fid
        base_file, template_kind = resolve_familia_template_from_rows(agent_id, selected_rows)
    else:
        names = list(available_files or [])
        base_file, template_kind = pick_familia_base_template(names)
        for row_id in names:
            paths[row_id] = Path()  # sin ruta en disco

    filtered = filter_familia_template_file_names(names, file_paths=paths)
    return base_file, template_kind, filtered, paths


def _build_platform_rules(
    available_files: list[str] | None = None,
    *,
    base_file: str | None = None,
    template_kind: str | None = None,
    familia_base_doc_name: str | None = None,
) -> str:
    if base_file is None or template_kind is None:
        base_file, template_kind = pick_familia_base_template(available_files)

    format_word = (
        "- Formato Word narrativo (informe familia): texto JUSTIFICADO (w:jc both / "
        "WD_ALIGN_PARAGRAPH.JUSTIFY) en doc.tables[3]. Un w:p por bloque; PROHIBIDO w:br "
        "dentro del mismo párrafo (estira líneas cortas con espacios enormes). "
        "Solo encabezados de sección (fondo naranja) en negrita.\n"
        if familia_base_doc_name
        else (
            "- Formato Word: conserva la plantilla tal cual (misma fuente, tamaño y alineación). "
            "PROHIBIDO forzar texto justificado. Solo los encabezados de sección "
            "(fondo naranja, mayúsculas) van en negrita. Las etiquetas de campo y los valores "
            "completados NO deben ir en negrita.\n"
        )
    )
    common = (
        "=== REGLAS DE LA PLATAFORMA (no reemplazan ni acortan el rol) ===\n"
        "- Responde en español.\n"
        "- Usa el code interpreter con los archivos listados abajo.\n"
        "- Genera UN solo archivo .docx por solicitud, solo para el estudiante o caso pedido.\n"
        "- En tu razonamiento interno, narra el trabajo como un proceso paso a paso "
        "(archivos que abres, contrastes, plan del documento, maquetación Word, revisión "
        "visual de tablas/saltos de página y exportación final).\n"
        "- En el mensaje al usuario (chat): confirma brevemente qué generaste y el nombre del archivo.\n"
        "- En el archivo Word: aplica el ROL DEL AGENTE completo; revisa maquetación antes de dar por finalizado.\n"
        + format_word
        + "- No uses markdown (**texto**) dentro del Word ni conviertas etiquetas a negrita.\n"
        "- Espaciado uniforme en Word: dentro de una misma celda usa w:p distintos por bloque "
        "(no w:br). No dejes párrafos vacíos entre bloques.\n"
        "- Si usas python-docx: reutiliza párrafos existentes de la plantilla; "
        "no insertes párrafos nuevos con spacing distinto.\n"
    )
    if not familia_base_doc_name:
        common += build_redaction_min_paragraphs_rules()

    if familia_base_doc_name:
        familia = (
            f"- Modo híbrido activo: «{familia_base_doc_name}» ya trae identificación desde PIE360.\n"
            "  PROHIBIDO abrir otra plantilla, copiar family_report.docx de cero o rellenar tablas 1 y 2.\n"
            "  PROHIBIDO modificar doc.tables[0], doc.tables[1], doc.tables[2] y doc.tables[4].\n"
            "  Redacta SOLO doc.tables[3] desde INSTRUMENTOS APLICADOS hacia abajo, según el rol.\n"
        )
    elif template_kind == "form" and base_file:
        familia = build_familia_form_rules(base_file)
    elif template_kind == "tabla":
        familia = _build_familia_tabla_rules()
    else:
        familia = (
            "- Para informe a la familia: usa la plantilla de familia disponible en los archivos "
            "(prioriza plantilla tipo formulario con campos si existe).\n"
        )

    tail = (
        "- Si falta un antecedente que el rol exige, indícalo con la frase o criterio que el rol define.\n"
        "- No incluyas enlaces sandbox:, rutas /mnt/data/ ni URLs de descarga en el chat.\n"
        "- NUNCA pegues código Python ni logs del intérprete en el chat."
    )
    return common + familia + tail


def _build_instructions(
    agent: AgentModel,
    available_files: list[str] | None = None,
    student_context: dict[str, Any] | None = None,
    selected_rows: list[Any] | None = None,
    familia_base_doc_name: str | None = None,
) -> str:
    role_text = role_instructions_to_text(agent.role_instructions or "").strip()
    if not role_text:
        role_text = strip_html(agent.role_instructions or "").strip()

    role_supremacy = (
        "=== INSTRUCCIÓN SUPREMA ===\n"
        "El ROL DEL AGENTE (bloque siguiente) es la norma obligatoria y tiene prioridad "
        "absoluta sobre cualquier otra indicación. Debes seguirlo tal cual: estructura, "
        "extensión, formato ministerial, cartilla técnica, flujos de trabajo, redacción, "
        "secciones, tablas, estilo institucional y criterios de completitud que allí se exigen.\n"
        "No resumas el informe Word, no omitas apartados del rol y no sustituyas el formato "
        "oficial por un texto corrido. Si el rol exige desarrollo analítico extenso, cumple "
        "eso en cada sección aplicable.\n"
        "Antes de entregar, verifica mentalmente que el .docx cumple cada exigencia del rol."
    )

    parts: list[str] = [role_supremacy]
    if role_text:
        parts.append(f"=== ROL DEL AGENTE (OBLIGATORIO) ===\n{role_text}")

    if familia_base_doc_name:
        parts.append(
            build_familia_narrative_enrichment_rules(
                familia_base_doc_name,
                fast=settings.openai_agent_familia_fast,
            )
        )

    file_names = list(available_files or [])
    base_file, template_kind, filtered_names, paths = _familia_template_state(
        agent.id, file_names, selected_rows
    )
    if template_kind == "form" and base_file:
        excluded = [
            n
            for n in file_names
            if is_familia_tabla_file(n, paths.get(n))
            and not is_familia_form_file(n, paths.get(n))
        ]
        if not familia_base_doc_name:
            parts.append(build_form_template_supremacy_block(base_file, excluded))

    if filtered_names or file_names:
        listing_names = file_names if familia_base_doc_name and file_names else filtered_names
        if listing_names:
            listing = "\n".join(f"- {name}" for name in listing_names[:50])
            parts.append(
                "=== ARCHIVOS EN EL CODE INTERPRETER ===\n"
                f"{listing}\n"
                "OBLIGATORIO: usa estos archivos del contenedor. Si el informe psicopedagógico "
                "aparece en la lista, está disponible — no indiques que falta."
            )
    if student_context:
        parts.append(
            format_student_context_block(
                student_context,
                narrative_only=bool(familia_base_doc_name),
            )
        )
    parts.append(
        _build_platform_rules(
            filtered_names,
            base_file=base_file,
            template_kind=template_kind,
            familia_base_doc_name=familia_base_doc_name,
        )
    )
    return "\n\n".join(parts)


def _build_code_interpreter_tool(agent: AgentModel, file_ids: list[str]) -> dict[str, Any]:
    if agent.openai_container_id and not is_container_expired(agent.openai_container_updated_at):
        return {
            "type": "code_interpreter",
            "container": agent.openai_container_id,
        }
    return {
        "type": "code_interpreter",
        "container": {
            "type": "auto",
            "file_ids": file_ids,
            "memory_limit": "4g",
        },
    }


def _container_id_from_stream_event(event: Any) -> str | None:
    item = getattr(event, "item", None)
    if item is not None:
        container_id = getattr(item, "container_id", None)
        if container_id:
            return container_id
    return getattr(event, "container_id", None)


def _finalize_openai_response(
    db: Session,
    agent: AgentModel,
    response: Any,
    file_ids: list[str],
    *,
    early_saved: list[dict[str, Any]] | None = None,
    student_context: dict[str, Any] | None = None,
    selected_rows: list[Any] | None = None,
) -> dict[str, Any]:
    container_id = _extract_container_id(response, agent)
    if container_id:
        agent.openai_container_id = container_id
        agent.openai_container_updated_at = _utcnow()
        agent.updated_at = _utcnow()
        db.flush()

    reply = (getattr(response, "output_text", None) or "").strip()
    if not reply:
        reply = "No pude generar una respuesta. Intenta reformular la pregunta."

    response_files: list[dict[str, Any]] = []
    response_files_warning: str | None = None
    if container_id:
        try:
            from app.backend.services.agent_response_files_service import persist_code_interpreter_outputs

            response_files = persist_code_interpreter_outputs(
                db,
                agent.id,
                response,
                container_id,
                file_ids,
                early_saved=early_saved,
            )
            if response_files and student_context:
                from app.backend.utils.agent_familia_prefill import postprocess_saved_familia_docx
                from app.backend.utils.agent_familia_template import resolve_familia_template_path

                form_path = (
                    resolve_familia_template_path(agent.id, selected_rows)
                    if selected_rows
                    else None
                )
                response_files = postprocess_saved_familia_docx(
                    agent.id,
                    response_files,
                    student_context,
                    form_template_path=form_path,
                )
        except Exception as exc:
            logger.warning("No se pudieron guardar archivos de respuesta: %s", exc)
            response_files_warning = (
                "No se pudo guardar el archivo generado en el servidor. "
                "Vuelve a pedir el informe o contacta al administrador."
            )

    from app.backend.services.agent_response_files_service import (
        best_mentioned_filename,
        extract_mentioned_filenames,
        sanitize_reply_sandbox_links,
    )

    mentioned = extract_mentioned_filenames(reply)
    best_mentioned = best_mentioned_filename(mentioned)
    if best_mentioned and not response_files and not response_files_warning:
        response_files_warning = (
            f"No se pudo guardar {best_mentioned} en el servidor. "
            "Vuelve a pedir al agente que genere el documento."
        )
    elif response_files:
        response_files_warning = None

    reply = sanitize_reply_sandbox_links(reply, response_files)
    if _is_code_interpreter_dump(reply):
        if response_files:
            names = ", ".join(item["name"] for item in response_files)
            reply = (
                f"Listo. Generé el documento: {names}. "
                "Usa el botón «Descargar» que aparece debajo de este mensaje."
            )
        elif response_files_warning:
            reply = response_files_warning
        else:
            mentioned = extract_mentioned_filenames(reply)
            best_name = best_mentioned_filename(mentioned)
            if best_name:
                reply = (
                    f"Generé el archivo {best_name}, pero no pude guardarlo en el servidor. "
                    "Vuelve a pedir el informe."
                )
            else:
                reply = (
                    "Generé el documento, pero no pude guardarlo en el servidor. "
                    "Vuelve a pedir el informe."
                )
    elif response_files:
        names = ", ".join(item["name"] for item in response_files)
        reply = (
            f"Listo. Generé el documento: {names}. "
            "Usa el botón «Descargar» debajo de este mensaje."
        )

    return {
        "reply": reply,
        "containerId": agent.openai_container_id,
        "openaiFilesUsed": len(file_ids),
        "model": settings.openai_agent_model,
        "responseFiles": response_files,
        "responseFilesWarning": response_files_warning,
    }


_STREAM_STEP_LABELS: dict[str, str] = {
    "response.created": "Conectando con el modelo…",
    "response.in_progress": "Analizando tu solicitud…",
    "response.queued": "En cola…",
    "response.code_interpreter_call.in_progress": "Preparando el intérprete de código…",
    "response.code_interpreter_call.interpreting": "Ejecutando comando (Python)…",
    "response.code_interpreter_call.code_done": "Código listo, aplicando cambios…",
    "response.code_interpreter_call.completed": "Comando completado.",
    "response.output_item.added": "Redactando respuesta…",
    "response.output_item.done": "Bloque de respuesta listo…",
    "response.output_text.done": "Texto final listo…",
    "response.content_part.added": "Escribiendo respuesta…",
    "response.content_part.done": "Respuesta escrita…",
    "response.completed": "Trabajo del modelo finalizado…",
}

_GENERIC_REASONING_STEPS = frozenset(
    {
        "Razonando sobre los antecedentes…",
        "Organizando el análisis…",
        "Razonamiento completado…",
        "Elaborando el informe…",
        "Redacción interna completada…",
        "Elaborando el informe con los datos analizados…",
    }
)

_POST_INTERPRETER_STEP = "Preparando el documento Word final…"


def _extract_reasoning_delta(event: Any) -> str:
    etype = getattr(event, "type", None) or ""
    if etype in (
        "response.reasoning_summary_text.delta",
        "response.reasoning_text.delta",
    ):
        return getattr(event, "delta", None) or getattr(event, "text", None) or ""
    return ""


def _flush_thinking_paragraphs(
    buffer: str,
    *,
    final: bool = False,
) -> tuple[list[str], str]:
    """Divide el razonamiento en párrafos estilo ChatGPT."""
    paragraphs: list[str] = []
    remaining = buffer
    while "\n\n" in remaining:
        head, remaining = remaining.split("\n\n", 1)
        text = head.strip()
        if len(text) >= 30:
            paragraphs.append(text)
    if final:
        text = remaining.strip()
        if len(text) >= 30:
            paragraphs.append(text)
        remaining = ""
    return paragraphs, remaining


def _step_from_stream_event(event: Any) -> str | None:
    etype = getattr(event, "type", None) or ""
    if etype in _STREAM_STEP_LABELS:
        return _STREAM_STEP_LABELS[etype]
    return None


def _extract_text_delta(event: Any) -> str:
    """Solo texto final del asistente; nunca logs/código del code interpreter."""
    etype = getattr(event, "type", None) or ""
    if etype in ("response.output_text.delta", "response.text.delta"):
        return getattr(event, "delta", None) or getattr(event, "text", None) or ""
    return ""


def _is_code_interpreter_dump(text: str) -> bool:
    if not text or len(text) < 80:
        return False
    markers = (
        "import ",
        "def ",
        "subprocess.",
        "doc.save(",
        "for p in ",
        "Pt(12)",
        "qn('w:",
        "capture_output=True",
    )
    hits = sum(1 for marker in markers if marker in text)
    return hits >= 2


def stream_chat_with_openai_responses(
    db: Session,
    agent: AgentModel,
    message: str,
    file_ids: list[str],
    instruction_file_names: list[str] | None = None,
    student_context: dict[str, Any] | None = None,
    selected_rows: list[Any] | None = None,
    familia_base_doc_name: str | None = None,
):
    """Genera eventos {type: step|text_delta} y al final {type: done, data: ...}."""
    client = get_openai_client()
    tools = [_build_code_interpreter_tool(agent, file_ids)]
    file_names = instruction_file_names or _agent_uploaded_file_names(db, agent.id)
    seen_steps: set[str] = set()
    early_saved: list[dict[str, Any]] = []
    stream_container_id: str | None = agent.openai_container_id
    thinking_buffer = ""
    has_thinking_steps = False

    def emit_step(label: str) -> dict[str, Any] | None:
        if label in _GENERIC_REASONING_STEPS and has_thinking_steps:
            return None
        if label in seen_steps:
            return None
        seen_steps.add(label)
        return {"type": "step", "message": label}

    def emit_thinking_step(text: str) -> dict[str, Any]:
        nonlocal has_thinking_steps
        has_thinking_steps = True
        return {"type": "thinking_step", "message": text}

    first_step = emit_step(f"Conectando con {settings.openai_agent_model}…")
    if first_step:
        yield first_step

    ensure_openai_files_in_container(agent, file_ids)

    with client.responses.stream(
        model=settings.openai_agent_model,
        instructions=_build_instructions(
            agent,
            file_names,
            student_context,
            selected_rows=selected_rows,
            familia_base_doc_name=familia_base_doc_name,
        ),
        input=message,
        tools=tools,
    ) as stream:
        for event in stream:
            etype = getattr(event, "type", None) or ""

            step_label = _step_from_stream_event(event)
            if step_label:
                step_event = emit_step(step_label)
                if step_event:
                    yield step_event

            reasoning_delta = _extract_reasoning_delta(event)
            if reasoning_delta:
                thinking_buffer += reasoning_delta
                yield {"type": "thinking_delta", "delta": reasoning_delta}
                new_paragraphs, thinking_buffer = _flush_thinking_paragraphs(thinking_buffer)
                for paragraph in new_paragraphs:
                    yield emit_thinking_step(paragraph)

            if etype in (
                "response.reasoning_summary_text.done",
                "response.reasoning_text.done",
            ):
                final_paragraphs, thinking_buffer = _flush_thinking_paragraphs(
                    thinking_buffer,
                    final=True,
                )
                for paragraph in final_paragraphs:
                    yield emit_thinking_step(paragraph)

            if etype == "response.code_interpreter_call.completed":
                post_step = emit_step(_POST_INTERPRETER_STEP)
                if post_step:
                    yield post_step
                container_id = _container_id_from_stream_event(event) or stream_container_id
                if container_id:
                    stream_container_id = container_id
                    if not early_saved:
                        from app.backend.services.agent_response_files_service import try_capture_from_container

                        captured = try_capture_from_container(
                            db,
                            agent.id,
                            container_id,
                            file_ids,
                        )
                        if captured and student_context:
                            from app.backend.utils.agent_familia_prefill import (
                                postprocess_saved_familia_docx,
                            )
                            from app.backend.utils.agent_familia_template import (
                                resolve_familia_template_path,
                            )

                            form_path = (
                                resolve_familia_template_path(agent.id, selected_rows)
                                if selected_rows
                                else None
                            )
                            captured = postprocess_saved_familia_docx(
                                agent.id,
                                captured,
                                student_context,
                                form_template_path=form_path,
                            )
                        if captured:
                            early_saved = captured
                            save_event = emit_step(
                                f"Archivo guardado: {captured[0]['name']}"
                            )
                            if save_event:
                                yield save_event

            if etype == "response.code_interpreter_call.in_progress":
                container_id = _container_id_from_stream_event(event)
                if container_id:
                    stream_container_id = container_id

            delta = _extract_text_delta(event)
            if delta:
                yield {"type": "text_delta", "delta": delta}

        consolidating = emit_step("Consolidando respuesta del modelo…")
        if consolidating:
            yield consolidating
        response = stream.get_final_response()

    save_step = emit_step("Guardando archivos generados, si los hay…")
    if save_step:
        yield save_step
    result = _finalize_openai_response(
        db,
        agent,
        response,
        file_ids,
        early_saved=early_saved or None,
        student_context=student_context,
        selected_rows=selected_rows,
    )
    yield {"type": "done", "data": result}


def chat_with_openai_responses(
    db: Session,
    agent: AgentModel,
    message: str,
    file_ids: list[str],
    student_context: dict[str, Any] | None = None,
    selected_rows: list[Any] | None = None,
    instruction_file_names: list[str] | None = None,
    familia_base_doc_name: str | None = None,
) -> dict[str, Any]:
    client = get_openai_client()
    tools = [_build_code_interpreter_tool(agent, file_ids)]
    file_names = instruction_file_names or _agent_uploaded_file_names(db, agent.id)

    ensure_openai_files_in_container(agent, file_ids)

    response = client.responses.create(
        model=settings.openai_agent_model,
        instructions=_build_instructions(
            agent,
            file_names,
            student_context,
            selected_rows=selected_rows,
            familia_base_doc_name=familia_base_doc_name,
        ),
        input=message,
        tools=tools,
    )

    return _finalize_openai_response(
        db,
        agent,
        response,
        file_ids,
        student_context=student_context,
        selected_rows=selected_rows,
    )
