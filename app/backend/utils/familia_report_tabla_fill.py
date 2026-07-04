"""Relleno de family_report.docx ministerial con tablas planas (sin campos Word)."""

from __future__ import annotations

import logging
import re
from copy import deepcopy
from pathlib import Path
from typing import Any, Literal

from app.backend.utils.familia_report_formtext import (
    _KEY_ALIASES,
    _NARRATIVE_SLOT_KEYS,
    _get_cell_element,
    _iter_cell_paragraphs,
    _resolve_replacement,
)

logger = logging.getLogger(__name__)

FillMode = Literal["replace", "append"]
# (tabla, fila, columna, clave, modo) — tabla 3 = narrativa; tablas 1-2 = identificación
FAMILIA_TABLA_SLOTS: tuple[tuple[int, int, int, str, FillMode], ...] = (
    (1, 3, 0, "student_full_name", "replace"),
    (1, 3, 3, "student_identification_number", "replace"),
    (1, 5, 0, "student_social_name", "replace"),
    (1, 5, 3, "student_birth_date", "replace"),
    (1, 6, 0, "student_age", "append"),
    (1, 6, 1, "student_course", "append"),
    (1, 6, 2, "student_school", "append"),
    (2, 3, 0, "professional_full_name", "replace"),
    (2, 3, 4, "professional_identification_number", "replace"),
    (2, 5, 0, "professional_social_name", "replace"),
    (2, 5, 4, "professional_job_position", "replace"),
    (2, 6, 0, "professional_phone", "append"),
    (2, 6, 1, "professional_email", "append"),
    (2, 6, 4, "professional_delivered_date_inform", "append"),
    (2, 10, 0, "person_full_name", "replace"),
    (2, 10, 4, "person_identification_number", "replace"),
    (2, 12, 0, "receiver_social_name", "replace"),
    (2, 12, 2, "person_phone", "replace"),
    (2, 14, 0, "person_email", "replace"),
    (3, 3, 0, "applied_instruments", "replace"),
    (3, 5, 0, "diagnostic", "replace"),
    (3, 8, 0, "pedagogical_strengths", "append"),
    (3, 8, 3, "pedagogical_support_needs", "append"),
    (3, 10, 0, "social_affective_strengths", "append"),
    (3, 10, 3, "social_affective_support_needs", "append"),
    (3, 12, 0, "collaborative_work", "replace"),
    (3, 14, 0, "home_based_description", "replace"),
    (3, 16, 0, "school_family_agreements", "replace"),
    (3, 17, 4, "evaluation_date_1", "replace"),
    (3, 17, 5, "evaluation_date_2", "replace"),
    (3, 17, 7, "evaluation_date_3", "replace"),
    (3, 17, 8, "evaluation_date_4", "replace"),
    (3, 17, 9, "evaluation_date_5", "replace"),
)

FAMILIA_TABLA_IDENTIFICATION_SLOTS = tuple(s for s in FAMILIA_TABLA_SLOTS if s[0] in (1, 2))
FAMILIA_TABLA_NARRATIVE_SLOTS = tuple(s for s in FAMILIA_TABLA_SLOTS if s[0] == 3)


def docx_is_familia_ministerial_tabla(path: Path) -> bool:
    """Plantilla ministerial actual: 5 tablas, sin SDT ni FORMTEXT."""
    try:
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document(str(path))
        if any(doc.element.body.iter(qn("w:sdt"))):
            return False
        for instr in doc.element.body.iter(qn("w:instrText")):
            if instr.text and "FORMTEXT" in instr.text:
                return False
        tbls = doc.element.body.findall(qn("w:tbl"))
        if len(tbls) != 5:
            return False
        header = "".join(t.text or "" for t in tbls[1].iter(qn("w:t")))
        return "IDENTIFICACIÓN DEL ESTUDIANTE" in header
    except Exception:
        return False


def _paragraph_text(p_el: Any, qn: Any) -> str:
    return "".join(t.text or "" for t in p_el.iter(qn("w:t")))


def _clear_paragraph_runs(p_el: Any, qn: Any) -> None:
    for child in list(p_el):
        if child.tag == qn("w:r"):
            p_el.remove(child)


def _append_run_text(p_el: Any, text: str, qn: Any, OxmlElement: Any, ref_r_pr: Any | None) -> None:
    r = OxmlElement("w:r")
    if ref_r_pr is not None:
        r.append(deepcopy(ref_r_pr))
    wt = OxmlElement("w:t")
    wt.text = text
    if text.startswith(" ") or text.endswith(" "):
        wt.set(qn("xml:space"), "preserve")
    r.append(wt)
    p_el.append(r)


def _set_cell_plain_text(
    tc_el: Any,
    value: str,
    qn: Any,
    OxmlElement: Any,
    *,
    mode: FillMode = "replace",
) -> bool:
    text = (value or "").strip()
    if not text:
        return False

    paragraphs = tc_el.findall(qn("w:p"))
    ref_r_pr = None
    if paragraphs:
        for r in paragraphs[0].iter(qn("w:r")):
            rpr = r.find(qn("w:rPr"))
            if rpr is not None:
                ref_r_pr = deepcopy(rpr)
                break

    if mode == "append" and paragraphs:
        label = _paragraph_text(paragraphs[0], qn).strip()
        if label:
            target = None
            for p_el in paragraphs[1:]:
                if not _paragraph_text(p_el, qn).strip():
                    target = p_el
                    break
            if target is None:
                target = OxmlElement("w:p")
                tc_el.append(target)
            else:
                _clear_paragraph_runs(target, qn)
            _append_run_text(target, text, qn, OxmlElement, ref_r_pr)
            return True

    # replace: solo rellena celda vacía o segundo párrafo; no borra etiquetas de fila superior
    if paragraphs:
        existing = _paragraph_text(paragraphs[0], qn).strip()
        if existing and not _is_label_like(existing):
            _clear_paragraph_runs(paragraphs[0], qn)
            _append_run_text(paragraphs[0], text, qn, OxmlElement, ref_r_pr)
            return True

    for p_el in paragraphs:
        tc_el.remove(p_el)
    new_p = OxmlElement("w:p")
    _append_run_text(new_p, text, qn, OxmlElement, ref_r_pr)
    tc_el.append(new_p)
    return True


def _clear_cell_placeholder_text(tc_el: Any, qn: Any) -> bool:
    from app.backend.utils.familia_report_prefill import _paragraph_has_placeholder

    changed = False
    for p_el in tc_el.findall(qn("w:p")):
        full = _paragraph_text(p_el, qn).strip()
        if not full or not _paragraph_has_placeholder(full):
            continue
        if _is_label_like(full):
            continue
        _clear_paragraph_runs(p_el, qn)
        changed = True
    return changed


def _is_label_like(text: str) -> bool:
    letters = [c for c in text if c.isalpha()]
    if len(letters) < 6:
        return True
    upper = sum(1 for c in letters if c.isupper())
    return (upper / len(letters)) > 0.75


def _fill_tabla_slot_rows(
    doc: Any,
    slots: tuple[tuple[int, int, int, str, FillMode], ...],
    replacements: dict[str, str],
    qn: Any,
    OxmlElement: Any,
) -> list[str]:
    seen: set[tuple[int, int, int]] = set()
    filled_keys: list[str] = []
    for table_idx, row_idx, col_idx, key, mode in slots:
        slot = (table_idx, row_idx, col_idx)
        if slot in seen:
            continue
        seen.add(slot)
        tc_el = _get_cell_element(doc, table_idx, row_idx, col_idx, qn)
        if tc_el is None:
            continue
        value = _resolve_replacement(key, replacements)
        if value:
            if _set_cell_plain_text(tc_el, value, qn, OxmlElement, mode=mode):
                filled_keys.append(key)
        else:
            if _clear_cell_placeholder_text(tc_el, qn):
                filled_keys.append(key)
    return filled_keys


def fix_familia_motivo_evaluacion_tabla(
    docx_path: Path,
    student_context: dict[str, Any] | None = None,
) -> None:
    """Marca ingreso o reevaluación reemplazando ☐ por x en la fila MOTIVO."""
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    from app.backend.utils.familia_report_prefill import (
        FAMILIA_CHECKBOX_CHECKED_MARK,
        _evaluation_type_flags,
    )

    doc = Document(str(docx_path))
    is_admission, is_reeval = _evaluation_type_flags(student_context)
    col = 1 if is_admission else (6 if is_reeval else None)
    if col is None:
        return

    tc_el = _get_cell_element(doc, 3, 1, col, qn)
    if tc_el is None:
        return

    for wt in tc_el.iter(qn("w:t")):
        if not wt.text:
            continue
        if "☐" in wt.text:
            wt.text = wt.text.replace("☐", FAMILIA_CHECKBOX_CHECKED_MARK, 1)
            doc.save(str(docx_path))
            return
        if wt.text.strip() in ("x", "X", FAMILIA_CHECKBOX_CHECKED_MARK):
            return

    doc.save(str(docx_path))


def apply_familia_arial_10_to_tabla(docx_path: Path) -> None:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    from app.backend.utils.familia_report_prefill import _apply_arial_10_to_run

    doc = Document(str(docx_path))
    append_slots = {
        (table_idx, row_idx, col_idx)
        for table_idx, row_idx, col_idx, _key, mode in FAMILIA_TABLA_SLOTS
        if mode == "append"
    }

    for table_idx, row_idx, col_idx, _key, _mode in FAMILIA_TABLA_SLOTS:
        tc_el = _get_cell_element(doc, table_idx, row_idx, col_idx, qn)
        if tc_el is None:
            continue
        skip_first = (table_idx, row_idx, col_idx) in append_slots
        for pi, p_el in enumerate(_iter_cell_paragraphs(tc_el, qn)):
            if skip_first and pi == 0:
                continue
            for r in p_el.iter(qn("w:r")):
                wt = r.find(qn("w:t"))
                if wt is not None and (wt.text or "").strip():
                    _apply_arial_10_to_run(r, qn, OxmlElement)

    doc.save(str(docx_path))


def compact_familia_tabla_narrative(docx_path: Path) -> None:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    from app.backend.utils.familia_report_prefill import (
        _append_narrative_paragraph,
        _apply_narrative_paragraph_format,
        _is_label_paragraph,
        _split_paragraph_segments,
    )

    doc = Document(str(docx_path))
    append_slots = {
        (table_idx, row_idx, col_idx)
        for table_idx, row_idx, col_idx, _key, mode in FAMILIA_TABLA_SLOTS
        if mode == "append"
    }
    narrative_slots = {
        (ti, ri, ci)
        for ti, ri, ci, key, _mode in FAMILIA_TABLA_SLOTS
        if key in _NARRATIVE_SLOT_KEYS
    }

    for table_idx, row_idx, col_idx in narrative_slots:
        tc_el = _get_cell_element(doc, table_idx, row_idx, col_idx, qn)
        if tc_el is None:
            continue
        paragraphs = tc_el.findall(qn("w:p"))
        if not paragraphs:
            continue

        skip_label = (table_idx, row_idx, col_idx) in append_slots
        body_paragraphs = [
            p for i, p in enumerate(paragraphs) if not (skip_label and i == 0)
        ]

        segments: list[str] = []
        has_br = False
        for p_el in body_paragraphs:
            text = _paragraph_text(p_el, qn).strip()
            if not text or _is_label_paragraph(text):
                continue
            if any(r.find(qn("w:br")) is not None for r in p_el.findall(qn("w:r"))):
                has_br = True
                segments.extend(_split_paragraph_segments(p_el, qn))
            else:
                for block in re.split(r"\n\s*\n", text):
                    block = block.strip()
                    if block:
                        segments.append(block)

        if segments and (len(segments) > 1 or has_br):
            for p_el in body_paragraphs:
                tc_el.remove(p_el)
            for seg in segments:
                _append_narrative_paragraph(tc_el, seg, qn, OxmlElement, None)

        for pi, p_el in enumerate(tc_el.findall(qn("w:p"))):
            if skip_label and pi == 0:
                continue
            text = _paragraph_text(p_el, qn).strip()
            if not text or _is_label_paragraph(text):
                continue
            ppr = p_el.find(qn("w:pPr"))
            if ppr is None:
                ppr = OxmlElement("w:pPr")
                p_el.insert(0, ppr)
            _apply_narrative_paragraph_format(ppr, qn, OxmlElement)

    doc.save(str(docx_path))


# Filas narrativas de doc.tables[3] con altura mínima fija en la plantilla ministerial
FAMILIA_TABLA_NARRATIVE_ROW_INDICES = (3, 5, 8, 10, 12, 14, 16, 17)


def relax_familia_tabla_layout(docx_path: Path) -> None:
    """
    Permite que el contenido narrativo fluya entre páginas sin huecos en blanco:
    quita keepNext/keepLines/pageBreakBefore, cantSplit y alturas fijas de fila.
    """
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    from app.backend.utils.familia_report_prefill import _clear_paragraph_pagination_locks

    if not docx_is_familia_ministerial_tabla(docx_path):
        return

    doc = Document(str(docx_path))
    tbls = doc.element.body.findall(qn("w:tbl"))

    for tbl_idx, tbl_el in enumerate(tbls):
        for tr_el in tbl_el.findall(qn("w:tr")):
            tr_pr = tr_el.find(qn("w:trPr"))
            if tr_pr is not None:
                for el in list(tr_pr.findall(qn("w:cantSplit"))):
                    tr_pr.remove(el)

        if tbl_idx == 3:
            rows = tbl_el.findall(qn("w:tr"))
            for row_idx in FAMILIA_TABLA_NARRATIVE_ROW_INDICES:
                if row_idx >= len(rows):
                    continue
                tr_pr = rows[row_idx].find(qn("w:trPr"))
                if tr_pr is not None:
                    for el in list(tr_pr.findall(qn("w:trHeight"))):
                        tr_pr.remove(el)
                for tc_el in rows[row_idx].findall(qn("w:tc")):
                    tc_pr = tc_el.find(qn("w:tcPr"))
                    if tc_pr is None:
                        continue
                    for el in list(tc_pr.findall(qn("w:vAlign"))):
                        tc_pr.remove(el)

        for tc_el in tbl_el.iter(qn("w:tc")):
            for p_el in tc_el.findall(qn("w:p")):
                ppr = p_el.find(qn("w:pPr"))
                if ppr is None:
                    ppr = OxmlElement("w:pPr")
                    p_el.insert(0, ppr)
                _clear_paragraph_pagination_locks(ppr, qn, OxmlElement)

    doc.save(str(docx_path))
