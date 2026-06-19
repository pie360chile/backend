"""Relleno de identificación en informe familia vía placeholders {{clave}} sin romper la plantilla."""

from __future__ import annotations

import logging
import re
import unicodedata
from copy import deepcopy
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

# Claves que van en tablas 1 y 2 (identificación). Narrativa sigue en tabla 3.
FAMILIA_IDENTIFICATION_KEYS: frozenset[str] = frozenset(
    {
        "student_full_name",
        "student_social_name",
        "student_identification_number",
        "student_birth_date",
        "student_born_date",
        "student_age",
        "student_course",
        "student_school",
        "professional_full_name",
        "professional_social_name",
        "professional_identification_number",
        "professional_job_position",
        "professional_phone",
        "professional_email",
        "professional_phone_email",
        "professional_delivered_date_inform",
        "person_full_name",
        "receiver_social_name",
        "person_identification_number",
        "person_relation_student",
        "person_presence",
        "person_phone",
        "person_email",
    }
)

_PLACEHOLDER_RE = re.compile(
    r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}|\{\s*([a-zA-Z0-9_]+)\s*\}|\[\s*([a-zA-Z0-9_]+)\s*\]"
)

# Etiquetas ministeriales → clave (celda con {{clave}} arriba y etiqueta abajo).
_IDENTIFICATION_LABEL_HINTS: dict[str, tuple[str, ...]] = {
    "student_full_name": ("nombre de identidad", "nombres y apellidos", "nombre completo"),
    "student_identification_number": ("run", "rut", "identificacion"),
    "student_social_name": ("nombre social",),
    "student_birth_date": ("fecha nacimiento", "fecha de nacimiento"),
    "student_age": ("edad",),
    "student_course": ("curso / nivel", "curso nivel"),
    "student_school": ("establecimiento",),
    "professional_full_name": ("nombre de identidad", "nombres y apellidos"),
    "professional_identification_number": ("run", "rut"),
    "professional_social_name": ("nombre social",),
    "professional_job_position": ("rol", "cargo", "funcion"),
    "professional_phone": ("telefono", "teléfono", "fono"),
    "professional_email": ("correo", "e mail", "email"),
    "professional_delivered_date_inform": ("fecha entrega", "fecha de entrega"),
    "person_full_name": ("nombre de identidad", "nombres y apellidos"),
    "receiver_social_name": ("nombre social",),
    "person_identification_number": ("run", "rut"),
    "person_phone": ("telefono", "teléfono", "fono"),
    "person_email": ("correo", "e mail", "email"),
    "person_relation_student": ("vinculo", "vínculo", "parentesco", "relacion"),
    "person_presence": ("presencia",),
}


def _normalize_label(text: str) -> str:
    folded = unicodedata.normalize("NFKD", text or "")
    asciiish = "".join(ch for ch in folded if not unicodedata.combining(ch))
    return re.sub(r"[^a-z0-9]+", " ", asciiish.lower()).strip()


def _paragraph_matches_label(text: str, key: str) -> bool:
    norm = _normalize_label(text)
    if not norm or len(norm) > 80:
        return False
    for hint in _IDENTIFICATION_LABEL_HINTS.get(key, ()):
        if norm == hint or norm.endswith(hint) or hint in norm:
            return True
    return False


def _formats_for_tag(tag: str) -> tuple[str, ...]:
    return (f"{{{{{tag}}}}}", f"{{{tag}}}", f"[{tag}]")


def _keys_in_text(text: str) -> set[str]:
    found: set[str] = set()
    for m in _PLACEHOLDER_RE.finditer(text or ""):
        key = m.group(1) or m.group(2) or m.group(3)
        if key in FAMILIA_IDENTIFICATION_KEYS:
            found.add(key)
    return found


def _paragraph_visible_text(p_el: Any, qn: Any) -> str:
    return "".join(t.text or "" for t in p_el.iter(qn("w:t")))


def _iter_body_paragraphs(doc: Any, qn: Any):
    for p_el in doc.element.body.iter(qn("w:p")):
        yield p_el
    for section in doc.sections:
        for hf in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
        ):
            if hf is None or hf._element is None:
                continue
            for p_el in hf._element.iter(qn("w:p")):
                yield p_el


def list_familia_placeholder_keys_in_doc(doc: Any) -> set[str]:
    """Claves {{...}} presentes (aunque Word las parta en varios w:t del mismo párrafo)."""
    from docx.oxml.ns import qn

    keys: set[str] = set()
    for p_el in _iter_body_paragraphs(doc, qn):
        keys |= _keys_in_text(_paragraph_visible_text(p_el, qn))
    return keys


def list_familia_placeholder_keys_in_path(path: Path) -> set[str]:
    try:
        from docx import Document

        doc = Document(str(path))
        return list_familia_placeholder_keys_in_doc(doc)
    except Exception:
        return set()


def docx_has_familia_placeholders(path: Path) -> bool:
    return bool(list_familia_placeholder_keys_in_path(path))


def _replace_text_placeholders(text: str, replacements: dict[str, str]) -> str:
    result = text
    for tag, value in replacements.items():
        val = str(value) if value is not None else ""
        for fmt in _formats_for_tag(tag):
            result = result.replace(fmt, val)
    return result


def _replace_placeholders_in_paragraph(
    p_el: Any,
    qn: Any,
    OxmlElement: Any,
    replacements: dict[str, str],
) -> list[str]:
    """Reemplaza placeholders en un w:p aunque estén partidos en varios runs."""
    full = _paragraph_visible_text(p_el, qn)
    if not full or "{{" not in full and "{" not in full and "[" not in full:
        return []

    keys_here = _keys_in_text(full)
    if not keys_here:
        return []

    applicable = {k: replacements[k] for k in keys_here if k in replacements}
    if not applicable:
        return []

    new_full = _replace_text_placeholders(full, applicable)
    if new_full == full:
        return []

    ref_r_pr = None
    for r in p_el.findall(qn("w:r")):
        rpr = r.find(qn("w:rPr"))
        if rpr is not None:
            ref_r_pr = deepcopy(rpr)
            break

    for child in list(p_el):
        if child.tag == qn("w:r"):
            p_el.remove(child)

    r = OxmlElement("w:r")
    if ref_r_pr is not None:
        r.append(deepcopy(ref_r_pr))
    wt = OxmlElement("w:t")
    wt.text = new_full
    if new_full.startswith(" ") or new_full.endswith(" "):
        wt.set(qn("xml:space"), "preserve")
    r.append(wt)
    p_el.append(r)

    filled: list[str] = []
    for tag in applicable:
        for fmt in _formats_for_tag(tag):
            if fmt in full:
                filled.append(tag)
                break
    return filled


def _iter_table_paragraphs(doc: Any, table_indices: tuple[int, ...], qn: Any):
    for table_idx in table_indices:
        if table_idx >= len(doc.tables):
            continue
        for row in doc.tables[table_idx].rows:
            for cell in row.cells:
                for p_el in cell._element.findall(qn("w:p")):
                    yield p_el


def _set_paragraph_plain_text(
    p_el: Any,
    value: str,
    qn: Any,
    OxmlElement: Any,
) -> None:
    ref_r_pr = None
    for r in p_el.findall(qn("w:r")):
        rpr = r.find(qn("w:rPr"))
        if rpr is not None:
            ref_r_pr = deepcopy(rpr)
            break
    for child in list(p_el):
        if child.tag == qn("w:r"):
            p_el.remove(child)
    r = OxmlElement("w:r")
    if ref_r_pr is not None:
        r.append(deepcopy(ref_r_pr))
    wt = OxmlElement("w:t")
    wt.text = value
    if value.startswith(" ") or value.endswith(" "):
        wt.set(qn("xml:space"), "preserve")
    r.append(wt)
    p_el.append(r)


def replace_placeholders_in_doc_tables(
    doc: Any,
    table_indices: tuple[int, ...],
    replacements: dict[str, str],
    *,
    keys_filter: frozenset[str] | None = None,
) -> list[str]:
    """Sustituye {{clave}} solo en tablas indicadas (identificación)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    filtered = replacements
    if keys_filter is not None:
        filtered = {k: v for k, v in replacements.items() if k in keys_filter}

    filled: list[str] = []
    for p_el in _iter_table_paragraphs(doc, table_indices, qn):
        filled.extend(
            _replace_placeholders_in_paragraph(p_el, qn, OxmlElement, filtered)
        )
    return sorted(set(filled))


def fill_identification_by_labels_in_tables(
    doc: Any,
    replacements: dict[str, str],
    table_indices: tuple[int, ...] = (1, 2),
) -> list[str]:
    """
    Rellena celdas con layout «valor arriba / etiqueta abajo» sin usar coordenadas.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    filled: list[str] = []
    for table_idx in table_indices:
        if table_idx >= len(doc.tables):
            continue
        for row in doc.tables[table_idx].rows:
            for cell in row.cells:
                tc_el = cell._element
                paragraphs = tc_el.findall(qn("w:p"))
                if not paragraphs:
                    continue
                texts = [_paragraph_visible_text(p, qn) for p in paragraphs]
                for key in FAMILIA_IDENTIFICATION_KEYS:
                    value = str(replacements.get(key) or "").strip()
                    if not value:
                        continue
                    label_idx: int | None = None
                    for i, text in enumerate(texts):
                        if _paragraph_matches_label(text, key):
                            label_idx = i
                            break
                    if label_idx is None:
                        continue
                    value_idx = label_idx - 1 if label_idx > 0 else 0
                    target = paragraphs[value_idx]
                    current = texts[value_idx].strip()
                    if "{{" in current or not current or _keys_in_text(current):
                        _set_paragraph_plain_text(target, value, qn, OxmlElement)
                        filled.append(key)
    return sorted(set(filled))


def fill_familia_identification_tables(
    doc: Any,
    replacements: dict[str, str],
) -> list[str]:
    """Placeholders + etiquetas en tablas 1-2 (sin coordenadas ministeriales)."""
    id_replacements = {k: v for k, v in replacements.items() if k in FAMILIA_IDENTIFICATION_KEYS}
    filled = replace_placeholders_in_doc_tables(
        doc, (1, 2), id_replacements, keys_filter=FAMILIA_IDENTIFICATION_KEYS
    )
    filled.extend(fill_identification_by_labels_in_tables(doc, id_replacements, (1, 2)))
    return sorted(set(filled))


def replace_familia_placeholders_in_doc(
    doc: Any,
    replacements: dict[str, str],
    *,
    keys_filter: frozenset[str] | None = None,
) -> list[str]:
    """
    Sustituye {{clave}} por párrafo (no por run suelto).
    Word suele partir {{student_full_name}} en varios w:t; por eso fallaba el reemplazo.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    filtered = replacements
    if keys_filter is not None:
        filtered = {k: v for k, v in replacements.items() if k in keys_filter}

    if not filtered:
        return []

    filled: list[str] = []
    for p_el in _iter_body_paragraphs(doc, qn):
        filled.extend(
            _replace_placeholders_in_paragraph(p_el, qn, OxmlElement, filtered)
        )

    return sorted(set(filled))


def copy_familia_tables_between_docs(
    source_path: Path,
    dest_path: Path,
    table_indices: tuple[int, ...] = (1, 2),
) -> None:
    """Copia tablas de identificación (1 y 2) de source a dest sin tocar narrativa."""
    from docx import Document
    from docx.oxml.ns import qn

    src_doc = Document(str(source_path))
    dst_doc = Document(str(dest_path))
    src_tbls = src_doc.element.body.findall(qn("w:tbl"))
    dst_tbls = dst_doc.element.body.findall(qn("w:tbl"))

    for idx in table_indices:
        if idx >= len(src_tbls) or idx >= len(dst_tbls):
            continue
        parent = dst_tbls[idx].getparent()
        if parent is None:
            continue
        parent.replace(dst_tbls[idx], deepcopy(src_tbls[idx]))

    dst_doc.save(str(dest_path))


def restore_familia_identification_from_template(
    output_path: Path,
    template_path: Path,
    replacements: dict[str, str],
) -> list[str]:
    """
    Tras GPT, el .docx ya no tiene {{...}}. Regenera tablas 1-2 desde la plantilla
    con datos de BD y las pega en el informe final (conserva tabla 3 del modelo).
    """
    import shutil
    import uuid

    template_path = Path(template_path)
    output_path = Path(output_path)
    if not template_path.is_file() or not output_path.is_file():
        return []

    if not docx_has_familia_placeholders(template_path):
        return []

    temp = output_path.parent / f"_familia_id_{uuid.uuid4().hex[:8]}.docx"
    try:
        result = fill_familia_identification_placeholders(
            template_path, replacements, temp
        )
        copy_familia_tables_between_docs(temp, output_path, (1, 2))
        filled = result.get("filled_keys") or []
        logger.info(
            "Identificación restaurada desde plantilla (%d claves) en %s",
            len(filled),
            output_path.name,
        )
        return filled
    finally:
        temp.unlink(missing_ok=True)


def fill_familia_identification_placeholders(
    template_path: str | Path,
    replacements: dict[str, str],
    output_path: str | Path,
) -> dict[str, Any]:
    """Copia plantilla y reemplaza placeholders de identificación."""
    import shutil

    from docx import Document

    template_path = Path(template_path)
    output_path = Path(output_path)
    if template_path.resolve() != output_path.resolve():
        shutil.copy2(template_path, output_path)

    doc = Document(str(output_path))
    id_replacements = {k: v for k, v in replacements.items() if k in FAMILIA_IDENTIFICATION_KEYS}
    filled = fill_familia_identification_tables(doc, id_replacements)
    doc.save(str(output_path))
    logger.info(
        "Placeholders identificación: %d claves en %s",
        len(filled),
        output_path.name,
    )
    return {"status": "success", "filled_keys": filled, "mode": "placeholders"}


def fill_familia_identification_only_docx(
    template_path: str | Path,
    replacements: dict[str, str],
    output_path: str | Path,
) -> dict[str, Any]:
    """Solo identificación (tablas 1-2). Sin narrativa ni coordenadas si hay {{...}}."""
    from app.backend.utils.agent_familia_tabla_fill import docx_is_familia_ministerial_tabla

    path = Path(template_path)
    if not docx_is_familia_ministerial_tabla(path):
        return fill_familia_identification_placeholders(template_path, replacements, output_path)

    if docx_has_familia_placeholders(path):
        return fill_familia_identification_placeholders(template_path, replacements, output_path)

    import shutil

    output_path = Path(output_path)
    shutil.copy2(path, output_path)
    from docx import Document

    from app.backend.utils.agent_familia_tabla_fill import (
        FAMILIA_TABLA_IDENTIFICATION_SLOTS,
        _fill_tabla_slot_rows,
    )
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document(str(output_path))
    filled = _fill_tabla_slot_rows(
        doc, FAMILIA_TABLA_IDENTIFICATION_SLOTS, replacements, qn, OxmlElement
    )
    doc.save(str(output_path))
    return {"status": "success", "filled_keys": filled, "mode": "coordinates-id-only"}


def restore_identification_from_context(docx_path: Path, student_context: dict[str, Any]) -> None:
    """Reaplica identificación desde BD (placeholders + coordenadas para el resto)."""
    from app.backend.utils.agent_familia_prefill import build_familia_identification_replacements
    from app.backend.utils.agent_familia_tabla_fill import refill_familia_identification_only

    replacements = build_familia_identification_replacements(student_context)
    refill_familia_identification_only(docx_path, replacements)
