"""Extrae texto de los archivos de contexto del agente (carpeta Files, no plantillas)."""

from __future__ import annotations

import re
from pathlib import Path

from app.backend.utils import agent_v2_storage as storage

# Límites para no saturar el contexto del modelo
MAX_FILES = 25
MAX_CHARS_PER_FILE = 24_000
MAX_TOTAL_CHARS = 120_000

_TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".json",
    ".csv",
    ".xml",
    ".html",
    ".htm",
    ".log",
    ".yaml",
    ".yml",
}
_BINARY_EXTENSIONS = {".docx", ".pdf"}
_SPREADSHEET_EXTENSIONS = {".xls", ".xlsx", ".xlsm"}

_INTERACTIVE_REPORT_MARKERS = (
    "reporte_interactivo",
    "reporte interactivo",
    "reporte-interactivo",
)


def _normalize_rut(value: str) -> str:
    return re.sub(r"[^0-9kK]", "", (value or "").strip()).upper()


def _is_context_file(path: Path, agent_root: Path) -> bool:
    if not path.is_file():
        return False
    rel = path.relative_to(agent_root).as_posix()
    if rel.startswith("documentos/"):
        return False
    ext = path.suffix.lower()
    return ext in _TEXT_EXTENSIONS or ext in _BINARY_EXTENSIONS or ext in _SPREADSHEET_EXTENSIONS


def _context_file_sort_key(path: Path) -> tuple[int, str]:
    """Prioriza reportes interactivos y Excel antes del orden alfabético."""
    name = path.name.lower()
    rel = path.as_posix().lower()
    if any(marker in name or marker in rel for marker in _INTERACTIVE_REPORT_MARKERS):
        return (0, rel)
    if path.suffix.lower() in _SPREADSHEET_EXTENSIONS:
        return (1, rel)
    return (2, rel)


def list_all_context_file_paths(agent_name: str) -> list[Path]:
    root = storage.agent_folder(agent_name)
    if not root.exists():
        return []
    files = [p for p in root.rglob("*") if _is_context_file(p, root)]
    return sorted(files, key=_context_file_sort_key)


def list_context_file_paths(agent_name: str) -> list[Path]:
    return list_all_context_file_paths(agent_name)[:MAX_FILES]


def list_spreadsheet_paths(agent_name: str) -> list[Path]:
    return [
        p
        for p in list_all_context_file_paths(agent_name)
        if p.suffix.lower() in _SPREADSHEET_EXTENSIONS
    ]


def _read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _read_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [((cell.text or "").strip()) for cell in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
    return "\n".join(parts)


def _read_pdf(path: Path) -> str:
    import fitz

    doc = fitz.open(str(path))
    try:
        parts: list[str] = []
        for page in doc:
            text = (page.get_text() or "").strip()
            if text:
                parts.append(text)
        return "\n\n".join(parts)
    finally:
        doc.close()


def _read_excel(path: Path) -> str:
    """Convierte hojas Excel (.xls / .xlsx) a texto tabular para el contexto del modelo."""
    import pandas as pd

    ext = path.suffix.lower()
    engine: str | None = None
    if ext == ".xls":
        engine = "xlrd"
    elif ext in {".xlsx", ".xlsm"}:
        engine = "openpyxl"

    read_kwargs: dict = {"sheet_name": None, "header": None, "dtype": str}
    if engine:
        read_kwargs["engine"] = engine

    try:
        sheets = pd.read_excel(path, **read_kwargs)
    except Exception as first_exc:
        try:
            sheets = pd.read_excel(path, sheet_name=None, header=None, dtype=str)
        except Exception as second_exc:
            raise RuntimeError(f"{first_exc}; reintento: {second_exc}") from first_exc

    if not isinstance(sheets, dict):
        sheets = {"Hoja1": sheets}

    parts: list[str] = []
    max_rows = 600
    for sheet_name, df in sheets.items():
        df = df.fillna("")
        label = str(sheet_name)
        parts.append(f"## Hoja: {label}")
        if len(df) > max_rows:
            df = df.head(max_rows)
            parts.append(f"[Mostrando primeras {max_rows} filas de {label}]")
        # TSV legible para el modelo (apoderados, reportes interactivos, etc.)
        parts.append(df.to_csv(index=False, sep="\t", lineterminator="\n"))
    return "\n".join(parts)


def _cell_matches_rut(cell_value: str, target_rut: str) -> bool:
    if not cell_value or not target_rut:
        return False
    cell_norm = _normalize_rut(cell_value)
    if not cell_norm:
        return False
    if cell_norm == target_rut:
        return True
    # RUT sin dígito verificador o con ceros a la izquierda
    if len(target_rut) >= 8 and cell_norm.startswith(target_rut[:-1]):
        return True
    if len(cell_norm) >= 8 and target_rut.startswith(cell_norm[:-1]):
        return True
    return False


def extract_spreadsheet_hint_for_rut(agent_name: str, student_rut: str) -> tuple[str, list[str]]:
    """
    Busca el RUT del estudiante en TODOS los Excel de Files (sin límite de 25 archivos).
    Devuelve (bloque de texto para el prompt, rutas de archivos donde hubo coincidencia).
    """
    target = _normalize_rut(student_rut)
    if len(target) < 8:
        return "", []

    root = storage.agent_folder(agent_name)
    paths = list_spreadsheet_paths(agent_name)
    if not paths:
        return "", []

    import pandas as pd

    sections: list[str] = []
    matched_files: list[str] = []

    for path in paths:
        rel = path.relative_to(root).as_posix()
        try:
            ext = path.suffix.lower()
            engine = "xlrd" if ext == ".xls" else ("openpyxl" if ext in {".xlsx", ".xlsm"} else None)
            read_kwargs: dict = {"sheet_name": None, "header": None, "dtype": str}
            if engine:
                read_kwargs["engine"] = engine
            sheets = pd.read_excel(path, **read_kwargs)
            if not isinstance(sheets, dict):
                sheets = {"Hoja1": sheets}
        except Exception as exc:
            sections.append(
                f"### Excel: {rel}\n[Error al leer: {exc}. "
                "Verifique openpyxl (xlsx) y xlrd (xls) en el servidor.]"
            )
            continue

        file_hits: list[str] = []
        for sheet_name, df in sheets.items():
            df = df.fillna("")
            for row_idx, row in df.iterrows():
                row_values = [str(v).strip() for v in row.tolist()]
                if not any(_cell_matches_rut(v, target) for v in row_values if v):
                    continue
                # Cabecera: fila anterior si existe (reportes interactivos suelen tener encabezados)
                header_vals: list[str] = []
                if int(row_idx) > 0:
                    header_vals = [
                        str(v).strip()
                        for v in df.iloc[int(row_idx) - 1].fillna("").tolist()
                    ]
                lines = [f"## Hoja: {sheet_name} — fila con RUT {student_rut}"]
                if header_vals and any(header_vals):
                    pairs = []
                    for col_idx, val in enumerate(row_values):
                        if not val:
                            continue
                        label = (
                            header_vals[col_idx]
                            if col_idx < len(header_vals) and header_vals[col_idx]
                            else f"col_{col_idx + 1}"
                        )
                        pairs.append(f"{label}: {val}")
                    lines.append("\n".join(pairs))
                else:
                    lines.append("\t".join(v for v in row_values if v))
                file_hits.append("\n".join(lines))
                break  # una fila por hoja

        if file_hits:
            matched_files.append(rel)
            sections.append(f"### Archivo: {rel}\n\n" + "\n\n".join(file_hits))

    if not sections:
        index = ", ".join(p.relative_to(root).as_posix() for p in paths)
        return (
            "EXCEL EN FILES (sin fila coincidente para el RUT indicado). "
            f"Archivos revisados ({len(paths)}): {index}",
            [],
        )

    header = (
        "DATOS EXTRAÍDOS DE EXCEL (búsqueda directa por RUT del estudiante). "
        "Usa esto para apoderado, contacto y datos del reporte interactivo.\n"
    )
    return header + "\n\n---\n\n".join(sections), matched_files


def extract_file_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in _TEXT_EXTENSIONS:
        return _read_text_file(path)
    if ext == ".docx":
        return _read_docx(path)
    if ext == ".pdf":
        return _read_pdf(path)
    if ext in _SPREADSHEET_EXTENSIONS:
        return _read_excel(path)
    return ""


def build_agent_files_context(
    agent_name: str,
    student_rut: str | None = None,
) -> tuple[str, int]:
    """
    Devuelve (bloque de texto para el prompt, cantidad de archivos incluidos).
    Solo archivos subidos en Files; excluye plantillas en documentos/.
    """
    root = storage.agent_folder(agent_name)
    paths = list_context_file_paths(agent_name)
    all_paths = list_all_context_file_paths(agent_name)
    spreadsheet_paths = [p for p in all_paths if p.suffix.lower() in _SPREADSHEET_EXTENSIONS]

    rut_block = ""
    if student_rut and student_rut.strip():
        rut_block, _matched = extract_spreadsheet_hint_for_rut(agent_name, student_rut)

    if not paths and not rut_block:
        return "", 0

    index_lines = [p.relative_to(root).as_posix() for p in all_paths]
    sections: list[str] = []
    total_chars = 0
    included = 0

    if rut_block:
        sections.append(rut_block)
        total_chars += len(rut_block)

    for path in paths:
        rel = path.relative_to(root).as_posix()
        ext = path.suffix.lower()
        try:
            content = extract_file_text(path).strip()
        except Exception as exc:
            content = f"[No se pudo leer el archivo: {exc}]"

        if not content:
            if ext in _TEXT_EXTENSIONS | _BINARY_EXTENSIONS | _SPREADSHEET_EXTENSIONS:
                content = "[Archivo sin texto extraíble o vacío]"
            else:
                continue

        if len(content) > MAX_CHARS_PER_FILE:
            content = content[:MAX_CHARS_PER_FILE] + "\n… [contenido truncado]"

        block = f"### Archivo: {rel}\n\n{content}"
        if total_chars + len(block) > MAX_TOTAL_CHARS:
            remaining = MAX_TOTAL_CHARS - total_chars
            if remaining > 200:
                sections.append(block[:remaining] + "\n… [contexto total truncado]")
            break

        sections.append(block)
        total_chars += len(block)
        included += 1

    if not sections:
        return "", 0

    excel_note = ""
    if spreadsheet_paths:
        excel_names = ", ".join(
            p.relative_to(root).as_posix() for p in spreadsheet_paths[:10]
        )
        extra = len(spreadsheet_paths) - 10
        if extra > 0:
            excel_names += f" … (+{extra} más)"
        excel_note = f"\nExcel en Files ({len(spreadsheet_paths)}): {excel_names}"

    header = (
        "ARCHIVOS DE CONTEXTO (subidos en Files del agente). "
        "Cuando las instrucciones pidan revisar, buscar o usar archivos adjuntos, "
        "utiliza ÚNICAMENTE el contenido de estos archivos como fuente.\n"
        f"Índice de archivos elegibles ({len(index_lines)}; "
        f"hasta {MAX_FILES} priorizados en el cuerpo): "
        + ", ".join(index_lines)
        + excel_note
        + "\n"
    )
    return header + "\n\n---\n\n".join(sections), included
