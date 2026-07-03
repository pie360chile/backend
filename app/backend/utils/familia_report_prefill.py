"""Relleno determinístico de campos de identificación en informe familia (formulario)."""

from __future__ import annotations

import logging
import re
import unicodedata
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

# Tags w:tag de identificación en family_report.docx (content controls)
FAMILIA_IDENTIFICATION_SDT_TAGS = (
    "student_full_name",
    "student_identification_number",
    "student_birth_date",
    "student_age",
    "student_course",
    "student_school",
    "professional_full_name",
    "professional_identification_number",
    "professional_job_position",
    "professional_phone_email",
    "professional_delivered_date_inform",
    "person_full_name",
    "person_identification_number",
    "person_relation_student",
    "person_presence",
)

_NARRATIVE_KEYS = (
    "evaluation_reason",
    "applied_instruments",
    "diagnosis",
    "diagnostic",
    "pedagogical_strengths",
    "pedagogical_support_needs",
    "social_affective_strengths",
    "social_affective_support_needs",
    "health_strengths",
    "health_support_needs",
    "collaborative_work",
    "home_based_description",
    "home_support",
    "school_family_agreements",
    "agreements_commitments",
    "strengths_1",
    "support_needs_1",
    "strengths_2",
    "support_needs_2",
    "strengths_3",
    "support_needs_3",
)


def _normalize_cc_tag(tag: str) -> str:
    t = (tag or "").strip().lower()
    t = unicodedata.normalize("NFKD", t)
    t = "".join(c for c in t if unicodedata.category(c) != "Mn")
    t = re.sub(r"[^a-z0-9]+", "_", t)
    return re.sub(r"_+", "_", t).strip("_")


_NARRATIVE_TAG_NORMS = frozenset(_normalize_cc_tag(k) for k in _NARRATIVE_KEYS)


def _paragraph_visible_text(p_el: Any, qn: Any) -> str:
    return "".join((t.text or "") for t in p_el.iter(qn("w:t")))


def _is_label_paragraph(text: str) -> bool:
    s = (text or "").strip()
    if len(s) < 10:
        return False
    letters = [c for c in s if c.isalpha()]
    if len(letters) < 8:
        return False
    upper = sum(1 for c in letters if c.isupper())
    return (upper / len(letters)) > 0.82


def _clear_paragraph_pagination_locks(ppr: Any, qn: Any, OxmlElement: Any) -> None:
    """Evita saltos forzados y bloques que Word empuja a la página siguiente."""
    for tag in ("w:keepNext", "w:keepLines", "w:pageBreakBefore"):
        for el in list(ppr.findall(qn(tag))):
            ppr.remove(el)
    for el in list(ppr.findall(qn("w:widowControl"))):
        ppr.remove(el)


def _zero_paragraph_spacing(ppr: Any, qn: Any, OxmlElement: Any) -> None:
    for sp in list(ppr.findall(qn("w:spacing"))):
        ppr.remove(sp)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    for ctx in list(ppr.findall(qn("w:contextualSpacing"))):
        ppr.remove(ctx)
    _clear_paragraph_pagination_locks(ppr, qn, OxmlElement)


def _paragraph_has_placeholder(text: str) -> bool:
    low = (text or "").lower()
    return any(
        ph in low
        for ph in (
            "haz clic",
            "pulse aqu",
            "escribir texto",
            "click here",
            "tap here",
            "click or tap",
            "clic aqu",
        )
    )


def _ppr_set_justify(ppr: Any, qn: Any, OxmlElement: Any) -> None:
    """Justificado estándar (w:jc both). La última línea de cada w:p queda a la izquierda."""
    for jc in list(ppr.findall(qn("w:jc"))):
        ppr.remove(jc)
    jc_el = OxmlElement("w:jc")
    jc_el.set(qn("w:val"), "both")
    ppr.append(jc_el)


def _apply_narrative_paragraph_format(ppr: Any, qn: Any, OxmlElement: Any) -> None:
    _ppr_set_justify(ppr, qn, OxmlElement)
    _zero_paragraph_spacing(ppr, qn, OxmlElement)


def _append_narrative_paragraph(
    parent: Any,
    text: str,
    qn: Any,
    OxmlElement: Any,
    ref_r_pr: Any | None,
) -> None:
    from copy import deepcopy

    new_p = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    _apply_narrative_paragraph_format(ppr, qn, OxmlElement)
    new_p.append(ppr)
    r = OxmlElement("w:r")
    if ref_r_pr is not None:
        r.append(deepcopy(ref_r_pr))
    wt = OxmlElement("w:t")
    wt.text = text
    if text.startswith(" ") or text.endswith(" "):
        wt.set(qn("xml:space"), "preserve")
    r.append(wt)
    new_p.append(r)
    parent.append(new_p)


def _split_paragraph_segments(p_el: Any, qn: Any) -> list[str]:
    """Parte un w:p en bloques de texto separados por w:br."""
    segments: list[str] = []
    buf: list[str] = []
    for r in p_el.findall(qn("w:r")):
        wt = r.find(qn("w:t"))
        if wt is not None and wt.text:
            buf.append(wt.text)
        if r.find(qn("w:br")) is not None:
            chunk = "".join(buf).strip()
            if chunk:
                segments.append(chunk)
            buf = []
    tail = "".join(buf).strip()
    if tail:
        segments.append(tail)
    return segments


def _reformat_narrative_sdt_content(sdt_content_el: Any, qn: Any, OxmlElement: Any) -> None:
    """
    Normaliza narrativa: un w:p por bloque, justificado, sin w:br.
    w:br + justify en el mismo párrafo estira líneas cortas con espacios enormes.
    """
    # Plantillas con w:tc dentro de w:sdtContent: reformatear solo dentro de la celda.
    direct_tc = sdt_content_el.find(qn("w:tc"))
    if direct_tc is not None:
        _reformat_narrative_paragraph_container(direct_tc, qn, OxmlElement)
        return

    _reformat_narrative_paragraph_container(sdt_content_el, qn, OxmlElement)


def _reformat_narrative_paragraph_container(container_el: Any, qn: Any, OxmlElement: Any) -> None:
    from copy import deepcopy

    segments: list[str] = []
    ref_r_pr = None

    for p in container_el.iter(qn("w:p")):
        raw = _paragraph_visible_text(p, qn).replace("\r\n", "\n").replace("\r", "\n")
        raw = raw.strip()
        if not raw or _paragraph_has_placeholder(raw) or _is_label_paragraph(raw):
            continue
        if ref_r_pr is None:
            for r in p.iter(qn("w:r")):
                rpr = r.find(qn("w:rPr"))
                if rpr is not None:
                    ref_r_pr = deepcopy(rpr)
                    break
        if any(r.find(qn("w:br")) is not None for r in p.findall(qn("w:r"))):
            segments.extend(_split_paragraph_segments(p, qn))
        else:
            for block in re.split(r"\n\s*\n", raw):
                block = block.strip()
                if block:
                    segments.append(block)

    for p in container_el.iter(qn("w:p")):
        ppr = p.find(qn("w:pPr"))
        if ppr is None:
            ppr = OxmlElement("w:pPr")
            p.insert(0, ppr)
        _apply_narrative_paragraph_format(ppr, qn, OxmlElement)

    if not segments:
        return

    for child in list(container_el):
        if child.tag == qn("w:tcPr"):
            continue
        container_el.remove(child)

    for seg in segments:
        _append_narrative_paragraph(container_el, seg, qn, OxmlElement, ref_r_pr)


def _compact_sdt_content_spacing(sdt_content_el: Any, qn: Any, OxmlElement: Any) -> None:
    """Reformatea párrafos narrativos: justificado, un w:p por bloque (sin w:br)."""
    _reformat_narrative_sdt_content(sdt_content_el, qn, OxmlElement)


def compact_familia_narrative_spacing(docx_path: Path) -> None:
    """Justifica narrativa, separa bloques en w:p distintos (sin w:br) y compacta espaciado."""
    from app.backend.utils.familia_report_formtext import (
        compact_familia_formtext_narrative,
        docx_has_legacy_formtext,
    )
    from app.backend.utils.familia_report_tabla_fill import (
        compact_familia_tabla_narrative,
        docx_is_familia_ministerial_tabla,
        relax_familia_tabla_layout,
    )

    if docx_is_familia_ministerial_tabla(docx_path):
        compact_familia_tabla_narrative(docx_path)
        relax_familia_tabla_layout(docx_path)
        return
    if docx_has_legacy_formtext(docx_path):
        compact_familia_formtext_narrative(docx_path)
        return

    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document(str(docx_path))

    def walk_sdts(parent: Any) -> None:
        for sdt in parent.iter(qn("w:sdt")):
            tag_el = sdt.find(".//" + qn("w:tag"))
            if tag_el is None:
                continue
            tag_n = _normalize_cc_tag(tag_el.get(qn("w:val")) or "")
            if tag_n not in _NARRATIVE_TAG_NORMS:
                continue
            sdt_content = sdt.find(qn("w:sdtContent"))
            if sdt_content is not None:
                _compact_sdt_content_spacing(sdt_content, qn, OxmlElement)

    walk_sdts(doc.element.body)
    for section in doc.sections:
        for hf in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
        ):
            if hf is not None and hf._element is not None:
                walk_sdts(hf._element)

    doc.save(str(docx_path))


W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"
W14_CHECKED_ATTR = f"{{{W14_NS}}}val"


def _evaluation_type_flags(context: dict[str, Any] | None) -> tuple[bool, bool]:
    """Devuelve (ingreso_marcado, reevaluacion_marcada)."""
    eval_type = str((context or {}).get("evaluation_type") or "").strip().lower()
    if eval_type in ("revaluation", "reevaluación", "reevaluacion", "2"):
        return False, True
    return True, False


def _set_sdt_checkbox_mark(sdt: Any, qn: Any, OxmlElement: Any, checked: bool) -> None:
    """Marca checkbox con 'x' sin cuadro (☐/☑)."""
    sdt_pr = sdt.find(qn("w:sdtPr"))
    if sdt_pr is not None:
        checkbox = None
        for child in sdt_pr:
            if child.tag.endswith("}checkbox"):
                checkbox = child
                break
        if checkbox is not None:
            check_el = None
            for child in checkbox:
                if child.tag.endswith("}checked"):
                    check_el = child
                    break
            val_str = "1" if checked else "0"
            if check_el is not None:
                check_el.set(W14_CHECKED_ATTR, val_str)
            else:
                check_el = OxmlElement("w14:checked")
                check_el.set(W14_CHECKED_ATTR, val_str)
                checkbox.append(check_el)

    sdt_content = sdt.find(qn("w:sdtContent"))
    if sdt_content is None:
        return
    wt_list = list(sdt_content.iter(qn("w:t")))
    mark = "x" if checked else ""
    if wt_list:
        wt_list[0].text = mark
        for wt in wt_list[1:]:
            wt.text = ""
    else:
        for p_el in sdt_content.iter(qn("w:p")):
            w_r = p_el.find(qn("w:r"))
            if w_r is None:
                w_r = OxmlElement("w:r")
                p_el.append(w_r)
            wt = w_r.find(qn("w:t"))
            if wt is None:
                wt = OxmlElement("w:t")
                w_r.append(wt)
            wt.text = mark
            break


def fix_familia_motivo_evaluacion_row(
    docx_path: Path,
    student_context: dict[str, Any] | None = None,
) -> None:
    """
    Restaura fila MOTIVO DE LA EVALUACIÓN como plantilla original:
    Evaluación: x de Ingreso   x Reevaluación (una x con cuadro).
    """
    from app.backend.utils.familia_report_formtext import (
        docx_has_legacy_formtext,
        fix_familia_motivo_evaluacion_formtext,
    )
    from app.backend.utils.familia_report_tabla_fill import (
        docx_is_familia_ministerial_tabla,
        fix_familia_motivo_evaluacion_tabla,
    )

    if docx_is_familia_ministerial_tabla(docx_path):
        fix_familia_motivo_evaluacion_tabla(docx_path, student_context)
        return
    if docx_has_legacy_formtext(docx_path):
        fix_familia_motivo_evaluacion_formtext(docx_path, student_context)
        return

    from copy import deepcopy

    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document(str(docx_path))
    if len(doc.tables) < 4 or len(doc.tables[3].rows) < 2:
        return

    is_admission, is_reeval = _evaluation_type_flags(student_context)
    cell = doc.tables[3].rows[1].cells[0]
    if not cell.paragraphs:
        return

    p = cell.paragraphs[0]._element
    p_pr = p.find(qn("w:pPr"))

    eval_sdt = reeval_sdt = None
    bold_r_pr = None
    for child in p:
        if child.tag.endswith("}r"):
            if bold_r_pr is None:
                found = child.find(qn("w:rPr"))
                if found is not None:
                    bold_r_pr = deepcopy(found)
        if child.tag.endswith("}sdt"):
            tag_el = child.find(".//" + qn("w:tag"))
            tag_val = tag_el.get(qn("w:val")) if tag_el is not None else None
            if tag_val == "evaluation":
                eval_sdt = child
            elif tag_val == "reevaluation":
                reeval_sdt = child

    for child in list(p):
        if child is not p_pr:
            p.remove(child)

    def add_run(text: str) -> None:
        r = OxmlElement("w:r")
        if bold_r_pr is not None:
            r.append(deepcopy(bold_r_pr))
        t = OxmlElement("w:t")
        t.text = text
        if text.startswith(" ") or text.endswith(" "):
            t.set(qn("xml:space"), "preserve")
        r.append(t)
        p.append(r)

    add_run("MOTIVO DE LA EVALUACIÓN")
    add_run("                                                      ")
    add_run("Evaluación:")
    add_run(" ")

    if eval_sdt is not None:
        _set_sdt_checkbox_mark(eval_sdt, qn, OxmlElement, is_admission)
        p.append(eval_sdt)
    elif is_admission:
        add_run("x")

    add_run("de Ingreso   ")

    if reeval_sdt is not None:
        _set_sdt_checkbox_mark(reeval_sdt, qn, OxmlElement, is_reeval)
        p.append(reeval_sdt)
    elif is_reeval:
        add_run("x")

    add_run(" Reevaluación")

    doc.save(str(docx_path))


FAMILIA_ANSWER_FONT = "Arial"
FAMILIA_ANSWER_SIZE_HALF_PT = "20"  # Word usa half-points → 10 pt
_CHECKBOX_SDT_TAGS = frozenset({"evaluation", "reevaluation"})


def _apply_arial_10_to_run(r: Any, qn: Any, OxmlElement: Any) -> None:
    r_pr = r.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        r.insert(0, r_pr)
    for tag in ("w:b", "w:bCs", "w:rFonts", "w:sz", "w:szCs"):
        for el in list(r_pr.findall(qn(tag))):
            r_pr.remove(el)
    r_fonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(attr), FAMILIA_ANSWER_FONT)
    r_pr.append(r_fonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), FAMILIA_ANSWER_SIZE_HALF_PT)
    r_pr.append(sz)
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), FAMILIA_ANSWER_SIZE_HALF_PT)
    r_pr.append(sz_cs)


def apply_familia_arial_10_font(docx_path: Path) -> None:
    """Todas las respuestas en content controls: Arial 10, sin negrita."""
    from app.backend.utils.familia_report_formtext import (
        apply_familia_arial_10_to_formtext,
        docx_has_legacy_formtext,
    )
    from app.backend.utils.familia_report_tabla_fill import (
        apply_familia_arial_10_to_tabla,
        docx_is_familia_ministerial_tabla,
    )

    if docx_is_familia_ministerial_tabla(docx_path):
        apply_familia_arial_10_to_tabla(docx_path)
        return
    if docx_has_legacy_formtext(docx_path):
        apply_familia_arial_10_to_formtext(docx_path)
        return

    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document(str(docx_path))

    def walk(parent: Any) -> None:
        for sdt in parent.iter(qn("w:sdt")):
            tag_el = sdt.find(".//" + qn("w:tag"))
            tag_val = tag_el.get(qn("w:val")) if tag_el is not None else None
            for r in sdt.iter(qn("w:r")):
                wt = r.find(qn("w:t"))
                if wt is None or not (wt.text or "").strip():
                    continue
                if tag_val in _CHECKBOX_SDT_TAGS and (wt.text or "").strip() in ("", "x", "X"):
                    _apply_arial_10_to_run(r, qn, OxmlElement)
                    continue
                if tag_val in _CHECKBOX_SDT_TAGS:
                    continue
                _apply_arial_10_to_run(r, qn, OxmlElement)

    walk(doc.element.body)
    for section in doc.sections:
        for hf in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
        ):
            if hf is not None and hf._element is not None:
                walk(hf._element)

    doc.save(str(docx_path))


def build_familia_identification_replacements(context: dict[str, Any]) -> dict[str, str]:
    """Rellena solo los content controls de identificación (tags exactos de la plantilla)."""
    student_full = str(context.get("student_full_name") or "").strip()
    student_social = str(context.get("student_social_name") or "").strip()
    student_id = str(context.get("student_identification_number") or "").strip()
    student_born = str(
        context.get("student_birth_date") or context.get("student_born_date") or ""
    ).strip()
    student_age = str(context.get("student_age") or "").strip()
    student_course = str(context.get("student_course") or "").strip()
    student_school = str(context.get("student_school") or "").strip()

    professional_full = str(
        context.get("professional_full_name")
        or context.get("professional_social_name")
        or ""
    ).strip()
    professional_social = str(context.get("professional_social_name") or "").strip()
    professional_id = str(context.get("professional_identification_number") or "").strip()
    professional_role = str(
        context.get("professional_job_position") or context.get("professional_role") or ""
    ).strip()
    ph = str(context.get("professional_phone") or "").strip()
    em = str(context.get("professional_email") or "").strip()
    professional_contact = f"{ph} / {em}".strip(" / ") if (ph or em) else ""
    report_delivery = str(
        context.get("professional_delivered_date_inform")
        or context.get("report_delivery_date")
        or ""
    ).strip()

    receiver_full = str(
        context.get("person_full_name") or context.get("receiver_full_name") or ""
    ).strip()
    receiver_social = str(context.get("receiver_social_name") or "").strip()
    receiver_id = str(
        context.get("person_identification_number")
        or context.get("receiver_identification_number")
        or ""
    ).strip()
    receiver_relation = str(
        context.get("person_relation_student") or context.get("receiver_relationship") or ""
    ).strip()
    receiver_presence = str(
        context.get("person_presence") or context.get("receiver_presence_of") or ""
    ).strip()
    receiver_phone = str(
        context.get("person_phone") or context.get("receiver_phone") or ""
    ).strip()
    receiver_email = str(
        context.get("person_email") or context.get("receiver_email") or ""
    ).strip()

    no_info = "No informado"

    def _fmt_rut(val: str) -> str:
        from app.backend.utils.agent_student_lookup import _format_rut_display

        raw = (val or "").strip()
        if not raw or raw.lower() == no_info.lower():
            return no_info
        return _format_rut_display(raw)

    replacements: dict[str, str] = {
        "student_full_name": student_full or no_info,
        "student_social_name": student_social or no_info,
        "student_identification_number": _fmt_rut(student_id),
        "student_birth_date": student_born or no_info,
        "student_age": student_age or no_info,
        "student_course": student_course or no_info,
        "student_school": student_school or no_info,
        "professional_full_name": professional_full or no_info,
        "professional_social_name": professional_social or professional_full or no_info,
        "professional_identification_number": _fmt_rut(professional_id),
        "professional_job_position": professional_role or no_info,
        "professional_phone_email": professional_contact or no_info,
        "professional_phone": ph or no_info,
        "professional_email": em or no_info,
        "professional_delivered_date_inform": report_delivery or no_info,
        "person_full_name": receiver_full or no_info,
        "receiver_social_name": receiver_social or receiver_full or no_info,
        "person_identification_number": _fmt_rut(receiver_id),
        "person_relation_student": receiver_relation or no_info,
        "person_presence": receiver_presence or no_info,
        "person_phone": receiver_phone or no_info,
        "person_email": receiver_email or no_info,
    }

    eval_type = str(context.get("evaluation_type") or "").strip().lower()
    if eval_type in ("admission", "admisión", "ingreso", "1"):
        replacements["evaluation"] = "1"
    elif eval_type in ("revaluation", "reevaluación", "reevaluacion", "2"):
        replacements["reevaluation"] = "1"

    for key in ("evaluation_date_1", "evaluation_date_2", "evaluation_date_3"):
        val = str(context.get(key) or "").strip()
        if val:
            replacements[key] = val

    return replacements


def build_familia_form_replacements(context: dict[str, Any]) -> dict[str, str]:
    """Identificación + narrativa desde BD (informe guardado en admin)."""
    return _append_narrative_replacements(
        build_familia_identification_replacements(context), context
    )


def _append_narrative_replacements(
    replacements: dict[str, str], context: dict[str, Any]
) -> dict[str, str]:
    diagnostic = str(context.get("diagnosis") or context.get("diagnostic") or "").strip()
    if diagnostic:
        replacements["diagnostic"] = diagnostic
        replacements["diagnosis"] = diagnostic

    evaluation_reason = str(
        context.get("evaluation_reason") or context.get("applied_instruments") or ""
    ).strip()
    if evaluation_reason:
        replacements["evaluation_reason"] = evaluation_reason

    alias_map = {
        "pedagogical_strengths": ("strengths_1", "fortalezas_1"),
        "pedagogical_support_needs": ("support_needs_1", "necesidades_apoyo_1"),
        "social_affective_strengths": ("strengths_2", "fortalezas_2", "fortalezas_social_afectivo"),
        "social_affective_support_needs": (
            "support_needs_2",
            "necesidades_apoyo_2",
            "necesidades_social_afectivo",
        ),
        "health_strengths": ("strengths_3",),
        "health_support_needs": ("support_needs_3",),
        "collaborative_work": ("educational_supports_school",),
        "home_based_description": ("home_support",),
        "school_family_agreements": ("agreements_commitments",),
    }
    for key in _NARRATIVE_KEYS:
        val = str(context.get(key) or "").strip()
        if not val:
            continue
        replacements[key] = val
        for alias in alias_map.get(key, ()):
            replacements[alias] = val

    return {k: v for k, v in replacements.items() if v}


def postprocess_saved_familia_docx(
    agent_id: str,
    saved: list[dict[str, Any]],
    student_context: dict[str, Any] | None,
    *,
    form_template_path: Path | None = None,
) -> list[dict[str, Any]]:
    """Corrige identificación en .docx generado usando fill_docx_form y datos de BD."""
    if not student_context or not saved:
        return saved

    entry = saved[0]
    file_id = entry.get("id")
    if not file_id:
        return saved

    path = agent_dir(agent_id) / file_id
    if path.suffix.lower() != ".docx":
        return saved

    if not path.is_file() and form_template_path and form_template_path.is_file():
        shutil.copy2(form_template_path, path)

    if not path.is_file():
        return saved

    uses_formtext = False
    try:
        from app.backend.utils.agent_familia_formtext import docx_has_legacy_formtext

        uses_formtext = docx_has_legacy_formtext(path)
    except Exception:
        pass

    if not docx_has_form_controls(path):
        if form_template_path and form_template_path.is_file():
            shutil.copy2(form_template_path, path)
            logger.info(
                "Informe familia: reemplazado formato tablas por plantilla formulario %s",
                form_template_path.name,
            )
            try:
                from app.backend.utils.agent_familia_formtext import docx_has_legacy_formtext

                uses_formtext = docx_has_legacy_formtext(path)
            except Exception:
                uses_formtext = False
        else:
            return saved

    if not docx_has_form_controls(path) and not uses_formtext:
        return saved

    try:
        replacements = build_familia_identification_replacements(student_context)
        from app.backend.utils.familia_report_tabla_fill import (
            docx_is_familia_ministerial_tabla,
            refill_familia_identification_only,
        )

        if docx_is_familia_ministerial_tabla(path):
            filled = refill_familia_identification_only(
                path,
                replacements,
                template_path=form_template_path,
            )
            if filled:
                logger.info(
                    "Identificación re-aplicada en %s (%d campos)",
                    path.name,
                    len(filled),
                )
        else:
            result = fill_familia_docx(path, replacements, path)
            if result.get("status") == "error":
                logger.warning(
                    "Prefill familia falló para %s: %s",
                    path.name,
                    result.get("message"),
                )
            else:
                logger.info("Prefill familia aplicado en %s", path.name)
    except Exception as exc:
        logger.warning("Prefill familia no aplicado en %s: %s", path.name, exc)

    try:
        compact_familia_narrative_spacing(path)
        logger.info("Espaciado narrativo compactado en %s", path.name)
    except Exception as exc:
        logger.warning("Compactación de espaciado falló en %s: %s", path.name, exc)

    try:
        fix_familia_motivo_evaluacion_row(path, student_context)
        logger.info("Fila motivo evaluación corregida en %s", path.name)
    except Exception as exc:
        logger.warning("Corrección motivo evaluación falló en %s: %s", path.name, exc)

    try:
        apply_familia_arial_10_font(path)
        logger.info("Arial 10 aplicado en %s", path.name)
    except Exception as exc:
        logger.warning("Arial 10 no aplicado en %s: %s", path.name, exc)

    return saved
