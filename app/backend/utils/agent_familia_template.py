"""Detección y reglas de plantilla para informe a la familia del agente."""

from __future__ import annotations

import re
import unicodedata
from pathlib import Path
from typing import Any


def _normalize(text: str) -> str:
    folded = unicodedata.normalize("NFKD", text or "")
    asciiish = "".join(ch for ch in folded if not unicodedata.combining(ch))
    return re.sub(r"[^a-z0-9]+", " ", asciiish.lower()).strip()


def is_familia_form_template(display_name: str) -> bool:
    """Plantilla PIE360 con campos de formulario (content controls)."""
    lower = _normalize(display_name.replace("_", " ").replace("-", " "))
    if not lower:
        return False
    if "family report" in lower or "family_report" in lower.replace(" ", "_"):
        return True
    if "family_report" in lower:
        return True
    if "formulario" in lower and "familia" in lower:
        return True
    if "informe familiar" in lower or "informe_familiar" in lower.replace(" ", "_"):
        return True
    if "con formulario" in lower or "tipo formulario" in lower:
        return True
    # Generado desde admin o plantilla PIE360 (no el formato ministerial vacío)
    if "informe familia" in lower or "informe_familia" in lower.replace(" ", "_"):
        if "formato" not in lower:
            return True
    return False


def agent_has_familia_form_template(
    file_names: list[str],
    *,
    file_paths: dict[str, Path] | None = None,
) -> bool:
    for name in file_names or []:
        if is_familia_form_file(name, (file_paths or {}).get(name)):
            return True
    return False


def filter_familia_template_file_names(
    file_names: list[str],
    *,
    file_paths: dict[str, Path] | None = None,
) -> list[str]:
    """Si hay plantilla formulario, excluye FORMATO INFORME DE FAMILIA (tablas)."""
    names = file_names or []
    if not agent_has_familia_form_template(names, file_paths=file_paths):
        return names
    return [
        n
        for n in names
        if not is_familia_tabla_file(n, (file_paths or {}).get(n))
    ]


def build_form_template_supremacy_block(base_filename: str, excluded: list[str]) -> str:
    excluded_txt = ", ".join(f"«{n}»" for n in excluded) if excluded else "(ninguno)"
    return (
        "=== PLANTILLA FORMULARIO — PRIORIDAD SOBRE FORMATO MINISTERIAL ===\n"
        f"OBLIGATORIO: genera el informe copiando y rellenando SOLO «{base_filename}».\n"
        f"PROHIBIDO abrir o completar plantillas de tablas ministeriales: {excluded_txt}.\n"
        "Pasos: (1) lee bloque DATOS EN BASE DE DATOS; (2) shutil.copy plantilla formulario; "
        "(3) rellena content controls / placeholders; (4) completa narrativa desde archivos del caso.\n"
    )


def docx_has_content_controls(path: Path) -> bool:
    """True si el .docx tiene content controls (w:sdt)."""
    try:
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document(str(path))
        count = sum(1 for _ in doc.element.body.iter(qn("w:sdt")))
        return count >= 3
    except Exception:
        return False


def docx_has_form_controls(path: Path) -> bool:
    """True si el .docx es rellenable: SDT, FORMTEXT o tablas ministeriales planas."""
    if docx_has_content_controls(path):
        return True
    try:
        from app.backend.utils.agent_familia_formtext import docx_has_legacy_formtext
        from app.backend.utils.agent_familia_tabla_fill import docx_is_familia_ministerial_tabla

        return docx_has_legacy_formtext(path) or docx_is_familia_ministerial_tabla(path)
    except Exception:
        return False


def is_familia_form_file(display_name: str, disk_path: Path | None = None) -> bool:
    if is_familia_form_template(display_name):
        return True
    if disk_path and disk_path.suffix.lower() == ".docx" and disk_path.is_file():
        return docx_has_form_controls(disk_path)
    return False


def is_familia_tabla_file(display_name: str, disk_path: Path | None = None) -> bool:
    if is_familia_form_file(display_name, disk_path):
        return False
    return is_familia_tabla_template(display_name)


def is_familia_tabla_template(display_name: str) -> bool:
    """FORMATO INFORME DE FAMILIA.docx (tablas etiqueta/valor)."""
    lower = _normalize(display_name.replace("_", " ").replace("-", " "))
    return ("formato" in lower and "familia" in lower) or "informe de familia" in lower


def familia_form_template_priority(display_name: str) -> int:
    lower = _normalize(display_name.replace("_", " ").replace("-", " "))
    if "family_report" in lower:
        return 0
    if is_familia_form_template(display_name):
        return 1
    if "cartilla" in lower:
        return 2
    if is_familia_tabla_template(display_name):
        return 10
    return 20


def pick_familia_base_template(
    file_names: list[str] | None,
    *,
    agent_id: str | None = None,
    file_paths: dict[str, Path] | None = None,
) -> tuple[str | None, str]:
    """Devuelve (nombre_archivo, tipo) con tipo form | tabla | none."""
    names = file_names or []
    form: list[str] = []
    tabla: list[str] = []
    for name in names:
        disk = (file_paths or {}).get(name)
        if is_familia_form_file(name, disk):
            form.append(name)
        elif is_familia_tabla_file(name, disk):
            tabla.append(name)
    if form:
        form.sort(key=familia_form_template_priority)
        return form[0], "form"
    if tabla:
        return tabla[0], "tabla"
    return None, "none"


def resolve_familia_template_from_rows(
    agent_id: str,
    rows: list[Any],
) -> tuple[str | None, str]:
    """Clasifica plantilla familia inspeccionando archivos en disco."""
    from app.backend.utils.agent_files import agent_dir

    names: list[str] = []
    paths: dict[str, Path] = {}
    for row in rows:
        name = getattr(row, "display_name", None) or ""
        if not name:
            continue
        names.append(name)
        fid = getattr(row, "id", None)
        if fid:
            paths[name] = agent_dir(agent_id) / fid
    return pick_familia_base_template(names, file_paths=paths)


def resolve_form_template_path(agent_id: str, rows: list[Any]) -> Path | None:
    """Ruta en disco de la plantilla formulario familia, si existe."""
    from app.backend.utils.agent_files import agent_dir

    base, kind = resolve_familia_template_from_rows(agent_id, rows)
    if kind != "form" or not base:
        return None
    for row in rows:
        name = getattr(row, "display_name", None) or ""
        fid = getattr(row, "id", None)
        if name == base and fid:
            path = agent_dir(agent_id) / fid
            if path.is_file():
                return path
    return None


def build_no_fabrication_rules() -> str:
    """Prohibición de inventar datos no presentes en BD o archivos."""
    return (
        "- NO INVENTAR DATOS (OBLIGATORIO): usa únicamente información que esté en "
        "«DATOS DEL ESTUDIANTE EN BASE DE DATOS», en los archivos del caso o en el rol. "
        "Prohibido fabricar fechas, calendarizaciones, nombres, RUT, diagnósticos o frases "
        "genéricas del tipo «Según calendarización del establecimiento», "
        "«antecedente no informado» como párrafo, o textos explicativos en celdas de fecha.\n"
        "- Si falta un dato narrativo en doc.tables[3]: indícalo según el rol o deja la celda vacía "
        "si es una fecha. Nunca sustituyas una fecha faltante por un párrafo.\n"
        "- NO rellenes ni modifiques identificación (doc.tables[1] y doc.tables[2]) si el .docx base "
        "ya viene con esos datos desde PIE360.\n"
        "- Fechas de seguimiento → doc.tables[3] fila 17: solo fechas DD/MM/YYYY "
        "en las celdas de valor (cols 4-9), una por celda asignada. "
        "Usa evaluation_date_1, evaluation_date_2, evaluation_date_3 de la BD si existen; "
        "si no existen, deja esas celdas vacías.\n"
    )


def build_typography_preserve_rules() -> str:
    """Conservar tipografía de la plantilla; narrativa justificada."""
    return (
        "- TIPOGRAFÍA RESPUESTAS (OBLIGATORIO): todo el texto que escribes en campos del formulario "
        "(identificación arriba y narrativa abajo) debe ser Arial 10 pt, sin negrita. "
        "Copia rPr de la plantilla pero fuerza w:rFonts ascii/hAnsi=Arial y w:sz val=20.\n"
        "- ALINEACIÓN NARRATIVA (OBLIGATORIO): todo párrafo narrativo del informe "
        "(instrumentos, diagnóstico, fortalezas, necesidades de apoyo, acuerdos, etc.) "
        "debe ir JUSTIFICADO (w:jc both / WD_ALIGN_PARAGRAPH.JUSTIFY). "
        "Campos cortos de identificación (nombre, RUT, fecha suelta) pueden ir a la izquierda.\n"
        "- TIPOGRAFÍA DE PLANTILLA (OBLIGATORIO): NO cambies fuente, tamaño, negrita, "
        "color ni alineación del párrafo/campo. Copia el formato del primer run existente "
        "en la plantilla (rPr: w:rFonts, w:sz, w:color, w:b). "
        "PROHIBIDO usar p.text = ... o cell.text = ... (destruye runs y cambia la letra).\n"
        "- Al rellenar content controls (w:sdt), escribe solo en w:t dentro del w:sdtContent "
        "preservando w:rPr del run; si no hay run, clónalo del párrafo vecino de la plantilla.\n"
        "- ESPACIADO NARRATIVO (OBLIGATORIO): los dos párrafos de cada campo van separados por "
        "UN solo salto (w:br con run.add_break() dentro del MISMO w:p), NO con add_paragraph() "
        "ni con \\n\\n. Si usas w:p distintos, space_before=Pt(0) y space_after=Pt(0) en ambos. "
        "PROHIBIDO dejar párrafos vacíos entre bloques (causan huecos grandes en Word).\n"
        "- Texto arriba en celda: space_before/space_after = 0; elimina párrafos vacíos sobrantes; "
        "no centres verticalmente.\n"
        + build_no_fabrication_rules()
    )


def build_redaction_min_paragraphs_rules() -> str:
    """Reglas de extensión mínima para textos narrativos del informe."""
    return (
        "- REDACCIÓN EXTENSA (OBLIGATORIO): cada apartado narrativo debe tener "
        "exactamente 2 párrafos LARGOS y desarrollados (no oraciones sueltas ni párrafos "
        "de 2-3 líneas). Aplica a motivo de evaluación, instrumentos, diagnóstico, "
        "fortalezas, necesidades de apoyo (pedagógico, social/afectivo y salud), "
        "trabajo colaborativo, apoyos en el hogar, acuerdos y secciones del rol.\n"
        "- Extensión mínima por párrafo: 6 a 10 oraciones completas (aprox. 120-200 palabras). "
        "Desarrolla contexto escolar, observaciones del expediente, implicancias pedagógicas "
        "y orientaciones concretas. Un párrafo breve de 1-4 oraciones es INACEPTABLE.\n"
        "- Párrafo 1: antecedentes, contexto e instrumentos/fuentes. "
        "Párrafo 2: hallazgos, interpretación e implicancias para apoyos y seguimiento.\n"
        "- Extrae contenido de la cartilla, evaluación psicopedagógica y archivos del caso; "
        "no resumas en frases genéricas.\n"
        "- Separa los dos párrafos con UN solo salto de línea (w:br), sin párrafo vacío intermedio "
        "ni espacio extra. Revisa que no queden huecos grandes entre párrafos antes de exportar.\n"
        "- Antes de exportar, revisa cada campo narrativo: si algún párrafo tiene menos "
        "de 6 oraciones, amplíalo antes de guardar.\n"
        + build_typography_preserve_rules()
    )


def build_familia_hybrid_scope_rules(base_doc_name: str) -> str:
    """Delimita qué tablas/filas puede editar GPT cuando PIE360 ya pre-llenó identificación."""
    return (
        "=== ÁMBITO DE EDICIÓN (OBLIGATORIO — NO SALIR DE AQUÍ) ===\n"
        f"Trabaja sobre «{base_doc_name}» tal como viene de PIE360.\n"
        "PROHIBIDO modificar, borrar o reescribir:\n"
        "  • doc.tables[0] — texto introductorio Decreto 170\n"
        "  • doc.tables[1] — IDENTIFICACIÓN DEL ESTUDIANTE (nombre, RUT, curso, establecimiento…)\n"
        "  • doc.tables[2] — IDENTIFICACIÓN del profesional y del apoderado/receptor\n"
        "  • doc.tables[4] — firmas\n"
        "  • Fila «MOTIVO DE LA EVALUACIÓN» (checkboxes ingreso/reevaluación — ya vienen de PIE360)\n"
        "TU ÚNICO ÁMBITO DE TRABAJO es doc.tables[3] «RESULTADOS DE LA EVALUACIÓN», "
        "desde INSTRUMENTOS APLICADOS hacia abajo:\n"
        "  • fila 3 col 0: instrumentos aplicados (applied_instruments)\n"
        "  • fila 5 col 0: diagnóstico NEE (diagnostic / diagnosis)\n"
        "  • fila 8: fortalezas pedagógicas (col 0) y necesidades de apoyo (col 3)\n"
        "  • fila 10: fortalezas y necesidades ámbito social/afectivo\n"
        "  • fila 12 col 0: trabajo colaborativo / apoyos en el establecimiento\n"
        "  • fila 14 col 0: apoyos en el hogar\n"
        "  • fila 16 col 0: acuerdos escuela-familia\n"
        "  • fila 17 cols 4-9: fechas de seguimiento (solo DD/MM/YYYY o vacío)\n"
        "No recorras doc.tables[1] ni doc.tables[2] con bucles genéricos. "
        "No uses cell.text = ... sobre celdas de identificación.\n"
    )


def build_familia_narrative_enrichment_rules(base_doc_name: str) -> str:
    """Instrucciones cuando PIE360 ya entregó base con identificación rellena."""
    return (
        "=== INFORME FAMILIA — BASE PIE360 + REDACCIÓN GPT (OBLIGATORIO) ===\n"
        f"El archivo «{base_doc_name}» ya tiene la IDENTIFICACIÓN superior rellena desde la base de datos "
        "(estudiante, profesional, apoderado, fechas, checkboxes de evaluación). NO modifiques esos campos.\n"
        + build_familia_hybrid_scope_rules(base_doc_name)
        + "Tu trabajo exclusivo: redactar los apartados narrativos de doc.tables[3] "
        "(desde instrumentos aplicados hacia abajo) según el ROL DEL AGENTE, usando la cartilla técnica, "
        "la evaluación psicopedagógica y los demás archivos del caso en el contenedor.\n"
        "El ROL DEL AGENTE define qué decir, cómo estructurar cada sección, extensión, tono técnico-pedagógico "
        "y criterios de completitud. Cada campo narrativo debe cumplir literalmente lo que el rol exige para "
        "ese apartado; no resumas ni omitas secciones que el rol mencione.\n"
        "Pasos:\n"
        f"  1) Abre «{base_doc_name}» (NO copies otra plantilla ni uses FORMATO INFORME DE FAMILIA.docx).\n"
        "  2) Lee cartilla, informe psicopedagógico y documentos del estudiante.\n"
        "  3) Completa SOLO las celdas narrativas vacías de doc.tables[3] según el rol.\n"
        "  4) Guarda UN solo .docx final y expórtalo.\n"
        "Mapeo rol → campos (solo parte inferior, doc.tables[3]):\n"
        "  - Instrumentos aplicados → applied_instruments (fila 3)\n"
        "  - Diagnóstico / diagnóstico NEE → diagnostic, diagnosis (fila 5)\n"
        "  - Fortalezas pedagógicas → pedagogical_strengths, strengths_1 (fila 8 col 0)\n"
        "  - Necesidades de apoyo pedagógico → pedagogical_support_needs, support_needs_1 (fila 8 col 3)\n"
        "  - Fortalezas social/afectivo → social_affective_strengths, strengths_2 (fila 10 col 0)\n"
        "  - Necesidades social/afectivo → social_affective_support_needs, support_needs_2 (fila 10 col 3)\n"
        "  - Trabajo colaborativo / apoyos en el establecimiento → collaborative_work (fila 12)\n"
        "  - Apoyos en el hogar → home_based_description, home_support (fila 14)\n"
        "  - Acuerdos escuela-familia → school_family_agreements, agreements_commitments (fila 16)\n"
        "OBLIGATORIO: ningún campo narrativo puede quedar vacío, con «Haz clic o pulse aquí…» ni con texto "
        "genérico. Extrae contenido concreto del expediente; si el rol exige citar cartilla técnica, Decreto 170 "
        "o criterios PIE, incorpóralos en la redacción.\n"
        "PROHIBIDO: borrar o reescribir identificación ya completada; tocar doc.tables[0], [1], [2] o [4]; "
        "usar formato ministerial de tablas distinto al base; entregar el documento sin redactar doc.tables[3]; "
        "modificar la fila «MOTIVO DE LA EVALUACIÓN».\n"
        + build_redaction_min_paragraphs_rules()
    )


def build_familia_form_rules(base_filename: str) -> str:
    return (
        f"- PLANTILLA FORMULARIO (OBLIGATORIO): usa «{base_filename}» como base del informe. "
        f"Copia ese archivo (shutil.copy) y trabaja sobre la copia. "
        "NO uses FORMATO INFORME DE FAMILIA.docx ni reconstruyas el documento.\n"
        "- La plantilla tiene campos de formulario Word (content controls; a veces muestran "
        "«Haz clic o pulse aquí para escribir texto»). Rellénalos sin mover tablas ni cambiar estilos.\n"
        "- Lee primero el expediente del estudiante (docx del caso, cartilla, etc.) y extrae los datos. "
        "Si el formulario ya trae identificación rellena, consérvala.\n"
        "- Identificación (tags w:tag exactos — NO modificar si ya vienen rellenos desde PIE360):\n"
        "    student_full_name, student_identification_number, student_birth_date, student_age, "
        "student_course, student_school, professional_full_name, professional_identification_number, "
        "professional_job_position, professional_phone_email, professional_delivered_date_inform, "
        "person_full_name, person_identification_number, person_relation_student, person_presence.\n"
        "    Contenido narrativo: evaluation_reason, diagnostic, strengths_1, support_needs_1, "
        "pedagogical_strengths / strengths_1; pedagogical_support_needs / support_needs_1; "
        "social_affective_strengths / strengths_2; social_affective_support_needs / support_needs_2; "
        "health_strengths / strengths_3; health_support_needs / support_needs_3; "
        "collaborative_work; home_based_description / home_support; "
        "school_family_agreements / agreements_commitments.\n"
        "- Nombre y Rut van en sus campos respectivos; el Rol/cargo NUNCA recibe el RUT.\n"
        "- Rellena con python-docx recorriendo w:sdt (w:tag) y asignando texto al w:sdtContent, "
        "o reemplazando {clave}, [clave] y <<clave>> en párrafos y celdas.\n"
        "- Patrón mínimo para content controls:\n"
        "    from docx import Document\n"
        "    from docx.oxml.ns import qn\n"
        "  def rellenar_formulario(ruta, reemplazos):\n"
        "      doc = Document(ruta)\n"
        "      def norm(s): return re.sub(r'[^a-z0-9]+','', (s or '').lower())\n"
        "      tags = {norm(k): str(v or '') for k,v in reemplazos.items()}\n"
        "      for sdt in doc.element.body.iter(qn('w:sdt')):\n"
        "          tag = sdt.find('.//'+qn('w:tag'))\n"
        "          if tag is None: continue\n"
        "          key = norm(tag.get(qn('w:val')))\n"
        "          if key not in tags: continue\n"
        "          content = sdt.find(qn('w:sdtContent'))\n"
        "          if content is None: continue\n"
        "          for t in content.iter(qn('w:t')): t.text = ''\n"
        "          ts = list(content.iter(qn('w:t')))\n"
        "          if ts: ts[0].text = tags[key]\n"
        "      for p in doc.paragraphs:\n"
        "          t = p.text\n"
        "          for k,v in reemplazos.items():\n"
        "              for fmt in ('{%s}','[%s]','<<%s>>'):\n"
        "                  t = t.replace(fmt % k, str(v or ''))\n"
        "          if t != p.text: p.text = t\n"
        "      for tbl in doc.tables:\n"
        "          for row in tbl.rows:\n"
        "              for cell in row.cells:\n"
        "                  for p in cell.paragraphs:\n"
        "                      t = p.text\n"
        "                      for k,v in reemplazos.items():\n"
        "                          for fmt in ('{%s}','[%s]','<<%s>>'):\n"
        "                              t = t.replace(fmt % k, str(v or ''))\n"
        "                      if t != p.text: p.text = t\n"
        "      doc.save(ruta)\n"
        "- Conserva tipografía de la plantilla; no uses cell.text sobre filas de rótulo.\n"
        "- Para narrativa: 2 párrafos largos (6-10 oraciones c/u) por campo; "
        "texto justificado (w:jc both), Arial 10 pt.\n"
        + build_typography_preserve_rules()
    )
