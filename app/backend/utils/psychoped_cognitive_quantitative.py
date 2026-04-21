"""
Evaluación psicopedagógica (doc 27): matriz cuantitativa tipo EVALÚA (PD/X/DT/E.T.M/PT × RE..RP)
e índices generales (IGC–IGM). Parseo flexible del JSON en BD y gráfico PT en PNG (Pillow).
"""
from __future__ import annotations

import json
import math
import os
import re
import unicodedata
from dataclasses import dataclass
from typing import Any, Dict, List, Optional, Tuple

INDICATORS: List[str] = ["RE", "PA", "OP", "MA", "CL", "PPM", "EF", "OR", "GR", "OF", "CN", "RP"]
GENERAL_LABELS: List[str] = ["IGC", "IGL", "IGE", "IGM"]

# Marcador en una línea propia; tras generar el DOCX se reemplaza por la imagen del gráfico.
CHART_PLACEHOLDER = "[[PSYCHOPED_PT_CHART]]"

# Valor no vacío para el CC `ac` hasta insertar la tabla (si ac="" Word desenvuelve el control).
AC_CONTENT_CONTROL_HOLD = "\ufeff"


def _safe_float(raw: Any) -> Optional[float]:
    if raw is None:
        return None
    if isinstance(raw, bool):
        return None
    if isinstance(raw, (int, float)):
        return float(raw)
    txt = str(raw).strip().replace(",", ".")
    if txt == "" or txt.lower() in ("null", "none", "n/a", "na", "-"):
        return None
    try:
        return float(txt)
    except (ValueError, TypeError):
        return None


def _loads_json_flex(raw: Any) -> Any:
    if raw is None:
        return None
    if isinstance(raw, (dict, list)):
        return raw
    s = str(raw).strip()
    if not s:
        return None
    try:
        v = json.loads(s)
        if isinstance(v, str):
            s2 = v.strip()
            if s2.startswith("{") or s2.startswith("["):
                try:
                    return json.loads(s2)
                except json.JSONDecodeError:
                    return v
        return v
    except json.JSONDecodeError:
        return None


def _norm_key(s: str) -> str:
    t = (s or "").strip().lower()
    t = "".join(c for c in unicodedata.normalize("NFKD", t) if unicodedata.category(c) != "Mn")
    t = re.sub(r"[^a-z0-9]+", "", t)
    return t


def _get_from_mapping(m: Dict[str, Any], *candidates: str) -> Any:
    if not isinstance(m, dict) or not m:
        return None
    norm_map = {_norm_key(str(k)): v for k, v in m.items()}
    for c in candidates:
        nk = _norm_key(c)
        if nk in norm_map:
            return norm_map[nk]
    return None


def _row_aliases(canonical: str) -> Tuple[str, ...]:
    if canonical == "PD":
        return ("PD", "pd", "puntuacion_directa", "puntuación_directa", "directa")
    if canonical == "X":
        return ("X", "x", "media", "media_baremo", "media_del_baremo", "baremo")
    if canonical == "DT":
        return ("DT", "dt", "desviacion_tipica", "desviación_típica", "desviacion", "desviación")
    if canonical == "ETM":
        return ("E.T.M", "E.T.M.", "ETM", "etm", "error_tipico", "error_tipico_medida", "error_típico")
    if canonical == "PT":
        return ("PT", "pt", "puntuacion_tipica", "puntuación_típica", "puntuacion_típica")
    return (canonical,)


def _empty_rows() -> Dict[str, List[Optional[float]]]:
    return {k: [None] * 12 for k in ("PD", "X", "DT", "ETM", "PT")}


def _fill_row_from_dict(row_key: str, source: Dict[str, Any], out: Dict[str, List[Optional[float]]]) -> None:
    row = out[row_key]
    for alias in _row_aliases(row_key):
        node = source.get(alias)
        if node is None and alias != row_key:
            node = _get_from_mapping(source, alias)
        if node is None:
            continue
        if isinstance(node, list):
            for i in range(min(12, len(node))):
                row[i] = _safe_float(node[i]) if row[i] is None else row[i]
            return
        if isinstance(node, dict):
            for i, ind in enumerate(INDICATORS):
                if row[i] is not None:
                    continue
                row[i] = _safe_float(
                    _get_from_mapping(node, ind, ind.lower(), ind.upper())
                )
            return


def _parse_transpose_tests(matrix_obj: Dict[str, Any], out: Dict[str, List[Optional[float]]]) -> bool:
    hit = False
    for ind in INDICATORS:
        node = matrix_obj.get(ind) or matrix_obj.get(ind.lower()) or matrix_obj.get(ind.upper())
        if not isinstance(node, dict):
            continue
        hit = True
        idx = INDICATORS.index(ind)
        for rk in ("PD", "X", "DT", "ETM", "PT"):
            for alias in _row_aliases(rk):
                v = _get_from_mapping(node, alias)
                if v is not None:
                    fv = _safe_float(v)
                    if out[rk][idx] is None:
                        out[rk][idx] = fv
                    break
    return hit


def _parse_tests_array(matrix_obj: Dict[str, Any], out: Dict[str, List[Optional[float]]]) -> bool:
    arr = matrix_obj.get("tests") or matrix_obj.get("pruebas") or matrix_obj.get("columns")
    if not isinstance(arr, list) or not arr:
        return False
    for item in arr:
        if not isinstance(item, dict):
            continue
        code = (
            item.get("code")
            or item.get("codigo")
            or item.get("abbr")
            or item.get("abreviatura")
            or item.get("test")
            or item.get("nombre")
        )
        if code is None:
            continue
        c = str(code).strip().upper()
        if c not in INDICATORS:
            continue
        idx = INDICATORS.index(c)
        for rk in ("PD", "X", "DT", "ETM", "PT"):
            if out[rk][idx] is not None:
                continue
            for alias in _row_aliases(rk):
                v = item.get(alias) if alias in item else _get_from_mapping(item, alias)
                if v is not None:
                    out[rk][idx] = _safe_float(v)
                    break
    return any(v is not None for row in out.values() for v in row)


def _parse_rows_container(matrix_obj: Dict[str, Any], out: Dict[str, List[Optional[float]]]) -> bool:
    rows_obj = matrix_obj.get("rows")
    if not isinstance(rows_obj, dict):
        return False
    any_hit = False
    for rk in ("PD", "X", "DT", "ETM", "PT"):
        for alias in _row_aliases(rk):
            nested = rows_obj.get(alias)
            if nested is None:
                continue
            any_hit = True
            if isinstance(nested, list):
                for i in range(min(12, len(nested))):
                    if out[rk][i] is None:
                        out[rk][i] = _safe_float(nested[i])
            elif isinstance(nested, dict):
                for i, ind in enumerate(INDICATORS):
                    if out[rk][i] is None:
                        out[rk][i] = _safe_float(
                            _get_from_mapping(nested, ind, ind.lower(), ind.upper())
                        )
            break
    return any_hit


def _parse_list_of_row_dicts(obj: List[Any], out: Dict[str, List[Optional[float]]]) -> bool:
    if not isinstance(obj, list):
        return False
    hit = False
    for item in obj:
        if not isinstance(item, dict):
            continue
        label = (
            item.get("row")
            or item.get("label")
            or item.get("metric")
            or item.get("tipo")
            or item.get("name")
        )
        if label is None:
            continue
        lk = _norm_key(str(label))
        rk_map = {
            "pd": "PD",
            "x": "X",
            "dt": "DT",
            "etm": "ETM",
            "pt": "PT",
            "e.t.m": "ETM",
            "etm.": "ETM",
        }
        rk = rk_map.get(lk)
        if rk is None:
            if "puntuaciondirecta" in lk or ("puntuacion" in lk and "directa" in lk):
                rk = "PD"
            elif "mediadelbaremo" in lk or lk == "media":
                rk = "X"
            elif "desviacion" in lk:
                rk = "DT"
            elif "error" in lk and "tipico" in lk:
                rk = "ETM"
            elif "puntuaciontipica" in lk or lk == "pt":
                rk = "PT"
        if rk is None:
            continue
        hit = True
        vals = item.get("values") or item.get("valores") or item.get("data") or item.get("cells")
        if isinstance(vals, list):
            for i in range(min(12, len(vals))):
                if out[rk][i] is None:
                    out[rk][i] = _safe_float(vals[i])
        elif isinstance(vals, dict):
            for i, ind in enumerate(INDICATORS):
                if out[rk][i] is None:
                    out[rk][i] = _safe_float(
                        _get_from_mapping(vals, ind, ind.lower(), ind.upper())
                    )
    return hit


def _normalize_matrix(matrix_raw: Any) -> Dict[str, List[Optional[float]]]:
    out = _empty_rows()
    obj = _loads_json_flex(matrix_raw)
    if obj is None:
        return out
    if isinstance(obj, list):
        _parse_list_of_row_dicts(obj, out)
        return out
    if not isinstance(obj, dict):
        return out

    inner = obj.get("matrix") or obj.get("matriz") or obj.get("cuantitativo") or obj.get("data")
    if isinstance(inner, (dict, list)) and not any(k in obj for k in INDICATORS + ["rows", "tests", "pruebas"]):
        return _normalize_matrix(inner)

    # Formato plano usado por frontend/BD:
    # quant_PD_RE, quant_X_PA, quant_DT_OP, quant_ETM_MA, quant_PT_CL, etc.
    # También soporta variantes de ETM: ETM / E_T_M / E.T.M
    if isinstance(obj, dict):
        flat_hit = False
        for i, ind in enumerate(INDICATORS):
            for rk in ("PD", "X", "DT", "ETM", "PT"):
                key_variants = (
                    f"quant_{rk}_{ind}",
                    f"quant_{rk.lower()}_{ind.lower()}",
                    f"quant_{rk}_{ind.lower()}",
                    f"quant_{rk.lower()}_{ind}",
                )
                if rk == "ETM":
                    key_variants += (
                        f"quant_E_T_M_{ind}",
                        f"quant_E_T_M_{ind.lower()}",
                        f"quant_E.T.M_{ind}",
                        f"quant_E.T.M_{ind.lower()}",
                    )
                val = None
                for k in key_variants:
                    if k in obj:
                        val = obj.get(k)
                        break
                if val is None:
                    val = _get_from_mapping(obj, *key_variants)
                fv = _safe_float(val)
                if fv is not None:
                    out[rk][i] = fv
                    flat_hit = True
        if flat_hit:
            # Completar PT si falta (PD - X)/DT
            for i in range(12):
                if out["PT"][i] is None:
                    pdv, xv, dtv = out["PD"][i], out["X"][i], out["DT"][i]
                    if pdv is not None and xv is not None and dtv is not None and abs(dtv) > 1e-12:
                        out["PT"][i] = (pdv - xv) / dtv
            return out

    for rk in ("PD", "X", "DT", "ETM", "PT"):
        _fill_row_from_dict(rk, obj, out)

    if not any(v is not None for row in out.values() for v in row):
        _parse_rows_container(obj, out)
    if not any(v is not None for row in out.values() for v in row):
        _parse_transpose_tests(obj, out)
    if not any(v is not None for row in out.values() for v in row):
        _parse_tests_array(obj, out)

    # PT = (PD - X) / DT si falta
    for i in range(12):
        if out["PT"][i] is None:
            pdv, xv, dtv = out["PD"][i], out["X"][i], out["DT"][i]
            if pdv is not None and xv is not None and dtv is not None and abs(dtv) > 1e-12:
                out["PT"][i] = (pdv - xv) / dtv

    return out


def _parse_general_raw(raw: Any) -> Dict[str, Optional[float]]:
    out: Dict[str, Optional[float]] = {k: None for k in GENERAL_LABELS}
    obj = _loads_json_flex(raw)
    if obj is None:
        return out
    if isinstance(obj, list):
        for item in obj:
            if not isinstance(item, dict):
                continue
            k = item.get("key") or item.get("clave") or item.get("code") or item.get("label")
            v = item.get("value") or item.get("valor") or item.get("pt") or item.get("puntuacion_tipica")
            if k is None:
                continue
            kk = str(k).strip().upper()
            if kk in out:
                out[kk] = _safe_float(v)
        return out
    if isinstance(obj, dict):
        for lab in GENERAL_LABELS:
            node = obj.get(lab) or obj.get(lab.lower()) or _get_from_mapping(obj, lab)
            if isinstance(node, dict):
                out[lab] = _safe_float(
                    _get_from_mapping(node, "pt", "PT", "puntuacion_tipica", "puntuación_típica", "value", "valor")
                )
            else:
                out[lab] = _safe_float(node)
        return out
    return out


def _fill_general_from_pt(pt: List[Optional[float]], gen: Dict[str, Optional[float]]) -> Dict[str, Optional[float]]:
    def _mean(vals: List[Optional[float]]) -> Optional[float]:
        xs = [float(v) for v in vals if v is not None]
        if not xs:
            return None
        return sum(xs) / len(xs)

    slices = {
        "IGC": (0, 4),
        "IGL": (4, 7),
        "IGE": (7, 10),
        "IGM": (10, 12),
    }
    out = dict(gen)
    for lab, (a, b) in slices.items():
        if out.get(lab) is None:
            out[lab] = _mean(pt[a:b])
    return out


@dataclass
class CognitiveQuantitativeParsed:
    rows: Dict[str, List[Optional[float]]]
    general: Dict[str, Optional[float]]

    def has_matrix_numbers(self) -> bool:
        return any(v is not None for row in self.rows.values() for v in row)

    def has_chart_numbers(self) -> bool:
        return any(v is not None for v in self.rows["PT"]) or any(
            v is not None for v in self.general.values()
        )

    def has_table_numbers(self) -> bool:
        if any(v is not None for v in self.general.values()):
            return True
        for rk in ("PD", "X", "DT", "ETM", "PT"):
            if any(v is not None for v in self.rows[rk]):
                return True
        return False


def parse_evalua_psychoped_matrices(eval_data: Dict[str, Any]) -> Optional[CognitiveQuantitativeParsed]:
    if not isinstance(eval_data, dict):
        return None
    rows = _normalize_matrix(eval_data.get("cognitive_quantitative_matrix"))
    general = _parse_general_raw(eval_data.get("cognitive_general_scales"))
    general = _fill_general_from_pt(rows["PT"], general)
    parsed = CognitiveQuantitativeParsed(rows=rows, general=general)
    if not parsed.has_matrix_numbers() and not any(v is not None for v in general.values()):
        return None
    return parsed


def _fmt_table_cell(v: Optional[float]) -> str:
    if v is None:
        return "-"
    fv = float(v)
    if math.isfinite(fv) and abs(fv - round(fv)) < 1e-9 and abs(fv) < 1e6:
        return str(int(round(fv)))
    s = f"{fv:.2f}"
    if "." in s:
        s = s.rstrip("0").rstrip(".")
    return s


def build_cognitive_quantitative_table_text(parsed: CognitiveQuantitativeParsed) -> str:
    lines: List[str] = []
    lines.append("CUADRO CUANTITATIVO (HABILIDADES COGNITIVAS Y COMUNICATIVAS)")
    lines.append("PRUEBAS: " + " | ".join(INDICATORS))
    for rk, label in (
        ("PD", "PD"),
        ("X", "X"),
        ("DT", "DT"),
        ("ETM", "E.T.M"),
        ("PT", "PT"),
    ):
        cells = " | ".join(_fmt_table_cell(v) for v in parsed.rows[rk])
        lines.append(f"{label}: {cells}")
    lines.append("")
    lines.append(
        "ESCALAS GENERALES: "
        + " | ".join(f"{k}: {_fmt_table_cell(parsed.general.get(k))}" for k in GENERAL_LABELS)
    )
    lines.append("")
    lines.append("Gráfico: Puntuación típica (PT) — pruebas e índices generales (EVALÚA).")
    return "\n".join(lines).strip()


def _find_sdt_element_by_tag(root_el: Any, tag: str) -> Optional[Any]:
    """Busca el primer w:sdt cuyo w:tag @w:val coincide (sin distinguir mayúsculas)."""
    try:
        from docx.oxml.ns import qn
    except ImportError:
        return None
    want = (tag or "").strip().lower()
    for sdt in root_el.iter(qn("w:sdt")):
        sdt_pr = sdt.find(qn("w:sdtPr"))
        if sdt_pr is None:
            continue
        tag_el = sdt_pr.find(qn("w:tag"))
        if tag_el is None:
            continue
        val = (tag_el.get(qn("w:val")) or "").strip().lower()
        if val == want:
            return sdt
    return None


def _set_table_borders(table: Any, color_hex: str = "4F81BD") -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    for el in list(tbl_pr.findall(qn("w:tblBorders"))):
        tbl_pr.remove(el)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        borders.append(el)
    tbl_pr.append(borders)


def _tbl_set_pct_width(table: Any, pct: int = 5000) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    for el in list(tbl_pr.findall(qn("w:tblW"))):
        tbl_pr.remove(el)
    tw = OxmlElement("w:tblW")
    tw.set(qn("w:w"), str(pct))
    tw.set(qn("w:type"), "pct")
    tbl_pr.append(tw)


def _set_cell_shading(cell: Any, fill_hex: str = "D9E2F3") -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    for el in list(tc_pr.findall(qn("w:shd"))):
        tc_pr.remove(el)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _style_cell_text(
    cell: Any,
    *,
    bold: bool = False,
    size_pt: float = 9,
    align_center: bool = False,
) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    for para in cell.paragraphs:
        if align_center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.bold = bold
            run.font.size = Pt(size_pt)


def inject_evalua_matrix_word_table(docx_path: str, parsed: CognitiveQuantitativeParsed) -> bool:
    """
    Sustituye el contenido del control de contenido `ac` por una tabla Word nativa
    (12 pruebas + 4 índices), estilo similar a EVALÚA (cabeceras y bordes azules).
    """
    if not parsed or not parsed.has_table_numbers():
        return False
    try:
        from docx import Document
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Pt
    except ImportError:
        return False

    doc = Document(docx_path)
    sdt = _find_sdt_element_by_tag(doc.element.body, "ac")
    if sdt is None:
        return False
    sdt_content = sdt.find(qn("w:sdtContent"))
    if sdt_content is None:
        return False

    for child in list(sdt_content):
        sdt_content.remove(child)

    # 7 filas: cabeceras (0–1) + PD, X, DT, E.T.M., PT (2–6). Sin fila vacía solo para composición.
    n_rows = 7
    n_cols = 17
    table = doc.add_table(rows=n_rows, cols=n_cols)
    tbl_el = table._tbl
    parent = tbl_el.getparent()
    if parent is not None:
        parent.remove(tbl_el)
    sdt_content.append(tbl_el)

    _tbl_set_pct_width(table, 5000)
    _set_table_borders(table, "4F81BD")

    # Tipografías compactas (cabeceras + cuerpo).
    sz_band = 8.0  # PRUEBAS / ESCALAS GENERALES*
    sz_head = 7.0  # RE, PA, … IGC, IGL, …
    sz_label = 7.0  # PD, X, … en primera columna
    sz_val = 7.0  # números y guiones
    sz_comp = 6.5  # escrito bajo índices (RE, PA, …)

    comp_text = {
        "IGC": "RE, PA, OP, MA",
        "IGL": "CL, PPM, EF",
        "IGE": "OR, GR, OF",
        "IGM": "CN, RP",
    }

    r0 = table.rows[0].cells
    r0[0].text = ""
    r0[1].merge(r0[12])
    r0[1].text = "PRUEBAS"
    r0[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _style_cell_text(r0[1], bold=True, size_pt=sz_band, align_center=True)
    r0[13].merge(r0[16])
    r0[13].text = "ESCALAS GENERALES*"
    r0[13].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _style_cell_text(r0[13], bold=True, size_pt=sz_band, align_center=True)
    _set_cell_shading(r0[1])
    _set_cell_shading(r0[13])

    r1 = table.rows[1].cells
    r1[0].text = ""
    for i, abbr in enumerate(INDICATORS):
        r1[1 + i].text = abbr
        r1[1 + i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _style_cell_text(r1[1 + i], bold=True, size_pt=sz_head, align_center=True)
        _set_cell_shading(r1[1 + i])
    for j, lab in enumerate(GENERAL_LABELS):
        c = r1[13 + j]
        c.text = lab
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _style_cell_text(c, bold=True, size_pt=sz_head, align_center=True)
        _set_cell_shading(c)

    # ESCALAS GENERALES (cols 13–16): una celda fusionada en vertical para las filas PD–E.T.M. (2–5);
    # el escrito (RE, PA, …) arriba y un guion para el bloque de métricas (sin «PD:» ni filas sueltas).
    row_defs: List[Tuple[str, str]] = [
        ("PD", "PD"),
        ("X", "X"),
        ("DT", "DT"),
        ("ETM", "E.T.M."),
        ("PT", "PT"),
    ]
    for r_idx, (rk, label) in enumerate(row_defs, start=2):
        row = table.rows[r_idx].cells
        row[0].text = label
        _style_cell_text(row[0], bold=True, size_pt=sz_label)
        for j in range(12):
            row[1 + j].text = _fmt_table_cell(parsed.rows[rk][j])
            row[1 + j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _style_cell_text(row[1 + j], size_pt=sz_val, align_center=True)
        if r_idx == 6:
            for j, lab in enumerate(GENERAL_LABELS):
                c = row[13 + j]
                c.text = _fmt_table_cell(parsed.general.get(lab))
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                _style_cell_text(c, size_pt=sz_val, align_center=True)
                _set_cell_shading(c, "E7EDF7")

    for j, lab in enumerate(GENERAL_LABELS):
        top = table.cell(2, 13 + j)
        bottom = table.cell(5, 13 + j)
        top.merge(bottom)
        top.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        _set_cell_shading(top, "E7EDF7")
        top.text = ""
        p0 = top.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_comp = p0.add_run(comp_text[lab])
        run_comp.bold = True
        run_comp.font.size = Pt(sz_comp)
        p_dash = top.add_paragraph()
        p_dash.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p_dash.add_run("-")
        rr.font.size = Pt(sz_val)

    doc.save(docx_path)
    return True


def inject_image_into_content_control_by_tag(
    docx_path: str,
    tag: str,
    image_path: str,
    *,
    width_inches: float = 6.0,
) -> bool:
    """
    Vacía el primer SDT con w:tag @w:val == `tag` e inserta una imagen centrada.
    Sustituye la tabla EVALÚA generada en el control `ac` del doc 27 cuando el usuario
    sube una captura en lugar de usar la matriz en el Word.
    """
    if not image_path or not os.path.isfile(image_path):
        return False
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches
        from docx.text.paragraph import Paragraph
    except ImportError:
        return False

    doc = Document(docx_path)
    sdt = _find_sdt_element_by_tag(doc.element.body, tag)
    if sdt is None:
        return False
    sdt_content = sdt.find(qn("w:sdtContent"))
    if sdt_content is None:
        return False

    for child in list(sdt_content):
        sdt_content.remove(child)

    new_p = OxmlElement("w:p")
    sdt_content.append(new_p)
    paragraph = Paragraph(new_p, doc._body)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    doc.save(docx_path)
    return True


def _diamond(draw: Any, xy: Tuple[float, float], r: float, fill: Tuple[int, int, int]) -> None:
    x, y = xy
    pts = [(x, y - r), (x + r, y), (x, y + r), (x - r, y)]
    draw.polygon(pts, outline=fill, fill=fill)


def render_evalua_pt_line_chart_png(parsed: CognitiveQuantitativeParsed, out_png_path: str) -> bool:
    if not parsed.has_chart_numbers():
        return False
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        return False

    W, H = 1500, 470
    img = Image.new("RGB", (W, H), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    font_path = os.path.join(os.environ.get("WINDIR", "C:\\Windows"), "Fonts", "arial.ttf")
    try:
        font = ImageFont.truetype(font_path, 20)
        font_title = ImageFont.truetype(font_path, 22)
        font_small = ImageFont.truetype(font_path, 16)
    except OSError:
        font = ImageFont.load_default()
        font_title = font
        font_small = font

    title = "Análisis cuantitativo — PT (pruebas e índices generales)"
    draw.text((80, 18), title, fill=(20, 20, 20), font=font_title)

    # Margen inferior extra para etiquetas del eje X + título "Puntuaciones" + leyenda sin solaparse.
    x0, y0, x1, y1 = 88, 72, W - 72, H - 148
    plot_w, plot_h = x1 - x0, y1 - y0

    def vy(val: Optional[float]) -> Optional[float]:
        if val is None or not math.isfinite(float(val)):
            return None
        vc = max(-3.0, min(3.0, float(val)))
        return y0 + (3.0 - vc) / 6.0 * plot_h

    # rejilla Y cada 0.5
    gy = -3.0
    while gy <= 3.001:
        yy = y0 + (3.0 - gy) / 6.0 * plot_h
        draw.line([(x0, yy), (x1, yy)], fill=(230, 235, 245), width=1)
        t = str(int(gy)) if abs(gy - round(gy)) < 1e-9 else f"{gy:.1f}".replace(".", ",")
        draw.text((x0 - 72, yy - 10), t, fill=(60, 60, 60), font=font_small)
        gy += 0.5

    draw.line([(x0, y0), (x0, y1), (x1, y1), (x1, y0)], fill=(40, 40, 40), width=1)

    gap = plot_w * 0.045
    w_tests_span = plot_w * 0.74
    w_gen_span = plot_w * 0.22
    xs_tests = [x0 + i * (w_tests_span / 11.0) for i in range(12)]
    divider = x0 + w_tests_span + gap / 2.0
    draw.line([(divider, y0), (divider, y1)], fill=(160, 170, 190), width=2)
    xs_gen = [divider + gap / 2.0 + j * (w_gen_span / 3.0) for j in range(4)]

    xs = xs_tests + xs_gen
    vals: List[Optional[float]] = list(parsed.rows["PT"]) + [parsed.general.get(k) for k in GENERAL_LABELS]
    labels = INDICATORS + GENERAL_LABELS

    blue = (0, 102, 204)
    for i in range(len(vals) - 1):
        v0, v1 = vals[i], vals[i + 1]
        y_a, y_b = vy(v0), vy(v1)
        if y_a is not None and y_b is not None:
            draw.line([(xs[i], y_a), (xs[i + 1], y_b)], fill=blue, width=3)

    for i, v in enumerate(vals):
        yp = vy(v)
        if yp is None:
            continue
        _diamond(draw, (xs[i], yp), 7, blue)

    y_cat = y1 + 10
    cat_h = 0
    for lab in labels:
        lb0 = draw.textbbox((0, 0), lab, font=font)
        cat_h = max(cat_h, lb0[3] - lb0[1])
    for i, lab in enumerate(labels):
        lb = draw.textbbox((0, 0), lab, font=font)
        lw = lb[2] - lb[0]
        draw.text((xs[i] - lw / 2, y_cat), lab, fill=(30, 30, 30), font=font)

    # Separación clara entre etiquetas del eje X (RE, OR, …) y el título «Puntuaciones».
    gap_below_categories = 26
    axis_title = "Puntuaciones"
    tb = draw.textbbox((0, 0), axis_title, font=font_small)
    tw = tb[2] - tb[0]
    y_axis_title = y_cat + cat_h + gap_below_categories
    draw.text(((x0 + x1 - tw) / 2, y_axis_title), axis_title, fill=(50, 50, 50), font=font_small)

    leg = "Puntuación típica (PT)"
    lb2 = draw.textbbox((0, 0), leg, font=font_small)
    lw2 = lb2[2] - lb2[0]
    y_leg = y_axis_title + (tb[3] - tb[1]) + 14
    draw.text(((x0 + x1 - lw2) / 2, y_leg), leg, fill=blue, font=font_small)

    img.save(out_png_path, format="PNG")
    return True


def insert_chart_placeholder_paragraph_image(docx_path: str, png_path: str, placeholder: str = CHART_PLACEHOLDER) -> bool:
    if not os.path.isfile(png_path):
        return False
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.text.paragraph import Paragraph
        from docx.shared import Inches
    except ImportError:
        return False

    doc = Document(docx_path)
    marker = placeholder
    replaced = False

    def _replace_in_paragraph_element(p_el: Any) -> None:
        nonlocal replaced
        full = "".join((t.text or "") for t in p_el.iter(qn("w:t")))
        if marker not in full:
            return
        p = Paragraph(p_el, doc._body)
        p.clear()
        r = p.add_run()
        # Reducido para que quepa dentro del cuadro `acg` del template.
        r.add_picture(png_path, width=Inches(5.9))
        replaced = True

    for p_el in doc.element.body.iter(qn("w:p")):
        _replace_in_paragraph_element(p_el)
        if replaced:
            break

    if not replaced:
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _replace_in_paragraph_element(para._element)
                        if replaced:
                            break
                    if replaced:
                        break
                if replaced:
                    break
            if replaced:
                break

    if replaced:
        doc.save(docx_path)
    return replaced


def strip_chart_placeholder_from_docx(docx_path: str, placeholder: str = CHART_PLACEHOLDER) -> None:
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.text.paragraph import Paragraph
    except ImportError:
        return

    doc = Document(docx_path)
    changed = False

    def _strip_p_el(p_el: Any) -> None:
        nonlocal changed
        full = "".join((t.text or "") for t in p_el.iter(qn("w:t")))
        if placeholder not in full:
            return
        p = Paragraph(p_el, doc._body)
        txt = p.text.replace(placeholder, "").strip()
        p.clear()
        if txt:
            p.add_run(txt)
        changed = True

    for p_el in doc.element.body.iter(qn("w:p")):
        _strip_p_el(p_el)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _strip_p_el(para._element)
    if changed:
        doc.save(docx_path)
