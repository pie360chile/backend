"""Relleno de campos legacy FORMTEXT (cuadros grises) en family_report.docx ministerial."""

from __future__ import annotations

import logging
import re
from copy import deepcopy
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

# (tabla, fila, columna, clave lógica) — una celda por grupo fusionado
FAMILIA_FORMTEXT_SLOTS: tuple[tuple[int, int, int, str], ...] = (
    (1, 3, 0, "student_full_name"),
    (1, 3, 3, "student_identification_number"),
    (1, 5, 3, "student_birth_date"),
    (2, 3, 0, "professional_full_name"),
    (2, 3, 4, "professional_identification_number"),
    (2, 5, 4, "professional_job_position"),
    (2, 10, 0, "person_full_name"),
    (2, 10, 4, "person_identification_number"),
    (2, 12, 4, "person_phone"),
    (2, 14, 0, "person_email"),
    (3, 3, 0, "applied_instruments"),
    (3, 5, 0, "diagnostic"),
    (3, 8, 0, "pedagogical_strengths"),
    (3, 8, 3, "pedagogical_support_needs"),
    (3, 10, 0, "social_affective_strengths"),
    (3, 10, 4, "social_affective_support_needs"),
    (3, 12, 0, "collaborative_work"),
    (3, 14, 0, "home_based_description"),
    (3, 16, 0, "school_family_agreements"),
    (3, 17, 3, "evaluation_date_1"),
    (3, 17, 4, "evaluation_date_2"),
    (3, 17, 5, "evaluation_date_3"),
    (3, 17, 6, "evaluation_date_4"),
    (3, 17, 7, "evaluation_date_5"),
    (3, 17, 8, "evaluation_date_6"),
    (3, 17, 9, "evaluation_date_7"),
    (3, 17, 10, "evaluation_date_8"),
)

_KEY_ALIASES: dict[str, tuple[str, ...]] = {
    "student_birth_date": ("student_born_date",),
    "professional_full_name": ("professional_social_name",),
    "professional_job_position": ("professional_role",),
    "person_full_name": ("receiver_full_name",),
    "person_identification_number": ("receiver_identification_number",),
    "person_phone": ("receiver_phone",),
    "person_email": ("receiver_email",),
    "diagnostic": ("diagnosis",),
    "applied_instruments": ("evaluation_reason",),
    "pedagogical_strengths": ("strengths_1",),
    "pedagogical_support_needs": ("support_needs_1",),
    "social_affective_strengths": ("strengths_2",),
    "social_affective_support_needs": ("support_needs_2",),
    "home_based_description": ("home_support",),
    "school_family_agreements": ("agreements_commitments",),
    "student_social_name": ("student_social_name",),
    "professional_social_name": ("professional_social_name",),
    "professional_email": ("professional_email",),
    "receiver_social_name": ("receiver_social_name", "person_social_name"),
}

_NARRATIVE_SLOT_KEYS = frozenset(
    {
        "applied_instruments",
        "diagnostic",
        "pedagogical_strengths",
        "pedagogical_support_needs",
        "social_affective_strengths",
        "social_affective_support_needs",
        "collaborative_work",
        "home_based_description",
        "school_family_agreements",
    }
)


def docx_has_legacy_formtext(path: Path) -> bool:
    """True si el .docx usa campos Word FORMTEXT (cuadros grises)."""
    try:
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document(str(path))
        count = 0
        for instr in doc.element.body.iter(qn("w:instrText")):
            if instr.text and "FORMTEXT" in instr.text:
                count += 1
                if count >= 3:
                    return True
        return False
    except Exception:
        return False


def _paragraph_has_formtext(p_el: Any, qn: Any) -> bool:
    for instr in p_el.iter(qn("w:instrText")):
        if instr.text and "FORMTEXT" in instr.text:
            return True
    return False


def _resolve_replacement(key: str, replacements: dict[str, str]) -> str:
    for candidate in (key, *_KEY_ALIASES.get(key, ())):
        val = replacements.get(candidate)
        if val is not None and str(val).strip():
            return str(val)
    return ""


def _set_formtext_paragraph_value(
    p_el: Any,
    value: str,
    qn: Any,
    OxmlElement: Any,
) -> bool:
    """Escribe valor en la zona de resultado de un campo FORMTEXT."""
    children = list(p_el)
    separate_idx = end_idx = None
    for idx, child in enumerate(children):
        if child.tag != qn("w:r"):
            continue
        fld = child.find(qn("w:fldChar"))
        if fld is None:
            continue
        ftype = fld.get(qn("w:fldCharType"))
        if ftype == "separate":
            separate_idx = idx
        elif ftype == "end" and separate_idx is not None:
            end_idx = idx
            break

    if separate_idx is None or end_idx is None:
        return False

    result_runs = [c for c in children[separate_idx + 1 : end_idx] if c.tag == qn("w:r")]
    text = value or ""

    ref_r_pr = None
    if result_runs:
        rpr = result_runs[0].find(qn("w:rPr"))
        if rpr is not None:
            ref_r_pr = deepcopy(rpr)
        for run in result_runs:
            p_el.remove(run)

    new_run = OxmlElement("w:r")
    if ref_r_pr is not None:
        new_run.append(ref_r_pr)
    wt = OxmlElement("w:t")
    wt.text = text
    if text.startswith(" ") or text.endswith(" "):
        wt.set(qn("xml:space"), "preserve")
    new_run.append(wt)
    p_el.insert(end_idx, new_run)
    return True


def _get_cell_element(doc: Any, table_idx: int, row_idx: int, col_idx: int, qn: Any) -> Any | None:
    """Obtiene w:tc por columna de cuadrícula (respeta celdas fusionadas)."""
    tbls = doc.element.body.findall(qn("w:tbl"))
    if table_idx >= len(tbls):
        return None
    rows = tbls[table_idx].findall(qn("w:tr"))
    if row_idx >= len(rows):
        return None

    grid_col = 0
    for tc_el in rows[row_idx].findall(qn("w:tc")):
        span_el = tc_el.find(qn("w:tcPr") + "/" + qn("w:gridSpan"))
        if span_el is None:
            span_el = tc_el.find(".//" + qn("w:gridSpan"))
        span = int(span_el.get(qn("w:val"))) if span_el is not None else 1
        if grid_col <= col_idx < grid_col + span:
            return tc_el
        grid_col += span
    return None


def _iter_cell_paragraphs(tc_el: Any, qn: Any):
    for p_el in tc_el.findall(qn("w:p")):
        yield p_el


def _fill_cell_formtext_element(
    tc_el: Any,
    value: str,
    qn: Any,
    OxmlElement: Any,
) -> int:
    filled = 0
    for p_el in _iter_cell_paragraphs(tc_el, qn):
        if _paragraph_has_formtext(p_el, qn):
            if _set_formtext_paragraph_value(p_el, value, qn, OxmlElement):
                filled += 1
    return filled


def _fill_cell_formtext(
    cell: Any,
    value: str,
    qn: Any,
    OxmlElement: Any,
) -> int:
    filled = 0
    for para in cell.paragraphs:
        if _paragraph_has_formtext(para._element, qn):
            if _set_formtext_paragraph_value(para._element, value, qn, OxmlElement):
                filled += 1
    return filled


def fill_familia_formtext_fields(
    template_path: str | Path,
    replacements: dict[str, str],
    output_path: str | Path,
) -> dict[str, Any]:
    """Rellena family_report.docx ministerial (FORMTEXT por posición en tabla)."""
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    template_path = Path(template_path)
    output_path = Path(output_path)
    if template_path.resolve() != output_path.resolve():
        import shutil

        shutil.copy2(template_path, output_path)

    doc = Document(str(output_path))
    tbl_count = len(doc.element.body.findall(qn("w:tbl")))
    if tbl_count < 4:
        return {"status": "error", "message": "Plantilla familia sin tablas esperadas"}

    seen_slots: set[tuple[int, int, int]] = set()
    filled_keys: list[str] = []

    for table_idx, row_idx, col_idx, key in FAMILIA_FORMTEXT_SLOTS:
        slot = (table_idx, row_idx, col_idx)
        if slot in seen_slots:
            continue
        seen_slots.add(slot)

        tc_el = _get_cell_element(doc, table_idx, row_idx, col_idx, qn)
        if tc_el is None:
            continue

        value = _resolve_replacement(key, replacements)
        count = _fill_cell_formtext_element(tc_el, value, qn, OxmlElement)
        if count:
            filled_keys.append(key)

    doc.save(str(output_path))
    logger.info("FORMTEXT familia: %d campos rellenados en %s", len(filled_keys), output_path.name)
    return {"status": "success", "filled_keys": filled_keys}


def fix_familia_motivo_evaluacion_formtext(
    docx_path: Path,
    student_context: dict[str, Any] | None = None,
) -> None:
    """Marca ingreso o reevaluación con «x» en la fila MOTIVO (plantilla FORMTEXT)."""
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    from app.backend.utils.familia_report_prefill import (
        FAMILIA_CHECKBOX_CHECKED_MARK,
        _evaluation_type_flags,
    )

    doc = Document(str(docx_path))
    if len(doc.tables) < 4 or len(doc.tables[3].rows) < 2:
        return

    is_admission, is_reeval = _evaluation_type_flags(student_context)
    from docx.oxml.ns import qn

    col = 2 if is_admission else (7 if is_reeval else None)
    if col is None:
        return

    tc_el = _get_cell_element(doc, 3, 1, col, qn)
    if tc_el is None:
        return

    for p_el in _iter_cell_paragraphs(tc_el, qn):
        parts = []
        for wt in p_el.iter(qn("w:t")):
            parts.append(wt.text or "")
        raw = "".join(parts).strip()
        if raw.startswith(FAMILIA_CHECKBOX_CHECKED_MARK) or raw.startswith("☒"):
            return
        if "☐" in raw:
            new_text = raw.replace("☐", FAMILIA_CHECKBOX_CHECKED_MARK, 1)
        elif raw.startswith("x ") or raw.startswith("x\n"):
            new_text = raw.replace("x", FAMILIA_CHECKBOX_CHECKED_MARK, 1)
        else:
            new_text = f"{FAMILIA_CHECKBOX_CHECKED_MARK} {raw}".strip() if raw else FAMILIA_CHECKBOX_CHECKED_MARK
        for wt in list(p_el.iter(qn("w:t"))):
            parent = wt.getparent()
            if parent is not None:
                parent.remove(wt)
        runs = [c for c in p_el if c.tag == qn("w:r")]
        if runs:
            wt = OxmlElement("w:t")
            wt.text = new_text
            runs[0].append(wt)
            for run in runs[1:]:
                p_el.remove(run)
        else:
            r = OxmlElement("w:r")
            wt = OxmlElement("w:t")
            wt.text = new_text
            r.append(wt)
            p_el.append(r)
        break

    doc.save(str(docx_path))


def apply_familia_arial_10_to_formtext(docx_path: Path) -> None:
    """Arial 10 en textos de resultado de campos FORMTEXT."""
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    from app.backend.utils.familia_report_prefill import (
        FAMILIA_ANSWER_FONT,
        FAMILIA_ANSWER_SIZE_HALF_PT,
        _apply_arial_10_to_run,
    )

    doc = Document(str(docx_path))
    tbls = doc.element.body.findall(qn("w:tbl"))

    def walk_paragraph(p_el: Any) -> None:
        if not _paragraph_has_formtext(p_el, qn):
            return
        in_result = False
        for child in p_el:
            if child.tag != qn("w:r"):
                continue
            fld = child.find(qn("w:fldChar"))
            if fld is not None:
                ftype = fld.get(qn("w:fldCharType"))
                if ftype == "separate":
                    in_result = True
                    continue
                if ftype == "end":
                    break
            if in_result:
                wt = child.find(qn("w:t"))
                if wt is not None and (wt.text or "").strip():
                    _apply_arial_10_to_run(child, qn, OxmlElement)

    for tbl_el in tbls:
        for tr_el in tbl_el.findall(qn("w:tr")):
            for tc_el in tr_el.findall(qn("w:tc")):
                for p_el in _iter_cell_paragraphs(tc_el, qn):
                    walk_paragraph(p_el)

    doc.save(str(docx_path))


def compact_familia_formtext_narrative(docx_path: Path) -> None:
    """Compacta narrativa en celdas FORMTEXT (justificado, sin w:br)."""
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    from app.backend.utils.familia_report_prefill import (
        _apply_narrative_paragraph_format,
        _is_label_paragraph,
        _paragraph_has_placeholder,
    )

    doc = Document(str(docx_path))
    tbls = doc.element.body.findall(qn("w:tbl"))
    if len(tbls) < 4:
        return

    slot_cells: set[tuple[int, int, int]] = set()
    for table_idx, row_idx, col_idx, key in FAMILIA_FORMTEXT_SLOTS:
        if key not in _NARRATIVE_SLOT_KEYS:
            continue
        slot_cells.add((table_idx, row_idx, col_idx))

    def _paragraph_result_text(p_el: Any) -> str:
        in_result = False
        parts: list[str] = []
        for child in p_el:
            if child.tag != qn("w:r"):
                continue
            fld = child.find(qn("w:fldChar"))
            if fld is not None:
                ftype = fld.get(qn("w:fldCharType"))
                if ftype == "separate":
                    in_result = True
                    continue
                if ftype == "end":
                    break
            if in_result:
                for wt in child.iter(qn("w:t")):
                    parts.append(wt.text or "")
        return "".join(parts)

    for table_idx, tbl_el in enumerate(tbls):
        for row_idx, tr_el in enumerate(tbl_el.findall(qn("w:tr"))):
            for col_idx, tc_el in enumerate(tr_el.findall(qn("w:tc"))):
                if (table_idx, row_idx, col_idx) not in slot_cells:
                    continue
                for p_el in _iter_cell_paragraphs(tc_el, qn):
                    if not _paragraph_has_formtext(p_el, qn):
                        continue
                    raw = _paragraph_result_text(p_el).replace("\r\n", "\n").replace("\r", "\n")
                    raw = raw.strip()
                    if not raw or _paragraph_has_placeholder(raw) or _is_label_paragraph(raw):
                        continue
                    segments = [b.strip() for b in re.split(r"\n\s*\n", raw) if b.strip()]
                    if len(segments) > 1:
                        _set_formtext_paragraph_value(p_el, segments[0], qn, OxmlElement)
                    ppr = p_el.find(qn("w:pPr"))
                    if ppr is None:
                        ppr = OxmlElement("w:pPr")
                        p_el.insert(0, ppr)
                    _apply_narrative_paragraph_format(ppr, qn, OxmlElement)

    doc.save(str(docx_path))
