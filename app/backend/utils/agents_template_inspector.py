"""Inspect fields in Word/PDF templates for Agents."""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from docx import Document

try:
    import fitz
except ImportError:
    fitz = None


_PLACEHOLDER_RE = re.compile(
    r"\{([^{}]+)\}|\[([^\[\]]+)\]|<<([^<>]+)>>|\{\{([^{}]+)\}\}"
)


def _docx_xml_roots(doc: Document) -> list[Any]:
    roots: list[Any] = [doc.element.body]
    for section in doc.sections:
        for hf in (section.header, section.footer):
            if hf is not None and hasattr(hf, "_element"):
                roots.append(hf._element)
    return roots


def _detect_docx_content_control_tags(doc: Document) -> set[str]:
    """Tags w:tag / w:alias de controles de contenido (SDT) en Word."""
    from docx.oxml.ns import qn

    found: set[str] = set()
    for root in _docx_xml_roots(doc):
        for sdt in root.iter(qn("w:sdt")):
            sdt_pr = sdt.find(qn("w:sdtPr"))
            if sdt_pr is None:
                continue
            tag_el = sdt_pr.find(qn("w:tag"))
            if tag_el is not None:
                tag_val = (tag_el.get(qn("w:val")) or "").strip()
                if tag_val:
                    found.add(tag_val)
                    continue
            alias_el = sdt_pr.find(qn("w:alias"))
            if alias_el is not None:
                alias_val = (alias_el.get(qn("w:val")) or "").strip()
                if alias_val:
                    found.add(alias_val)
    return found


def _detect_docx_legacy_form_fields(doc: Document) -> set[str]:
    """Campos de formulario heredados (w:ffData / w:name)."""
    from docx.oxml.ns import qn

    found: set[str] = set()
    for root in _docx_xml_roots(doc):
        for ff_data in root.iter(qn("w:ffData")):
            name_el = ff_data.find(qn("w:name"))
            if name_el is None:
                continue
            name = (name_el.get(qn("w:val")) or "").strip()
            if name:
                found.add(name)
    return found


def _detect_docx_text_placeholders(doc: Document) -> set[str]:
    found: set[str] = set()
    chunks: list[str] = []

    for paragraph in doc.paragraphs:
        chunks.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)

    for text in chunks:
        for match in _PLACEHOLDER_RE.finditer(text):
            field = next(g for g in match.groups() if g)
            cleaned = field.strip()
            if cleaned:
                found.add(cleaned)
    return found


def detect_docx_fields(path: Path) -> list[str]:
    doc = Document(str(path))
    found: set[str] = set()
    found.update(_detect_docx_text_placeholders(doc))
    found.update(_detect_docx_content_control_tags(doc))
    found.update(_detect_docx_legacy_form_fields(doc))
    return sorted(found)


def detect_pdf_fields(path: Path) -> list[str]:
    if fitz is None:
        return []
    found: set[str] = set()
    pdf = fitz.open(str(path))
    try:
        for page in pdf:
            widgets = page.widgets()
            if widgets:
                for widget in widgets:
                    name = (widget.field_name or "").strip()
                    if name:
                        found.add(name)
            text = page.get_text("text") or ""
            for match in _PLACEHOLDER_RE.finditer(text):
                field = next(g for g in match.groups() if g)
                cleaned = field.strip()
                if cleaned:
                    found.add(cleaned)
    finally:
        pdf.close()
    return sorted(found)


def inspect_template_fields(path: Path, format_type: str) -> dict[str, Any]:
    if not path.exists():
        raise ValueError("Template not found.")
    fmt = format_type.lower()
    if fmt == "docx":
        fields = detect_docx_fields(path)
    elif fmt == "pdf":
        fields = detect_pdf_fields(path)
    else:
        raise ValueError("Unsupported format. Use .docx or .pdf.")
    return {
        "formatType": fmt,
        "fields": fields,
        "fieldCount": len(fields),
    }


def fields_to_json(fields: list[str]) -> str:
    return json.dumps(fields, ensure_ascii=False)

def fields_from_json(raw: str | None) -> list[str]:
    if not raw:
        return []
    try:
        parsed = json.loads(raw)
        return parsed if isinstance(parsed, list) else []
    except json.JSONDecodeError:
        return []
