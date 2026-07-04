"""Prueba: checkboxes vacíos del Informe Familia deben conservar ☐."""

from __future__ import annotations

import shutil
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from app.backend.utils.agent_v2_familia_fill import fill_familia_template
from app.backend.utils.familia_report_prefill import (
    FAMILIA_CHECKBOX_CHECKED_MARK,
    FAMILIA_CHECKBOX_UNCHECKED,
    apply_familia_checkbox_states,
    ensure_familia_checkbox_boxes_visible,
)


def _add_sdt_w14(paragraph, tag: str, checked: bool = False) -> None:
    val = "1" if checked else "0"
    sdt = parse_xml(
        '<w:sdt xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml">'
        f"<w:sdtPr>"
        f'<w:tag w:val="{tag}"/>'
        f'<w14:checkbox><w14:checked w14:val="{val}"/></w14:checkbox>'
        f"</w:sdtPr>"
        f"<w:sdtContent><w:p><w:r><w:t></w:t></w:r></w:p>"
        f"</w:sdtContent></w:sdt>"
    )
    paragraph._p.addnext(sdt)


def _add_sdt(paragraph, tag: str, text: str = "☐") -> None:
    sdt = parse_xml(
        f'<w:sdt xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f"<w:sdtPr><w:tag w:val=\"{tag}\"/></w:sdtPr>"
        f"<w:sdtContent><w:p><w:r><w:t>{text}</w:t></w:r></w:p>"
        f"</w:sdtContent></w:sdt>"
    )
    paragraph._p.addnext(sdt)


def _count_symbol(docx_path: Path, symbol: str) -> int:
    doc = Document(str(docx_path))
    return sum((t.text or "").count(symbol) for t in doc.element.body.iter(qn("w:t")))


def _build_fixture(path: Path) -> None:
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Apoderado/a titular"
    p0 = table.cell(0, 1).paragraphs[0]
    _add_sdt(p0, "primary", "☐")
    table.cell(1, 0).text = "Evaluación de Ingreso"
    p1 = table.cell(1, 1).paragraphs[0]
    _add_sdt(p1, "evaluation", "☐")
    p1b = table.cell(1, 1).add_paragraph()
    _add_sdt_w14(p1b, "reevaluation")
    p1c = table.cell(1, 1).add_paragraph()
    _add_sdt_w14(p1c, "primary")
    doc.add_paragraph("Campo narrativo")
    p2 = doc.paragraphs[-1]
    sdt2 = OxmlElement("w:sdt")
    sdt_pr = OxmlElement("w:sdtPr")
    tag = OxmlElement("w:tag")
    tag.set(qn("w:val"), "diagnostic")
    sdt_pr.append(tag)
    sdt2.append(sdt_pr)
    sdt_content = OxmlElement("w:sdtContent")
    p_el = OxmlElement("w:p")
    r_el = OxmlElement("w:r")
    t_el = OxmlElement("w:t")
    t_el.text = "Haz clic o pulse aquí para escribir texto."
    r_el.append(t_el)
    p_el.append(r_el)
    sdt_content.append(p_el)
    sdt2.append(sdt_content)
    p2._p.addnext(sdt2)
    doc.save(str(path))


def main() -> int:
    tmp = Path(tempfile.mkdtemp())
    template = tmp / "template.docx"
    output = tmp / "output.docx"
    _build_fixture(template)

    before = _count_symbol(template, FAMILIA_CHECKBOX_UNCHECKED)
    assert before >= 2, f"fixture debe tener ☐, tiene {before}"

    # w14 con w:t vacío debe recuperar ☐ tras postproceso
    w14_only = tmp / "w14.docx"
    doc_w14 = Document()
    pw = doc_w14.add_paragraph()
    _add_sdt_w14(pw, "substitute")
    doc_w14.save(str(w14_only))
    apply_familia_checkbox_states(w14_only, {})
    ensure_familia_checkbox_boxes_visible(w14_only)
    w14_count = _count_symbol(w14_only, FAMILIA_CHECKBOX_UNCHECKED)
    print(f"checkbox_empty w14_vacio={w14_count}")
    if w14_count < 1:
        print("FALLO: w14 vacio sin cuadro visible")
        shutil.rmtree(tmp, ignore_errors=True)
        return 1

    result = fill_familia_template(
        template,
        output,
        {
            "diagnostic": "Diagnóstico de prueba con texto suficiente.",
            "evaluation_type": "admission",
            "evaluation": "1",
            "reevaluation": "",
            "primary": "",
            "substitute": "",
            "yes": "",
            "no": "",
        },
        student_context={"student_fullname": "Isabella Test", "identification_number": "23.442.145-K"},
    )
    assert result.get("status") == "success", result

    after = _count_symbol(output, FAMILIA_CHECKBOX_UNCHECKED)
    checked = _count_symbol(output, FAMILIA_CHECKBOX_CHECKED_MARK)
    print(f"checkbox_empty antes={before} despues={after}")
    print(f"checkbox_checked simbolo={checked}")

    if after < 1:
        print("FALLO: desaparecieron checkboxes vacios")
        shutil.rmtree(tmp, ignore_errors=True)
        return 1
    if checked < 1:
        print("FALLO: ingreso marcado sin simbolo dentro del recuadro")
        shutil.rmtree(tmp, ignore_errors=True)
        return 1

    # Postproceso aislado
    iso = tmp / "iso.docx"
    shutil.copy2(template, iso)
    apply_familia_checkbox_states(iso, {})
    ensure_familia_checkbox_boxes_visible(iso)
    iso_count = _count_symbol(iso, FAMILIA_CHECKBOX_UNCHECKED)
    print(f"checkbox_empty postproceso_aislado={iso_count}")
    if iso_count < before:
        print("FALLO: postproceso borra ☐")
        shutil.rmtree(tmp, ignore_errors=True)
        return 1

    print("OK: checkboxes vacios conservan cuadro")
    shutil.rmtree(tmp, ignore_errors=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
